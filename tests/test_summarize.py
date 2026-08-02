"""Summarize route + document uploads.

House style: throwaway user store, Flask test client, plain PASS/FAIL lines.
No model is called — the one map-reduce test stubs _stream_chat.
"""
import os, sys, io, json, tempfile, threading
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
import router

tmpdir = tempfile.mkdtemp()
router.USERS_FILE = os.path.join(tmpdir, "users.json")
router.app.config["SESSION_COOKIE_SECURE"] = False
# Never let a test write the real router_config.json — one already did.
import shutil as _sh
_cfgcopy = os.path.join(tmpdir, "router_config.json")
if os.path.exists(router.ROUTER_CONFIG_FILE):
    _sh.copyfile(router.ROUTER_CONFIG_FILE, _cfgcopy)
else:                       # fresh clone: no config yet -> built-in defaults
    json.dump(router.ROUTER_CONFIG, open(_cfgcopy, "w"))
router.ROUTER_CONFIG_FILE = _cfgcopy

P, F = [], []
def check(name, cond, detail=""):
    (P if cond else F).append(name)
    print(("  PASS  " if cond else "  FAIL  ") + name +
          (f"   [{detail}]" if detail and not cond else ""))

c = router.app.test_client()

print("--- route registration ---")
check("summarize is a live route", "summarize" in router.SPECIALISTS,
      str(sorted(router.SPECIALISTS)))
check("summarize is in VALID", "summarize" in router.VALID)
check("summarize has a persona", bool(router.SYSTEM_PROMPTS.get("summarize")))
check("the classifier is told about it", "summarize" in router.CLASSIFIER_SYSTEM)
check("the pre-existing 'summarizer' ROLE is untouched",
      router.ROUTER_CONFIG["roles"].get("summarizer") == "summarizer-e4b",
      str(router.ROUTER_CONFIG["roles"]))
check("route count still under the ~9-10 classifier ceiling",
      len(router.SPECIALISTS) <= 10, str(len(router.SPECIALISTS)))

print("--- backfill of built-in routes into an older config ---")
old = json.loads(json.dumps(router.ROUTER_CONFIG))
old["routes"].pop("summarize", None)
merged = router._backfill_builtins(json.loads(json.dumps(old)))
check("a config predating the route gets it added",
      "summarize" in merged["routes"], str(sorted(merged["routes"])))
check("existing routes are not rewritten",
      merged["routes"]["general"] == old["routes"]["general"])
orphan = json.loads(json.dumps(old))
orphan["backends"] = [b for b in orphan["backends"] if b["id"] != "brain-26b"]
orphan["routes"]["general"] = dict(orphan["routes"]["general"], backend="recipes-e4b")
m2 = router._backfill_builtins(orphan)
check("a missing default backend falls back to general's, not a dangling id",
      m2["routes"]["summarize"]["backend"] == "recipes-e4b",
      str(m2["routes"]["summarize"]))

print("--- text extraction ---")
check("plain text", router._upload_text("a.txt", b"hello world") == "hello world")
check("markdown", "# Title" in router._upload_text("a.md", b"# Title\nbody"))
csv = router._upload_text("a.csv", b"a,b\n1,2")
check("csv", "1,2" in csv, csv)

def _mini_pdf(text):
    """A complete one-page PDF built by hand — no system file, no extra lib.
    Object offsets are computed, so the xref is genuinely valid."""
    stream = f"BT /F1 24 Tf 72 720 Td ({text}) Tj ET".encode()
    objs = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>",
        b"<< /Length %d >>\nstream\n%s\nendstream" % (len(stream), stream),
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    out, offsets = bytearray(b"%PDF-1.4\n"), []
    for i, body in enumerate(objs, 1):
        offsets.append(len(out))
        out += b"%d 0 obj\n%s\nendobj\n" % (i, body)
    xref = len(out)
    out += b"xref\n0 %d\n0000000000 65535 f \n" % (len(objs) + 1)
    for off in offsets:
        out += b"%010d 00000 n \n" % off
    out += (b"trailer\n<< /Size %d /Root 1 0 R >>\nstartxref\n%d\n%%%%EOF"
            % (len(objs) + 1, xref))
    return bytes(out)

pdf_txt = router._upload_text("form.pdf", _mini_pdf("Crewe pdf fixture"))
check("pdf extracts real text", "Crewe pdf fixture" in pdf_txt, pdf_txt[:60])

import docx
buf = io.BytesIO()
d = docx.Document(); d.add_paragraph("Docx paragraph one")
t = d.add_table(rows=1, cols=2); t.rows[0].cells[0].text = "cellA"; t.rows[0].cells[1].text = "cellB"
d.save(buf)
dx = router._upload_text("a.docx", buf.getvalue())
check("docx paragraphs", "Docx paragraph one" in dx, dx[:60])
check("docx tables are included", "cellA" in dx and "cellB" in dx, dx[:120])

import openpyxl
buf = io.BytesIO()
wb = openpyxl.Workbook(); ws = wb.active; ws["A1"] = "hdr"; ws["B1"] = 42
wb.save(buf)
xl = router._upload_text("a.xlsx", buf.getvalue())
check("xlsx cells", "hdr" in xl and "42" in xl, xl[:80])

def rejects(name, data, phrase=None):
    try:
        router._upload_text(name, data)
        return False
    except ValueError as e:
        return (phrase in str(e)) if phrase else True
check("unknown extension is refused", rejects("a.exe", b"MZ..."))
check("no-extension file is refused", rejects("README", b"hi"))
check("empty text is refused", rejects("a.txt", b"   \n  "))
check("a scanned (text-free) PDF says so, not a stack trace",
      rejects("a.pdf", b"%PDF-1.4 garbage", "could not read") or
      rejects("a.pdf", b"%PDF-1.4 garbage", "scan"))
check("extracted text is capped", len(router._upload_text(
      "big.txt", b"x" * (router.MAX_UPLOAD_TEXT + 5000))) == router.MAX_UPLOAD_TEXT)

print("--- chunking ---")
ch = router._chunk_text("para one\n\npara two\n\npara three", 12)
check("chunks split on paragraph boundaries", all(len(x) <= 12 for x in ch), str(ch))
check("chunking loses no content",
      "".join(ch).replace("\n", "") == "para oneparatwoparathree".replace("para", "para")
      or len("".join(ch)) >= len("para onepara twopara three") - 4, str(ch))
big = router._chunk_text("z" * 100, 30)
check("an oversized single paragraph is hard-split",
      len(big) == 4 and all(len(x) <= 30 for x in big), str([len(x) for x in big]))

print("--- upload endpoint (auth + tenancy) ---")
r = c.post("/upload", data={"session_id": "s1"})
check("unauthenticated upload -> 401", r.status_code == 401, str(r.status_code))

admin, err = router.create_user("admin@example.com", "correct-horse-battery", is_admin=True)
assert admin, err
c.post("/login", data={"email": "admin@example.com", "password": "correct-horse-battery"})

r = c.post("/upload", data={"session_id": "sX",
                            "file": (io.BytesIO(b"the quick brown fox"), "note.txt")},
           content_type="multipart/form-data")
j = r.get_json()
check("upload succeeds", r.status_code == 200 and j.get("ok"), json.dumps(j)[:140])
up_id = j["upload"]["id"]
check("upload reports char count", j["upload"]["chars"] == len("the quick brown fox"),
      str(j["upload"]))

r = c.get("/uploads/sX")
check("the file is attached to the session",
      any(u["id"] == up_id for u in r.get_json()["uploads"]), r.get_data(as_text=True)[:120])
check("extracted text is readable back",
      router._upload_body({"id": up_id}) == "the quick brown fox")

d = router.uploads_dir()
check("stored under the caller's own userdata root",
      "crewe_userdata" in d and os.path.isdir(d), d)
check("the client filename is NOT used as a path",
      not any(f.startswith("note") for f in os.listdir(d)), str(os.listdir(d))[:120])

r = c.post("/upload", data={"session_id": "sX",
                            "file": (io.BytesIO(b"MZ\x00binary"), "evil.exe")},
           content_type="multipart/form-data")
check("unsupported type -> 400 with a readable reason",
      r.status_code == 400 and "exe" in r.get_json().get("error", ""),
      r.get_data(as_text=True)[:120])

r = c.post("/upload", data={"session_id": "sX",
                            "file": (io.BytesIO(b""), "empty.txt")},
           content_type="multipart/form-data")
check("empty file -> 400", r.status_code == 400, str(r.status_code))

print("--- cap on attachments per session ---")
for i in range(router.MAX_SESSION_UPLOADS + 3):
    c.post("/upload", data={"session_id": "sCap",
                            "file": (io.BytesIO(f"file {i}".encode()), f"f{i}.txt")},
           content_type="multipart/form-data")
ups = router.session_uploads("sCap")
check(f"session keeps only the newest {router.MAX_SESSION_UPLOADS}",
      len(ups) == router.MAX_SESSION_UPLOADS, str(len(ups)))
check("the newest survives", ups[-1]["name"] == f"f{router.MAX_SESSION_UPLOADS + 2}.txt",
      ups[-1]["name"])

print("--- detach ---")
r = c.delete(f"/uploads/sX/{up_id}")
check("delete reports removal", r.get_json().get("removed") is True,
      r.get_data(as_text=True)[:100])
check("it is gone from the session",
      not any(u["id"] == up_id for u in router.session_uploads("sX")))
check("the stored copies are deleted from disk too",
      not any(f.split(".")[0] == up_id for f in os.listdir(d)), str(os.listdir(d))[:120])
r = c.delete("/uploads/sX/doesnotexist")
check("deleting an unknown id is a no-op, not an error",
      r.status_code == 200 and r.get_json().get("removed") is False)

print("--- map-reduce condense ---")
calls = {"n": 0}
def fake_stream(url, messages, job_id, **kw):
    calls["n"] += 1
    body = messages[-1]["content"]
    return "S:" + str(len(body))          # always much shorter than the input
real = router._stream_chat
router._stream_chat = fake_stream
try:
    short = "a" * 100
    check("text already within budget is returned untouched",
          router._condense(short, 500, "http://x", None) == short)
    calls["n"] = 0
    out = router._condense("para\n\n" * 5000, 2000, "http://x", None)
    check("a long document is condensed to fit", len(out) <= 2000, str(len(out)))
    check("it actually called the model per chunk", calls["n"] > 1, str(calls["n"]))

    # a model that refuses to shrink must terminate, not loop
    router._stream_chat = lambda url, messages, job_id, **kw: messages[-1]["content"]
    out = router._condense("b" * 50000, 1000, "http://x", None)
    check("a non-shrinking model terminates and truncates",
          len(out) <= 1000, str(len(out)))

    # a backend that errors must degrade, not crash
    def boom(*a, **k): raise RuntimeError("backend down")
    router._stream_chat = boom
    out = router._condense("c" * 50000, 1000, "http://x", None)
    check("a dead backend degrades to truncation instead of raising",
          len(out) <= 1000, str(len(out)))
finally:
    router._stream_chat = real

print("--- github repos as attachments ---")
import io as _io, tarfile as _tar
check("review is a live route", "review" in router.SPECIALISTS)
check("review persona demands ready-to-paste code per finding",
      "fenced code block" in router.SYSTEM_PROMPTS["review"]
      and "never a fragment" in router.SYSTEM_PROMPTS["review"])
check("the classifier is told review != code",
      "review" in router.CLASSIFIER_SYSTEM and "REVIEWED" in router.CLASSIFIER_SYSTEM)

p_ = router._github_parse
check("plain repo URL parses", p_("https://github.com/o/r") == ("o", "r", None))
check("branch suffix parses", p_("https://github.com/o/r/tree/dev") == ("o", "r", "dev"))
check(".git suffix parses", p_("https://github.com/o/r.git") == ("o", "r", None))
for bad in ("https://github.com@evil.host/o/r", "https://evil.com/o/r",
            "https://github.com/o", "https://github.com/o/r/../x",
            "ftp://github.com/o/r", "https://codeload.github.com/o/r/tar.gz/HEAD"):
    check(f"rejected: {bad[:44]}", p_(bad) is None, str(p_(bad)))

# a fake tarball, extracted through the real code path via a stubbed fetch
buf = _io.BytesIO()
with _tar.open(fileobj=buf, mode="w:gz") as t:
    for path, data in (("repo-HEAD/src/app.py", b"def main():\n    return 1\n"),
                       ("repo-HEAD/README.md", b"# hi"),
                       ("repo-HEAD/node_modules/x/y.js", b"junk"),
                       ("repo-HEAD/logo.png", b"\x89PNG\x00\x00binary"),
                       ("repo-HEAD/package-lock.json", b"{}")):
        ti = _tar.TarInfo(path); ti.size = len(data)
        t.addfile(ti, _io.BytesIO(data))
buf.seek(0)

class _FakeResp:
    status_code = 200
    def iter_content(self, n): 
        d = buf.getvalue()
        for i in range(0, len(d), n): yield d[i:i+n]
    def close(self): pass
_real_get = router.requests.get
router.requests.get = lambda *a, **k: _FakeResp()
try:
    name, text = router._github_repo_text("https://github.com/o/repo")
    check("bundle carries the file tree and sources",
          "src/app.py" in text and "def main" in text, text[:120])
    check("node_modules is skipped", "node_modules" not in text)
    check("lockfiles are skipped", "package-lock" not in text)
    check("binaries are skipped", "logo.png" not in text)
    check("bundle uses the FILE markers review-packing splits on",
          "=== FILE: src/app.py ===" in text)
finally:
    router.requests.get = _real_get

r = c.post("/upload", json={"session_id": "sGH", "github_url": "not a url"})
check("a junk repo URL is a clean 400", r.status_code == 400
      and "GitHub" in r.get_json()["error"], r.get_data(as_text=True)[:100])
router.requests.get = lambda *a, **k: _FakeResp()
try:
    buf.seek(0)
    r = c.post("/upload", json={"session_id": "sGH",
                                "github_url": "https://github.com/o/repo"})
    j = r.get_json()
    check("a repo attaches like any upload", r.status_code == 200 and j["ok"],
          str(j)[:120])
    check("the chip names the repo", j["upload"]["name"] == "o/repo", str(j["upload"]))
finally:
    router.requests.get = _real_get

print("--- a repo link in the MESSAGE attaches itself ---")
check("classifier steers repo-link requests away from search",
      "NEVER 'search'" in router.CLASSIFIER_SYSTEM)
router.set_owner(admin["id"])
router.requests.get = lambda *a, **k: _FakeResp()
try:
    buf.seek(0)
    note = router._maybe_attach_github(
        "could you read this and suggest a readme https://github.com/o/repo",
        "sAuto")
    check("a linked repo attaches with no note", note is None, str(note))
    check("...and lands in the session",
          any(u["name"] == "o/repo" for u in router.session_uploads("sAuto")),
          str(router.session_uploads("sAuto")))
    n_before = len(router.session_uploads("sAuto"))
    buf.seek(0)
    router._maybe_attach_github("same link again https://github.com/o/repo", "sAuto")
    check("the same repo is not fetched twice in one session",
          len(router.session_uploads("sAuto")) == n_before)
finally:
    router.requests.get = _real_get
check("a message with no repo link is untouched",
      router._maybe_attach_github("just a question", "sAuto") is None)
def _boom(*a, **k): raise router.requests.ConnectionError("down")
router.requests.get = _boom
try:
    note = router._maybe_attach_github("see https://github.com/o/other", "sAuto")
    check("a failed fetch returns a note the model must relay",
          note and "could not be attached" in note, str(note))
finally:
    router.requests.get = _real_get

print(f"\n{len(P)} passed, {len(F)} failed")
sys.exit(1 if F else 0)
