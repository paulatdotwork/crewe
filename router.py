#!/usr/bin/env python3
#
# Crewe — a self-hosted junction for your local LLMs.
# Copyright (C) 2026  Crewe contributors
#
# This program is free software: you can redistribute it and/or modify it under
# the terms of the GNU Affero General Public License as published by the Free
# Software Foundation, either version 3 of the License, or (at your option) any
# later version.
#
# This program is distributed in the hope that it will be useful, but WITHOUT
# ANY WARRANTY; without even the implied warranty of MERCHANTABILITY or FITNESS
# FOR A PARTICULAR PURPOSE.  See the GNU Affero General Public License for more
# details.  You should have received a copy of the GNU Affero General Public
# License along with this program.  If not, see <https://www.gnu.org/licenses/>.
#
# AGPL SECTION 13 — READ THIS BEFORE YOU DEPLOY A MODIFIED COPY.
# Crewe is served over a network, so if you modify it and let anyone else use
# it, you must offer those users the source of YOUR version. The "source" link
# in the page header is how Crewe does that: set CREWE_SOURCE_URL to point at
# your fork. Leaving it pointing at upstream while running modified code does
# not satisfy the licence.
"""
router.py — a small Flask app that:
  1. serves a single web page (chat box)
  2. sends each question to the E2B classifier (port 8083) -> one of recipes/creative/code
  3. forwards the question to the chosen specialist (8080/8081/8082)
  4. returns the answer + which specialist handled it (shown as a badge)

While the specialist works, the page shows a spinner and a live token counter
(polled from a lightweight progress endpoint) so you know it's actually working.

Run:  python3 router.py
Then open http://localhost:5000
"""

import array
import base64
import difflib
import functools
import io
import ipaddress
import json
import os
import queue
import re
import shutil
import socket
import subprocess
import tempfile
import threading
import time
import uuid
import wave
from http.server import ThreadingHTTPServer, SimpleHTTPRequestHandler
from urllib.parse import urlparse
import requests
from flask import Flask, request, jsonify, Response

app = Flask(__name__)

# ==== accounts / auth (Phase 1, 2026-07-20) ==================================
# Authentication + a DEFAULT-DENY gate. Deliberately NOT Cloudflare Access:
# Crewe is meant to be a hosted product with its own login page, and this is the
# same system that later flips on public signup (Phase 3).
#
# ⚠ SCOPE: this is auth only. Multi-tenancy is Phase 2 — every logged-in user
# still shares ONE data pool (docs, sheets, cookbook, chat history are global).
# Until Phase 2 lands, only invite people who may see each other's data.
import secrets as _secrets
from datetime import datetime as _dt
from functools import wraps as _wraps
from flask import session as flask_session, redirect  # aliased: router.py uses
                                                      # `session_id` everywhere
from werkzeug.security import generate_password_hash, check_password_hash

USERS_FILE  = os.path.expanduser("~/crewe_users.json")
SECRET_FILE = os.path.expanduser("~/crewe_secret")
USERS_LOCK  = threading.Lock()

def _load_secret_key():
    """Persist the session key. If this regenerates, every user is logged out —
    which matters here because the router gets restarted often."""
    try:
        with open(SECRET_FILE, "rb") as f:
            k = f.read().strip()
            if len(k) >= 32:
                return k
    except FileNotFoundError:
        pass
    k = _secrets.token_bytes(48)
    fd = os.open(SECRET_FILE, os.O_WRONLY | os.O_CREAT | os.O_TRUNC, 0o600)
    with os.fdopen(fd, "wb") as f:
        f.write(k)
    return k

app.secret_key = _load_secret_key()
# Secure cookies are correct in production (served over HTTPS via the tunnel) but
# block login over plain http://localhost. Set CREWE_DEV=1 for local testing.
_DEV = os.environ.get("CREWE_DEV") == "1"
app.config.update(
    SESSION_COOKIE_HTTPONLY=True,
    SESSION_COOKIE_SAMESITE="Lax",
    SESSION_COOKIE_SECURE=not _DEV,
    PERMANENT_SESSION_LIFETIME=60 * 60 * 24 * 30,   # 30 days
    # Cap request bodies so a giant POST can't exhaust memory. 8 MB covers the
    # largest real payload (code pipeline seed_code / doc writes); override with
    # CREWE_MAX_BODY (bytes). Flask returns 413 past this.
    MAX_CONTENT_LENGTH=int(os.environ.get("CREWE_MAX_BODY", 8 * 1024 * 1024)),
)

def _load_users():
    with USERS_LOCK:
        try:
            with open(USERS_FILE) as f:
                return json.load(f).get("users", [])
        except (FileNotFoundError, json.JSONDecodeError):
            return []

def _save_users(users):
    with USERS_LOCK:
        tmp = USERS_FILE + ".tmp"
        with open(tmp, "w") as f:
            json.dump({"users": users}, f, indent=2)
        os.replace(tmp, USERS_FILE)          # atomic; no half-written user file

def _user_by_email(email):
    email = (email or "").strip().lower()
    return next((u for u in _load_users() if u.get("email") == email), None)

def create_user(email, password, is_admin=False):
    """Returns (user, None) or (None, error). Used by create_user.py."""
    email = (email or "").strip().lower()
    if not email or "@" not in email:
        return None, "invalid email"
    if len(password or "") < 10:
        return None, "password must be at least 10 characters"
    users = _load_users()
    if any(u.get("email") == email for u in users):
        return None, "that email already has an account"
    u = {
        "id": _secrets.token_hex(8),
        "email": email,
        "pw_hash": generate_password_hash(password),
        "is_admin": bool(is_admin),
        "created": _dt.now().isoformat(timespec="seconds"),
    }
    users.append(u)
    _save_users(users)
    return u, None

def current_user():
    uid = flask_session.get("uid")
    if not uid:
        return None
    return next((u for u in _load_users() if u.get("id") == uid), None)

def set_password(uid, new_password):
    """Returns (True, None) or (False, error)."""
    if len(new_password or "") < 10:
        return False, "password must be at least 10 characters"
    with USERS_LOCK:
        # re-read inside the lock so we don't clobber a concurrent write
        try:
            with open(USERS_FILE) as f:
                users = json.load(f).get("users", [])
        except (FileNotFoundError, json.JSONDecodeError):
            users = []
        u = next((x for x in users if x.get("id") == uid), None)
        if not u:
            return False, "user not found"
        u["pw_hash"] = generate_password_hash(new_password)
        tmp = USERS_FILE + ".tmp"
        with open(tmp, "w") as f:
            json.dump({"users": users}, f, indent=2)
        os.replace(tmp, USERS_FILE)
    return True, None

def delete_user(uid):
    """Remove an account. Does NOT delete their data dir (kept for recovery)."""
    users = [u for u in _load_users() if u.get("id") != uid]
    _save_users(users)

# --- invites: the one sanctioned path to a new account before public signup.
# An admin mints a single-use, expiring token; the invitee sets their own
# password at /invite/<token>. This is also the seed of Phase 3: flipping on
# public signup is "let anyone reach the acceptance page without a token".
INVITES_FILE = os.path.expanduser("~/crewe_invites.json")
INVITES_LOCK = threading.Lock()
INVITE_TTL_S = 7 * 24 * 3600

def _load_invites():
    with INVITES_LOCK:
        try:
            with open(INVITES_FILE) as f:
                return json.load(f).get("invites", [])
        except (FileNotFoundError, json.JSONDecodeError):
            return []

def _save_invites(invites):
    with INVITES_LOCK:
        tmp = INVITES_FILE + ".tmp"
        with open(tmp, "w") as f:
            json.dump({"invites": invites}, f, indent=2)
        os.replace(tmp, INVITES_FILE)

def create_invite(email, is_admin, invited_by):
    email = (email or "").strip().lower()
    if not email or "@" not in email:
        return None, "invalid email"
    if _user_by_email(email):
        return None, "that email already has an account"
    invites = _load_invites()
    # supersede any outstanding invite for the same email
    invites = [i for i in invites if i.get("email") != email or i.get("used")]
    inv = {
        "token": _secrets.token_urlsafe(32),
        "email": email,
        "is_admin": bool(is_admin),
        "invited_by": invited_by,
        "created": _dt.now().isoformat(timespec="seconds"),
        "expires": time.time() + INVITE_TTL_S,
        "used": False,
    }
    invites.append(inv)
    _save_invites(invites)
    return inv, None

def _invite_by_token(token):
    inv = next((i for i in _load_invites() if i.get("token") == token), None)
    if not inv or inv.get("used"):
        return None
    if time.time() > inv.get("expires", 0):
        return None
    return inv

def _consume_invite(token):
    invites = _load_invites()
    for i in invites:
        if i.get("token") == token:
            i["used"] = True
            i["used_at"] = _dt.now().isoformat(timespec="seconds")
    _save_invites(invites)

# --- login throttle: in-memory, per-IP. Resets on restart, which is fine for a
# --- small beta; swap for a persistent store if this ever gets real traffic.
_LOGIN_FAILS, _LOGIN_LOCK = {}, threading.Lock()
_MAX_FAILS, _LOCKOUT_S = 8, 300

def _throttled(ip):
    with _LOGIN_LOCK:
        fails, until = _LOGIN_FAILS.get(ip, (0, 0))
        if until and time.time() < until:
            return int(until - time.time())
        return 0

def _note_fail(ip):
    with _LOGIN_LOCK:
        fails, _ = _LOGIN_FAILS.get(ip, (0, 0))
        fails += 1
        _LOGIN_FAILS[ip] = (fails, time.time() + _LOCKOUT_S if fails >= _MAX_FAILS else 0)

def _clear_fails(ip):
    with _LOGIN_LOCK:
        _LOGIN_FAILS.pop(ip, None)

# --- the gate -----------------------------------------------------------------
# Default-deny by endpoint allowlist rather than @login_required on 50 routes,
# so a route added later is protected on the day it's written. Adding a new
# PUBLIC route is the thing that now requires a deliberate edit here.
# 'invite_accept' is the ONLY public write path — it is token-gated (256-bit,
# single-use, expiring), which is what makes that safe. Do not add to this set
# without the same care.
PUBLIC_ENDPOINTS = {"login", "logout", "static", "healthz", "invite_accept"}
ADMIN_PREFIXES   = ("/settings", "/admin", "/agent")   # settings/users + the code agent

@app.before_request
def _require_login():
    # Clear first, unconditionally. Flask reuses worker threads, so a value left
    # over from a previous request on this thread would be a cross-user leak.
    set_owner(None)
    if request.endpoint in PUBLIC_ENDPOINTS:
        return None
    u = current_user()
    if not u:
        # Page loads get a redirect; XHR gets JSON, so the frontend shows an
        # error instead of silently rendering a login page into a <div>.
        wants_html = "text/html" in (request.headers.get("Accept") or "")
        if request.method == "GET" and wants_html:
            return redirect("/login")
        return jsonify({"error": "not authenticated"}), 401
    if request.path.startswith(ADMIN_PREFIXES) and not u.get("is_admin"):
        return jsonify({"error": "admin only"}), 403
    set_owner(u["id"])
    return None

@app.route("/healthz")
def healthz():
    return jsonify({"ok": True})

@app.route("/whoami")
def whoami():
    u = current_user()   # gate guarantees this is set for a protected route
    return jsonify({"email": u["email"], "is_admin": bool(u.get("is_admin"))})

LOGIN_PAGE = """<!DOCTYPE html>
<html lang="en"><head><meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>Sign in &middot; Crewe</title>
<link rel="icon" href="data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 64 64'%3E%3Crect width='64' height='64' rx='14' fill='%238b4a18'/%3E%3Cpath d='M32 57V34M32 34 16 13M32 34 48 13' stroke='%23faf6f0' stroke-width='7' stroke-linecap='round' fill='none'/%3E%3Ccircle cx='32' cy='34' r='6' fill='%23faf6f0'/%3E%3C/svg%3E">
<link href="/static/fonts.css" rel="stylesheet">
<script>try{var t=localStorage.getItem('creweTheme');if(t)document.documentElement.setAttribute('data-theme',t)}catch(e){}</script>
<style>
:root{--bg:#faf6f0;--card:#fff;--ink:#2b211a;--muted:#7d6b5d;--line:#e4d9cc;--gold:#8b4a18;--btn-ink:#faf6f0;--err:#a3301c}
[data-theme=dark]{--bg:#1e1815;--card:#2a2220;--ink:#f0e7dc;--muted:#a89684;--line:#3d322c;--gold:#c97a3d;--btn-ink:#1e1815;--err:#e08a76}
*{box-sizing:border-box}
body{margin:0;min-height:100vh;display:flex;align-items:center;justify-content:center;
     background:var(--bg);color:var(--ink);font-family:'DM Sans',system-ui,sans-serif;padding:24px}
.card{background:var(--card);border:1px solid var(--line);border-radius:16px;padding:38px 34px;
      width:100%;max-width:390px;box-shadow:0 10px 34px rgba(0,0,0,.09)}
.brand{display:flex;align-items:center;gap:11px;margin-bottom:6px}
.brand svg{width:38px;height:38px;border-radius:9px;flex:none}
h1{font-family:'Lora',Georgia,serif;font-size:27px;margin:0;font-weight:600}
.sub{color:var(--muted);font-size:14px;margin:0 0 26px}
label{display:block;font-size:13px;font-weight:500;margin:15px 0 6px}
input{width:100%;padding:11px 13px;border:1px solid var(--line);border-radius:9px;
      background:var(--bg);color:var(--ink);font-size:15px;font-family:inherit}
input:focus{outline:2px solid var(--gold);outline-offset:-1px;border-color:transparent}
button{width:100%;margin-top:22px;padding:12px;border:0;border-radius:9px;background:var(--gold);
       color:var(--btn-ink);font-size:15px;font-weight:600;font-family:inherit;cursor:pointer}
button:hover{filter:brightness(1.07)}
.err{background:color-mix(in srgb,var(--err) 12%,transparent);border:1px solid var(--err);
     color:var(--err);padding:10px 13px;border-radius:9px;font-size:14px;margin-bottom:18px}
.foot{margin-top:26px;padding-top:19px;border-top:1px solid var(--line);
      text-align:center;font-size:13.5px;color:var(--muted)}
.soon{display:inline-block;margin-left:6px;padding:2px 8px;border-radius:99px;
      background:var(--line);color:var(--muted);font-size:11.5px;font-weight:600;letter-spacing:.03em}
</style></head><body>
<div class="card">
  <div class="brand">
    <svg viewBox="0 0 64 64" xmlns="http://www.w3.org/2000/svg"><rect width="64" height="64" rx="14" fill="#8b4a18"/><path d="M32 57V34M32 34 16 13M32 34 48 13" stroke="#faf6f0" stroke-width="7" stroke-linecap="round" fill="none"/><circle cx="32" cy="34" r="6" fill="#faf6f0"/></svg>
    <h1>Crewe</h1>
  </div>
  <p class="sub">Sign in to continue.</p>
  __ERROR__
  <form method="POST" action="/login">
    <label for="email">Email</label>
    <input id="email" name="email" type="email" autocomplete="username" required autofocus>
    <label for="password">Password</label>
    <input id="password" name="password" type="password" autocomplete="current-password" required>
    <button type="submit">Sign in</button>
  </form>
  <div class="foot">Don't have an account?<span class="soon">COMING SOON</span></div>
</div></body></html>"""

# Computed once at import so the unknown-email path costs the same as a real one.
_DUMMY_HASH = generate_password_hash("not-a-real-password-" + _secrets.token_hex(8))

def _login_page(error=""):
    block = f'<div class="err">{error}</div>' if error else ""
    return LOGIN_PAGE.replace("__ERROR__", block)

@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "GET":
        return redirect("/") if current_user() else _login_page()
    ip = request.headers.get("Cf-Connecting-Ip") or request.remote_addr or "?"
    wait = _throttled(ip)
    if wait:
        return _login_page(f"Too many attempts. Try again in {wait}s."), 429
    email = (request.form.get("email") or "").strip().lower()
    pw = request.form.get("password") or ""
    u = _user_by_email(email)
    # Hash a dummy when the email is unknown, so a bad email and a bad password
    # cost the same time and don't reveal which accounts exist.
    try:
        ok = check_password_hash(u["pw_hash"] if u else _DUMMY_HASH, pw)
    except Exception:
        ok = False
    if not u or not ok:
        _note_fail(ip)
        return _login_page("Incorrect email or password."), 401
    _clear_fails(ip)
    flask_session.permanent = True
    flask_session["uid"] = u["id"]
    return redirect("/")

@app.route("/logout", methods=["GET", "POST"])
def logout():
    flask_session.clear()
    return redirect("/login")

# ---- account + admin pages ---------------------------------------------------
# Shared shell so /account, /admin and /invite match the login page's theme.
_PAGE_CSS = """
:root{--bg:#faf6f0;--card:#fff;--ink:#2b211a;--muted:#7d6b5d;--line:#e4d9cc;--gold:#8b4a18;--btn-ink:#faf6f0;--err:#a3301c;--ok:#3f7d3f}
[data-theme=dark]{--bg:#1e1815;--card:#2a2220;--ink:#f0e7dc;--muted:#a89684;--line:#3d322c;--gold:#c97a3d;--btn-ink:#1e1815;--err:#e08a76;--ok:#8fc98f}
*{box-sizing:border-box}
body{margin:0;min-height:100vh;background:var(--bg);color:var(--ink);
     font-family:'DM Sans',system-ui,sans-serif;padding:34px 18px}
.wrap{max-width:640px;margin:0 auto}
.top{display:flex;align-items:center;gap:11px;margin-bottom:26px}
.top svg{width:34px;height:34px;border-radius:8px;flex:none}
.top h1{font-family:'Lora',Georgia,serif;font-size:24px;margin:0;font-weight:600;flex:1}
.top a{color:var(--gold);text-decoration:none;font-size:14px;font-weight:500}
.card{background:var(--card);border:1px solid var(--line);border-radius:15px;padding:26px 24px;margin-bottom:20px}
.card h2{font-family:'Lora',Georgia,serif;font-size:18px;margin:0 0 4px}
.card .hint{color:var(--muted);font-size:13.5px;margin:0 0 18px}
label{display:block;font-size:13px;font-weight:500;margin:14px 0 6px}
input[type=text],input[type=email],input[type=password]{width:100%;padding:10px 12px;border:1px solid var(--line);
     border-radius:9px;background:var(--bg);color:var(--ink);font-size:15px;font-family:inherit}
input:focus{outline:2px solid var(--gold);outline-offset:-1px;border-color:transparent}
.row{display:flex;gap:9px;align-items:center;margin-top:14px}
.row label{margin:0;font-weight:400}
button,.btn{display:inline-block;margin-top:18px;padding:10px 18px;border:0;border-radius:9px;background:var(--gold);
       color:var(--btn-ink);font-size:14.5px;font-weight:600;font-family:inherit;cursor:pointer;text-decoration:none}
button:hover{filter:brightness(1.07)}
button.ghost,.btn.ghost{background:transparent;color:var(--gold);border:1px solid var(--line);margin:0;padding:6px 12px;font-size:13px}
.msg{padding:10px 13px;border-radius:9px;font-size:14px;margin-bottom:16px}
.msg.err{background:color-mix(in srgb,var(--err) 12%,transparent);border:1px solid var(--err);color:var(--err)}
.msg.ok{background:color-mix(in srgb,var(--ok) 14%,transparent);border:1px solid var(--ok);color:var(--ok)}
table{width:100%;border-collapse:collapse;font-size:14px;margin-top:6px}
th,td{text-align:left;padding:9px 8px;border-bottom:1px solid var(--line)}
th{color:var(--muted);font-weight:500;font-size:12.5px;text-transform:uppercase;letter-spacing:.03em}
.tag{display:inline-block;padding:1px 8px;border-radius:99px;background:var(--line);color:var(--muted);font-size:11.5px;font-weight:600}
.tag.admin{background:color-mix(in srgb,var(--gold) 22%,transparent);color:var(--gold)}
.linkbox{display:flex;gap:8px;margin-top:8px}
.linkbox input{font-family:ui-monospace,monospace;font-size:12.5px}
.muted{color:var(--muted);font-size:13px}
form.inline{display:inline;margin:0}
"""
_BRAND_SVG = ('<svg viewBox="0 0 64 64" xmlns="http://www.w3.org/2000/svg">'
    '<rect width="64" height="64" rx="14" fill="#8b4a18"/>'
    '<path d="M32 57V34M32 34 16 13M32 34 48 13" stroke="#faf6f0" stroke-width="7" stroke-linecap="round" fill="none"/>'
    '<circle cx="32" cy="34" r="6" fill="#faf6f0"/></svg>')

def _shell(title, body, back="/"):
    return f"""<!DOCTYPE html><html lang="en"><head><meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>{title} &middot; Crewe</title>
<link rel="icon" href="data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 64 64'%3E%3Crect width='64' height='64' rx='14' fill='%238b4a18'/%3E%3Cpath d='M32 57V34M32 34 16 13M32 34 48 13' stroke='%23faf6f0' stroke-width='7' stroke-linecap='round' fill='none'/%3E%3Ccircle cx='32' cy='34' r='6' fill='%23faf6f0'/%3E%3C/svg%3E">
<link href="/static/fonts.css" rel="stylesheet">
<script>try{{var t=localStorage.getItem('creweTheme');if(t)document.documentElement.setAttribute('data-theme',t)}}catch(e){{}}</script>
<style>{_PAGE_CSS}</style></head><body><div class="wrap">
<div class="top">{_BRAND_SVG}<h1>{title}</h1><a href="{back}">&larr; Back to Crewe</a></div>
{body}</div></body></html>"""

def _esc(s):
    return (str(s).replace("&", "&amp;").replace("<", "&lt;")
            .replace(">", "&gt;").replace('"', "&quot;"))

HELP_BODY = """
<div class="card">
<h2>How Crewe works</h2>
<p>Crewe does not have one model. It has <b>routes</b> &mdash; recipes, code,
search, and so on &mdash; and each route points at a <b>backend</b>, which is
any OpenAI-compatible server: llama.cpp, Ollama, LM Studio, vLLM, or a hosted
API. A small, fast <b>classifier</b> reads every question and picks the route.</p>
<p>So there are two separate questions to answer when you set Crewe up:
<i>what servers do I have?</i> (backends) and <i>what should each kind of
question go to?</i> (routes). You do both on the
<a href="/settings">&#9881; Settings</a> page, and changes apply live &mdash;
no restart.</p>
</div>

<div class="card">
<h2>1. Adding your model</h2>
<p>Open <a href="/settings">Settings</a> and either:</p>
<ul>
<li><b>Scan localhost</b> &mdash; finds servers already running on this machine
(it checks the usual ports, including Ollama's 11434), or</li>
<li><b>Add backend</b> and type the base URL yourself, e.g.
<code>http://192.168.1.50:8080</code> or <code>http://localhost:11434</code>.</li>
</ul>
<p>Give it a name you will recognise, then press <b>Check</b>. Crewe asks the
server what models it has; if that works, you are connected.</p>
<p class="note"><b>The "default model" field.</b> llama.cpp serves one model and
ignores this. Ollama and hosted APIs serve many and need to be told which
&mdash; put the exact name there, e.g. <code>qwen3:8b</code>. If Check lists
models but answers fail, this is almost always why.</p>
</div>

<div class="card">
<h2>2. Pointing routes at backends</h2>
<p>Under <b>Routes</b>, each route has a backend dropdown. Pick one per route.
Several routes can share a backend &mdash; that is normal, they just queue.</p>
<p>Below that, <b>system roles</b> are the jobs that are not routes:</p>
<ul>
<li><b>classifier</b> &mdash; picks the route for every question. Wants small and
fast; it runs on every single request.</li>
<li><b>summarizer</b> &mdash; condenses older conversation turns in the
background. Small and fast is fine.</li>
<li><b>easy coder</b> / <b>hard coder</b> &mdash; the two coders the Effort
switch chooses between.</li>
</ul>
</div>

<div class="card">
<h2>3. Effort: three coders</h2>
<p>Coding questions are the expensive ones, so they get a switch. The
<b>Effort</b> dropdown next to the message box picks between your
<b>fast</b>, <b>normal</b> and <b>extra</b> coders — quickest, everyday, and
strongest. It is ignored for every other kind of question. A tier you have
not wired up simply falls back to the one above it.</p>
<p>Effort changes more than the model. Crewe asks each coder how big its
context window is and sizes the prompt budgets to match, so a small model is
never handed a prompt it cannot hold &mdash; the classic way a capable little
model is made to look stupid.</p>
<p class="note">Whichever level you pick, planning, writing and checking all
happen on <b>one</b> model. Mixing two models inside a single build produces
worse results than either one alone.</p>
</div>

<div class="card">
<h2>4. Custom routes</h2>
<p><b>Add route</b> creates your own. Three fields matter, and they do
completely different jobs:</p>
<ul>
<li><b>subject</b> &mdash; the <i>only</i> thing the classifier ever sees about
your route. It decides whether a question comes here. Describe the territory
plainly: <i>"D&amp;D 5e character builds, combat encounters, monster statistics,
and rules clarifications"</i>.</li>
<li><b>persona</b> &mdash; the system prompt used to <i>answer</i>, once the
question has arrived. It has no effect on routing.</li>
<li><b>backend</b> &mdash; which server answers.</li>
</ul>
<p>A brilliant persona with a vague subject means the route never fires. A sharp
subject with no persona means it fires and then answers like a generic
assistant.</p>
<p>Two buttons help: <b>&#10024; draft</b> writes a subject and persona from a
one-line description, and <b>&#129514; test</b> runs your <i>real</i>
classifier over example questions with the candidate route in place, so you can
see what it actually decides before you commit.</p>
<p class="note"><b>Test a counter-example, not just the obvious hits.</b> The
classifier is a small model and its accuracy degrades past roughly nine or ten
routes. The useful test for a D&amp;D route is not "build me a rogue" &mdash;
it is "beef stew for my character's tavern", which should still go to
recipes.</p>
</div>

<div class="card">
<h2>5. Paid models, and choosing when to spend</h2>
<p>A backend can be a hosted API instead of your own hardware. Add it like any
other: base URL, plus an <b>API key</b>, plus the model name. Anything
OpenAI-compatible works &mdash; OpenAI, OpenRouter, DeepSeek, Groq, Together.
OpenRouter is the easiest single key if you want to reach many providers'
models at once.</p>
<p><b>Pick the right type when you add it.</b> llama.cpp and Ollama speak the
same wire format as OpenAI, so they differ only in small details (Ollama needs
a model name; llama.cpp doesn't). <b>Claude is genuinely different</b> &mdash;
a different endpoint, different auth headers, and a different request and
response shape &mdash; so select <b>Claude (Anthropic)</b> and Crewe translates
for you. Use <code>https://api.anthropic.com</code> as the base URL, and name a
model such as <code>claude-opus-5</code>. One quirk worth knowing: Claude
ignores <code>temperature</code>, so steer it with the route's persona instead.</p>
<p>In every case the base URL has <b>no <code>/v1</code></b> &mdash; Crewe
appends the right path for the type you picked. A doubled <code>/v1/v1/</code>
is the most common reason a new backend 404s.</p>
<p>Tick <b>this backend costs money</b> and enter the per-million-token prices
from your provider's pricing page. Then Crewe can tell you what you are
spending, and the Effort dropdown will mark which option bills you.</p>
<p>No tier means "paid" &mdash; each of the three can be local or hosted,
whichever suits you. A common arrangement is a tiny local model on fast, a
solid local model on normal, and a paid API on extra, so nothing bills unless
you choose it per question. Crewe marks any tier that costs money with a
&#128181; in the dropdown.</p>
<p class="note"><b>Spend figures are estimates.</b> They use the prices you
typed and the token counts your provider reported &mdash; or a rough
four-characters-per-token guess if it reported none. Useful as a running
indication. Never treat them as a bill; your provider's dashboard is the
truth.</p>
</div>

<div class="card">
<h2>6. Comparing models &mdash; the Compare page</h2>
<p>One page puts answers side by side, with a <b>models / routes</b> switch that
changes what is being compared:</p>
<ul>
<li><b>models</b> gives every entrant the <i>same</i> system prompt, so the only
thing that differs is the model. Use it for "is this model better than that
one?" &mdash; and note it is the only way to test a backend that no route points
at, such as one wired only to a coder role.</li>
<li><b>routes</b> gives each entrant its own persona, so you are judging your
configured setups rather than raw models. That is closer to what your users
actually experience.</li>
</ul>
<p>Tick two or more, ask once, and every entrant answers the same prompt side by
side. Your selection is remembered per browser, separately for each mode.</p>
<p class="note"><b>Every entrant needs a working backend.</b> They run
simultaneously and the judge only starts once they have <i>all</i> finished, so
one dead or very slow backend holds up the whole comparison. If it seems to
hang, check the route health dots in the header &mdash; and remember a model on
CPU can legitimately take a minute or two for a long answer.</p>
<p>Both pages finish with a judge, which runs on your <b>reasoning</b> route.
Give that route a capable model &mdash; it reads every answer and has to justify
a verdict.</p>
</div>

<div class="card">
<h2>7. Attaching documents</h2>
<p>The &#128206; button attaches a file to the conversation: PDF, Word, Excel,
Markdown, CSV, plain text or source code. Crewe extracts the text &mdash; the
model never sees the original file.</p>
<p>The &#128230; button attaches a <b>public GitHub repository</b>: paste
<code>https://github.com/owner/repo</code> (a <code>/tree/branch</code> suffix
works too) and Crewe fetches the source, skips binaries, lockfiles and
<code>node_modules</code>, and attaches it like any document. Then ask for a
<b>code review</b> &mdash; "suggest improvements", "audit this" &mdash; and the
<b>review</b> route reads whole files and reports findings, each with
ready-to-paste suggested code, telling you which files it could not fit. Private repositories are not supported. One thing to
know: <i>"review this code"</i> gets a review, while <i>"fix the login bug"</i>
goes to the code route and builds &mdash; ask for findings when you want
findings.</p>
<p>An attachment belongs to the <i>conversation</i>, not to one question, so you
can keep asking about the same document. Ask for a summary and the
<b>summarize</b> route handles it; ask a specific question and it goes wherever
it belongs.</p>
<p>Documents too large for the model's context window are summarised in passes
rather than cut off, so the end of a long report still counts.</p>
</div>

<div class="card">
<h2>Where your data lives</h2>
<p>Everything is on your own machine, in your home directory:
<code>~/crewe_userdata/&lt;your-id&gt;/</code> holds documents, sheets, recipes,
uploads and conversation memory. Back up that folder and you have backed up
everything of yours. Your API keys live in <code>~/router_config.json</code>
&mdash; never commit that file anywhere.</p>
</div>
"""


@app.route("/help")
def help_page():
    return Response(_shell("Help", HELP_BODY), mimetype="text/html")


@app.route("/account", methods=["GET"])
def account():
    u = current_user()
    msg = ""
    flashed = flask_session.pop("account_msg", None)
    if flashed:
        cls, text = flashed
        msg = f'<div class="msg {cls}">{_esc(text)}</div>'
    admin_link = ('<p class="muted" style="margin-top:18px">You are an admin. '
                  '<a href="/admin" style="color:var(--gold)">Manage people &rarr;</a></p>'
                  ) if u.get("is_admin") else ""
    body = f"""{msg}
<div class="card">
  <h2>Your account</h2>
  <p class="hint">{_esc(u['email'])}
     {'<span class="tag admin">ADMIN</span>' if u.get('is_admin') else ''}
     &middot; joined {_esc(u.get('created',''))}</p>
  <form method="POST" action="/account/password">
    <label for="cur">Current password</label>
    <input id="cur" name="current" type="password" autocomplete="current-password" required>
    <label for="new">New password (min 10 characters)</label>
    <input id="new" name="new" type="password" autocomplete="new-password" required>
    <label for="conf">Confirm new password</label>
    <input id="conf" name="confirm" type="password" autocomplete="new-password" required>
    <button type="submit">Change password</button>
  </form>
  {admin_link}
</div>
<div class="card">
  <h2>Session</h2>
  <p class="hint">Signed in on this device.</p>
  <a class="btn ghost" href="/logout">Sign out</a>
</div>"""
    return _shell("Account", body)

@app.route("/account/password", methods=["POST"])
def account_password():
    u = current_user()
    cur = request.form.get("current") or ""
    new = request.form.get("new") or ""
    conf = request.form.get("confirm") or ""
    try:
        ok = check_password_hash(u["pw_hash"], cur)
    except Exception:
        ok = False
    if not ok:
        flask_session["account_msg"] = ("err", "Current password is incorrect.")
    elif new != conf:
        flask_session["account_msg"] = ("err", "New passwords do not match.")
    else:
        done, err = set_password(u["id"], new)
        flask_session["account_msg"] = ("ok", "Password changed.") if done else ("err", err)
    return redirect("/account")

def _base_url():
    # Prefer an explicit public URL; else infer from the request (correct when
    # reached via the tunnel, which forwards the real Host).
    env = os.environ.get("CREWE_BASE_URL")
    return (env or request.url_root).rstrip("/")

@app.route("/admin", methods=["GET"])
def admin():
    users = _load_users()
    invites = [i for i in _load_invites()
               if not i.get("used") and time.time() <= i.get("expires", 0)]
    me = current_user()

    flashed = flask_session.pop("admin_msg", None)
    msg = f'<div class="msg {flashed[0]}">{flashed[1]}</div>' if flashed else ""

    urows = ""
    for u in users:
        is_me = u["id"] == me["id"]
        badge = '<span class="tag admin">ADMIN</span>' if u.get("is_admin") else '<span class="tag">user</span>'
        del_btn = "" if is_me else (
            f'<form class="inline" method="POST" action="/admin/user/delete" '
            f'onsubmit="return confirm(\'Delete {_esc(u["email"])}? Their data dir is kept.\')">'
            f'<input type="hidden" name="uid" value="{_esc(u["id"])}">'
            f'<button class="ghost" type="submit">Remove</button></form>')
        who = _esc(u["email"]) + (' <span class="muted">(you)</span>' if is_me else "")
        urows += f"<tr><td>{who}</td><td>{badge}</td><td class='muted'>{_esc(u.get('created',''))}</td><td>{del_btn}</td></tr>"

    irows = ""
    for i in invites:
        irows += (f"<tr><td>{_esc(i['email'])}</td>"
                  f"<td>{'<span class=\"tag admin\">admin</span>' if i.get('is_admin') else '<span class=\"tag\">user</span>'}</td>"
                  f"<td class='muted'>expires {_esc(_dt.fromtimestamp(i['expires']).strftime('%Y-%m-%d'))}</td>"
                  f"<td><form class='inline' method='POST' action='/admin/invite/revoke'>"
                  f"<input type='hidden' name='token' value='{_esc(i['token'])}'>"
                  f"<button class='ghost' type='submit'>Revoke</button></form></td></tr>")
    invites_block = (f"<table><tr><th>Email</th><th>Role</th><th></th><th></th></tr>{irows}</table>"
                     if irows else '<p class="muted">No pending invites.</p>')

    new_inv = flask_session.pop("new_invite", None)
    link_block = ""
    if new_inv:
        link_block = f"""<div class="msg ok">Invite ready for <b>{_esc(new_inv['email'])}</b> — send them this link (valid 7 days, one use):</div>
<div class="linkbox">
  <input type="text" id="invlink" value="{_esc(new_inv['link'])}" readonly onclick="this.select()">
  <button type="button" class="ghost" onclick="navigator.clipboard.writeText(document.getElementById('invlink').value);this.textContent='Copied'">Copy</button>
</div>"""

    body = f"""{msg}
<div class="card">
  <h2>Invite someone</h2>
  <p class="hint">Generates a one-time link. They set their own password — you never see it.</p>
  {link_block}
  <form method="POST" action="/admin/invite">
    <label for="iemail">Email</label>
    <input id="iemail" name="email" type="email" required placeholder="person@example.com">
    <div class="row"><input type="checkbox" id="iadmin" name="is_admin" value="1" style="width:auto">
      <label for="iadmin">Make them an admin (can manage people and settings)</label></div>
    <button type="submit">Create invite link</button>
  </form>
</div>
<div class="card">
  <h2>People <span class="muted">({len(users)})</span></h2>
  <table><tr><th>Email</th><th>Role</th><th>Joined</th><th></th></tr>{urows}</table>
</div>
<div class="card">
  <h2>Pending invites</h2>
  {invites_block}
</div>"""
    return _shell("People", body)

@app.route("/admin/invite", methods=["POST"])
def admin_invite():
    inv, err = create_invite(request.form.get("email"),
                             bool(request.form.get("is_admin")),
                             current_user()["email"])
    if err:
        flask_session["admin_msg"] = ("err", _esc(err))
    else:
        flask_session["new_invite"] = {
            "email": inv["email"],
            "link": f"{_base_url()}/invite/{inv['token']}",
        }
    return redirect("/admin")

@app.route("/admin/invite/revoke", methods=["POST"])
def admin_invite_revoke():
    _consume_invite(request.form.get("token") or "")   # mark used = effectively revoked
    flask_session["admin_msg"] = ("ok", "Invite revoked.")
    return redirect("/admin")

@app.route("/admin/user/delete", methods=["POST"])
def admin_user_delete():
    uid = request.form.get("uid") or ""
    me = current_user()
    if uid == me["id"]:
        flask_session["admin_msg"] = ("err", "You can't remove your own account here.")
        return redirect("/admin")
    users = _load_users()
    target = next((u for u in users if u["id"] == uid), None)
    admins = [u for u in users if u.get("is_admin")]
    if target and target.get("is_admin") and len(admins) <= 1:
        flask_session["admin_msg"] = ("err", "Can't remove the last admin.")
        return redirect("/admin")
    if target:
        delete_user(uid)
        flask_session["admin_msg"] = ("ok", f"Removed {_esc(target['email'])}. Their data dir was kept.")
    return redirect("/admin")

@app.route("/invite/<token>", methods=["GET", "POST"])
def invite_accept(token):
    inv = _invite_by_token(token)
    if not inv:
        return _shell("Invite", '<div class="card"><h2>This invite link is invalid or expired.</h2>'
                      '<p class="hint">Ask whoever invited you for a fresh link.</p></div>',
                      back="/login"), 410
    err = ""
    if request.method == "POST":
        pw = request.form.get("password") or ""
        conf = request.form.get("confirm") or ""
        if pw != conf:
            err = "Passwords do not match."
        else:
            u, e = create_user(inv["email"], pw, is_admin=inv.get("is_admin"))
            if e:
                err = e
            else:
                _consume_invite(token)
                flask_session.permanent = True
                flask_session["uid"] = u["id"]
                return redirect("/")
    emsg = f'<div class="msg err">{_esc(err)}</div>' if err else ""
    body = f"""{emsg}
<div class="card">
  <h2>Welcome to Crewe</h2>
  <p class="hint">Creating your account for <b>{_esc(inv['email'])}</b>. Choose a password to finish.</p>
  <form method="POST" action="/invite/{_esc(token)}">
    <label for="pw">Password (min 10 characters)</label>
    <input id="pw" name="password" type="password" autocomplete="new-password" required autofocus>
    <label for="conf">Confirm password</label>
    <input id="conf" name="confirm" type="password" autocomplete="new-password" required>
    <button type="submit">Create account &amp; sign in</button>
  </form>
</div>"""
    return _shell("Accept invite", body, back="/login")

# ---- tenancy: per-user data roots (Phase 2) ---------------------------------
# The owning user is carried in a thread-local rather than threaded through ~76
# call sites, so the per-store accessors below take no arguments and the call
# sites stay a plain rename. Two rules make that safe:
#   1. _require_login() sets it on EVERY request (Flask reuses worker threads,
#      so a stale value from a previous request would be a cross-user leak).
#   2. Background threads get it via spawn_owned(), which captures the uid at
#      launch time — a bare threading.Thread would see None and blow up loudly
#      in _owner(), which is the intended failure mode. Never paper over that
#      by defaulting to a random user.
USERDATA_ROOT = os.path.expanduser("~/crewe_userdata")
_OWNER = threading.local()

def set_owner(uid):
    _OWNER.uid = uid

def _owner():
    uid = getattr(_OWNER, "uid", None)
    if not uid:
        raise RuntimeError(
            "no owning user in context — a background thread was started with "
            "threading.Thread instead of spawn_owned(), or a store was touched "
            "outside a request")
    return uid

def spawn_owned(target, args=(), daemon=True, **kw):
    """threading.Thread, but the worker inherits the caller's owning user."""
    uid = _owner()
    def _run(*a, **k):
        set_owner(uid)
        return target(*a, **k)
    return threading.Thread(target=_run, args=args, daemon=daemon, **kw)

def user_root(uid=None):
    p = os.path.join(USERDATA_ROOT, uid or _owner())
    os.makedirs(p, exist_ok=True)
    return p

def _ustore(name, uid=None):
    p = os.path.join(user_root(uid), name)
    os.makedirs(p, exist_ok=True)
    return p

def _ufile(name, uid=None):
    return os.path.join(user_root(uid), name)

# Per-user stores. Previously ~/router_<name>; now ~/crewe_userdata/<uid>/<name>.
def docs_dir(uid=None):     return _ustore("docs", uid)
def sheets_dir(uid=None):   return _ustore("sheets", uid)
def books_dir(uid=None):    return _ustore("books", uid)
def photos_dir(uid=None):   return _ustore("photos", uid)
def audio_dir(uid=None):    return _ustore("audio", uid)
def uploads_dir(uid=None):  return _ustore("uploads", uid)
def spend_file(uid=None):   return _ufile("spend.jsonl", uid)
def checks_dir(uid=None):   return _ustore("checks", uid)
def cookbook_file(uid=None): return _ufile("cookbook.json", uid)
def memory_file(uid=None):   return _ufile("memory.json", uid)

def owned_job(jobs, job_id):
    """Fetch a job only if the caller owns it. Foreign jobs return None so the
    caller's existing 'unknown job' 404 fires — a probe for someone else's job
    is then indistinguishable from a probe for one that never existed."""
    j = jobs.get(job_id)
    if not j:
        return None
    try:
        if j.get("owner") != _owner():
            return None
    except RuntimeError:
        return None
    return j

# Deliberately GLOBAL, not per-user:
#   ROUTER_CONFIG_FILE — fleet config + API keys, admin-only via ADMIN_PREFIXES
#   PREFS_FILE         — standing build standards, operator-set
#   ~/llama_logs/*     — operator debugging
# ==== end accounts / auth ====================================================

# ---- endpoint map -----------------------------------------------------------
# Backend hosts — overridable via env so a fork never has to edit this file:
#   CREWE_HOST        host of the local specialist llama.cpp servers (default 127.0.0.1)
#   CREWE_CODER_HOST  host of the dedicated coder machine
# AGPL s13: users interacting over a network must be offered this version's
# source. Point this at YOUR repository if you are running a modified Crewe.
SOURCE_URL = os.environ.get("CREWE_SOURCE_URL", "https://github.com/paulatdotwork/crewe")

HOST = os.environ.get("CREWE_HOST", "127.0.0.1")
CODER_HOST = os.environ.get("CREWE_CODER_HOST", HOST)   # dedicated coder box; set CREWE_CODER_HOST if it is a different machine
# TEMP: all routes → single model on 8080 for testing
# To restore multi-model setup, uncomment the original block below and remove this one.
#SPECIALISTS = {
#    "recipes":   f"http://{HOST}:8080/v1/chat/completions",
#    "creative":  f"http://{HOST}:8080/v1/chat/completions",
#    "code":      f"http://{HOST}:8080/v1/chat/completions",
#    "general":   f"http://{HOST}:8080/v1/chat/completions",
#    "reasoning": f"http://{HOST}:8080/v1/chat/completions",
#}
#CLASSIFIER = f"http://{HOST}:8080/v1/chat/completions"
#SUMMARIZER = f"http://{HOST}:8080/v1/chat/completions"
# Original multi-model config:
SPECIALISTS = {
    "recipes":   f"http://{HOST}:8080/v1/chat/completions",
    "creative":  f"http://{HOST}:8081/v1/chat/completions",
    "code":      f"http://{CODER_HOST}:8080/v1/chat/completions",
    "general":   f"http://{HOST}:8087/v1/chat/completions",
    "reasoning": f"http://{HOST}:8087/v1/chat/completions",
    # web search: results are fetched first, then synthesized by the 26B brain
    "search":    f"http://{HOST}:8087/v1/chat/completions",
}
CLASSIFIER = f"http://{HOST}:8083/v1/chat/completions"
SUMMARIZER = f"http://{HOST}:8084/v1/chat/completions"
VALID = set(SPECIALISTS.keys()) | {"audio"}

# ---- web search (SearXNG) ---------------------------------------------------
SEARXNG_URL       = f"http://{HOST}:8888/search"  # JSON API enabled in settings.yml
SEARCH_RESULTS    = 5     # snippets handed to the model
SEARCH_FETCH      = 3     # of those, how many pages to fetch + extract in full
SEARCH_FETCH_CHARS = 3500 # chars of extracted page text kept per fetched page

# ---- conversation memory ----------------------------------------------------
# Per-user since Phase 2. Previously one global dict loaded at import, which —
# combined with a client-supplied session_id — let anyone read anyone else's
# history just by passing a different id. Namespacing by owner closes that.
SESSIONS_LOCK = threading.Lock()      # guards MUTATION of a user's dict
_SESS_CACHE: dict = {}                # uid -> that user's sessions
_SESS_CACHE_LOCK = threading.Lock()   # guards the cache map ONLY.
# ^ Must be a different lock from SESSIONS_LOCK: callers hold SESSIONS_LOCK while
# calling sessions(), and threading.Lock is not reentrant — sharing one would
# deadlock the request thread.

def _load_memory(uid=None) -> dict:
    try:
        with open(memory_file(uid)) as f:
            return json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        return {}

def sessions() -> dict:
    """The calling user's sessions, loaded on first touch. Raises outside a
    request or an owned thread rather than guessing whose data this is."""
    uid = _owner()
    with _SESS_CACHE_LOCK:
        if uid not in _SESS_CACHE:
            _SESS_CACHE[uid] = _load_memory(uid)
        return _SESS_CACHE[uid]

def _save_memory():
    # snapshot under the lock (a consistent view), write atomically — a
    # concurrent summarize() thread mid-dump used to corrupt the file
    uid = _owner()
    mine = sessions()          # resolve before taking SESSIONS_LOCK (see above)
    with SESSIONS_LOCK:
        snap = json.dumps(mine, indent=2)
    tmp = f"{memory_file(uid)}.tmp{os.getpid()}-{threading.get_ident()}"
    with open(tmp, "w") as f:
        f.write(snap)
    os.replace(tmp, memory_file(uid))

# Verbatim recent turns, replayed to the specialists as real chat history.
# The rolling summary covers the long tail; this covers "the user replied
# ten seconds later" — it is written synchronously when an answer completes,
# so the very next request is guaranteed to see it (the summary is not).
HISTORY_TURNS = 6      # exchanges kept per session
HISTORY_CHARS = 2000   # max chars stored per question/answer


def _append_history(session_id: str, question: str, answer: str, route: str):
    with SESSIONS_LOCK:
        sess = sessions().setdefault(session_id, {})
        hist = sess.setdefault("history", [])
        hist.append({"q": question[:HISTORY_CHARS], "a": answer[:HISTORY_CHARS]})
        del hist[:-HISTORY_TURNS]
        sess["last_route"] = route
    _save_memory()


def _recent_context(session_id: str):
    """(history list, last_route) snapshot for prompt building and routing."""
    with SESSIONS_LOCK:
        sess = sessions().get(session_id, {})
        return list(sess.get("history") or []), sess.get("last_route", "")

# ---- cookbook ---------------------------------------------------------------
# Per-user since Phase 2, same lazy-cache shape as sessions() above.
COOKBOOK_LOCK = threading.Lock()      # guards MUTATION of a user's list
_COOK_CACHE: dict = {}                # uid -> that user's cookbook
_COOK_CACHE_LOCK = threading.Lock()   # guards the cache map ONLY (not reentrant)

def _load_cookbook(uid=None) -> list:
    try:
        with open(cookbook_file(uid)) as f:
            return json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        return []

def cookbook() -> list:
    """The calling user's cookbook, loaded on first touch."""
    uid = _owner()
    with _COOK_CACHE_LOCK:
        if uid not in _COOK_CACHE:
            _COOK_CACHE[uid] = _load_cookbook(uid)
        return _COOK_CACHE[uid]

def _save_cookbook():
    uid = _owner()
    mine = cookbook()          # resolve before taking COOKBOOK_LOCK
    with COOKBOOK_LOCK:
        snap = json.dumps(mine, indent=2)
    tmp = f"{cookbook_file(uid)}.tmp{os.getpid()}-{threading.get_ident()}"
    with open(tmp, "w") as f:
        f.write(snap)
    os.replace(tmp, cookbook_file(uid))


# Headings with these names are sections WITHIN a recipe, not new dishes —
# so they never start a new cookbook card.
_RECIPE_SECTIONS = re.compile(
    r"^(ingredients?|instructions?|steps?|method|directions?|preparation|prep"
    r"|notes?|tips?|nutrition|servings?|yield|equipment|to serve|garnish"
    r"|for the .+|optional.*)\b", re.I)


def split_recipes(answer, question=""):
    """Split a recipes answer into individual {title, body} dicts, one per dish.
    Dishes are delimited by markdown headings; headings named like recipe
    sections (Ingredients, Instructions, …) stay inside the current recipe.
    Leading prose before the first dish heading is dropped as preamble. Falls
    back to a single untitled entry when there are no dish headings."""
    recipes, title, buf = [], None, []

    def flush():
        body = "\n".join(buf).strip()
        if body or title:
            recipes.append({"title": title, "body": body})

    for line in answer.split("\n"):
        m = re.match(r"^\s{0,3}#{1,4}\s+(.+?)\s*#*\s*$", line)
        if m:
            t = re.sub(r"[*_`]", "", m.group(1)).strip()
            t = re.sub(r"^(recipe\s*)?\d+[.):]\s*", "", t, flags=re.I)
            if _RECIPE_SECTIONS.match(t):   # subsection — keep with current dish
                buf.append(line)
                continue
            flush()
            title, buf = t, []
        else:
            buf.append(line)
    flush()

    # Drop a leading title-less preamble chunk if real titled recipes follow.
    if len(recipes) > 1 and recipes[0]["title"] is None:
        recipes = recipes[1:]
    if not recipes:
        return [{"title": None, "body": answer.strip()}]
    return recipes


# ---- audio (Piper TTS) ------------------------------------------------------
PIPER_BIN   = os.path.expanduser("~/piper_tts/piper/piper")
PIPER_VOICE = os.path.expanduser("~/piper_tts/voices/en_US-amy-medium.onnx")
PIPER_LIB   = os.path.expanduser("~/piper_tts/piper")   # bundled .so files

# The 'audio' route writes a script meant to be SPOKEN, then Piper records it.
AUDIO_SCRIPT_SYSTEM = (
    "You are writing a short script that you will then read aloud yourself. "
    "Output ONLY the spoken words — natural, conversational sentences meant to be heard, not read. "
    "No markdown, no headings, no bullet points, no stage directions, no emoji, no parentheticals, "
    "no '[music]' style cues. Write numbers and symbols as words to be spoken. "
    "Keep it engaging and tight. Unless the user asks for a specific length, aim for 60-150 words."
)

def _clean_for_speech(text: str) -> str:
    """Strip markdown so Piper doesn't literally pronounce '*', '#', backticks, etc."""
    text = re.sub(r"[*_`#>|]+", "", text)          # inline markdown punctuation
    text = re.sub(r"^\s*[-•]\s+", "", text, flags=re.M)  # list bullets
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()

AUDIO_OUT_RATE = 44100  # target sample rate for saved recordings

def _upsample_wav(src: str, dst: str, target_rate: int = AUDIO_OUT_RATE):
    """Resample a 16-bit mono WAV to target_rate via linear interpolation.

    Piper's voice is natively 22050 Hz. Going to 44100 is an exact 2x, but this
    handles any integer-ish ratio by interpolating. Pure stdlib (wave + array),
    since this box has no ffmpeg/sox/numpy and Py3.14 dropped audioop.
    Note: this changes the container's rate for compatibility; it does not add
    real fidelity beyond the source's ~11 kHz ceiling.
    """
    with wave.open(src, "rb") as w:
        nch, sw, src_rate, n = (w.getnchannels(), w.getsampwidth(),
                                w.getframerate(), w.getnframes())
        raw = w.readframes(n)

    # Only 16-bit mono is expected from Piper; anything else, copy unchanged.
    if sw != 2 or nch != 1 or src_rate >= target_rate and target_rate % src_rate != 0:
        with open(src, "rb") as f, open(dst, "wb") as g:
            g.write(f.read())
        return

    samples = array.array("h")
    samples.frombytes(raw)
    m = len(samples)
    if m == 0:
        with open(src, "rb") as f, open(dst, "wb") as g:
            g.write(f.read())
        return

    ratio = target_rate / src_rate          # 2.0 for 22050 -> 44100
    out_len = int(m * ratio)
    out = array.array("h", bytes(out_len * 2))
    for i in range(out_len):
        pos = i / ratio
        j = int(pos)
        frac = pos - j
        a = samples[j]
        b = samples[j + 1] if j + 1 < m else a
        out[i] = int(a + (b - a) * frac)

    with wave.open(dst, "wb") as w:
        w.setnchannels(1)
        w.setsampwidth(2)
        w.setframerate(target_rate)
        w.writeframes(out.tobytes())

# system prompts injected per route (llama-server no longer takes them as a CLI flag)
SYSTEM_PROMPTS = {
    "recipes":  "You are a concise, practical cooking assistant. Begin EVERY recipe with a markdown heading naming the dish, e.g. '## Honey Garlic Chicken Thighs'. If the user asks for several recipes, give each its own '## Dish Name' heading. Under each heading, return a clear ingredients list with quantities followed by numbered steps. Note prep/cook time and servings. Keep commentary minimal.",
    "creative": "You are an imaginative creative-writing partner. Write vivid, characterful prose. Favor strong imagery, varied sentence rhythm, and emotional texture. Continue collaboratively and keep the established tone, characters, and plot consistent across turns.",
    "code":     "You are an expert software engineer. Return correct, idiomatic, production-quality code with minimal preamble. Prefer complete, runnable solutions. Explain only when asked or when a non-obvious decision needs justifying. Default to the user's language/framework if evident.",
    "general":   "You are a knowledgeable, helpful general assistant. Answer questions clearly and accurately across any topic. Be concise but complete, and say when you are unsure rather than inventing facts.",
    "review": "You are a senior software engineer performing a code review. Ground every finding in the actual code shown — name the file path and quote the offending lines. Every finding MUST end with working suggested code: the file path on its own line, then a fenced code block containing the corrected version, complete enough to paste — a whole function or coherent block, never a fragment with '...' elisions. Keep each suggestion minimal: change what the finding requires and preserve the surrounding style. Order findings by impact: correctness first, then security, then performance, then clarity. Say clearly when something is already good, and when you could not review a file because it was omitted. Never invent files, lines, or APIs that were not shown to you; if a fix depends on code you cannot see, say so instead of guessing.",
    "summarize": "You condense material the user already has. Be faithful above all: keep names, numbers, dates, decisions and conclusions exactly as written, never infer a fact that is not present, and say so plainly when the source does not cover something. Lead with the single most important point, then structure the rest to match the material — bullets for lists of findings, short prose for an argument. Match the requested length if one is given; otherwise aim for roughly a tenth of the original. Do not add opinion, praise, or advice unless asked.",
    "reasoning": "You are a precise analytical reasoner. Work through problems step by step, showing your reasoning clearly. Handle math, logic, science, and multi-step analysis. Check your work before giving a final answer. Acknowledge uncertainty rather than guessing.",
    "search":    "You are a research assistant answering with the help of live web search results supplied in the user message. Base your answer on those results and synthesize across them — do not rely on a single source when several are given. Cite sources inline as [n] matching the numbered results, and attach a date to any fact that can change over time (prices, rankings, counts, 'latest' anything). When the question compares things across several items, present the comparison as a markdown table with a source/date column. Prefer recent, agreeing sources; when sources disagree or a figure is uncertain, say so explicitly rather than averaging silently. If the results don't actually answer the question, state that plainly instead of guessing. Never fabricate URLs, numbers, or facts beyond what the sources support.",
}

# in-memory progress tracking: job_id -> {"tokens": int, "done": bool, "route": str, "answer": str}
JOBS: dict = {}
JOBS_LOCK = threading.Lock()


def _prune_jobs(jobs, keep_done=50):
    """Drop the oldest finished jobs — answers are large (a code build is
    routinely 50-200KB of markdown) and nothing else ever evicts them."""
    done = [k for k, v in jobs.items() if v.get("done")]
    for k in done[:-keep_done]:
        jobs.pop(k, None)


# ---- agentic code pipeline ----------------------------------------------------
# ONE model does everything: planning, coding, and the per-step checks.
# Two models rewriting the same files produced code neither would write
# alone — that rule is structural now. Each step runs with a fresh, minimal
# context — the goal, the plan outline, the step, and only the files that
# step touches — so late steps can't drift and the model never sees enough
# clutter to get lazy.
#
# The coder wants a machine to itself: builds are long, and a coder that shares
#   a backend with chat makes both feel broken. Point CODER_HOST at a dedicated
#   box if you have one — otherwise it defaults to the same host as everything
#   else and simply queues. Routed via SPECIALISTS["code"].
# Trade-off worth knowing: a big-context coder lets a step see the whole
#   project, which is what stops steps redeclaring each other's globals — but
#   the budgets must follow the model, which is what _budgets_for() does.
#   Docs: code-pipeline-A.md / code-pipeline-B.md.
CODE_AGENT_URL      = SPECIALISTS["code"]        # remote Qwen3.6-35B coder, 128k ctx
CODE_AGENT_URL_FAST  = CODE_AGENT_URL   # "normal" tier; _apply_router_config sets it
CODE_AGENT_URL_QUICK = CODE_AGENT_URL   # "fast" tier; falls back down the chain
_BUDGETS_READY = False   # True once _refresh_budgets() exists (defined below)
# CODE_AGENT_URL    = SPECIALISTS["reasoning"]   # local 26B on 8087 (drop STEP_FILE_CHARS back to 24000)
STEP_TEMP           = 0.2    # explicit low temp — never write code at the server's default (1.0)
MAX_PLAN_STEPS      = 24     # sanity bound only — a real request can need many steps
EXTRA_STEPS         = 8      # bonus steps the delivery inspector may add for missing features
STEP_FIX_ROUNDS     = 2      # foreman fix attempts per step
# Run the per-step LLM foreman ONLY when a free deterministic check fails.
# Set 0 to always run it (the old behaviour).
FOREMAN_ONLY_ON_ISSUES = os.environ.get("CREWE_FOREMAN_ALWAYS") != "1"
# Env-tunable so slow-backend TESTS (e.g. the laguna TB-RPC rig) can shrink the
# per-call prompt without code edits — prompt-eval cost scales with these.
STEP_FILE_CHARS     = int(os.environ.get("CREWE_STEP_FILE_CHARS", "160000"))  # ~40k tokens of files
STEP_FILE_HARD      = int(os.environ.get("CREWE_STEP_FILE_HARD", "320000"))  # a step's own target file may use this much alone — with diff
                             # edits the OUTPUT stays small, so a big file still fits the ctx
# Below this, EVERY step sees the WHOLE project. Was 12000 — written when the
# coder had 8-16k of context. The coder now has 131072 tokens (~500k chars), so
# starving a step of the project it is editing bought nothing and cost plenty:
# steps redeclared each other's globals, asked for files they were refused, and
# edited against code they could not see. 160k chars is ~40k tokens, a third of
# the window, leaving ample room for the plan, working memory and the output.
SMALL_PROJECT_CHARS = int(os.environ.get("CREWE_WHOLE_PROJECT_CHARS", "160000"))

# Standing quality standards, injected into every pipeline call. Same idea as
# coding-bot's preferences.md — edit the file, no code change needed.
PREFS_FILE = os.path.expanduser("~/router_prefs.md")

def _load_prefs():
    try:
        with open(PREFS_FILE) as f:
            return f.read().strip()
    except OSError:
        return ""

# Snap Firefox can't read /tmp or hidden dirs, so browser checks run here.

PLANNER_SYSTEM = (
    "You are a software architect writing the build spec and plan for another "
    "engineer. First decide the project's TRUE scope: what the request "
    "IMPLIES, not the literal minimum. A request 'like' a known game or app "
    "implies its core loop and signature parts — a JRPG implies an overworld, "
    "enterable buildings with doors, NPCs with dialogue, styled menus, "
    "turn-based battles, and real sprite art. Shipping a hollow skeleton or "
    "placeholder-rectangle art is a FAILURE, not a small version.\n"
    "Respond with ONLY a JSON object — no markdown fences, no commentary:\n"
    '{"goal": "one sentence",\n'
    ' "features": ["every concrete feature the finished build must have — '
    'specific and checkable, e.g. \'player enters the town through a door '
    'tile and a distinct town map loads\'"],\n'
    ' "design": "2-4 sentences of visual direction: art method (layered SVG / '
    'canvas pixel sprites and tile size), palette, UI style",\n'
    ' "steps": [\n'
    '   {"title": "short imperative title",\n'
    '    "detail": "exactly what to implement in this step and nothing more",\n'
    '    "files": ["files this step creates or edits"],\n'
    '    "done_when": "objective, checkable completion criterion"}]}\n'
    "Rules:\n"
    "- The steps must collectively deliver EVERY feature in the list.\n"
    "- 1 to 10 steps. A genuinely small request (one fix, one widget) is ONE "
    "step; an ambitious build uses as many as it needs.\n"
    "- Each step is one focused change touching at most 2-3 files.\n"
    "- Design the file layout so NO file grows beyond ~200 lines — many "
    "small modules beat one giant script. Large files become uneditable.\n"
    "- Building from scratch: step 1 creates the full file skeleton (page "
    "layout, empty functions, element ids) so every later step edits "
    "existing files. Later steps may only depend on earlier ones.\n"
    "- When existing project files are provided, plan ONLY the changes the "
    "request needs — do not rebuild or restyle what already works."
)

# Kept identical across a build's step calls (plus the standards block) so
# llama.cpp can reuse the cached prompt prefix; everything step-specific
# travels in the user message.
STEP_SYSTEM = (
    "You are an expert software engineer executing ONE step of an approved "
    "build plan. Implement exactly what the step says — nothing more. Do not "
    "add features and do not touch files outside the step's file list — but "
    "within the step's scope, build to the project's design standards: real "
    "art and real layout. A plain colored rectangle standing in for a "
    "character, object, or scene is never acceptable unless the step "
    "explicitly says placeholder. Keep the code clean and idiomatic. "
    "Files you do not output are kept exactly as they are.\n"
    "OUTPUT RULES (absolute):\n"
    "- Do NOT deliberate, narrate, or think out loud. No 'Let me check', no "
    "'Actually, wait', no reasoning about what to do.\n"
    "- Do NOT quote existing code back before changing it.\n"
    "- Your ENTIRE reply is file output: a path line, then one fenced block "
    "per file. Nothing before the first block, nothing between blocks.\n"
    "- The ONLY permitted non-file reply is a single 'NEED FILES: ...' line.\n"
    "- At most one short sentence of explanation, and only if a decision "
    "genuinely needs it."
)

FOREMAN_SYSTEM = (
    "You are the build foreman checking ONE completed step of a build plan. "
    "You get the step's instructions, its done-when criterion, automated "
    "syntax-check results, and the files the step produced. Judge ONLY this "
    "step: is it fully implemented, correct, within scope, and up to the "
    "design standards? Placeholder-grade visuals (single-color boxes standing "
    "in for characters, scenery, or UI) fail the step even when the logic "
    "works. Ignore work that belongs to later steps.\n"
    "If the step is correct and complete, reply with exactly: ON TRACK\n"
    "Otherwise return every file that needs fixing — corrected and COMPLETE, "
    "the file path on its own line then one fenced code block per file — and "
    "end with one line starting 'FIXED:' that summarizes the correction. "
    "Never output partial files, diffs, or placeholder comments.\n"
    "OUTPUT RULES (absolute):\n"
    "- Do NOT deliberate, narrate, or reason out loud. No 'Let me check', no "
    "'Looking at the code', no walking through the files.\n"
    "- Do NOT quote or echo existing code back. The only fenced blocks in "
    "your reply are COMPLETE corrected files you are replacing.\n"
    "- Your whole reply is either the two words ON TRACK, or file blocks "
    "followed by one FIXED: line. Nothing else."
)

VISION_SYSTEM = (
    "You are reviewing a SCREENSHOT of a web app that was just built. Your job "
    "is to catch what automated selector checks CANNOT see: a page that renders "
    "blank or unstyled, a layout that collapsed, regions that overlap or sit on "
    "top of each other, controls that are invisible or cut off, default browser "
    "styling where a designed interface was expected.\n"
    "Judge ONLY what is visible in the image. Never guess about behaviour you "
    "cannot see (clicks, saving, drag), and never ask for anything that was not "
    "requested.\n"
    "If the page looks built, styled and coherent, reply with exactly: LOOKS GOOD\n"
    "Otherwise list ONLY what is visibly wrong or missing, one per line, most "
    "important first, at most 5 lines. No commentary, no code, no headings."
)

COMPLETE_SYSTEM = (
    "You are the delivery inspector for a finished build. You get the "
    "feature list the build promised, the design standards, and the project "
    "files. Decide whether every promised feature is genuinely present and "
    "usable — not stubbed, not hollow, not placeholder art.\n"
    "If everything is delivered, reply with exactly: COMPLETE\n"
    "Otherwise list the missing or hollow features, one per line, most "
    "important first, max 4 lines, no commentary. Do not output code."
)

_TRUNC_RETRY_NOTE = (
    "YOUR PREVIOUS OUTPUT WAS TRUNCATED or contained no complete files. "
    "Output every file for this step in full, first line to last.\n\n"
)


# Rejections the model can actually fix by trying again. A blind-rewrite
# refusal is policy, not a mistake, so retrying it just burns a call.
_RECOVERABLE_REJECT = re.compile(
    r"fragment|didn'?t match|cut off|unlabeled|no complete file|shorter than"
    r"|doesn'?t exist",
    re.I)


def _retry_note(why, files_to_redo=None):
    """Second-attempt preamble that tells the model WHY attempt 1 was thrown
    away. A bare 'try again' repeats the same mistake."""
    if files_to_redo:
        extra = ("\nIf a file does NOT exist yet, send it as a COMPLETE NEW "
                 "FILE — SEARCH/REPLACE cannot patch a file that is not there."
                 if "doesn't exist" in why or "does not exist" in why else "")
        return (f"PART OF YOUR PREVIOUS REPLY WAS REJECTED — {why}.{extra}\n"
                f"These file(s) did NOT get applied and must be re-sent: "
                f"{', '.join(files_to_redo)}.\n"
                "Send each of them COMPLETE from first line to last (or as "
                "SEARCH/REPLACE sections quoting the current file EXACTLY). "
                "Do not resend files that already succeeded.\n\n")
    return ("YOUR PREVIOUS REPLY PRODUCED NOTHING USABLE — " + why + ".\n"
            "Do not repeat that mistake. Either output each file for this step "
            "COMPLETE from first line to last, or use SEARCH/REPLACE sections "
            "whose SEARCH text is copied EXACTLY, character for character, from "
            "the current file shown above. Keep each reply focused: fewer files, "
            "smaller edits.\n\n")


def _vision_url():
    """Where to send screenshots.

    NOT tied to the coder: swapping in a coder without a vision encoder used to
    silently disable visual verification. Prefer an explicit CREWE_VISION_URL,
    else the first backend that actually reports a vision encoder — normally
    the local chat model, which stays put while coders come and go."""
    envu = os.environ.get("CREWE_VISION_URL")
    if envu:
        return envu
    for u in (SPECIALISTS.get("general"), SPECIALISTS.get("reasoning"),
              coder_url()):
        if u and _vision_ready(u):
            return u
    return SPECIALISTS.get("general") or coder_url()


@functools.lru_cache(maxsize=8)
def _vision_ready(url):
    """True when that backend actually has a vision encoder loaded. Cached —
    but keyed on url, so pointing at a different backend re-probes."""
    try:
        base = url.rsplit("/v1/", 1)[0]
        r = requests.get(f"{base}/props", timeout=10)
        return bool((r.json().get("modalities") or {}).get("vision"))
    except Exception:
        return False


def _vision_critique(shot, goal, features, job_id=None):
    """Show the RENDERED page to a vision model and return the visible defects
    ([] means it looks right).

    This exists because selector checks pass on a page that renders as an
    unstyled skeleton: asserting '#workspace exists' says nothing about whether
    anything was drawn. Never raises — a critic that is down or blind must
    never fail a build that is otherwise fine."""
    url = _vision_url()
    if not shot or not _vision_ready(url):
        return []
    try:
        b64 = base64.b64encode(open(shot, "rb").read()).decode()
    except Exception:
        return []
    want = "\n".join(f"- {f}" for f in (features or [])[:8]) or f"- {goal}"
    payload = {"messages": [
        {"role": "system", "content": VISION_SYSTEM},
        {"role": "user", "content": [
            {"type": "text",
             "text": f"The build was asked for:\n{want}\n\n"
                     f"Overall goal: {goal}\n\nReview the screenshot."},
            {"type": "image_url",
             "image_url": {"url": f"data:image/png;base64,{b64}"}}]}],
        "stream": False, "temperature": 0.2, "max_tokens": 600,
        "chat_template_kwargs": {"enable_thinking": False}}
    try:
        r = requests.post(url, json=_inject_model(payload, url),
                          headers=_hdrs(url), timeout=300)
        clean = _strip_think(
            _resp_json(r)["choices"][0]["message"]["content"] or "").strip()
    except Exception as e:
        print(f"[code] vision critic error: {e}")
        return []
    head = next((l for l in clean.splitlines() if l.strip()), "")
    if re.search(r"\bLOOKS GOOD\b", head, re.I):
        return []
    out = []
    for l in clean.splitlines():
        m = l.strip("-•* \t").strip()
        if len(m) <= 12 or m.endswith(":") or _INSPECT_NOISE.search(m):
            continue
        if re.search(r"\bLOOKS GOOD\b", m, re.I):
            continue
        out.append(m)
    return out[:5]


def _drop_truncated(new, out, idx, warnings, mem):
    """Discard the LAST file of a generation that was cut off mid-file.

    An unclosed final fence is definitive proof the last file is half-written,
    and the elision guard only rejects losses over 50% — so a file truncated at
    60% would be ACCEPTED and would overwrite a working file with a fragment.
    Dropping it keeps the previous version and lets the retry rebuild it.

    Keyed on fence parity, NOT on "the call hit the cap": a capped reply that
    still closed its last fence is complete, and a stale cap flag from an
    earlier attempt must never discard a good file."""
    if not new or out.count("```") % 2 == 0:
        return new
    new = dict(new)
    last = list(new)[-1]
    new.pop(last)
    warnings.append(f"step {idx}: {last} was cut off mid-file — discarded, "
                    f"kept the previous version")
    _mem_note(mem, f"Step {idx}: {last} was cut off mid-file by the output "
                   f"limit and was DISCARDED. Write fewer files per reply, or "
                   f"use targeted SEARCH/REPLACE edits instead of full "
                   f"rewrites, so the reply fits.")
    return new

# Full-file rewrites of large files reliably come back as fragments (which
# the elision guard rightly rejects) — the fix then never lands. Targeted
# SEARCH/REPLACE edits are how big files get changed. Ported from coding-bot.
_EDIT_FORMAT_NOTE = (
    "\n\nTo EDIT an existing file shown above, you may return targeted edits "
    "instead of the whole file: put the file path on the line before a fenced "
    "block containing one or more sections in EXACTLY this format:\n"
    "<<<<<<< SEARCH\n"
    "[exact lines copied verbatim from the current file]\n"
    "=======\n"
    "[replacement lines]\n"
    ">>>>>>> REPLACE\n"
    "The SEARCH text must match the current file character-for-character, "
    "including indentation, and must appear EXACTLY ONCE in the file — "
    "include enough surrounding lines to make it unique. Use several small "
    "sections rather than one big one — but note that if ANY section fails to "
    "match, the whole file is discarded, so quote conservatively and prefer a "
    "few reliable sections over many fragile ones. Prefer targeted edits for "
    "any file over ~100 lines. NEW files (that do not exist yet) must be "
    "complete files, never edits. NEVER output a file that exists but was not "
    "shown to you — it will be rejected."
)

# The trailing word is OPTIONAL: models close the block with a bare ">>>>>>>"
# about as often as ">>>>>>> REPLACE". Requiring the word meant the whole block
# was not recognised as an edit at all — it fell through to the full-file path,
# where a 1.2 KB patch "replacing" a 39 KB file is rejected as a fragment. The
# user sees "rejecting partial files"; the truth is a missing terminator word.
# Terminator = the arrows alone. The trailing word is optional and simply not
# consumed: models close with a bare ">>>>>>>" as often as ">>>>>>> REPLACE",
# and requiring the word made the whole block parse as a full file (a 1.2 KB
# "index.html" then loses to the 39 KB original as a "fragment").
# Deliberately NO lookahead for a following newline — with it, two adjacent
# blocks separated by no blank line merged into one and spliced the literal
# marker text into the file.
_DIFF_BLOCK = re.compile(
    r"<{5,} *SEARCH *\n([\s\S]*?)\n={5,} *\n([\s\S]*?)\n>{5,}", re.I)


def _apply_diff(code, text):
    """Apply SEARCH/REPLACE sections to code. Returns (new_code, applied,
    total, failed) — failed is the SEARCH text of every block that did not
    match uniquely. Callers must reject the result unless applied == total > 0
    — a partially-applied edit is worse than no edit."""
    applied = total = 0
    failed = []
    for m in _DIFF_BLOCK.finditer(text):
        total += 1
        search, replace = m.group(1), m.group(2)
        if not search.strip():
            # an empty SEARCH matches everywhere (and normalizes to [''] in
            # the flex tier) — it can only corrupt the file
            failed.append(search)
            continue
        # an ambiguous match would edit the wrong occurrence silently —
        # demand uniqueness and let the all-blocks-applied contract reject
        if code.count(search) == 1:
            code = code.replace(search, replace, 1)
            applied += 1
            continue
        # flex tier: models routinely drop/add TRAILING whitespace when
        # quoting. Match per-line with rstrip, still demanding uniqueness.
        c_lines = code.split("\n")
        s_lines = [l.rstrip() for l in search.split("\n")]
        idxs = [i for i in range(len(c_lines) - len(s_lines) + 1)
                if [l.rstrip() for l in c_lines[i:i + len(s_lines)]] == s_lines]
        if len(idxs) == 1:
            i = idxs[0]
            c_lines[i:i + len(s_lines)] = replace.split("\n")
            code = "\n".join(c_lines)
            applied += 1
        else:
            failed.append(search)
    return code, applied, total, failed


def _merge_output_file(name, content, files, allowed, warnings, tag, mem=None):
    """Turn one extracted output 'file' into its final content: SEARCH/REPLACE
    edits are applied against the current version, full files pass the
    blind-rewrite and elision guards. Returns the new content or None.

    `allowed` is the exact set of EXISTING files this call's model context
    included — an existing file outside it can only be a blind guess, no
    matter which bucket (unseen/others/foreman-omitted) kept it hidden."""
    if name in files and name not in allowed:
        warnings.append(f"{tag}: refused a blind rewrite of {name} "
                        f"(not shown to the model)")
        return None
    body = re.sub(r"^\s*(?://|#|/\*|<!--)[^\n]*$", "", content or "",
                  flags=re.M).strip()
    head = (content or "")[:200].strip()
    if head and name not in files:
        twin = next((k for k, v in files.items()
                     if k != name and (v or "")[:200].strip() == head), None)
        if twin:
            warnings.append(f"{tag}: refused {name} — its content is a copy of "
                            f"{twin} (mislabelled block, would clobber)")
            _mem_note(mem, f"{tag}: a block labelled {name} was byte-identical "
                           f"to {twin} and was REFUSED — label each block with "
                           f"the file it actually belongs to.")
            return None
    if name not in files and not body:
        # Nothing but comments/whitespace = an echoed format example, not a
        # file. Keyed on "no code at all", NOT a length threshold: a genuine
        # one-line new file (`let a=1;`) must still be accepted.
        warnings.append(f"{tag}: refused {name} — it contains no actual code "
                        f"(looks like an echoed format example)")
        _mem_note(mem, f"{tag}: a block named {name} held no code (an echoed "
                       f"format example) and was REFUSED. Never repeat the "
                       f"format example back — emit only real project files.")
        return None
    if files and re.fullmatch(r"file\d+\.\w+", name):
        # an auto-named fallback file means the model forgot the path line —
        # in an existing project that's a mislabeled duplicate, never a file.
        # This used to reject SILENTLY, so the model repeated the mistake in
        # the same step — twice in one observed run. Tell it what went wrong.
        warnings.append(f"{tag}: rejected an unlabeled code block "
                        f"(would have become junk file '{name}')")
        _mem_note(mem, f"{tag}: a code block had NO filename line above its "
                       f"fence and was THROWN AWAY as '{name}'. Every fenced "
                       f"block must have its file path on the line directly "
                       f"above the fence — re-emit that work with the path.")
        return None
    if _DIFF_BLOCK.search(content):
        base = files.get(name)
        if base is None:
            warnings.append(f"{tag}: {name} arrived as edits to a file "
                            f"that doesn't exist")
            _mem_note(mem, f"{tag}: you sent SEARCH/REPLACE edits for {name}, "
                           f"which does not exist, so they were DROPPED. New "
                           f"files must be written IN FULL — path line, then a "
                           f"fenced block with the whole content.")
            return None
        patched, ok, tot, failed = _apply_diff(base, content)
        if tot and ok == tot:
            return patched
        # PARTIAL APPLY. Discarding 8 correct edits because 2 missed threw the
        # step's whole answer away and pushed the model into a worse retry
        # (observed: 8/10 and 3/4 rejected, then a 14k-token junk retry). Keep
        # what matched IF the file is no more broken than before — syntax is
        # checked deterministically, so a partial edit can't ship a file that
        # stopped parsing. The unmatched blocks are reported, not silently lost.
        if ok and tot:
            before_issues = _static_check_files({name: base})
            after_issues = _static_check_files({name: patched})
            if len(after_issues) <= len(before_issues):
                warnings.append(f"{tag}: {name} applied {ok}/{tot} edits — "
                                f"{tot - ok} did not match and were SKIPPED "
                                f"(file still parses)")
                _mem_note(mem, f"{tag}: {name} took {ok} of {tot} edits; "
                               f"{tot - ok} did NOT match the real file and "
                               f"were skipped, so that part of the change is "
                               f"still missing. Quote the CURRENT file exactly "
                               f"if you need to finish it.")
                return patched
            warnings.append(f"{tag}: {name} partial edit ({ok}/{tot}) would "
                            f"have broken the file; rejected entirely")
        warnings.append(f"{tag}: {name} edits didn't match the current "
                        f"file ({ok}/{tot} applied); rejected")
        # Aider-style feedback: show the model the closest ACTUAL text so the
        # next attempt quotes the real file instead of re-guessing from memory
        detail = ""
        if failed:
            first = next((l for l in failed[0].splitlines() if l.strip()), "")
            lines = base.splitlines()
            close = difflib.get_close_matches(first, lines, n=1, cutoff=0.3)
            if close:
                i = lines.index(close[0])
                ctx = "\n".join(lines[max(0, i - 2):i + 3])[:400]
                detail = (f" Its first line was '{first[:90]}' — the closest "
                          f"ACTUAL text in {name} is:\n{ctx}\n")
        _mem_note(mem, f"{tag}: a SEARCH/REPLACE for {name} was REJECTED — "
                       f"only {ok}/{tot} blocks matched the real file.{detail}"
                       f"Next edit must quote the current file EXACTLY.")
        return None
    if _looks_elided(content, files.get(name)):
        warnings.append(f"{tag}: {name} came back as a fragment; "
                        f"kept the previous version")
        _mem_note(mem, f"{tag}: {name} came back SHORTER than the current "
                       f"version (a fragment or lazy rewrite) and was "
                       f"REJECTED. Output the file complete, or edit it with "
                       f"SEARCH/REPLACE sections.")
        return None
    return content

# ---- project file memory ----------------------------------------------------
# The code route keeps the latest version of every file the models produce, per
# session, so follow-ups ("here's the error", "change the header") are edits
# with full context instead of blind regeneration. Lives inside SESSIONS and is
# persisted to router_memory.json with the summaries.

# Extensions Crewe will treat as project source. This gates BOTH which files a
# step may write and — via _plausible_filename — whether a name written above a
# fence is recognised as a filename at all. A language missing here doesn't
# degrade: the label is rejected, the block gets a generated name like
# "index.xml", and the build ends with "produced no files". That is exactly how
# a correct C# reply was thrown away, so keep this list generous.
CODE_EXTS = {
    # web
    "html", "htm", "css", "scss", "less", "js", "mjs", "cjs", "ts", "tsx",
    "jsx", "vue", "svelte", "astro",
    # general purpose
    "py", "rb", "php", "go", "rs", "java", "kt", "kts", "scala", "swift",
    "cs", "fs", "vb", "c", "h", "cpp", "cc", "hpp", "m", "mm", "dart", "lua",
    "pl", "r", "jl", "ex", "exs", "erl", "hs", "clj", "cljs", "zig", "nim",
    # data / config / project files a build legitimately writes
    "json", "sql", "sh", "bash", "ps1", "bat",
    "csproj", "fsproj", "vbproj", "sln", "gradle", "sbt", "cmake",
    "toml", "yaml", "yml", "xml", "ini", "cfg", "properties", "gemspec",
}
_CMD_LANGS = {"bash", "shell", "sh", "zsh", "console", "terminal", "cmd"}
# Fence tag -> extension, used to name a block the model didn't label. An
# unmapped tag becomes the filename extension verbatim, which is how a
# ```csharp block became "file2.csharp".
_EXT_BY_LANG = {"html": "html", "htm": "html", "css": "css", "scss": "scss",
                "javascript": "js", "js": "js", "jsx": "jsx",
                "typescript": "ts", "ts": "ts", "tsx": "tsx",
                "python": "py", "py": "py", "json": "json",
                "bash": "sh", "shell": "sh", "sh": "sh", "powershell": "ps1",
                "sql": "sql", "c": "c", "cpp": "cpp", "c++": "cpp",
                "objectivec": "m", "rust": "rs", "rs": "rs",
                "csharp": "cs", "c#": "cs", "cs": "cs", "fsharp": "fs",
                "java": "java", "kotlin": "kt", "kt": "kt", "scala": "scala",
                "go": "go", "golang": "go", "ruby": "rb", "rb": "rb",
                "php": "php", "swift": "swift", "dart": "dart", "lua": "lua",
                "perl": "pl", "r": "r", "julia": "jl", "elixir": "ex",
                "erlang": "erl", "haskell": "hs", "clojure": "clj",
                "xml": "xml", "yaml": "yml", "yml": "yml", "toml": "toml",
                "ini": "ini", "vue": "vue", "svelte": "svelte"}
_LANG_BY_EXT = {"html": "html", "htm": "html", "css": "css", "js": "javascript",
                "mjs": "javascript", "ts": "typescript", "py": "python",
                "json": "json", "sh": "bash", "sql": "sql", "c": "c",
                "cpp": "cpp", "rs": "rust"}
_ERRORISH = re.compile(
    r"traceback|exception|syntax\s*error|type\s*error|reference\s*error|"
    r"is not defined|undefined|stack trace|errno|error:|uncaught|"
    r"doesn'?t work|not working|broken|fails?\b|"
    r"truncat|cut\s*off|didn'?t finish|incomplete|unfinished", re.I)

# Short replies that continue the previous exchange (game answers, yes/no,
# "keep going") — these stay on the previous route without asking the
# classifier, which can't be trusted with a bare "no".
_CONTINUATION = re.compile(
    r"^(yes|yeah|yep|no|nope|nah|maybe|sometimes|correct|right|wrong|"
    r"i don'?t know|idk|not sure|ok(ay)?|sure|continue|go on|next|"
    r"keep going|again|more|why|how so|what else)[\s!.?]*$", re.I)


# Unambiguous PROSE markers of an elided file. Safe to match anywhere: no
# language writes these in real code.
_ELIDE_PHRASE = re.compile(
    r"\(lines? \d|rest (?:of|is) (?:the )?(?:file|code)|"
    r"remains? (?:the )?same|unchanged below|omitted for brevity|"
    r"code (?:here )?(?:is )?unchanged|same as (?:before|above)", re.I)

# A bare "..." only means elision inside a COMMENT or alone on its line.
# It must NEVER match JavaScript spread syntax — `[...(state.ids || [])]`,
# `{...(opts ?? {})}`, `foo(...args)` are ordinary code, and treating them as
# placeholders silently rejected complete files and kept the stale version.
_ELIDE_DOTS = re.compile(
    r"^\s*(?://|#|/\*|\*|<!--)[^\n]*\.\.\."      # ... inside a comment
    r"|^\s*\.\.\.\s*$"                            # ... alone on a line
    r"|^\s*\{?\s*/\*\s*\.\.\.[^\n]*\*/\s*\}?\s*$",  # /* ... */ placeholder line
    re.M)


def _scrub_tokens(text):
    """Remove Gemma unused/special tokens that a degenerating backend can
    emit — they are never legitimate output and must not reach storage."""
    return re.sub(r"<unused\d+>", "", text or "")


def _strip_think(text):
    """Remove reasoning spans. Closed <think>...</think> pairs, and ALSO an
    unclosed trailing <think> — a truncated reasoning stream must never be
    parsed as answer text (its draft code fences become phantom files and
    its prose becomes garbage build steps)."""
    text = re.sub(r"<think>[\s\S]*?</think>", "", text, flags=re.I)
    m = re.search(r"<think>", text, flags=re.I)
    if m:
        head = text[:m.start()]
        # If the unclosed reasoning IS the whole reply, returning "" is
        # indistinguishable from a dead backend and kills the step with no
        # diagnostic. Keep the text when it still holds a fenced block (the
        # model wrote code inside its reasoning — salvage it); drop it when it
        # is pure prose, so rambling can never become build steps.
        if not head.strip():
            return text if "```" in text else head
        text = head
    return text


# Extensions a model-emitted file may plausibly have. A prose-scanned
# "filename" outside this list is a code expression (eq.art, obj.type),
# not a file — accepting one plants a junk file the whole pipeline then
# carries as context forever.
_KNOWN_EXTS = CODE_EXTS | {
    "md", "txt", "svg", "xml", "csv", "yml", "yaml", "toml", "ini", "cfg",
    "conf", "env", "scss", "less", "vue", "svelte", "frag", "vert", "glsl",
    "webmanifest", "map"}


def _plausible_filename(name):
    ext = name.rsplit(".", 1)[-1].lower() if "." in name else ""
    return ext in _KNOWN_EXTS


def _normalize_name(c, existing):
    """Map a candidate onto the path the project already uses, if any.
    exact -> unique basename -> case-insensitive -> unique case-insensitive
    basename. Looser tiers demand a UNIQUE hit so a project holding both
    app.js and js/app.js is never guessed at."""
    if not existing:
        return None
    if c in existing:
        return c
    base = {}
    for k in existing:
        base.setdefault(os.path.basename(k), []).append(k)
    if len(base.get(c, [])) == 1:
        return base[c][0]
    lower = {k.lower(): k for k in existing}
    if c.lower() in lower:
        return lower[c.lower()]
    lower_base = {}
    for k in existing:
        lower_base.setdefault(os.path.basename(k).lower(), []).append(k)
    cb = os.path.basename(c).lower()
    if len(lower_base.get(cb, [])) == 1:
        return lower_base[cb][0]
    return None


def _resolve_label(cands, existing, near=None):
    """Pick a filename for a code block.

    THE NEAREST LABEL WINS. `near` holds the candidates from the line closest
    to the fence; if it names exactly one plausible file, that is the answer —
    mapped onto an existing path when it matches one, otherwise taken as a NEW
    file. Only when the nearest line names nothing do we scan earlier prose.

    This ordering is load-bearing. Preferring "any existing filename anywhere
    in the preceding prose" silently DELETED files: a step whose preamble read
    "creating js/connections.js and integrating it into js/canvas.js" filed the
    new connections.js block under canvas.js (which existed), and the next
    block then overwrote it. The file vanished with no warning, and the foreman
    burned ten minutes trying to rewrite what the parser had eaten."""
    for c in (near or []):
        if _plausible_filename(c) or _normalize_name(c, existing):
            return _normalize_name(c, existing) or c
    if existing:
        # Match against files the project ALREADY has, progressively looser:
        # exact -> basename -> case-insensitive. Looser tiers require a UNIQUE
        # hit, so a project holding both app.js and js/app.js never guesses.
        # Without this, a model writing "App.js" or "connections.js" for an
        # existing "js/connections.js" spawned a parallel duplicate file, and
        # the build ended up with App.js + app.js + js/app.js all diverging.
        def _uniq(mapping, key):
            hits = [v for k, v in mapping.items() if k == key]
            return hits[0] if len(hits) == 1 else None

        base = {}
        for k in existing:
            base.setdefault(os.path.basename(k), []).append(k)
        lower = {k.lower(): k for k in existing}
        lower_base = {}
        for k in existing:
            lower_base.setdefault(os.path.basename(k).lower(), []).append(k)
        for c in cands:
            if c in existing:
                return c
            if len(base.get(c, [])) == 1:
                return base[c][0]
            if c.lower() in lower:
                return lower[c.lower()]
            cb = os.path.basename(c).lower()
            if len(lower_base.get(cb, [])) == 1:
                return lower_base[cb][0]
    for c in cands:
        if _plausible_filename(c):
            return c
    return None


# Lines in an inspector verdict that are code or narration, not feature
# statements. One of these becoming a build step sends the job in circles
# (observed: steps literally titled '```javascript' and 'const svgString =').
_INSPECT_NOISE = re.compile(
    # braces, a trailing semicolon, or a fence marker = code. A semicolon or
    # backtick ANYWHERE was too broad — it deleted real feature lines that
    # merely used one — but a bare ```lang line must still never become a step.
    r"[{}]|;\s*$|^\s*```"
    r"|^\s*(const|let|var|function|def|class|return|import|if|for|while|else)\b"
    r"|^(looking|let'?s|let me|i need|i will|i'?ll|i am|i'm|okay|ok[,.]|"
    r"next,|checking|to (verify|check|evaluate)|based on|overall|in summary|"
    r"wait\b|hmm|first,? (i|let))", re.I)


def _inspector_missing(clean):
    """Distill an inspector verdict into real missing-feature lines.
    Returns [] when the verdict is COMPLETE or contains nothing usable —
    reasoning-style models wrap the list in prose and stray code, and a
    noise line must never become a build step."""
    if clean.upper().startswith("COMPLETE"):
        return []
    for l in clean.splitlines():
        # a short line ENDING in COMPLETE is the verdict, however wrapped
        # ("COMPLETE", "The build is COMPLETE.") — never a missing feature.
        # \b before COMPLETE keeps "INCOMPLETE" out; the NOT test keeps
        # "NOT COMPLETE" from being read as the opposite of what it says.
        s = l.strip()
        if (re.fullmatch(r"[A-Za-z ,'—-]{0,40}\bCOMPLETE[.!]?", s)
                and not re.search(r"\bNOT\b", s.upper())):
            return []
    out = []
    for l in clean.splitlines():
        m = l.strip("-•* \t").strip()
        if len(m) <= 12 or m.endswith(":"):
            continue
        if _INSPECT_NOISE.search(m):
            continue
        # kept whole — a verbose model was having its real missing-feature
        # findings discarded for being wordy, then truncated for being long
        out.append(m)
    return out[:extra_steps()]


def _looks_elided(code, old=None):
    """True when a 'new version' of a file is really a fragment: placeholder
    comments standing in for code, or drastically shorter than the version it
    claims to replace."""
    if _ELIDE_PHRASE.search(code) or _ELIDE_DOTS.search(code):
        return True
    return bool(old) and len(code) < 0.5 * len(old)


def _extract_files(md_text, existing=None, default_name=None):
    """Parse fenced code blocks into {path: content}. Same heuristics as the
    compare page's extractFiles(): filename from the prose right before the
    fence, else an inline comment on the first line, else a generated name.
    Terminal-command blocks are skipped. Pass the project's current files as
    `existing` when parsing edits to a live project — unlabeled blocks then
    get junk-pattern names (which the merge guard rejects) instead of an
    'index.*' guess that could clobber a real file."""
    blocks, last = [], 0
    for m in re.finditer(r"```(\w*)[ \t]*\n([\s\S]*?)```", md_text):
        blocks.append(((m.group(1) or "").lower(), m.group(2), md_text[last:m.start()]))
        last = m.end()
    # A trailing fence that never closed means the generation was cut off
    # mid-file. Keep what arrived — losing it makes recovery impossible.
    tail = re.search(r"```(\w*)[ \t]*\n([\s\S]+)", md_text[last:])
    if tail:
        blocks.append(((tail.group(1) or "").lower(), tail.group(2),
                       md_text[last:last + tail.start()]))

    parsed = []
    for lang, code, before in blocks:
        if lang in _CMD_LANGS and not code.strip().startswith("#!/"):
            continue
        cands, near = [], []
        for line in reversed(before.rstrip().split("\n")[-4:]):
            line = re.sub(r"[`*_#>]+", " ", line).strip()
            hits = []
            for fm in re.finditer(
                    r"([A-Za-z0-9_\-]+(?:/[A-Za-z0-9_\-.]+)*\.[A-Za-z0-9]+)", line):
                cand = fm.group(1)
                if not cand.startswith("http") and len(cand) < 80:
                    hits.append(cand)
            # the FIRST line (nearest the fence) that names anything is the label
            if hits and not near and not cands:
                near = hits
            cands.extend(hits)
        first = code.split("\n", 1)[0]
        cm = re.search(r"(?://|#|<!--|/\*)\s*([A-Za-z0-9_\-./]+\.[A-Za-z0-9]+)", first)
        if cm and not cm.group(1).startswith("http"):
            cands.append(cm.group(1))
        parsed.append((_resolve_label(cands, existing, near), lang, code))

    # Name the leftovers in a SECOND pass so block ORDER is preserved (callers
    # rely on the last file being the one a truncated reply cut off).
    unresolved = sum(1 for n, _, _ in parsed if n is None)
    files = {}
    last_named = None
    for name, lang, code in parsed:
        # Models label a file ONCE and then emit a run of further edit blocks
        # for it: "**`js/canvas.js`**" then ```…``` ```…``` ```…```. The lines
        # above blocks 2..N are the previous block's ">>>>>>> REPLACE", not a
        # name, so they looked unlabeled and 12 blocks of real edits were
        # dropped as junk. A diff-only block with no label of its own belongs
        # to the file most recently named. Safe: if that guess is wrong the
        # SEARCH text simply won't match and the merge guard rejects it.
        is_diff = bool(_DIFF_BLOCK.search(code))
        if name is None and is_diff:
            # An edit block is self-identifying: its SEARCH text must already
            # exist in the file it patches. Look it up instead of guessing —
            # this beats label inheritance and is impossible to get wrong,
            # because a unique hit IS the proof.
            if existing:
                owners = set()
                for srch, _rep in _DIFF_BLOCK.findall(code):
                    if not srch.strip():
                        continue
                    hits = [k for k, v in existing.items() if srch in v]
                    if len(hits) == 1:
                        owners.add(hits[0])
                if len(owners) == 1:
                    name = owners.pop()
            if name is None and last_named:
                name = last_named
        if name is None:
            ext = _EXT_BY_LANG.get(lang, lang or "txt")
            if unresolved == 1 and default_name:
                # one unnamed block + one obvious target = that target.
                # Refusing it (as a junk fileN.<ext> the merge guard drops)
                # silently threw away entire steps.
                name = default_name
            elif existing:
                # An unlabeled block whose language matches EXACTLY ONE file in
                # the project is that file — but only when it is big enough to
                # be a real rewrite of it, so a brand-new file never clobbers
                # an existing one. Ambiguous languages stay junk on purpose.
                same = [k for k in existing
                        if k.rsplit(".", 1)[-1].lower() == ext.lower()]
                if len(same) == 1 and len(code) >= 0.5 * len(existing[same[0]]):
                    name = same[0]
            # Junk fallback LAST and unconditional: an earlier elif that failed
            # its own test must still end up with a name, or the block is
            # stored under the key None and every later lookup breaks.
            if name is None:
                name = (f"index.{ext}" if not files and not existing
                        else f"file{len(files) + 1}.{ext}")
        # Several edit blocks for one file ACCUMULATE — they are different
        # sections of the same patch, and overwriting would keep only the last.
        # For full files the old rule still holds: the same path emitted twice
        # is draft-then-correction, so the LAST one wins.
        if is_diff and name in files and _DIFF_BLOCK.search(files[name]):
            files[name] = files[name] + "\n" + code
        else:
            files[name] = code
        if not re.fullmatch(r"file\d+\.\w+", name):
            last_named = name

    # THE STEP-6 CLIFF: early steps CREATE files (fenced code blocks); later
    # steps EDIT them and the model switches to SEARCH/REPLACE — often with no
    # fence at all. Fenced-only parsing then returns {} and the entire reply is
    # discarded silently. Recover those blocks by attributing each to the
    # nearest filename above it, exactly as fenced blocks are labelled.
    if not files and _DIFF_BLOCK.search(md_text):
        groups = {}
        for m in _DIFF_BLOCK.finditer(md_text):
            cands = []
            for line in reversed(md_text[:m.start()].rstrip().split("\n")[-6:]):
                line = re.sub(r"[`*_#>]+", " ", line).strip()
                for fm in re.finditer(
                        r"([A-Za-z0-9_\-]+(?:/[A-Za-z0-9_\-.]+)*\.[A-Za-z0-9]+)",
                        line):
                    c = fm.group(1)
                    if not c.startswith("http") and len(c) < 80:
                        cands.append(c)
            name = _resolve_label(cands, existing) or default_name
            if name:
                groups.setdefault(name, []).append(m.group(0))
        for name, blocks in groups.items():
            files[name] = "\n".join(blocks)
    return files


def _files_context(files, limit=80000,
                   header="Current project files (complete, latest versions):"):
    parts, total, omitted = [], 0, []
    for name, code in files.items():
        ext = name.rsplit(".", 1)[-1].lower() if "." in name else ""
        block = f"**`{name}`**\n```{_LANG_BY_EXT.get(ext, '')}\n{code}\n```"
        if total + len(block) > limit:
            omitted.append(name)
            continue
        parts.append(block)
        total += len(block)
    out = header + "\n\n" + "\n\n".join(parts)
    if omitted:
        out += f"\n\n(omitted for length: {', '.join(omitted)})"
    return out


def _store_files(session_id, new_files, force=False):
    if not new_files:
        return
    with SESSIONS_LOCK:
        sess = sessions().setdefault(session_id, {})
        cur = sess.setdefault("files", {})
        for name, code in new_files.items():
            # never let a fragment or lazy rewrite clobber a complete version.
            # force=True is for pipeline output that already survived its own
            # merge guards: re-guarding it against the stale session copy made
            # legitimate shrinks (file splits, dead-code removal) silently
            # diverge session memory from the answer the user was shown.
            if not force and _looks_elided(code, cur.get(name)):
                print(f"[files] skipped elided/truncated {name} (session {session_id})")
                continue
            cur[name] = code
    _save_memory()


def _files_to_markdown(files, prose=None, unchanged=None, fixed_note=None, warnings=None):
    """Canonical answer format: optional prose, then one complete fenced block
    per file with its path on the line above."""
    parts = []
    if prose:
        parts.append(prose)
    for name, code in files.items():
        ext = name.rsplit(".", 1)[-1].lower() if "." in name else ""
        parts.append(f"**`{name}`**\n```{_LANG_BY_EXT.get(ext, '')}\n{code.rstrip()}\n```")
    if unchanged:
        parts.append(f"_Unchanged from earlier: {', '.join(unchanged)}_")
    if fixed_note:
        parts.append(f"🔧 Reviewer: {fixed_note}")
    if warnings:
        parts.append("⚠ Automated checks still flag:\n"
                     + "\n".join(f"- {w}" for w in warnings))
    return "\n\n".join(parts)


# ---- verification: cheap deterministic syntax checks --------------------------

def _node_check(label, code, node):
    suffix = ".mjs" if re.search(r"^\s*(import|export)\s", code, re.M) else ".js"
    path = None
    try:
        with tempfile.NamedTemporaryFile("w", suffix=suffix, delete=False) as f:
            f.write(code)
            path = f.name
        p = subprocess.run([node, "--check", path],
                           capture_output=True, text=True, timeout=20)
        if p.returncode != 0:
            err = (p.stderr or p.stdout).strip().splitlines()
            # node prints location/source/caret first and the actual
            # 'SyntaxError: ...' line several lines later — keep location
            # AND message; the first three lines alone amputate the message
            # and leave every fix prompt guessing at an unnamed error
            msg = [l for l in err if re.search(r"\w+Error\b", l)]
            keep = ([err[0]] if err else []) + [m for m in msg if m != (err[0] if err else "")]
            return [f"{label}: " + " | ".join(keep or err[:3])[:300]]
    except Exception:
        pass
    finally:
        if path:
            try:
                os.unlink(path)
            except OSError:
                pass
    return []


def _static_check_files(files):
    """Deterministic syntax checks run before the reviewer model sees the
    draft — catches what an LLM reviewer glosses over. JS checks need node on
    PATH and are skipped silently without it."""
    issues = []
    node = shutil.which("node")
    for name, code in files.items():
        ext = name.rsplit(".", 1)[-1].lower() if "." in name else ""
        # ES-module syntax in a classic-script project is a guaranteed
        # browser SyntaxError — and node --check alone can't catch it here,
        # because _node_check validates import/export files AS modules.
        if ext == "js" and re.search(r"^\s*(import|export)\s", code, re.M):
            issues.append(f"{name}: uses ES module import/export — builds "
                          f"load classic <script> tags; rewrite without "
                          f"module syntax (see design standards)")
        if ext == "py":
            try:
                compile(code, name, "exec")
            except SyntaxError as e:
                issues.append(f"{name}: Python syntax error line {e.lineno}: {e.msg}")
        elif ext in ("js", "mjs") and node:
            issues.extend(_node_check(name, code, node))
        elif ext in ("html", "htm") and node:
            for i, script in enumerate(re.findall(
                    r"<script(?![^>]*\bsrc=)[^>]*>([\s\S]*?)</script>", code, re.I)):
                if script.strip():
                    issues.extend(_node_check(f"{name} <script #{i + 1}>", script, node))
    return issues

def _symbol_manifest(files):
    """Mechanical one-line-per-file map of top-level names — cheap cross-step
    memory. Every step sees WHAT exists project-wide (and what it's called)
    without paying to see every file's contents. Extracted by regex, so it
    can't hallucinate and can't miss the way retrieval could."""
    lines = []
    for name in sorted(files):
        code = files[name]
        if name.endswith((".js", ".mjs")):
            pairs = re.findall(
                r"^(?:class|function)\s+([A-Za-z_$][\w$]*)"
                r"|^(?:const|let|var)\s+([A-Za-z_$][\w$]*)", code, re.M)
            syms = ", ".join(dict.fromkeys(a or b for a, b in pairs))[:220]
            lines.append(f"{name}: {syms}" if syms else f"{name}: (no top-level names)")
        elif name.endswith((".html", ".htm")):
            ids = list(dict.fromkeys(re.findall(r'\bid=["\']([\w-]+)', code)))[:20]
            srcs = re.findall(r'<script[^>]*\bsrc=["\']([^"\']+)', code, re.I)
            lines.append(f"{name}: ids: {', '.join(ids)}; loads: {', '.join(srcs)}")
    return "\n".join(lines)


def _dup_symbol_check(files):
    """Two files declaring the same top-level name is a guaranteed
    SyntaxError (let/class/const) or a silent clobber (var/function) once
    classic scripts share the page's global scope. Deterministic — this is
    the check that would have caught the Player saga at the step that
    caused it."""
    owners, issues = {}, []
    for name in sorted(files):
        if not name.endswith((".js", ".mjs")):
            continue
        pairs = re.findall(
            r"^(?:class|function)\s+([A-Za-z_$][\w$]*)"
            r"|^(?:const|let|var)\s+([A-Za-z_$][\w$]*)", files[name], re.M)
        for a, b in pairs:
            s = a or b
            if s in owners and owners[s] != name:
                issues.append(f"{name}: top-level '{s}' is already declared "
                              f"in {owners[s]} — one owner per name")
            else:
                owners.setdefault(s, name)
    return issues


def _bundle_check(files):
    """Cross-file JS check: for each html, node-check the concatenation of its
    classic <script src> files in include order. Catches what per-file checks
    can't — duplicate top-level declarations across files ('redeclaration of
    let X') and load-order breaks — plus double-included scripts."""
    node = shutil.which("node")
    if not node:
        return []
    issues = []
    for hname, html in files.items():
        if not hname.endswith((".html", ".htm")):
            continue
        srcs = re.findall(r"<script[^>]*\bsrc=[\"']([^\"']+)[\"']", html, re.I)
        for s in set(srcs):
            if srcs.count(s) > 1:
                issues.append(f"{hname}: includes {s} more than once")
        parts = []
        hdir = os.path.dirname(hname)
        for s in srcs:
            s2 = s.lstrip("./")
            # resolve the way a browser would: relative to the including html
            # first — a root-level game.js must not shadow pages/game.js
            rel = os.path.normpath(os.path.join(hdir, s2)) if hdir else s2
            tgt = (rel if rel in files else
                   s2 if s2 in files else next(
                       (n for n in files
                        if n.endswith("/" + s2)
                        or os.path.basename(n) == os.path.basename(s2)), None))
            if tgt and tgt.rsplit(".", 1)[-1].lower() in ("js", "mjs"):
                parts.append(f"// ==== {tgt} ====\n{files[tgt]}")
            elif not tgt and not s.startswith(("http://", "https://", "//")):
                # deterministic, free, and catches the exact failure the
                # foreman was burning whole minutes to spot: a step that
                # referenced a file it never wrote.
                issues.append(f"{hname}: <script src=\"{s}\"> does not exist "
                              f"in the project — the page will not run")
        if len(parts) > 1:
            issues.extend(_node_check(
                f"{hname} scripts loaded together", "\n".join(parts), node))
    return issues


SUMMARIZER_SYSTEM = (
    "You maintain a concise conversation memory. "
    "Given the current memory (if any) and a new exchange, produce an updated summary "
    "in 3-5 sentences max. Capture key topics discussed, any preferences or context the "
    "user revealed, and important facts established. Be factual and terse. "
    "Output only the updated summary, nothing else."
)

# ---- router settings: backends / routes / roles ------------------------------
# Public-config layer: everything the classifier and routing table know comes
# from ROUTER_CONFIG (backends with URL+key+model, routes with subject/persona,
# system roles). Saved to router_config.json, editable at /settings, applied
# LIVE — module globals (SPECIALISTS, CLASSIFIER, VALID, CLASSIFIER_SYSTEM…)
# are rebuilt on every save; no restart needed. Defaults mirror the original
# hardcoded fleet so first run changes nothing.
ROUTER_CONFIG_FILE = os.path.expanduser("~/router_config.json")
CONFIG_LOCK = threading.Lock()
ROUTE_MODELS: dict = {}      # route -> per-route model override
BACKEND_INFO: dict = {}      # chat-url -> {"key","model"} for request time
ROUTER_CONFIG: dict = {}

# Tuned classifier rule lines for the built-in routes; custom routes get a
# generated line from their 'subject'. Priority order preserved from the
# original prompt: audio, search, recipes, code, then customs, then the rest.
_BUILTIN_RULES = {
    "audio": ("- 'audio': the user wants to HEAR something — a voiceover, "
              "narration, podcast, spoken recording, audio clip, or wants you "
              "to read/say/record something aloud or 'in your voice'. Choose "
              "audio whenever the desired output is sound/speech, even if a "
              "script is involved.\n"),
    "search": ("- 'search': the answer needs current, real-world, or external "
               "information — news, recent events, prices, weather, sports "
               "scores, release dates, 'look up', 'search for', 'what's the "
               "latest', or any fact likely to have changed since training or "
               "that you can't be sure of. Choose search whenever fresh or "
               "verifiable web information would matter.\n"),
    "recipes": "- 'recipes': food, cooking, baking, drinks, meal planning.\n",
    "code": ("- 'code': ANYTHING asking to build, create, make, generate, "
             "write, or fix something that produces a working artifact — "
             "website, app, program, function, HTML, CSS, JSON, SQL, etc. "
             "If the user wants something BUILT or CODED, always choose code, "
             "even if the topic sounds creative. Also choose code when the "
             "user reports errors, stack traces, or that something previously "
             "built is broken or needs fixing.\n"),
    "review": ("- 'review': the user wants EXISTING code REVIEWED — 'suggest "
               "improvements', 'code review', 'critique this code', 'audit', "
               "'how would you refactor this', especially about an attached "
               "repository or pasted code. They want FINDINGS about code that "
               "already exists, not something new built and not a specific "
               "reported bug fixed in a live build (those are 'code').\n"),
    "summarize": ("- 'summarize': the user wants EXISTING material condensed — "
                  "'summarise this', 'tl;dr', 'key points', 'what does this "
                  "say', 'brief me on this'. Choose summarize when the user "
                  "is pointing at text they already have (pasted in, attached "
                  "as a file, or the conversation so far) and wants it made "
                  "shorter. If they want NEW information researched or a "
                  "question answered from knowledge, that is not summarize.\n"),
    "reasoning": ("- 'reasoning': math, logic puzzles, science explanations, "
                  "step-by-step analysis, comparisons.\n"),
    "creative": ("- 'creative': writing prose, stories, poems, roleplay, or "
                 "describing/imagining something where NO working code or "
                 "artifact is being requested.\n"),
    "general": ("- 'general': factual questions, advice, definitions, "
                "conversational exchanges.\n"),
}
_BUILTIN_ROUTES = ("recipes", "creative", "code", "general", "reasoning",
                   "search", "summarize", "review")
_ROUTE_NAME_RE = re.compile(r"^[a-z][a-z0-9]{1,15}$")
_CUSTOM_COLOR_POOL = ["#7a4ba0", "#8c2a5e", "#4a6b1f", "#1f6f8c",
                      "#96431c", "#5e5a1f", "#265e8c", "#8c264d"]


# Backend kinds. "openai" is the generic OpenAI-compatible shape and the
# default, so a config written before kinds existed behaves identically.
# llama.cpp and Ollama both speak that shape — they differ only in how you ask
# them their context size and whether a model name is required. Anthropic is
# the genuinely different one: a different path, different auth headers, and a
# different request AND response body.
BACKEND_KINDS = ("openai", "llamacpp", "ollama", "anthropic")

ANTHROPIC_VERSION = "2023-06-01"
# Anthropic requires max_tokens on every request; OpenAI treats it as optional
# and Crewe often omits it (or uses 0 to mean "uncapped"). Neither is legal
# there, so an omitted cap becomes this.
# 32768, not 8192: this fills in when Crewe would send no cap (its 0 means
# "uncapped", which Anthropic does not allow). A whole-file rewrite can pass
# 8k output tokens, and this project already learned the hard way that a cap
# binding MID-FILE destroys work — the truncated file passes the elision
# guard and overwrites the good version. Anthropic bills actual tokens, not
# the cap, so the higher ceiling costs nothing; every current Claude model
# allows at least 64k out.
ANTHROPIC_DEFAULT_MAX_TOKENS = int(os.environ.get("CREWE_ANTHROPIC_MAX_TOKENS", "32768"))


def _kind(backend_or_url):
    if isinstance(backend_or_url, dict):
        k = backend_or_url.get("kind") or "openai"
    else:
        k = (BACKEND_INFO.get(backend_or_url, {}) or {}).get("kind") or "openai"
    return k if k in BACKEND_KINDS else "openai"


def _chat_url(backend):
    base = backend["url"].rstrip("/")
    if _kind(backend) == "anthropic":
        return base + "/v1/messages"
    return base + "/v1/chat/completions"


def _default_router_config():
    """Mirrors the classic hardcoded fleet — first run behaves identically."""
    backends = [
        {"id": "recipes-e4b",    "url": f"http://{HOST}:8080",       "key": "", "model": ""},
        {"id": "creative-e4b",   "url": f"http://{HOST}:8081",       "key": "", "model": ""},
        {"id": "code-e4b",       "url": f"http://{HOST}:8082",       "key": "", "model": ""},
        {"id": "brain-26b",      "url": f"http://{HOST}:8087",       "key": "", "model": ""},
        {"id": "classifier-e2b", "url": f"http://{HOST}:8083",       "key": "", "model": ""},
        {"id": "summarizer-e4b", "url": f"http://{HOST}:8084",       "key": "", "model": ""},
        {"id": "coder-qwen",     "url": f"http://{CODER_HOST}:8080", "key": "", "model": ""},
    ]
    rb = {"recipes": "recipes-e4b", "creative": "creative-e4b",
          "code": "coder-qwen", "general": "brain-26b",
          "reasoning": "brain-26b", "search": "brain-26b",
          # the biggest local window — documents are long
          "summarize": "brain-26b",
          # code review wants the strongest reader you have
          "review": "brain-26b"}
    routes = {n: {"backend": rb[n], "model": "", "subject": "", "persona": "",
                  "custom": False, "color": ""} for n in _BUILTIN_ROUTES}
    return {"backends": backends, "routes": routes,
            "roles": {"classifier": "classifier-e2b",
                      "summarizer": "summarizer-e4b",
                      "coder": "coder-qwen",
                      "coder_fast": "code-e4b",
                      "coder_quick": "code-e4b"}}


def _load_router_config():
    try:
        with open(ROUTER_CONFIG_FILE) as f:
            cfg = json.load(f)
        assert isinstance(cfg.get("backends"), list)
        assert isinstance(cfg.get("routes"), dict)
        assert isinstance(cfg.get("roles"), dict)
        return _backfill_builtins(cfg, save=True)
    except Exception:
        return _default_router_config()


def _backfill_builtins(cfg, save=False):
    """Add built-in routes the saved config predates.

    A config written before a route existed simply lacks it, and saved routes
    win over defaults — so without this, shipping a new built-in route would be
    invisible to every existing install until someone hand-edited their JSON.
    Only ADDS — an existing entry is never touched. (The settings UI only
    offers delete for CUSTOM routes, so this cannot resurrect something the
    user deliberately removed.) The default backend may not exist in this
    user's config, so fall back to whatever `general` uses before adding a
    route that would silently be skipped for pointing at nothing."""
    defaults = _default_router_config()
    have = {b.get("id") for b in cfg["backends"]}
    fallback = (cfg["routes"].get("general") or {}).get("backend")
    added = [n for n in _BUILTIN_ROUTES if n not in cfg["routes"]]
    for n in added:
        r = dict(defaults["routes"][n])
        if r.get("backend") not in have:
            if fallback not in have:
                continue                 # nothing sane to point it at
            r["backend"] = fallback
        cfg["routes"][n] = r
    added = [n for n in added if n in cfg["routes"]]
    if added:
        print(f"[config] added new built-in route(s): {', '.join(added)}")
        # Only the loader persists. Callers inspecting a HYPOTHETICAL config
        # must not have it written to disk as a side effect — doing so once
        # overwrote the live config with a test fixture.
        if save:
            _save_router_config(cfg)
    return cfg


def _save_router_config(cfg):
    tmp = f"{ROUTER_CONFIG_FILE}.tmp{os.getpid()}"
    with open(tmp, "w") as f:
        json.dump(cfg, f, indent=2)
    os.replace(tmp, ROUTER_CONFIG_FILE)


def _build_classifier_system(cfg):
    names = [n for n in cfg["routes"]] + ["audio"]
    # Priority-ordered, then customs, then EVERY remaining configured route.
    # This used to be a fixed list, so a built-in added later (summarize,
    # review) appeared in the answer-word list but its RULE line was silently
    # dropped — the classifier was guessing from the route's name alone.
    head = ["audio", "search", "recipes", "code"]
    customs = [n for n in cfg["routes"] if cfg["routes"][n].get("custom")]
    tail = ["reasoning", "creative", "general"]
    middle = [n for n in cfg["routes"]
              if n not in head + customs + tail]
    order = head + customs + sorted(middle) + tail
    lines = ""
    for n in order:
        if n != "audio" and n not in cfg["routes"]:
            continue
        r = cfg["routes"].get(n, {})
        if n in _BUILTIN_RULES and not r.get("custom"):
            lines += _BUILTIN_RULES[n]
        else:
            subj = (r.get("subject") or n).strip().rstrip(".")
            lines += f"- '{n}': {subj}.\n"
    return (
        "You are a routing classifier. Read the user's message and respond "
        "with EXACTLY ONE word, lowercase, no punctuation: "
        + ", ".join(names[:-1]) + ", or " + names[-1] + ". "
        "Rules in priority order:\n" + lines +
        "Context rule (overrides all of the above): if a previous exchange "
        "is shown and the new message continues it — answers a question the "
        "assistant asked, says yes/no, gives the next move in a game, or "
        "refers back with 'it/that/again' — repeat the previous route. Only "
        "switch routes when the new message clearly starts a different kind "
        "of task.\nOutput only the single word."
    )


def _apply_router_config(cfg):
    """Swap the live routing tables to match cfg. Called at import and on
    every settings save — readers pick up the new dicts atomically."""
    global SPECIALISTS, VALID, CLASSIFIER, SUMMARIZER, CODE_AGENT_URL, \
        CODE_AGENT_URL_FAST, CODE_AGENT_URL_QUICK, \
        CLASSIFIER_SYSTEM, BACKEND_INFO, ROUTER_CONFIG, ROUTE_MODELS
    with CONFIG_LOCK:
        b_by_id = {b["id"]: b for b in cfg["backends"]}
        spec, binfo, rmodels = {}, {}, {}
        for b in cfg["backends"]:
            binfo[_chat_url(b)] = {"key": b.get("key", ""),
                                   "model": b.get("model", ""),
                                   "kind": b.get("kind") or "openai",
                                   "paid": bool(b.get("paid")),
                                   "price_in": b.get("price_in", 0),
                                   "price_out": b.get("price_out", 0)}
        for name, r in cfg["routes"].items():
            b = b_by_id.get(r.get("backend"))
            if not b:
                continue
            spec[name] = _chat_url(b)
            if r.get("model"):
                rmodels[name] = r["model"]
            if r.get("persona"):
                SYSTEM_PROMPTS[name] = r["persona"]
            elif r.get("custom") and name not in SYSTEM_PROMPTS:
                SYSTEM_PROMPTS[name] = (
                    "You are a helpful, knowledgeable specialist assistant "
                    f"for: {r.get('subject') or name}. Answer clearly and "
                    "practically.")
        roles = cfg.get("roles", {})
        def role_url(role, fallback):
            b = b_by_id.get(roles.get(role, ""))
            return _chat_url(b) if b else fallback
        SPECIALISTS = spec
        VALID = set(spec) | {"audio"}
        CLASSIFIER = role_url("classifier", CLASSIFIER)
        SUMMARIZER = role_url("summarizer", SUMMARIZER)
        CODE_AGENT_URL = role_url("coder", CODE_AGENT_URL)
        # Easy-effort coder. Absent in configs written before effort existed,
        # in which case both levels use the deep coder -> identical behaviour.
        CODE_AGENT_URL_FAST = role_url("coder_fast", CODE_AGENT_URL)
        # The fast tier falls back to normal, which falls back to extra — an
        # older config with fewer coder roles keeps working at every level.
        CODE_AGENT_URL_QUICK = role_url("coder_quick", CODE_AGENT_URL_FAST)
        CLASSIFIER_SYSTEM = _build_classifier_system(cfg)
        BACKEND_INFO = binfo
        ROUTE_MODELS = rmodels
        ROUTER_CONFIG = cfg
    # A settings save can repoint either coder, so budgets must follow them.
    # Done outside CONFIG_LOCK: it probes each backend over the network.
    if _BUDGETS_READY:
        _refresh_budgets()


def _hdrs(url):
    """Auth headers for a chat URL, from its backend's stored API key.

    Anthropic does not use bearer auth — it wants x-api-key plus a version
    header, and sending Authorization instead is a 401 with a confusing
    message."""
    k = BACKEND_INFO.get(url, {}).get("key", "")
    if _kind(url) == "anthropic":
        h = {"anthropic-version": ANTHROPIC_VERSION}
        if k:
            h["x-api-key"] = k
        return h
    return {"Authorization": f"Bearer {k}"} if k else {}


# ---------------------------------------------------------------------------
# Cost tracking for paid backends.
#
# Crewe has always been able to talk to hosted APIs — a backend has a key and a
# model name. What it could not do is tell you that a question was about to
# spend money, or how much it had spent. That is the whole point of this
# module: make "is this request free?" a visible property rather than something
# you have to remember about your own config.
#
# Prices are per MILLION tokens, in whatever currency you type, because every
# provider quotes them that way. Crewe does not know or fetch prices — you set
# them per backend, and they are only ever used for an ESTIMATE.
SPEND_LOCK = threading.Lock()


def _bmeta(url):
    return BACKEND_INFO.get(url, {}) or {}


def _is_paid(url):
    return bool(_bmeta(url).get("paid"))


def _prices(url):
    b = _bmeta(url)
    try:
        return float(b.get("price_in") or 0), float(b.get("price_out") or 0)
    except (TypeError, ValueError):
        return 0.0, 0.0


def _with_usage(payload, url):
    """Ask a paid backend to report token usage on the final stream chunk.

    Only sent to paid backends: llama.cpp ignores unknown fields, but there is
    no reason to change the request shape for a local server we bill nothing
    for."""
    # Anthropic reports usage unconditionally and rejects stream_options.
    if _is_paid(url) and payload.get("stream") and _kind(url) != "anthropic":
        payload.setdefault("stream_options", {"include_usage": True})
    return payload


def _rough_tokens(text):
    """~4 chars per token. Only used when a provider reports no usage."""
    return max(1, len(text or "") // 4)


def record_spend(url, route, usage, prompt_text="", answer_text=""):
    """Append one spend record. Returns the estimated cost (0.0 if free).

    Never raises: a billing-log failure must not lose the user's answer."""
    try:
        if not _is_paid(url):
            return 0.0
        p_in, p_out = _prices(url)
        estimated = not usage
        if usage:
            tin = int(usage.get("prompt_tokens") or 0)
            tout = int(usage.get("completion_tokens") or 0)
        else:
            tin, tout = _rough_tokens(prompt_text), _rough_tokens(answer_text)
        cost = (tin * p_in + tout * p_out) / 1_000_000.0
        rec = {"ts": time.time(), "route": route, "url": url,
               "model": _bmeta(url).get("model", ""), "in": tin, "out": tout,
               "cost": round(cost, 6), "estimated": estimated}
        with SPEND_LOCK:
            with open(spend_file(), "a", encoding="utf-8") as f:
                f.write(json.dumps(rec) + "\n")
        return cost
    except Exception as e:
        print(f"[spend] could not record: {e}")
        return 0.0


def spend_summary(uid=None):
    """Totals for the current user. Cheap enough to call per page load."""
    now = time.time()
    day, month = now - 86400, now - 30 * 86400
    out = {"day": 0.0, "month": 0.0, "all": 0.0, "requests": 0,
           "estimated_any": False, "recent": []}
    try:
        with open(spend_file(uid), encoding="utf-8") as f:
            for line in f:
                try:
                    r = json.loads(line)
                except Exception:
                    continue
                c = float(r.get("cost") or 0)
                out["all"] += c
                out["requests"] += 1
                if r.get("estimated"):
                    out["estimated_any"] = True
                if r["ts"] >= month:
                    out["month"] += c
                if r["ts"] >= day:
                    out["day"] += c
                out["recent"].append(r)
    except OSError:
        return out
    out["recent"] = out["recent"][-20:][::-1]
    for k in ("day", "month", "all"):
        out[k] = round(out[k], 4)
    return out


# Fields Crewe sends that Anthropic rejects outright. temperature/top_p/top_k
# were REMOVED on the current Claude models (Opus 5, Opus 4.7/4.8, Sonnet 5,
# Fable 5) — sending one is a 400, not a warning. Crewe sets temperature on
# almost every call, so stripping these is what makes the route work at all.
# Anthropic's guidance is to steer with prompting instead.
_ANTHROPIC_DROP = ("temperature", "top_p", "top_k", "chat_template_kwargs",
                   "stream_options", "presence_penalty", "frequency_penalty",
                   "n", "stop", "logit_bias", "seed", "response_format")


def _to_anthropic(payload):
    """Translate an OpenAI-shaped chat payload into Anthropic's Messages shape.

    Three things actually differ, and all three break the request if missed:
      * system prompts are a TOP-LEVEL string, not a message with role=system
      * max_tokens is REQUIRED (OpenAI treats it as optional; Crewe uses 0 to
        mean uncapped, which is not a legal value here)
      * sampling parameters are rejected — see _ANTHROPIC_DROP
    """
    out = {k: v for k, v in payload.items() if k not in _ANTHROPIC_DROP}

    systems, msgs = [], []
    for m in payload.get("messages") or []:
        role = m.get("role")
        content = m.get("content", "")
        if role == "system":
            if content:
                systems.append(content if isinstance(content, str) else str(content))
        else:
            # Anthropic accepts only user/assistant here; anything else would
            # 400, so an unknown role is treated as the user speaking.
            msgs.append({"role": "assistant" if role == "assistant" else "user",
                         "content": content})
    # A conversation must start with a user turn.
    while msgs and msgs[0]["role"] != "user":
        msgs.pop(0)
    if not msgs:
        msgs = [{"role": "user", "content": " "}]
    out["messages"] = msgs
    if systems:
        out["system"] = "\n\n".join(systems)

    cap = payload.get("max_tokens")
    out["max_tokens"] = int(cap) if cap else ANTHROPIC_DEFAULT_MAX_TOKENS
    return out


def _inject_model(payload, url, route=None):
    """Add the model name a backend/route needs, then shape the body for the
    backend's kind. Every backend POST in Crewe passes through here, which is
    why the Anthropic translation hangs off it rather than off each caller."""
    m = (ROUTE_MODELS.get(route) if route else None) \
        or BACKEND_INFO.get(url, {}).get("model", "")
    if m and "model" not in payload:
        payload["model"] = m
    if _kind(url) == "anthropic":
        payload = _to_anthropic(payload)
    return payload


CLASSIFIER_SYSTEM = ""   # set by _apply_router_config below
ROUTER_CONFIG = _load_router_config()
_apply_router_config(ROUTER_CONFIG)


def _backend_ctx(url, timeout=6):
    """The backend's REAL context window, straight from the server.

    Every kind answers this differently: llama.cpp has /props, Ollama has
    /api/show, and Anthropic publishes max_input_tokens on /v1/models. Asking
    the wrong one just returns None, which leaves the default budgets in
    place — wrong numbers would be worse than no numbers."""
    base = url.rsplit("/v1/", 1)[0]
    kind = _kind(url)
    try:
        if kind == "anthropic":
            model = (BACKEND_INFO.get(url, {}) or {}).get("model", "")
            r = requests.get(f"{base}/v1/models", headers=_hdrs(url), timeout=timeout)
            for m in (r.json().get("data") or []):
                if not model or m.get("id") == model:
                    n = m.get("max_input_tokens")
                    if n:
                        return int(n)
            return None
        if kind == "ollama":
            model = (BACKEND_INFO.get(url, {}) or {}).get("model", "")
            if model:
                r = requests.post(f"{base}/api/show", json={"model": model},
                                  timeout=timeout)
                info = r.json().get("model_info") or {}
                for k, v in info.items():
                    if k.endswith(".context_length") and v:
                        return int(v)
            return None
        d = requests.get(f"{base}/props", timeout=timeout).json()
        n = (d.get("default_generation_settings") or {}).get("n_ctx") or d.get("n_ctx")
        return int(n) if n else None
    except Exception:
        return None


BUDGET_CAP_CHARS = int(os.environ.get("CREWE_BUDGET_CAP_CHARS", "160000"))


def _budgets_for(url):
    """Size the file budgets to whatever coder is plugged into `url`.

    Hard-coded budgets are a trap when the backend can be swapped from the
    settings page: numbers tuned for a 128k model silently overflow a 32k one,
    llama.cpp truncates the prompt FROM THE FRONT (losing the system prompt and
    the project), and the model looks stupid for reasons that are entirely our
    fault. An explicit CREWE_* env var always wins.

    Returns (budgets_dict, ctx_or_None). Pure -- callers store the result, so
    each effort level gets budgets matched to ITS OWN coder."""
    default = {"small_project_chars": SMALL_PROJECT_CHARS,
               "step_file_chars":     STEP_FILE_CHARS,
               "step_file_hard":      STEP_FILE_HARD}
    if any(os.environ.get(k) for k in
           ("CREWE_WHOLE_PROJECT_CHARS", "CREWE_STEP_FILE_CHARS",
            "CREWE_STEP_FILE_HARD")):
        print("[code] file budgets pinned by env — not autosizing")
        return default, None
    n = _backend_ctx(url)
    if not n:
        print(f"[code] coder ctx unknown ({url}) — keeping "
              f"{SMALL_PROJECT_CHARS:,}-char budgets")
        return default, None
    # ~30% of the window for project files (~4 chars/token), leaving room for
    # the plan, working memory, the step text and the whole generation.
    chars = int(n * 0.30) * 4
    # CAPPED. The 30% rule was written for local models, where context is
    # free. Fed Claude's 1,000,000-token window it produced a 1.2M-char
    # budget — ~300k tokens of prompt PER STEP on a metered API. Past ~40k
    # tokens of project text more context stops helping anyway. The cap is
    # the long-standing 160k default, tunable for whoever disagrees.
    if chars > BUDGET_CAP_CHARS:
        print(f"[code] coder ctx {n:,} tokens -> autosized budget "
              f"{chars:,} chars CAPPED at {BUDGET_CAP_CHARS:,}")
        chars = BUDGET_CAP_CHARS
    else:
        print(f"[code] coder ctx {n:,} tokens -> whole-project budget "
              f"{chars:,} chars (~{chars // 4:,} tokens), hard {chars * 2:,}")
    return {"small_project_chars": chars,
            "step_file_chars":     chars,
            "step_file_hard":      chars * 2}, n


CODER_CTX = None


def _classify_with(system, valid, question, timeout=60):
    """One classification with an explicit prompt+vocabulary — used by the
    settings page's live routing test."""
    payload = {"messages": [{"role": "system", "content": system},
                            {"role": "user", "content": question}],
               "temperature": 0.0, "max_tokens": 1024, "stream": False}
    try:
        r = requests.post(CLASSIFIER, json=_inject_model(payload, CLASSIFIER),
                          headers=_hdrs(CLASSIFIER), timeout=timeout)
        r.raise_for_status()
        content = _resp_json(r)["choices"][0]["message"]["content"] or ""
        content = _strip_think(content)
        for word in reversed(re.findall(r"[a-z]+", content.lower())):
            if word in valid:
                return word
    except Exception as e:
        print(f"[classify-test] error: {e}")
    return "(no answer)"


def classify(question: str, prev_q: str = "", prev_a: str = "", last_route: str = "") -> str:
    """Ask the classifier to pick a single-word route. Falls back to 'general'."""
    user_content = question
    if prev_q or prev_a:
        user_content = (
            f"Previous exchange (routed to '{last_route or 'unknown'}'):\n"
            f"User: {prev_q[:300]}\n"
            f"Assistant: {prev_a[:300]}\n\n"
            f"New message to classify:\n{question}"
        )
    payload = {
        "messages": [
            {"role": "system", "content": CLASSIFIER_SYSTEM},
            {"role": "user", "content": user_content},
        ],
        "temperature": 0.0,
        # Generous budget: reasoning models think before emitting the route word,
        # and that thinking counts against max_tokens. The old cap of 8 starved
        # them into returning empty content -> silent 'general' fallback.
        "max_tokens": 1024,
        "stream": False,
    }
    try:
        # The classifier shares the 12B with generation work. If a generation is
        # mid-flight the classify call queues behind it, so allow generous time —
        # waiting and routing correctly beats timing out into a silent 'general'.
        r = requests.post(CLASSIFIER, json=_inject_model(payload, CLASSIFIER),
                          headers=_hdrs(CLASSIFIER), timeout=120)
        r.raise_for_status()
        content = _resp_json(r)["choices"][0]["message"]["content"] or ""
        # Reasoning models may leak think-text into content; the route is
        # whatever valid word appears last.
        content = _strip_think(content)
        for word in reversed(re.findall(r"[a-z]+", content.lower())):
            if word in VALID:
                return word
    except Exception as e:
        print(f"[classify] error: {e}")
    # fallback bucket if the classifier output can't be parsed
    return "general"


def summarize(session_id: str, question: str, answer: str):
    """Update the rolling session summary using the E4B summarizer. Fire-and-forget."""
    with SESSIONS_LOCK:
        current = sessions().get(session_id, {}).get("summary", "")
        turns   = sessions().get(session_id, {}).get("turns", 0)

    user_content = (
        f"Current memory:\n{current}\n\nNew exchange:\nUser: {question}\nAssistant: {answer[:800]}"
        if current else
        f"New exchange:\nUser: {question}\nAssistant: {answer[:800]}"
    )
    payload = {
        "messages": [
            {"role": "system",  "content": SUMMARIZER_SYSTEM},
            {"role": "user",    "content": user_content},
        ],
        "temperature": 0.3,
        "max_tokens": 200,
        "stream": False,
    }
    try:
        r = requests.post(SUMMARIZER, json=_inject_model(payload, SUMMARIZER),
                          headers=_hdrs(SUMMARIZER), timeout=60)
        r.raise_for_status()
        new_summary = _resp_json(r)["choices"][0]["message"]["content"].strip()
        with SESSIONS_LOCK:
            # update in place — the session dict also carries project files
            sess = sessions().setdefault(session_id, {})
            sess["summary"] = new_summary
            sess["turns"] = turns + 1
        _save_memory()
    except Exception as e:
        print(f"[summarize] error: {e}")


def web_search(query: str) -> list:
    """Query the local SearXNG JSON API. Returns the raw result list (possibly empty)."""
    try:
        r = requests.get(SEARXNG_URL,
                         params={"q": query, "format": "json"},
                         timeout=10)
        r.raise_for_status()
        return r.json().get("results", []) or []
    except Exception as e:
        print(f"[search] searxng error: {e}")
        return []


_TAG_RE   = re.compile(r"(?s)<[^>]+>")
_DROP_RE  = re.compile(r"(?is)<(script|style|noscript|template|head)[^>]*>.*?</\1>")
_WS_RE    = re.compile(r"\s+")

def _pdf_to_text(data: bytes) -> str:
    """Extract text from a PDF via poppler's pdftotext (read from stdin)."""
    try:
        p = subprocess.run(["pdftotext", "-q", "-", "-"], input=data,
                          capture_output=True, timeout=15)
        return _WS_RE.sub(" ", p.stdout.decode("utf-8", "ignore")).strip()
    except Exception as e:
        print(f"[search] pdftotext error: {e}")
        return ""


def _url_is_public(url: str) -> bool:
    """SSRF guard for the search fetcher: only http(s) to a host that resolves
    to a PUBLIC address. Blocks a crafted search result from reaching internal
    services — the llama ports on this box, or a cloud metadata endpoint
    (169.254.169.254) if this ever runs on a VM. Checks EVERY resolved address,
    so a name that returns one public and one private IP is still rejected."""
    try:
        p = urlparse(url)
        if p.scheme not in ("http", "https") or not p.hostname:
            return False
        infos = socket.getaddrinfo(p.hostname, p.port or (443 if p.scheme == "https" else 80),
                                   proto=socket.IPPROTO_TCP)
        if not infos:
            return False
        for *_, sockaddr in infos:
            ip = ipaddress.ip_address(sockaddr[0])
            if (ip.is_private or ip.is_loopback or ip.is_link_local
                    or ip.is_reserved or ip.is_multicast or ip.is_unspecified):
                return False
        return True
    except Exception:
        return False


def _fetch_page_text(url: str) -> str:
    """Best-effort fetch + readable-text extraction. Handles HTML and PDF.

    Redirects are followed manually so each hop is re-checked against the SSRF
    guard — otherwise a public page could 302 the fetcher to an internal host."""
    try:
        hops = 0
        while True:
            if not _url_is_public(url):
                return ""      # internal/loopback/link-local target — refuse (SSRF)
            r = requests.get(url, timeout=8, allow_redirects=False, headers={
                "User-Agent": "Mozilla/5.0 (compatible; router-search/1.0)"})
            if r.is_redirect and r.headers.get("location") and hops < 4:
                url = requests.compat.urljoin(url, r.headers["location"])
                hops += 1
                continue
            break
        r.raise_for_status()
        ctype = r.headers.get("content-type", "").lower()
        if "pdf" in ctype or url.lower().endswith(".pdf"):
            return _pdf_to_text(r.content)
        if "html" not in ctype and "text" not in ctype:
            return ""
        html = _DROP_RE.sub(" ", r.text)
        text = _TAG_RE.sub(" ", html)
        text = text.replace("&nbsp;", " ").replace("&amp;", "&")
        return _WS_RE.sub(" ", text).strip()
    except Exception as e:
        print(f"[search] fetch error {url}: {e}")
        return ""


# ---------------------------------------------------------------------------
# Uploaded documents.
#
# A file attaches to the SESSION, not to one question: once uploaded it stays
# available to whatever route the NEXT question lands on, so "summarise this"
# and "what does clause 4 say?" both work off a single upload.
#
# Extraction is deliberately boring — plain text out, formatting discarded. The
# model never sees the original bytes, and the original is kept only so the
# user can re-download what they sent.
_TEXTISH_EXTS = {"txt", "md", "markdown", "rst", "csv", "tsv", "json", "yaml",
                 "yml", "log", "ini", "cfg", "toml", "xml", "html", "htm",
                 "py", "js", "ts", "sh", "sql", "c", "h", "cpp", "java", "go"}
_DOC_EXTS = {"pdf", "docx", "xlsx"}
UPLOAD_EXTS = _TEXTISH_EXTS | _DOC_EXTS
# A hard ceiling on extracted text. MAX_CONTENT_LENGTH already bounds the
# upload itself; this bounds what a pathological file can expand INTO.
MAX_UPLOAD_TEXT = 4_000_000
MAX_SESSION_UPLOADS = 5      # newest N stay attached to a session


def _upload_text(name: str, raw: bytes) -> str:
    """Extract plain text from an uploaded file.

    Raises ValueError with a message meant for the user — every caller surfaces
    it verbatim, so it must never leak a path or a stack trace."""
    ext = name.rsplit(".", 1)[-1].lower() if "." in name else ""
    if ext not in UPLOAD_EXTS:
        raise ValueError(f"can't read .{ext or '?'} files — "
                         f"supported: {', '.join(sorted(UPLOAD_EXTS))}")
    try:
        if ext in _TEXTISH_EXTS:
            text = raw.decode("utf-8", errors="replace")
        elif ext == "pdf":
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(raw))
            pages = [(p.extract_text() or "") for p in reader.pages]
            text = "\n\n".join(pages)
            if not text.strip():
                raise ValueError(
                    "this PDF has no extractable text — it is probably a scan. "
                    "OCR it first, then upload the result.")
        elif ext == "docx":
            import docx
            d = docx.Document(io.BytesIO(raw))
            parts = [p.text for p in d.paragraphs]
            for t in d.tables:
                for row in t.rows:
                    parts.append("\t".join(c.text for c in row.cells))
            text = "\n".join(parts)
        elif ext == "xlsx":
            import openpyxl
            wb = openpyxl.load_workbook(io.BytesIO(raw), data_only=True)
            parts = []
            for ws in wb.worksheets:
                parts.append(f"--- sheet: {ws.title} ---")
                for row in ws.iter_rows(values_only=True):
                    if any(v is not None for v in row):
                        parts.append("\t".join(
                            "" if v is None else str(v) for v in row))
            text = "\n".join(parts)
        else:                                    # unreachable, kept honest
            raise ValueError(f"no extractor for .{ext}")
    except ValueError:
        raise
    except Exception as e:
        raise ValueError(f"could not read this .{ext} file ({e.__class__.__name__})")
    if not text.strip():
        raise ValueError("that file contains no readable text")
    return text[:MAX_UPLOAD_TEXT]


# ---------------------------------------------------------------------------
# GitHub repositories as attachments.
#
# A pasted repo URL is NEVER fetched directly. The owner/repo/ref are parsed
# out with a strict pattern and the tarball URL is built from constants, so
# the only host this feature can ever contact is codeload.github.com — a
# crafted "github.com@evil.host/..." URL has nothing to grab onto. Public
# repositories only; that is a feature statement, not a TODO.
_GH_URL = re.compile(
    r"^https?://github\.com/([A-Za-z0-9][A-Za-z0-9-]{0,38})/"
    r"([A-Za-z0-9._-]{1,100}?)(?:\.git)?"
    r"(?:/(?:tree|commit|releases/tag)/([A-Za-z0-9._/-]{1,120}))?/?"
    r"(?:[?#].*)?$")
GH_TARBALL_CAP = int(os.environ.get("CREWE_GH_TARBALL_CAP", str(60 * 1024 * 1024)))
GH_FILE_CHARS = 60_000          # per file; a 200k-char bundle is not "a file"
GH_MAX_FILES = 400
_GH_SKIP_DIRS = {"node_modules", ".git", "dist", "build", "vendor", "target",
                 "__pycache__", ".next", ".venv", "venv", "coverage"}
_GH_SKIP_FILES = ("package-lock.json", "yarn.lock", "pnpm-lock.yaml",
                  "poetry.lock", "Cargo.lock", "go.sum", ".min.js", ".min.css",
                  ".map", ".lock")


def _github_parse(url):
    """(owner, repo, ref-or-None) from a github.com URL, else None."""
    m = _GH_URL.match((url or "").strip())
    if not m:
        return None
    owner, repo, ref = m.group(1), m.group(2), m.group(3)
    if repo.startswith(".") or owner.startswith("-"):
        return None
    return owner, repo, ref


def _github_repo_text(url):
    """Fetch a public repo as one reviewable text bundle.

    Returns (display_name, text). Raises ValueError with a user-facing message
    — same contract as _upload_text. The bundle is a file tree followed by
    each kept file under an '=== FILE: path ===' header; that marker is what
    the review route's whole-file packing splits on, so change them together.
    """
    parsed = _github_parse(url)
    if not parsed:
        raise ValueError("that doesn't look like a GitHub repository URL — "
                         "expected https://github.com/owner/repo")
    owner, repo, ref = parsed
    tar_url = (f"https://codeload.github.com/{owner}/{repo}/tar.gz/"
               f"{ref or 'HEAD'}")
    try:
        r = requests.get(tar_url, stream=True, timeout=60,
                         headers={"User-Agent": "crewe-repo-attach"})
    except Exception as e:
        raise ValueError(f"could not reach GitHub ({e.__class__.__name__})")
    if r.status_code == 404:
        raise ValueError(f"GitHub returned 404 for {owner}/{repo}"
                         + (f"@{ref}" if ref else "")
                         + " — private repos and typos both look like this; "
                           "only public repositories can be attached")
    if r.status_code != 200:
        raise ValueError(f"GitHub returned {r.status_code} for {owner}/{repo}")

    buf, total = io.BytesIO(), 0
    for chunk in r.iter_content(65536):
        total += len(chunk)
        if total > GH_TARBALL_CAP:
            r.close()
            raise ValueError(f"repository archive exceeds "
                             f"{GH_TARBALL_CAP // (1024*1024)} MB — too large "
                             "to attach")
        buf.write(chunk)
    buf.seek(0)

    import tarfile
    kept, skipped, tree = [], 0, []
    used = 0
    try:
        with tarfile.open(fileobj=buf, mode="r:gz") as tar:
            for m in tar:                       # members read in memory only —
                if not m.isfile():              # nothing is extracted to disk,
                    continue                    # so tar paths can't traverse
                parts = m.name.split("/")[1:]   # strip the repo-ref top dir
                if not parts:
                    continue
                rel = "/".join(parts)
                if any(p in _GH_SKIP_DIRS for p in parts[:-1]):
                    continue
                base = parts[-1]
                ext = base.rsplit(".", 1)[-1].lower() if "." in base else ""
                if any(base.endswith(sfx) for sfx in _GH_SKIP_FILES):
                    continue
                if ext not in _KNOWN_EXTS and base not in (
                        "Dockerfile", "Makefile", "LICENSE", "README"):
                    continue
                tree.append(rel)
                if len(kept) >= GH_MAX_FILES or used >= MAX_UPLOAD_TEXT:
                    skipped += 1
                    continue
                raw = tar.extractfile(m).read(GH_FILE_CHARS * 4)
                text = raw.decode("utf-8", errors="replace")
                if "\x00" in text[:1000]:
                    continue                    # binary wearing a text suffix
                if len(text) > GH_FILE_CHARS:
                    text = text[:GH_FILE_CHARS] + "\n… (file truncated)\n"
                kept.append((rel, text))
                used += len(text)
    except tarfile.TarError:
        raise ValueError("GitHub's archive could not be read — try again, and "
                         "check the branch name if you gave one")

    if not kept:
        raise ValueError("no readable source files found in that repository")
    name = f"{owner}/{repo}" + (f"@{ref}" if ref else "")
    head = (f"GitHub repository {name} — {len(tree)} source files, "
            f"{len(kept)} included below"
            + (f", {skipped} omitted for size" if skipped else "") + ".\n\n"
            "FILE TREE:\n" + "\n".join(f"  {t}" for t in tree[:GH_MAX_FILES]))
    body = "\n\n".join(f"=== FILE: {rel} ===\n{text}" for rel, text in kept)
    return name, (head + "\n\n" + body)[:MAX_UPLOAD_TEXT]


def session_uploads(session_id: str) -> list:
    """Metadata for files attached to a session (no text — that lives on disk)."""
    with SESSIONS_LOCK:
        return list(sessions().get(session_id, {}).get("uploads") or [])


def _upload_body(up: dict) -> str:
    """Read back the extracted text for one upload. Missing file -> ''."""
    p = os.path.join(uploads_dir(), up["id"] + ".txt")
    try:
        with open(p, encoding="utf-8") as f:
            return f.read()
    except OSError:
        return ""


def _chunk_text(text: str, size: int) -> list:
    """Split on paragraph boundaries where possible, hard-split when not."""
    chunks, cur = [], ""
    for para in text.split("\n\n"):
        if len(cur) + len(para) + 2 <= size:
            cur += (("\n\n" if cur else "") + para)
            continue
        if cur:
            chunks.append(cur)
            cur = ""
        while len(para) > size:                 # one enormous paragraph
            chunks.append(para[:size])
            para = para[size:]
        cur = para
    if cur.strip():
        chunks.append(cur)
    return chunks


CONDENSE_SYSTEM = (
    "You condense documents. Return a faithful, information-dense summary of "
    "the text given: keep names, numbers, dates, decisions and conclusions "
    "verbatim where they matter. Do not add commentary, do not speculate, and "
    "do not omit a section merely because it seems minor. Prose or bullets, "
    "whichever suits the material."
)


def _condense(text: str, budget: int, url: str, job_id: str = None) -> str:
    """Map-reduce a document down to `budget` characters.

    A 300-page PDF does not fit any window we have, so summarise each chunk,
    then summarise the summaries, until it fits. Bounded: if a pass fails to
    shrink the text, stop and truncate rather than loop forever."""
    if len(text) <= budget:
        return text
    chunk_size = max(4000, budget)
    for _ in range(3):                       # depth bound
        chunks = _chunk_text(text, chunk_size)
        outs = []
        for i, ch in enumerate(chunks, 1):
            try:
                outs.append(_stream_chat(
                    url,
                    [{"role": "system", "content": CONDENSE_SYSTEM},
                     {"role": "user", "content":
                      f"Part {i} of {len(chunks)} of a document. Condense it.\n\n{ch}"}],
                    job_id, temperature=0.2, timeout=600))
            except Exception as e:
                print(f"[summarize] chunk {i}/{len(chunks)} failed: {e}")
                outs.append(ch[:budget // max(1, len(chunks))])
        new = "\n\n".join(o.strip() for o in outs if o and o.strip())
        if not new.strip():
            break
        if len(new) >= len(text):            # not shrinking — stop, don't loop
            text = new
            break
        text = new
        if len(text) <= budget:
            return text
    return text[:budget]


def build_search_context(results: list):
    """Turn raw SearXNG results into a numbered context block + a sources list.

    Fetches the first SEARCH_FETCH pages in parallel for fuller excerpts; the
    rest contribute snippet-only entries. Returns (context_str, sources)."""
    top = results[:SEARCH_RESULTS]
    if not top:
        return "", []

    fetched: dict = {}
    threads = []
    for r in top[:SEARCH_FETCH]:
        url = r.get("url", "")
        if not url:
            continue
        t = threading.Thread(target=lambda u=url: fetched.__setitem__(u, _fetch_page_text(u)),
                             daemon=True)
        t.start(); threads.append(t)
    for t in threads:
        t.join(timeout=10)

    blocks, sources = [], []
    for i, r in enumerate(top, 1):
        title   = (r.get("title") or "").strip()
        url     = r.get("url") or ""
        snippet = (r.get("content") or "").strip()
        block = f"[{i}] {title}\nURL: {url}"
        if snippet:
            block += f"\nSnippet: {snippet}"
        page = fetched.get(url, "")
        if page:
            block += f"\nExcerpt: {page[:SEARCH_FETCH_CHARS]}"
        blocks.append(block)
        sources.append((i, title, url))
    return "\n\n".join(blocks), sources


def run_audio_job(job_id: str, question: str, session_id: str):
    """Audio route: write a spoken script with the 12B, then record it with Piper."""
    # ---- stage 1: write the spoken script (12B on GPU, thinking off) ----
    with JOBS_LOCK:
        JOBS[job_id]["stage"] = "scripting"
        JOBS[job_id]["stage_ts"] = time.time()
    with SESSIONS_LOCK:
        sess = sessions().get(session_id, {})
        summary = sess.get("summary", "")
        history = list(sess.get("history") or [])
    system = AUDIO_SCRIPT_SYSTEM
    if summary:
        system += f"\n\nConversation context (older exchanges, summarized):\n{summary}"

    messages = [{"role": "system", "content": system}]
    for turn in history:
        messages.append({"role": "user",      "content": turn.get("q", "")})
        messages.append({"role": "assistant", "content": turn.get("a", "")})
    messages.append({"role": "user", "content": question})
    payload = {"messages": messages, "stream": True}
    parts = []
    try:
        _gurl = SPECIALISTS["general"]
        with requests.post(_gurl, json=_inject_model(payload, _gurl, "general"),
                           headers=_hdrs(_gurl), stream=True, timeout=600) as r:
            r.raise_for_status()
            for line in _cancellable_lines(r, job_id):
                if not line:
                    continue
                line = line.decode("utf-8")
                if line.startswith("data: "):
                    data = line[6:]
                    if data.strip() == "[DONE]":
                        break
                    try:
                        chunk = json.loads(data)
                        delta = chunk["choices"][0]["delta"].get("content", "")
                        if delta:
                            parts.append(delta)
                            with JOBS_LOCK:
                                JOBS[job_id]["tokens"] += 1
                    except Exception:
                        pass
    except Exception as e:
        with JOBS_LOCK:
            JOBS[job_id]["answer"] = f"[router error writing script: {e}]"
            JOBS[job_id]["done"] = True
        return

    script = _clean_for_speech("".join(parts))

    # ---- stage 2: record the script to a wav with Piper ----
    with JOBS_LOCK:
        JOBS[job_id]["stage"] = "recording"
        JOBS[job_id]["stage_ts"] = time.time()
    wav_path = os.path.join(audio_dir(), f"{job_id}.wav")          # final (44.1 kHz)
    raw_path = os.path.join(audio_dir(), f"{job_id}.raw22k.wav")   # piper's native 22 kHz
    audio_url = None
    if script:
        try:
            env = dict(os.environ)
            env["LD_LIBRARY_PATH"] = PIPER_LIB + ":" + env.get("LD_LIBRARY_PATH", "")
            proc = subprocess.run(
                [PIPER_BIN, "--model", PIPER_VOICE, "--output_file", raw_path],
                input=script.encode("utf-8"),
                env=env, capture_output=True, timeout=300,
            )
            if proc.returncode == 0 and os.path.exists(raw_path):
                # upsample 22.05 kHz -> 44.1 kHz for project compatibility
                try:
                    _upsample_wav(raw_path, wav_path)
                    os.remove(raw_path)
                except Exception as e:
                    print(f"[audio] resample failed, serving 22k: {e}")
                    os.replace(raw_path, wav_path)
                audio_url = f"/audio/{job_id}.wav"
            else:
                print(f"[audio] piper rc={proc.returncode}: "
                      f"{proc.stderr.decode('utf-8', 'ignore')[-300:]}")
        except Exception as e:
            print(f"[audio] piper error: {e}")

    _append_history(session_id, question, script, "audio")

    with JOBS_LOCK:
        JOBS[job_id]["answer"] = script
        JOBS[job_id]["audio"]  = audio_url
        JOBS[job_id]["done"]   = True

    spawn_owned(target=summarize, args=(session_id, question, script), daemon=True).start()


# Code-pipeline stream read timeout. Raised 600→1200 (2026-07-21): the coder on
# .234 is single-slot at ~223 tok/s prefill, so a big multi-file step can spend
# most of 10 min just prefilling before the first token — 600s was cutting those
# builds off as "specialist unreachable". Tune with CREWE_CODE_TIMEOUT (seconds).
CODE_STREAM_TIMEOUT = int(os.environ.get("CREWE_CODE_TIMEOUT", "1200"))
# Hard per-call output ceiling for pipeline calls that set no explicit cap —
# ~8k tokens ≈ 24 KB of code, generous for any real step, fatal to runaways.
# 0 = UNCAPPED, the default. A cap on step output only ever destroyed work
# here: cut mid-file, a rewrite that kept >50% still passes the elision guard
# and silently overwrites the good file. Loops are caught by the repetition
# detector; runaway wall-clock by CODE_STREAM_TIMEOUT; the user by Stop.
# Set CREWE_STEP_MAX_TOKENS to a number only to debug a specific backend.
STEP_MAX_TOKENS = int(os.environ.get("CREWE_STEP_MAX_TOKENS", "0"))


# ---------------------------------------------------------------------------
# Effort levels.
#
# The constants above are tuned for the DEEP coder. Pointing the pipeline at a
# small model without changing them overflows its context on any real project
# -- the same class of bug as the old 12k cap starving the 131k coder, just in
# the other direction. So effort selects a whole PROFILE: coder URL *and*
# budgets, never one without the other. Budgets are not guessed: each level is
# sized from its own backend's advertised n_ctx by _budgets_for().
#
# NOTE: whichever level is active, plan/execute/check all run on ONE model.
# Never split them across the two coders -- mixing models inside a single
# pipeline has produced worse output than either model alone, twice.
EFFORTS        = ("fast", "normal", "extra")
# Old names stay valid forever: browsers persist the choice in localStorage
# and older clients send it in requests, so a rename must never break them.
EFFORT_ALIASES = {"easy": "normal", "hard": "extra"}
# "normal", NOT "extra": an absent parameter must never silently select the
# paid tier. (Before the three-tier rename the absent default was the deep
# coder; anyone relying on that sends the value explicitly anyway — the UI
# always does.)
DEFAULT_EFFORT = "normal"

# Non-budget knobs. Put stage toggles here if the easy path should later SKIP
# verification stages rather than just run them smaller -- then it stays a
# constant change, not a code change.
EFFORT_TUNING = {
    "extra": {"label": "Extra",
              "max_plan_steps":  MAX_PLAN_STEPS,
              "extra_steps":     EXTRA_STEPS,
              "step_max_tokens": STEP_MAX_TOKENS,
              "step_fix_rounds": STEP_FIX_ROUNDS},
    "normal": {"label": "Normal",
               "max_plan_steps":  16,
               "extra_steps":     6,
               # 0 = uncapped on EVERY tier: a cap that binds mid-file lets a
               # truncated rewrite pass the elision guard and destroy the file.
               "step_max_tokens": 0,
               "step_fix_rounds": 2},
    "fast": {"label": "Fast",
             "max_plan_steps":  10,
             "extra_steps":     4,
             "step_max_tokens": 0,
             "step_fix_rounds": 2},
}
EFFORT_BUDGETS = {}    # level -> budget dict; filled by _refresh_budgets()


def _refresh_budgets():
    """Re-probe both coders and size each level's budgets to its own window."""
    global EFFORT_BUDGETS, CODER_CTX
    urls = {"extra": CODE_AGENT_URL, "normal": CODE_AGENT_URL_FAST,
            "fast": CODE_AGENT_URL_QUICK}
    probed = {}                                  # one probe per distinct URL
    out = {}
    for lvl, u in urls.items():
        if u not in probed:
            probed[u] = _budgets_for(u)
        out[lvl] = probed[u]
    EFFORT_BUDGETS = {lvl: b for lvl, (b, _n) in out.items()}
    CODER_CTX = out["extra"][1]
    print("[code] effort budgets — " + ", ".join(
        f"{lvl}: {b['small_project_chars']:,} chars (ctx {n or '?'})"
        for lvl, (b, n) in out.items()))


# Thread-local, mirroring set_owner/_owner: the pipeline reads these budgets at
# ~30 call sites, and threading a parameter through all of them is exactly the
# problem that pattern already solved here. Background threads must capture the
# level at launch the same way they capture the owner.
#
# Deliberate DIFFERENCE from _owner(): this does NOT raise when unset. A missing
# owner is a cross-user data leak and must fail loudly; a missing effort just
# means "hard", which is what the pipeline did before effort existed.
_EFFORT = threading.local()


def set_effort(level):
    level = EFFORT_ALIASES.get(level, level)
    _EFFORT.level = level if level in EFFORTS else DEFAULT_EFFORT


def _effort():
    return getattr(_EFFORT, "level", DEFAULT_EFFORT)


def coder_url():
    """The coder backend for the ACTIVE effort level."""
    lvl = _effort()
    if lvl == "fast":
        return CODE_AGENT_URL_QUICK
    if lvl == "normal":
        return CODE_AGENT_URL_FAST
    return CODE_AGENT_URL


def _budget(key):
    lvl = _effort()
    if lvl not in EFFORT_BUDGETS:            # probe failed / not yet run
        return {"small_project_chars": SMALL_PROJECT_CHARS,
                "step_file_chars":     STEP_FILE_CHARS,
                "step_file_hard":      STEP_FILE_HARD}[key]
    return EFFORT_BUDGETS[lvl][key]


def small_project_chars(): return _budget("small_project_chars")
def step_file_chars():     return _budget("step_file_chars")
def step_file_hard():      return _budget("step_file_hard")
def max_plan_steps():      return EFFORT_TUNING[_effort()]["max_plan_steps"]
def extra_steps():         return EFFORT_TUNING[_effort()]["extra_steps"]
def step_max_tokens():     return EFFORT_TUNING[_effort()]["step_max_tokens"]
def step_fix_rounds():     return EFFORT_TUNING[_effort()]["step_fix_rounds"]


_refresh_budgets()
_BUDGETS_READY = True


TRACE_DIR = os.path.expanduser(
    os.environ.get("CREWE_TRACE_DIR", "~/llama_logs/trace"))
TRACE_ON = os.environ.get("CREWE_TRACE", "1") != "0"
TRACE_TICK = int(os.environ.get("CREWE_TRACE_TICK", "2000"))  # tokens per tick
_TRACE_SEQ = {}
_TRACE_T0 = {}
_TRACE_LOCK = threading.Lock()


def _trace_path(job_id, name):
    d = os.path.join(TRACE_DIR, job_id or "nojob")
    os.makedirs(d, exist_ok=True)
    return os.path.join(d, name)


def _trace(job_id, event, **fields):
    """Append ONE event and flush immediately.

    Everything here is written the instant it happens — the jobmem file is
    only serialized when a job ENDS, which meant a running job was completely
    unobservable and every diagnosis had to be reconstructed afterwards.
    Read live with:  tail -f ~/llama_logs/trace/<job_id>/trace.jsonl"""
    if not TRACE_ON:
        return
    try:
        t0 = _TRACE_T0.get(job_id)
        rec = {"t": time.strftime("%H:%M:%S"), "event": event}
        # seconds since the job began — the field that answers
        # "where did the 50 minutes actually go"
        if t0:
            rec["el"] = round(time.time() - t0, 1)
        rec.update(fields)
        with open(_trace_path(job_id, "trace.jsonl"), "a") as f:
            f.write(json.dumps(rec, default=str) + "\n")
            f.flush()
    except Exception:
        pass


def _trace_blob(job_id, seq, label, kind, text):
    """Persist a full prompt/reply so a rejected or ballooning generation can
    be re-read and re-parsed offline instead of re-run."""
    if not TRACE_ON:
        return None
    try:
        name = f"{seq:03d}-{re.sub(r'[^a-zA-Z0-9_-]', '', label or 'call')}.{kind}.txt"
        with open(_trace_path(job_id, name), "w") as f:
            f.write(text or "")
        return name
    except Exception:
        return None


def _output_shape(text):
    """Explain WHY a generation is the size it is.

    A token count alone cannot distinguish 'wrote nine files' from 'repeated
    one line four hundred times'. These few numbers do: a low unique-line
    ratio with a high top-repeat count is degeneration; a high diff_blocks
    count is a fragile mega-edit; many fences is a multi-file dump."""
    lines = [l.strip() for l in (text or "").splitlines() if l.strip()]
    n = len(lines)
    counts = {}
    for l in lines:
        counts[l] = counts.get(l, 0) + 1
    # Rank only SUBSTANTIVE lines. Structural punctuation ("}", "},", "});")
    # repeats constantly in healthy code, so ranking it reported "REPEAT x20"
    # on a perfectly good generation — noise that makes the whole signal
    # untrustworthy. A repeated line with real content is the actual tell.
    top = sorted(((l, c) for l, c in counts.items()
                  if len(re.sub(r"[^A-Za-z0-9_]", "", l)) >= 4),
                 key=lambda kv: -kv[1])[:3]
    return {
        "chars": len(text or ""),
        "lines": n,
        "unique_line_ratio": round(len(counts) / n, 3) if n else 1.0,
        "top_repeats": [{"n": c, "line": l[:70]} for l, c in top if c > 2],
        "fences": (text or "").count("```"),
        "diff_blocks": len(_DIFF_BLOCK.findall(text or "")),
        "tail_period": _min_period(text[-1200:]) if len(text or "") >= 1200 else 0,
    }


def _job_notes_len(job_id):
    with JOBS_LOCK:
        return len(JOBS.get(job_id, {}).get("notes", []))


def _job_notes_since(job_id, k):
    with JOBS_LOCK:
        return list(JOBS.get(job_id, {}).get("notes", [])[k:])


def _min_period(t):
    """Minimal linear period of t (KMP failure): p such that t[i]==t[i+p]."""
    n = len(t)
    f = [0] * (n + 1)
    k = 0
    for i in range(1, n):
        while k and t[i] != t[k]:
            k = f[k]
        if t[i] == t[k]:
            k += 1
        f[i + 1] = k
    return n - f[n]


# ---------------------------------------------------------------------------
# Anthropic -> OpenAI stream adapter.
#
# Crewe has eight streaming loops, all of which parse OpenAI's shape
# (`choices[0].delta.content`). Rather than rewrite every one of them, the
# translation happens at the wire: Anthropic's events are converted into
# OpenAI-shaped `data:` lines, so everything downstream — token counting,
# cancellation, usage capture, the repetition detector — works unchanged.
#
# Anthropic's stream is a sequence of typed events rather than uniform chunks:
#   message_start        carries input_tokens
#   content_block_delta  text_delta (visible) or thinking_delta (reasoning)
#   message_delta        carries stop_reason and output_tokens
#   message_stop         end of stream
# Only message_delta knows the output count, and only message_start knows the
# input count, so usage has to be accumulated across the stream — which is why
# this is a stateful generator rather than a per-line function.


def _anthropic_to_openai_lines(raw_lines):
    """Yield OpenAI-shaped `data: {...}` byte lines from an Anthropic stream."""
    usage = {"prompt_tokens": 0, "completion_tokens": 0}

    def pack(delta=None, finish=None, use=None):
        chunk = {"choices": [{"delta": delta or {}, "finish_reason": finish,
                              "index": 0}]}
        if use:
            chunk["usage"] = dict(use)
        return b"data: " + json.dumps(chunk).encode()

    for ln in raw_lines:
        if not ln:
            continue
        s = ln.decode("utf-8", errors="replace") if isinstance(ln, bytes) else ln
        if not s.startswith("data: "):
            continue                      # `event:` lines carry no payload
        body = s[6:].strip()
        if body == "[DONE]":
            yield b"data: [DONE]"
            return
        try:
            ev = json.loads(body)
        except Exception:
            continue
        t = ev.get("type")

        if t == "message_start":
            u = ((ev.get("message") or {}).get("usage") or {})
            usage["prompt_tokens"] = int(u.get("input_tokens") or 0)
        elif t == "content_block_delta":
            d = ev.get("delta") or {}
            dt = d.get("type")
            if dt == "text_delta":
                yield pack({"content": d.get("text", "")})
            elif dt == "thinking_delta":
                # keep the live token counter moving during thinking, exactly
                # as it does for llama.cpp's reasoning_content
                yield pack({"reasoning_content": d.get("thinking", "")})
        elif t == "message_delta":
            u = ev.get("usage") or {}
            if u.get("output_tokens"):
                usage["completion_tokens"] = int(u["output_tokens"])
            if u.get("input_tokens"):
                usage["prompt_tokens"] = int(u["input_tokens"])
            stop = (ev.get("delta") or {}).get("stop_reason")
            # A safety decline is a normal 200 with stop_reason "refusal" — it
            # must not look like an empty successful answer.
            if stop == "refusal":
                yield pack({"content": "\n\n*(the model declined this request)*"})
            yield pack(finish=stop, use=usage)
        elif t == "error":
            msg = (ev.get("error") or {}).get("message", "stream error")
            yield pack({"content": f"\n\n[backend error: {msg}]"})
        elif t == "message_stop":
            yield b"data: [DONE]"
            return


def _resp_json(r):
    """Normalise a non-streaming reply to the OpenAI shape.

    Anthropic returns {"content": [{"type": "text", "text": ...}], "usage": ...}
    where OpenAI returns {"choices": [{"message": {"content": ...}}]}. Callers
    read the OpenAI shape, so translate here rather than at each call site."""
    d = r.json()
    try:
        if _kind(r.request.url) != "anthropic":
            return d
    except Exception:
        return d
    parts = [b.get("text", "") for b in (d.get("content") or [])
             if b.get("type") == "text"]
    text = "".join(parts)
    if d.get("stop_reason") == "refusal":
        text = text or "*(the model declined this request)*"
    u = d.get("usage") or {}
    return {"choices": [{"message": {"role": "assistant", "content": text},
                         "finish_reason": d.get("stop_reason"), "index": 0}],
            "usage": {"prompt_tokens": u.get("input_tokens", 0),
                      "completion_tokens": u.get("output_tokens", 0)},
            "model": d.get("model", "")}


def _maybe_translate(r, raw_lines):
    """Wrap a response's line iterator in the adapter when the backend needs it."""
    try:
        if _kind(r.request.url) == "anthropic":
            return _anthropic_to_openai_lines(raw_lines)
    except Exception:
        pass
    return raw_lines


def _lines(r):
    """Line iterator for non-cancellable streams, kind-aware."""
    return _maybe_translate(r, r.iter_lines())


def _cancellable_lines(r, job_id):
    """Iterate a streaming response while honoring job cancellation within
    ~1 s even when the server sends NOTHING (long prompt processing). A
    blocking iter_lines() only notices the cancel flag when bytes arrive, so
    a silent backend made the Stop button useless — this moves the read to a
    thread and polls the flag. On cancel the response is closed, which drops
    the connection so the backend abandons the request when it next checks."""
    q = queue.Queue(maxsize=4096)
    DONE = object()
    ERR = object()

    def _reader():
        try:
            for ln in _maybe_translate(r, r.iter_lines()):
                q.put(ln)
            q.put(DONE)
        except Exception as e:
            # a mid-stream failure MUST reach the caller — swallowing it here
            # once turned dropped connections into silently-truncated answers
            # that the pipeline built garbage on top of
            q.put((ERR, e))

    threading.Thread(target=_reader, daemon=True).start()
    cancelled_here = False
    while True:
        if job_id and _job_cancelled(job_id):
            cancelled_here = True
            try:
                r.close()
            except Exception:
                pass
            return
        try:
            item = q.get(timeout=0.8)
        except queue.Empty:
            continue
        if item is DONE:
            return
        if isinstance(item, tuple) and item[0] is ERR:
            if cancelled_here:
                return
            raise item[1]
        yield item

def _stream_chat(url, messages, job_id, temperature=None, max_tokens=None,
                 timeout=CODE_STREAM_TIMEOUT, label=None):
    """POST a streaming chat request, bumping the job's live token counter per
    delta (visible and hidden reasoning alike), and return the full text.
    Raises on transport errors — callers decide how to degrade.

    Only the code pipeline uses this, and every one of its roles (plan, write,
    check) wants thinking OFF: measured on the 26B, the same planner prompt
    took 524 tokens / 20.6s with thinking vs 101 tokens / 3.6s without, with
    identical output quality. Templates without the switch ignore the kwarg."""
    payload = {"messages": messages, "stream": True,
               "chat_template_kwargs": {"enable_thinking": False}}
    if temperature is not None:
        payload["temperature"] = temperature
    # Steps are NOT capped by default. A cap that binds mid-file silently
    # corrupts the build, and starving a step is never the right answer to a
    # long one. Pathology is caught by what it actually is: the repetition
    # detector below (loops) and CODE_STREAM_TIMEOUT (wall clock). Callers with
    # a genuinely bounded job (classifier, planner, inspector) pass max_tokens.
    cap = max_tokens or step_max_tokens()
    if cap:
        payload["max_tokens"] = cap
    with _TRACE_LOCK:
        seq = _TRACE_SEQ[job_id] = _TRACE_SEQ.get(job_id, 0) + 1
    label = label or "call"
    prompt_txt = "\n\n".join(
        f"--- {m.get('role')} ---\n{m.get('content')}" for m in messages)
    _trace_blob(job_id, seq, label, "prompt", prompt_txt)
    _trace(job_id, "call_start", seq=seq, label=label, url=url,
           prompt_chars=len(prompt_txt), messages=len(messages),
           max_tokens=cap or None, temperature=temperature)
    t0 = time.time()
    finish_reason = None
    parts = []
    think_n = content_n = 0
    payload = _with_usage(_inject_model(payload, url), url)
    _usage = None
    with requests.post(url, json=payload, headers=_hdrs(url),
                       stream=True, timeout=timeout) as r:
        r.raise_for_status()
        for line in _cancellable_lines(r, job_id):
            if not line:
                continue
            line = line.decode("utf-8")
            if not line.startswith("data: "):
                continue
            data = line[6:]
            if data.strip() == "[DONE]":
                break
            try:
                chunk = json.loads(data)
                # usage arrives on a final chunk whose choices list is EMPTY
                if chunk.get("usage"):
                    _usage = chunk["usage"]
                if not chunk.get("choices"):
                    continue
                finish_reason = (chunk["choices"][0].get("finish_reason")
                                 or finish_reason)
                d = chunk["choices"][0]["delta"]
                delta = d.get("content", "")
                if delta:
                    parts.append(delta)
                    content_n += 1
                    # LIVE balloon telemetry: every TRACE_TICK tokens, record
                    # how fast it is going AND what it is actually emitting
                    # right now. A token count alone never explained a runaway;
                    # the tail excerpt does, while the call is still running.
                    if TRACE_TICK and content_n % TRACE_TICK == 0:
                        el = max(0.1, time.time() - t0)
                        so_far = "".join(parts)
                        tail = so_far[-200:].replace("\n", "\\n")
                        shape = _output_shape(so_far)
                        _trace(job_id, "tick", seq=seq, label=label,
                               tokens=content_n, elapsed=round(el, 1),
                               tok_s=round(content_n / el, 1),
                               unique_line_ratio=shape["unique_line_ratio"],
                               top_repeat=(shape["top_repeats"] or [{}])[0],
                               fences=shape["fences"],
                               diff_blocks=shape["diff_blocks"], tail=tail)
                        print(f"[code] {label} @{content_n}tok "
                              f"{round(content_n / el, 1)}/s "
                              f"uniq={shape['unique_line_ratio']} "
                              f"fences={shape['fences']} "
                              f"diffs={shape['diff_blocks']} … {tail[-90:]}")
                    # repetition-loop detector — STRICT periodicity only:
                    # fire iff the ENTIRE last 1200 chars are one exact unit
                    # (<=300 chars, so >=4 repeats) tiling the tail. Real code
                    # is repetitive-but-VARYING and is never exactly periodic
                    # over a full 1200-char window; only degeneration is.
                    # Deliberation-spiral abort. Measured on the corpus:
                    # every legitimate generation held unique_line_ratio
                    # >= 0.48 past 10k tokens; the observed spirals read
                    # 0.38 -> 0.27 while the model narrated "Actually,
                    # wait." for 17 minutes. Both signals must agree.
                    if content_n and content_n % 2048 == 0 and content_n >= 10240:
                        u = _output_shape("".join(parts))["unique_line_ratio"]
                        if u < 0.35:
                            msg = (f"deliberation spiral detected "
                                   f"(unique-line ratio {u} after "
                                   f"{content_n} tokens) — call aborted; "
                                   f"the model is reasoning in circles, "
                                   f"not writing code")
                            print(f"[code] {msg}")
                            _trace(job_id, "spiral_abort", seq=seq, label=label,
                                   tokens=content_n, uniq=u)
                            if job_id:
                                with JOBS_LOCK:
                                    j = JOBS.get(job_id)
                                    if j is not None:
                                        j.setdefault("notes", []).append(msg)
                            try:
                                r.close()
                            except Exception:
                                pass
                            break
                    if content_n and content_n % 512 == 0 and content_n >= 2048:
                        tail = "".join(parts)[-1200:]
                        p = _min_period(tail) if len(tail) == 1200 else 0
                        if 0 < p <= 300:
                            msg = (f"repetition loop detected after "
                                   f"{content_n} tokens (period {p} chars) — "
                                   f"call aborted")
                            print(f"[code] {msg}")
                            _trace(job_id, "loop_abort", seq=seq, label=label,
                                   tokens=content_n, period=p,
                                   repeated=tail[-p:] if p else "")
                            if job_id:
                                with JOBS_LOCK:
                                    j = JOBS.get(job_id)
                                    if j is not None:
                                        j.setdefault("notes", []).append(msg)
                            try:
                                r.close()
                            except Exception:
                                pass
                            break
                if d.get("reasoning_content"):
                    think_n += 1
                if delta or d.get("reasoning_content"):
                    with JOBS_LOCK:
                        JOBS[job_id]["tokens"] += 1
            except Exception:
                pass
    total = think_n + content_n
    if cap and total >= cap - 8:
        msg = (f"call hit the {cap}-token cap "
               f"(reasoning={think_n}, content={content_n})"
               + (" — most of the budget went to THINKING; the model's "
                  "template may ignore enable_thinking:false"
                  if think_n > content_n * 3 else ""))
        print(f"[code] {msg}")
        if job_id:
            with JOBS_LOCK:
                j = JOBS.get(job_id)
                if j is not None:
                    j.setdefault("notes", []).append(msg)
    out = "".join(parts)
    # Every pipeline step is a separate billable call, so record each one —
    # a code build on a paid coder is many requests, not one.
    record_spend(url, label or "code", _usage,
                 "".join(m.get("content", "") for m in messages), out)
    el = max(0.1, time.time() - t0)
    shape = _output_shape(out)
    _trace_blob(job_id, seq, label, "reply", out)
    _trace(job_id, "call_end", seq=seq, label=label,
           finish_reason=finish_reason, capped=bool(cap and total >= cap - 8),
           think_tokens=think_n, content_tokens=content_n,
           elapsed=round(el, 1), tok_s=round(total / el, 1), shape=shape)
    print(f"[code] {label} done: {content_n}tok content / {think_n} think "
          f"in {el:.0f}s ({total / el:.1f}/s) finish={finish_reason} "
          f"lines={shape['lines']} uniq={shape['unique_line_ratio']} "
          f"fences={shape['fences']} diffs={shape['diff_blocks']}"
          + (f" TOP-REPEAT x{shape['top_repeats'][0]['n']}: "
             f"{shape['top_repeats'][0]['line'][:50]}"
             if shape["top_repeats"] else ""))
    return out


def _set_stage(job_id, stage):
    with JOBS_LOCK:
        JOBS[job_id]["stage"] = stage
        # lets /progress report how long the current stage has run — a big
        # prompt legitimately produces no tokens for minutes while the model
        # reads it, and without stage age that's indistinguishable from a hang
        JOBS[job_id]["stage_ts"] = time.time()


def _parse_plan(text, question):
    """Extract the JSON build plan from planner output. Any parse failure
    degrades to a single catch-all step, which is the old one-shot behavior."""
    text = _strip_think(text)
    plan = None
    why = "no_json_span"
    i = text.find("{")
    if i >= 0:
        # raw_decode reads ONE JSON value and stops at its natural end, so a
        # stray trailing bracket cannot destroy the plan. A greedy
        # `\{[\s\S]*\}` grabbed through the LAST brace instead: one extra
        # "]}" from the model turned a valid 9-step plan into "Extra data" and
        # silently collapsed the build to a single catch-all step.
        try:
            plan, _end = json.JSONDecoder().raw_decode(text, i)
            why = "json"
        except Exception as e:
            why = f"json_decode: {e}"[:120]
            # last resort: trim back to the final closing brace and retry
            j = text.rfind("}")
            if j > i:
                try:
                    plan = json.loads(text[i:j + 1])
                    why = "json_trimmed"
                except Exception:
                    pass
    if not isinstance(plan, dict):
        print(f"[code] planner JSON unusable ({why}) — falling back to one step")
    # Local models produce every wrong shape imaginable (steps as a dict,
    # files as a bare string or number) — validate every level; any surprise
    # collapses to the fallback rather than killing the worker thread.
    steps = []
    try:
        raw_steps = plan.get("steps") if isinstance(plan, dict) else None
        if not isinstance(raw_steps, list):
            raw_steps = []
        for s in raw_steps[:max_plan_steps()]:
            if not isinstance(s, dict) or not (s.get("detail") or s.get("title")):
                continue
            fl = s.get("files")
            if isinstance(fl, str):
                fl = [fl]
            elif not isinstance(fl, list):
                fl = []
            steps.append({
                "title": str(s.get("title") or f"step {len(steps) + 1}")[:80],
                "detail": str(s.get("detail") or s.get("title")),
                "files": [str(f) for f in fl if isinstance(f, str)][:4],
                "done_when": str(s.get("done_when") or "the step's change works"),
            })
    except Exception as e:
        print(f"[code] plan validation error: {e}")
        steps = []
    features, design = [], ""
    if isinstance(plan, dict):
        raw_f = plan.get("features")
        if isinstance(raw_f, str):
            raw_f = [raw_f]
        if isinstance(raw_f, list):
            features = [str(f).strip()[:200] for f in raw_f
                        if isinstance(f, str) and f.strip()][:12]
        if isinstance(plan.get("design"), str):
            design = plan["design"].strip()[:600]
    if not steps:
        print(f"[code] plan fallback: could not parse a plan from "
              f"{len(text)} chars of planner output")
        return {"goal": question[:200], "features": features, "design": design,
                "steps": [{
                    "title": "build the request", "detail": question, "files": [],
                    "done_when": "the user's request is fulfilled"}]}
    goal = str(plan.get("goal") or question)[:200]
    return {"goal": goal, "features": features, "design": design, "steps": steps}


_BROWSER_ERR_CAPTURE = """<script>
window.__router_errors = [];
window.addEventListener('error', function (e) {
  if (e.target && (e.target.src || e.target.href)) {
    window.__router_errors.push('Failed to load: ' + (e.target.src || e.target.href));
  } else {
    window.__router_errors.push('JS error: ' + e.message +
      (e.filename ? ' @' + e.filename.split('/').pop() + ':' + e.lineno : ''));
  }
}, true);
window.addEventListener('unhandledrejection', function (e) {
  window.__router_errors.push('Unhandled promise rejection: ' + e.reason);
});
</script>"""

# "Loads without errors" is not "works": a draw loop that silently renders
# nothing produces a black screen and zero exceptions. Sample the canvas —
# one flat color and no visible text means nothing was actually drawn.
_BLANK_DETECT_JS = """
const out = {blank: false, size: ''};
const text = (document.body && document.body.innerText || '').trim();
const c = document.querySelector('canvas');
if (c) {
  out.size = c.width + 'x' + c.height;
  try {
    const ctx = c.getContext('2d');
    if (ctx) {
      const d = ctx.getImageData(0, 0, c.width, c.height).data;
      const seen = new Set();
      for (let i = 0; i < d.length; i += 397 * 4)
        seen.add((d[i] << 16) | (d[i + 1] << 8) | d[i + 2]);
      out.blank = seen.size <= 2 && text.length < 20;
    }
  } catch (e) {}
  if (!out.blank) {
    const gl = c.getContext('webgl') || c.getContext('webgl2');
    if (gl && !c.getContext('2d')) out.webgl = true; /* pixel check impossible without preserveDrawingBuffer */
  }
} else {
  out.blank = text.length < 20 && !document.querySelector('img, svg, video');
}
return out;
"""


def _inject_html(html, snippet):
    """Insert a snippet as early as VALIDLY possible: inside <head>, else at
    the top of <body>, else right after the doctype — never before it (a
    token ahead of <!doctype> flips the parser into quirks mode and can void
    the snippet entirely)."""
    for pat in (r"<head[^>]*>", r"<body[^>]*>", r"<!doctype[^>]*>"):
        m = re.search(pat, html, flags=re.I)
        if m:
            return html[:m.end()] + snippet + html[m.end():]
    return snippet + html


def _browser_check(files):
    """Write the project to a run dir, load its entry html in headless
    Firefox, and return (runtime errors, screenshot path). Degrades to
    ([], None) whenever selenium/geckodriver are unavailable — the pipeline
    must work without a browser, just with one less safety net."""
    entry = ("index.html" if "index.html" in files
             else next((n for n in files if n.endswith((".html", ".htm"))), None))
    if not entry:
        return [], None
    try:
        from selenium import webdriver
        from selenium.webdriver.firefox.options import Options
        from selenium.webdriver.firefox.service import Service
    except Exception:
        return [], None
    run_dir = os.path.join(checks_dir(), uuid.uuid4().hex[:12])
    driver = None
    httpd = None
    try:
        runs = sorted(os.listdir(checks_dir()),
                      key=lambda d: os.path.getmtime(
                          os.path.join(checks_dir(), d)))
        for old in runs[:-60]:
            shutil.rmtree(os.path.join(checks_dir(), old),
                          ignore_errors=True)
    except OSError:
        pass
    try:
        for name, code in files.items():
            path = os.path.abspath(os.path.join(run_dir, name))
            if not path.startswith(os.path.abspath(run_dir) + os.sep):
                continue  # no path escapes from model-named files
            os.makedirs(os.path.dirname(path), exist_ok=True)
            if name == entry:
                # error hook must run before any page script
                code = _inject_html(code, _BROWSER_ERR_CAPTURE)
            with open(path, "w") as f:
                f.write(code)
        # Serve over HTTP, the way builds actually run: file:// silently
        # blocks ES modules and local fetch(), which fails working builds
        # (Pal diagnosed this by spinning up a python web host by hand).
        class _Quiet(SimpleHTTPRequestHandler):
            def log_message(self, *args):
                pass
        httpd = ThreadingHTTPServer(
            ("127.0.0.1", 0), functools.partial(_Quiet, directory=run_dir))
        threading.Thread(target=httpd.serve_forever, daemon=True).start()

        opts = Options()
        opts.add_argument("--headless")
        gecko = "/snap/bin/geckodriver"
        service = Service(executable_path=gecko) if os.path.exists(gecko) else Service()
        driver = webdriver.Firefox(options=opts, service=service)
        driver.set_page_load_timeout(25)
        try:
            driver.get(f"http://127.0.0.1:{httpd.server_address[1]}/{entry}")
        except Exception as e:
            # a page that never finishes loading is a FAILURE — the old
            # catch-all reported it as 'loads clean'
            return ([f"page never finished loading within 25s — likely an "
                     f"infinite loop or blocked startup [{type(e).__name__}]"],
                    None)
        time.sleep(2.5)  # let init scripts and first frames run
        errors = driver.execute_script("return window.__router_errors || []")
        blank = None
        try:
            blank = driver.execute_script(_BLANK_DETECT_JS)
            if blank and blank.get("blank"):
                # might be a title screen waiting for input — nudge first
                from selenium.webdriver.common.keys import Keys
                driver.find_element("tag name", "body").send_keys(Keys.ENTER, " ")
                try:
                    driver.find_element("tag name", "canvas").click()
                except Exception:
                    pass
                time.sleep(1.5)
                blank = driver.execute_script(_BLANK_DETECT_JS)
        except Exception:
            blank = None
        if blank and blank.get("blank"):
            # a minimal-but-valid page (some text, a button) is NOT blank —
            # the pixel heuristic false-positives on tiny pages and used to
            # burn fix rounds restyling them
            try:
                vis = driver.execute_script(
                    "return {t:(document.body.innerText||'').trim().length,"
                    " i:!!document.querySelector("
                    "'button,a[href],input,select,textarea')}")
                # visible text or any interactive control = a real page, not
                # a blank screen ('0' + an 'Add' button is a valid tiny app)
                if vis and (vis.get("i") or (vis.get("t") or 0) >= 8):
                    blank = None
            except Exception:
                pass
        shot = os.path.join(run_dir, "screenshot.png")
        try:
            driver.save_screenshot(shot)
        except Exception:
            shot = None
        errs = [str(e)[:200] for e in errors][:10]
        if blank and blank.get("blank"):
            errs.append(
                "BLANK SCREEN: the page renders as a single flat color — the "
                f"canvas ({blank.get('size') or 'none'}) draws nothing and the "
                "page shows no visible text, even after pressing Enter and "
                "clicking. Initialization or the first draw never produces "
                "visible output; find why and fix it.")
        return errs, shot
    except Exception as e:
        print(f"[code] browser check unavailable: {e}")
        return [], None
    finally:
        if driver:
            try:
                driver.quit()
            except Exception:
                pass
        if httpd:
            try:
                httpd.shutdown()
            except Exception:
                pass


# ---- behavior checks: prove the build WORKS, not just loads ------------------
# The model writes a tiny declarative check script (clicks/typing + assertions);
# the harness executes it in the same headless browser as _browser_check and
# feeds concrete pass/fail results back. Patch mode runs the checks BEFORE
# fixing (reproduce the reported bug — SWE-agent style) and again after, so
# "fixed" means a passing check, not the model's own opinion of its diff.

CHECKWRITER_SYSTEM = (
    "You write automated behavior checks for a small web build. Given the "
    "goal (or a bug report) and the project's HTML/JS, respond with ONLY a "
    "JSON array — no commentary, no fences — of at most 8 check objects, "
    "using EXACTLY these forms:\n"
    '  {"do":"click","sel":"<css>"}\n'
    '  {"do":"type","sel":"<css>","text":"..."}\n'
    '  {"do":"key","key":"ENTER"}\n'
    '  {"do":"wait","ms":400}\n'
    '  {"do":"snap"}                        (remember how the page looks now)\n'
    '  {"do":"goto","page":"other.html"}    (open another page of the project)\n'
    '  {"assert":"text","sel":"<css>","equals":"..."}   (or "contains":"...")\n'
    '  {"assert":"exists","sel":"<css>"}\n'
    '  {"assert":"js","expr":"<js expression that must be true>","desc":"short name"}\n'
    '  {"do":"js","expr":"<js statement run for side effects — seed state, dispatch events>"}\n'
    '  {"assert":"no_errors"}\n'
    '  {"assert":"screen_changed"}          (page LOOKS different than at the last snap)\n'
    '  {"assert":"canvas_pixels","sel":"canvas"}  (the canvas actually draws content)\n'
    "Rules: use ONLY selectors, pages, and globals that exist in the provided "
    "files; keep the sequence deterministic (no animation-timing bets); "
    "prefer one precise assertion over many vague ones. For saved state use "
    "assert js (e.g. localStorage.getItem('save') !== null). For games and "
    "graphics: canvas_pixels proves something is drawn; snap → interact → "
    "screen_changed proves the UI responds to input."
)


def _check_desc(c):
    if "do" in c:
        return f"do:{c['do']} " \
               f"{c.get('sel', c.get('key', c.get('page', c.get('ms', ''))))}".strip()
    a = c.get("assert")
    if a == "text":
        want = c.get("equals", c.get("contains", ""))
        return f"text of {c.get('sel', '?')} ~ '{str(want)[:40]}'"
    if a == "exists":
        return f"exists {c.get('sel', '?')}"
    if a == "js":
        return c.get("desc") or f"js: {str(c.get('expr', ''))[:60]}"
    if a == "screen_changed":
        return "screen changed since snap"
    if a == "canvas_pixels":
        return f"canvas {c.get('sel', 'canvas')} draws content"
    return "no page errors"


def _parse_checks(text):
    """Parse the behaviour-check script.

    Accepts a JSON array OR newline-delimited JSON objects. Requiring an array
    meant the whole verification layer silently did nothing: the model reliably
    emitted one {"do":...} object per line, this returned [], and every job
    skipped behaviour checks entirely without a word in any log."""
    text = _strip_think(text)
    raw = []
    m = re.search(r"\[[\s\S]*\]", text)
    if m:
        try:
            raw = json.loads(m.group(0))
        except Exception:
            raw = []
    def _valid(objs):
        return [c for c in objs if isinstance(c, dict) and (
            c.get("do") in ("click", "type", "key", "wait", "snap", "goto",
                            "js")
            or c.get("assert") in ("text", "exists", "js", "no_errors",
                                   "screen_changed", "canvas_pixels"))]

    # Fall through to JSONL when the array parsed but held nothing usable —
    # a stray ["a","b"] earlier in the text must not mask the real script.
    if not _valid(raw if isinstance(raw, list) else []):
        raw = []
    if not raw:
        # JSONL form — one object per line, optionally fenced or comma-tailed
        for line in text.splitlines():
            line = line.strip().rstrip(",").strip()
            if not line.startswith("{") or not line.endswith("}"):
                continue
            try:
                obj = json.loads(line)
            except Exception:
                continue
            if isinstance(obj, dict):
                raw.append(obj)
    out = []
    for c in raw if isinstance(raw, list) else []:
        if not isinstance(c, dict):
            continue
        if (c.get("do") in ("click", "type", "key", "wait", "snap", "goto",
                        "js")
                or c.get("assert") in ("text", "exists", "js", "no_errors",
                                       "screen_changed", "canvas_pixels")):
            out.append(c)
    # 14, not 10: do:js seeding/event steps consume slots, and the first real
    # interaction script needed exactly 10 — a repro must never lose its
    # assertions to its own setup.
    return out[:14]


def _write_checks(job_id, goal_text, files, patch_bug=None, mem=None):
    """One model call → parsed check list ([] on any failure; never fatal)."""
    ctx = {n: c for n, c in files.items()
           if n.endswith((".html", ".htm", ".js", ".mjs"))}
    if not ctx:
        return []
    user = ""
    if patch_bug:
        user += ("Write checks that REPRODUCE this bug report. The assertions "
                 "must describe the CORRECT behavior — what SHOULD happen "
                 "once the bug is fixed. Because the bug currently exists, "
                 "those assertions will FAIL on the current build; that "
                 "failure is the reproduction. NEVER write an assertion that "
                 "expects the buggy behavior itself. The report:\n"
                 + patch_bug + "\n\n")
    elif goal_text:
        user += goal_text + "\n\n"
    user += _files_context(ctx, limit=24000, header="Project files:")
    user += "\n\nRespond with ONLY the JSON array of checks."
    try:
        out = _stream_chat(coder_url(),
                           [{"role": "system", "content": CHECKWRITER_SYSTEM},
                            {"role": "user", "content": user}],
                           job_id, temperature=0.1, max_tokens=4096,
                           label="checkwriter")
    except Exception as e:
        print(f"[code] check writer error: {e}")
        return []
    checks = _parse_checks(out)
    if checks:
        # observability: the exact checks belong in working memory so every
        # later call (and the human reading jobmem) can question them
        _mem_note(mem, "Behavior checks written: "
                       + json.dumps(checks, separators=(",", ":"))[:900])
    return checks


def _run_behavior_checks(files, checks):
    """Execute checks against the served build in headless Firefox.
    Returns (results, fatal): results = [{desc, ok, detail, action_error}];
    fatal is a string when the ENVIRONMENT failed (no selenium, page never
    loaded) — callers must treat that as 'checks unavailable', never as a
    build failure."""
    entry = ("index.html" if "index.html" in files
             else next((n for n in files if n.endswith((".html", ".htm"))), None))
    if not entry or not checks:
        return [], "no entry html or no checks"
    try:
        from selenium import webdriver
        from selenium.webdriver.firefox.options import Options
        from selenium.webdriver.firefox.service import Service
        from selenium.webdriver.common.by import By
        from selenium.webdriver.common.keys import Keys
    except Exception:
        return [], "selenium unavailable"
    run_dir = os.path.join(checks_dir(), "chk-" + uuid.uuid4().hex[:12])
    driver = httpd = None
    try:
        for name, code in files.items():
            path = os.path.abspath(os.path.join(run_dir, name))
            if not path.startswith(os.path.abspath(run_dir) + os.sep):
                continue
            os.makedirs(os.path.dirname(path), exist_ok=True)
            if name == entry:
                code = _inject_html(code, _BROWSER_ERR_CAPTURE)
            with open(path, "w") as f:
                f.write(code)

        class _Quiet(SimpleHTTPRequestHandler):
            def log_message(self, *args):
                pass
        httpd = ThreadingHTTPServer(
            ("127.0.0.1", 0), functools.partial(_Quiet, directory=run_dir))
        threading.Thread(target=httpd.serve_forever, daemon=True).start()

        opts = Options()
        opts.add_argument("--headless")
        gecko = "/snap/bin/geckodriver"
        service = Service(executable_path=gecko) if os.path.exists(gecko) else Service()
        driver = webdriver.Firefox(options=opts, service=service)
        driver.set_page_load_timeout(25)
        base = f"http://127.0.0.1:{httpd.server_address[1]}/"
        try:
            driver.get(base + entry)
        except Exception as e:
            return [], f"page never finished loading [{type(e).__name__}]"
        time.sleep(1.5)

        _CANVAS_PROBE = """
var c=document.querySelector(arguments[0]);
if(!c||!c.toDataURL) return {ok:false,d:'no canvas found for selector'};
var w=c.width||0,h=c.height||0;
if(!w||!h) return {ok:false,d:'zero-size canvas'};
var ctx=null; try{ctx=c.getContext('2d');}catch(e){}
if(ctx){
  var d=ctx.getImageData(0,0,w,h).data,seen={},n=0;
  var step=Math.max(1,Math.floor(d.length/4/2000));
  for(var i=0;i<d.length;i+=step*4){
    var k=(d[i]<<16)|(d[i+1]<<8)|d[i+2];
    if(!seen[k]){seen[k]=1;n++;} if(n>16)break;}
  return {ok:n>=3,d:'distinct sampled colors: '+n};
}
var b=document.createElement('canvas'); b.width=w; b.height=h;
return {ok:c.toDataURL()!==b.toDataURL(),d:'webgl canvas vs blank compare'};
"""
        results = []
        last_snap = None
        for c in checks:
            desc = _check_desc(c)
            try:
                if "do" in c:
                    if c["do"] == "click":
                        driver.find_element(By.CSS_SELECTOR, c["sel"]).click()
                    elif c["do"] == "type":
                        driver.find_element(By.CSS_SELECTOR, c["sel"]).send_keys(
                            str(c.get("text", "")))
                    elif c["do"] == "key":
                        k = str(c.get("key", ""))
                        driver.find_element(By.TAG_NAME, "body").send_keys(
                            getattr(Keys, k.upper(), k))
                    elif c["do"] == "wait":
                        time.sleep(min(int(c.get("ms", 300)), 3000) / 1000.0)
                    elif c["do"] == "snap":
                        last_snap = driver.get_screenshot_as_png()
                    elif c["do"] == "js":
                        # action form of the js assert: run a statement for
                        # its side effects (seed state, dispatch synthetic
                        # events). The checkwriter was already emitting this
                        # verb; dropping it silently meant interaction checks
                        # never had their preconditions set up.
                        driver.execute_script(str(c.get("expr", "")))
                    elif c["do"] == "goto":
                        page = str(c.get("page", ""))
                        if page not in files:
                            raise ValueError(f"'{page}' is not a project file")
                        driver.get(base + page)
                        time.sleep(1.0)
                    time.sleep(0.15)
                    results.append({"desc": desc, "ok": True, "detail": "",
                                    "action_error": False})
                elif c.get("assert") == "text":
                    el = driver.find_element(By.CSS_SELECTOR, c["sel"])
                    actual = (el.text or "").strip()
                    if "equals" in c:
                        ok = actual == str(c["equals"]).strip()
                    else:
                        ok = str(c.get("contains", "")) in actual
                    results.append({"desc": desc, "ok": ok,
                                    "detail": f"actual text: '{actual[:80]}'",
                                    "action_error": False})
                elif c.get("assert") == "exists":
                    n_found = len(driver.find_elements(By.CSS_SELECTOR, c["sel"]))
                    results.append({"desc": desc, "ok": n_found > 0,
                                    "detail": f"matches found: {n_found}",
                                    "action_error": False})
                elif c.get("assert") == "js":
                    val = bool(driver.execute_script(
                        f"return !!({c.get('expr', 'false')})"))
                    results.append({"desc": desc, "ok": val,
                                    "detail": f"expression was {val}",
                                    "action_error": False})
                elif c.get("assert") == "no_errors":
                    errs = driver.execute_script(
                        "return window.__router_errors || []")
                    results.append({"desc": desc, "ok": not errs,
                                    "detail": (str(errs[0])[:120] if errs else ""),
                                    "action_error": False})
                elif c.get("assert") == "screen_changed":
                    if last_snap is None:
                        results.append({"desc": desc, "ok": False,
                                        "detail": "no prior {\"do\":\"snap\"} "
                                                  "to compare against",
                                        "action_error": True})
                    else:
                        cur = driver.get_screenshot_as_png()
                        # Byte inequality is far too weak: a focus ring on the
                        # button "changed the screen" and this passed on a build
                        # where the click did nothing. Require a MEANINGFUL
                        # fraction of pixels to differ. Falls back to the byte
                        # compare only if PIL is unavailable.
                        changed_frac = None
                        try:
                            from PIL import Image, ImageChops
                            import io
                            a = Image.open(io.BytesIO(last_snap)).convert("RGB")
                            b = Image.open(io.BytesIO(cur)).convert("RGB")
                            if a.size == b.size:
                                diff = ImageChops.difference(a, b).convert("L")
                                # count pixels differing by more than noise
                                hist = diff.point(lambda v: 255 if v > 12 else 0).histogram()
                                changed_frac = hist[-1] / float(a.size[0] * a.size[1])
                        except Exception:
                            changed_frac = None
                        if changed_frac is not None:
                            # LIVENESS, not behaviour. Measured on a real build:
                            # adding a list item moved 0.191% of pixels while a
                            # mere focus ring moved 0.167% — far too close to
                            # separate. So this only catches "the page did not
                            # react at all"; proving WHAT happened is the job of
                            # {"assert":"js"}, which does discriminate.
                            visibly = changed_frac >= 0.0005   # 0.05%
                            results.append({"desc": desc, "ok": visibly,
                                            "detail": (f"{changed_frac:.3%} of pixels changed"
                                                       + ("" if visibly else
                                                          " — the page did not visibly "
                                                          "react at all")),
                                            "action_error": False})
                            last_snap = cur
                            continue
                        results.append({"desc": desc, "ok": cur != last_snap,
                                        "detail": ("screenshot differs from snap"
                                                   if cur != last_snap else
                                                   "screenshot is IDENTICAL to "
                                                   "the snap — nothing visibly "
                                                   "changed"),
                                        "action_error": False})
                elif c.get("assert") == "canvas_pixels":
                    probe = driver.execute_script(
                        _CANVAS_PROBE, c.get("sel", "canvas"))
                    probe = probe or {}
                    results.append({"desc": desc, "ok": bool(probe.get("ok")),
                                    "detail": str(probe.get("d", ""))[:120],
                                    "action_error": False})
            except Exception as e:
                results.append({"desc": desc, "ok": False,
                                "detail": f"{type(e).__name__}: {str(e)[:100]}",
                                "action_error": True})
        return results, None
    except Exception as e:
        print(f"[code] behavior checks unavailable: {e}")
        return [], f"environment error [{type(e).__name__}]"
    finally:
        if driver:
            try:
                driver.quit()
            except Exception:
                pass
        if httpd:
            try:
                httpd.shutdown()
            except Exception:
                pass


# ---- job working memory ------------------------------------------------------
# A per-job markdown log of what has been tried and what survived each fix,
# injected into every model call in the job. Without it, retry loops rebuild an
# identical prompt each round and the model (temp 0.1-0.2) regenerates the same
# failed fix until retries run out. The digest is kept per-session so the NEXT
# follow-up job starts knowing what the last one already tried, and the full
# log is written to ~/llama_logs/jobmem/<job_id>.md for inspection.
JOB_MEM_CHARS = 7000     # rendered budget inside prompts

PATCH_PLANNER_NOTE = (
    "\n\nPATCH MODE: existing project files are provided and this request is "
    "a follow-up on that build. Plan the MINIMAL steps that satisfy the "
    "request — fix or change ONLY what it names. Do NOT re-plan or re-promise "
    "the original app: the features list must contain ONLY the specific "
    "changes this request asks for (it may even be empty), never the app's "
    "original feature set. Never rebuild or restyle working code the request "
    "didn't mention."
)


def _mem_new(question, followup):
    return {"question": question, "followup": followup, "events": [],
            "prev": None}


def _mem_note(mem, text):
    if mem is not None:
        text = text.strip()
        if not mem["events"] or mem["events"][-1] != text:
            mem["events"].append(text)


def _mem_render(mem, limit=JOB_MEM_CHARS):
    """Markdown block for prompts: verbatim request + newest-events-first log."""
    if mem is None:
        return ""
    head = ("USER REQUEST (verbatim — THIS is what must be satisfied):\n"
            f"{mem['question']}\n")
    if mem.get("prev"):
        head += f"\nFrom the PREVIOUS job on this project:\n{mem['prev']}\n"
    if not mem["events"]:
        return head
    lines, total = [], 0
    for ev in reversed(mem["events"]):        # newest events survive the cap
        if total + len(ev) > limit:
            lines.append("(older events trimmed)")
            break
        lines.append(ev)
        total += len(ev)
    return (head + "\nWORKING MEMORY — what already happened in this job. "
            "Never repeat an approach recorded as failed or reverted; form a "
            "NEW hypothesis instead:\n"
            + "\n".join(f"- {l}" for l in reversed(lines)) + "\n")


def _mem_digest(mem, limit=2000):
    """Compact cross-job carryover: request + the last few events."""
    if mem is None:
        return ""
    out = f"request: {mem['question'][:200]}"
    tail, total = [], 0
    for ev in reversed(mem["events"]):
        if total + len(ev) > limit:
            break
        tail.append(ev)
        total += len(ev)
    if tail:
        out += "\nwhat happened: " + " | ".join(reversed(tail))
    return out


def _execute_step(job_id, step, idx, total, goal, outline, files,
                  step_notes, warnings, step_system, foreman_system,
                  mem=None, temp=None):
    """Run one build step end to end: minimal-context generation, guarded
    apply, then the foreman check/fix loop. Mutates files/step_notes/warnings
    in place. Returns False only when the specialist is unreachable (callers
    should abort the build)."""
    _set_stage(job_id, f"step {idx}/{total}: {step['title']}"[:60])

    # Context: this step's files in full. A small project rides along
    # whole, so cross-file references (ids, function names) stay visible.
    step_ctx = {n: files[n] for n in step["files"] if n in files}
    if files and sum(len(v) for v in files.values()) <= small_project_chars():
        step_ctx = dict(files)
    if files and not step_ctx:
        # Synthetic steps (inspector catch-ups, browser fixes) often name no
        # files. A step must NEVER run blind on an existing project — pull in
        # the files the step text mentions, else fall back to the whole set
        # and let the packer keep what fits.
        detail_lc = step["detail"].lower()
        step_ctx = {n: files[n] for n in files
                    if os.path.basename(n).lower() in detail_lc
                    or n.lower() in detail_lc} or dict(files)

    # Pack explicitly so we KNOW which files the model saw — set(shown)
    # becomes the write allowlist for this call. Explicit step targets go
    # first and may use the hard budget, so a single file bigger than the
    # soft cap stays visible and diff-editable instead of locked out forever.
    # Then entry html (it wires everything), then smallest-first.
    shown, unseen, ctx_total, packed = {}, [], 0, set()
    order = [(n, step_ctx[n]) for n in step["files"] if n in step_ctx]
    order += sorted(((n, c) for n, c in step_ctx.items()
                     if n not in step["files"]),
                    key=lambda kv: (not kv[0].endswith((".html", ".htm")),
                                    len(kv[1])))
    for n, code in order:
        if n in packed:
            continue
        packed.add(n)
        cap = step_file_hard() if n in step["files"] else step_file_chars()
        if ctx_total + len(code) > cap:
            unseen.append(n)
            continue
        shown[n] = code
        ctx_total += len(code)
    others = [n for n in files if n not in shown and n not in unseen]

    manifest = _symbol_manifest(files)
    user = f"Project goal: {goal}\n\nFull plan:\n{outline}\n\n"
    if mem:
        user += _mem_render(mem) + "\n"
    if manifest:
        user += ("PROJECT MAP — every file's top-level names. NEVER redeclare "
                 "a name another file owns; use the exact names shown:\n"
                 f"{manifest}\n\n")
    user += (f"YOUR STEP — {idx} of {total}: {step['title']}\n"
            f"{step['detail']}\n"
            f"Done when: {step['done_when']}\n")
    if step["files"]:
        user += f"Files for this step: {', '.join(step['files'])}\n"
    if shown:
        user += "\n" + _files_context(
            shown, limit=step_file_chars() * 2,
            header="Current contents of the step's files:")
    if unseen:
        user += ("\n\nToo large to show — do NOT output these files, "
                 "they are kept as-is: " + ", ".join(unseen))
    if others:
        user += ("\n\nOther project files exist but are not shown: "
                 + ", ".join(others))
    if unseen or others:
        user += ("\n\nIf this step cannot be done correctly without seeing "
                 "one of the hidden files, reply with ONLY one line — "
                 "NEED FILES: <comma-separated names> — and you will be "
                 "called again with them included.")
    user += _FILE_FORMAT_NOTE
    if shown:
        user += _EDIT_FORMAT_NOTE

    if _job_cancelled(job_id):
        step_notes.append(f"{idx}. {step['title']} — skipped (stopped)")
        return True
    _trace(job_id, "step_start", step=idx, title=step["title"],
           shown=list(shown), shown_chars=ctx_total,
           unseen=unseen, others=others,
           project_files=len(files), project_chars=sum(len(v) for v in files.values()),
           whole_project=(len(shown) == len(files) and not unseen and not others))
    changed = {}
    # a step with ONE target file lets _extract_files rescue an unlabeled block
    step_target = step["files"][0] if len(step["files"]) == 1 else (
        list(shown)[0] if len(shown) == 1 else None)
    call_mark = _job_notes_len(job_id)
    try:
        out = _stream_chat(
            coder_url(),
            [{"role": "system", "content": step_system},
             {"role": "user", "content": user}], job_id,
            temperature=temp or STEP_TEMP, label=f"step{idx}")
        # a reasoning model's think-text can carry stray code fences
        out = _strip_think(out)
        # NEED-FILES protocol (Aider-style): the model may ask once for
        # hidden files instead of guessing blind; we re-call with them shown
        nf = (re.match(r"\s*NEED FILES:\s*([^\n]+?)\s*$", out.strip(), re.I)
              if "```" not in out else None)
        if nf:
            req = [x.strip().strip("`'\"") for x in nf.group(1).split(",")]
            extra, budget = {}, ctx_total
            for nm in req:
                if (nm in files and nm not in shown
                        and budget + len(files[nm]) <= step_file_hard() * 2):
                    extra[nm] = files[nm]
                    budget += len(files[nm])
            if extra:
                _mem_note(mem, f"Step {idx}: model asked for hidden files "
                               f"({', '.join(extra)}) — provided")
                shown.update(extra)
                unseen = [n for n in unseen if n not in extra]
                user += ("\n\n" + _files_context(
                             extra, limit=step_file_hard() * 2,
                             header="The files you requested (the 'do not "
                                    "output' restriction is lifted for these):")
                         + "\nYou now have everything — complete the step; "
                           "do not reply NEED FILES again.")
            else:
                _mem_note(mem, f"Step {idx}: model asked for files that "
                               f"don't exist or don't fit "
                               f"({nf.group(1)[:100]}) — refused")
                user += ("\n\nThe files you asked for do not exist or cannot "
                         "be shown. Complete the step with what you have; "
                         "do not reply NEED FILES again.")
            call_mark = _job_notes_len(job_id)
            out = _stream_chat(
                coder_url(),
                [{"role": "system", "content": step_system},
                 {"role": "user", "content": user}], job_id,
                temperature=temp or STEP_TEMP, label=f"step{idx}-needfiles")
            out = _strip_think(out)
        new = _extract_files(out, existing=files, default_name=step_target)
        # Merge, and if NOTHING survives the guards, retry ONCE with the real
        # reason fed back. Testing only `new` (files PARSED) was a bug: a
        # cap-truncated file parses fine and is then rejected at merge, so the
        # retry never fired while the step reported "produced no usable files".
        # The condition that matters is `changed` — files ACCEPTED.
        for attempt in range(2):
            # only the MOST RECENT call's notes count — a cap from attempt 1
            # must not inflate attempt 2's budget or judgement
            capped = any("token cap" in s
                         for s in _job_notes_since(job_id, call_mark))
            new = _drop_truncated(new, out, idx, warnings, mem)
            wmark = len(warnings)
            _trace(job_id, "parse", step=idx, attempt=attempt,
                   files=list(new), count=len(new))
            rejected = []
            for name, code in new.items():
                wpre = len(warnings)
                merged = _merge_output_file(name, code, files, set(shown),
                                            warnings, f"step {idx}", mem)
                _trace(job_id, "merge", step=idx, file=name,
                       verdict="accepted" if merged is not None else "REJECTED",
                       in_chars=len(code), prev_chars=len(files.get(name, "")),
                       out_chars=len(merged or ""),
                       reason=(warnings[wpre:] or [""])[0][:200])
                if merged is not None:
                    changed[name] = merged
                else:
                    rejected.append((name, (warnings[wpre:] or [""])[0]))
            # A PARTIAL step used to report plain success: one file accepted,
            # another silently dropped, `changed` non-empty, nobody told.
            # Retry whenever a rejection is something the model can fix.
            fixable = [n for n, w_ in rejected if _RECOVERABLE_REJECT.search(w_)]
            if (changed and not fixable) or attempt or _job_cancelled(job_id):
                break
            why = "; ".join(w.split(": ", 1)[-1] for w in warnings[wmark:]) \
                  or "the reply contained no complete file"
            if changed and fixable:
                _mem_note(mem, f"Step {idx}: {', '.join(sorted(changed))} landed "
                               f"but {', '.join(fixable)} was REJECTED ({why}) "
                               f"— retrying for the rejected file(s)")
            else:
                _mem_note(mem, f"Step {idx}: attempt 1 produced nothing usable "
                               f"({why}) — retrying once")
            call_mark = _job_notes_len(job_id)
            out = _stream_chat(
                coder_url(),
                [{"role": "system", "content": step_system},
                 {"role": "user", "content": _retry_note(why, fixable) + user}],
                job_id, temperature=temp or STEP_TEMP,
                max_tokens=step_max_tokens() * 2 if capped else None,
                label=f"step{idx}-retry")
            out = _strip_think(out)
            new = _extract_files(out, existing=files, default_name=step_target)
    except Exception as e:
        warnings.append(f"step {idx} ({step['title']}): aborted — {e}")
        step_notes.append(f"{idx}. {step['title']} — FAILED (specialist unreachable)")
        return False
    still_bad = [n for n, _ in rejected if n not in changed]
    _trace(job_id, "step_end", step=idx, title=step["title"],
           changed=sorted(changed), usable=bool(changed),
           rejected=still_bad, partial=bool(changed and still_bad))
    if changed and still_bad:
        warnings.append(f"step {idx} ({step['title']}): PARTIAL — "
                        f"{', '.join(still_bad)} rejected and NOT applied")
        _mem_note(mem, f"Step {idx} was PARTIAL: {', '.join(sorted(changed))} "
                       f"landed, {', '.join(still_bad)} did NOT — that work is "
                       f"still missing.")
    if not changed:
        warnings.append(f"step {idx} ({step['title']}): produced no usable files")
        step_notes.append(f"{idx}. {step['title']} — skipped (no output)")
        _mem_note(mem, f"Step {idx} '{step['title']}': produced NO usable "
                       f"files (output empty, truncated, or rejected)")
        return True
    files.update(changed)
    _mem_note(mem, f"Step {idx} '{step['title']}': edited "
                   + ", ".join(sorted(changed)))

    # ---- foreman: keep the step in line before the next one builds on it ----
    note = None
    prev_issues: set = set()
    for _ in range(step_fix_rounds()):
        issues = _static_check_files(changed) + [
            d for d in _dup_symbol_check(files)
            if any(n in d for n in changed)] + [
            b for b in _bundle_check(files)
            if "does not exist" in b]
        # Only pay for an LLM review when something deterministic is wrong.
        # 18 foreman calls across two runs cost 39 minutes and merged zero
        # fixes; every one of them was handed "No issues found".
        if not issues and FOREMAN_ONLY_ON_ISSUES:
            break
        survived = [i for i in issues if i in prev_issues]
        if survived:
            _mem_note(mem, f"Step {idx}: these check failures SURVIVED the "
                           f"previous foreman fix (that approach was wrong — "
                           f"do not repeat it): " + "; ".join(survived[:4]))
        prev_issues = set(issues)
        _set_stage(job_id, f"checking step {idx}/{total}")
        checks = ("\n".join(f"- {i}" for i in issues) if issues
                  else "No issues found by automated syntax checks.")
        f_shown, f_total = {}, 0
        for n, c in sorted(changed.items(), key=lambda kv: len(kv[1])):
            if f_total + len(c) > step_file_hard():
                continue
            f_shown[n] = c
            f_total += len(c)
        check_user = (
            (_mem_render(mem) + "\n" if mem else "")
            + f"Step {idx} of {total}: {step['title']}\n{step['detail']}\n"
            f"Done when: {step['done_when']}\n\n"
            f"Automated syntax-check results:\n{checks}\n\n"
            + _files_context(f_shown, limit=step_file_hard() * 2,
                             header="Files produced by this step:"))
        f_skipped = [n for n in changed if n not in f_shown]
        if f_skipped:
            check_user += ("\n\nAlso produced but too large to show — "
                           "judge only what you can see, do NOT rewrite "
                           "these: " + ", ".join(f_skipped))
        check_user += _EDIT_FORMAT_NOTE
        try:
            verdict = _stream_chat(
                coder_url(),
                [{"role": "system", "content": foreman_system},
                 {"role": "user", "content": check_user}],
                job_id, temperature=0.1, label=f"step{idx}-foreman",
                # A check is a BOUNDED job: "ON TRACK", or the corrected files
                # it was shown. The prose rule holds for single-file steps but
                # not when a step touched several files at once — observed live
                # at 12k tokens and climbing. This ceiling is ~2x the largest
                # legitimate fix; a truncated tail is discarded by the existing
                # unclosed-fence guard, and the step has already banked its work.
                max_tokens=int(os.environ.get("CREWE_FOREMAN_MAX_TOKENS", "12288")))
        except Exception as e:
            print(f"[code] foreman error: {e}")
            warnings.append(f"step {idx}: foreman check skipped "
                            f"(specialist unreachable)")
            break
        clean = _strip_think(verdict).strip()
        # verdict may arrive wrapped in prose ("The step is on track and…") —
        # accept it from the first real line, but never past a negation
        head = next((l for l in clean.splitlines() if l.strip()), "").upper()
        # \b on NOT so NOTE / NOTHING / ANOTHER / CANNOT / NOTICED don't read
        # as negations and reject a genuine ON TRACK verdict
        approves = ("ON TRACK" in head
                    or re.search(r"\b(CORRECT AND COMPLETE|COMPLETE AND CORRECT"
                                 r"|FULLY IMPLEMENTED|NO CHANGES? (?:ARE )?NEEDED"
                                 r"|NOTHING TO FIX|LOOKS CORRECT)\b", head))
        if (approves and not re.search(r"\bNOT\b|N'T\b", head)
                and "```" not in clean):
            break
        fixes = {}
        f_target = list(f_shown)[0] if len(f_shown) == 1 else None
        for n, c in _extract_files(clean, existing=files,
                                   default_name=f_target).items():
            merged = _merge_output_file(n, c, files, set(f_shown), warnings,
                                        f"step {idx} foreman", mem)
            if merged is not None:
                fixes[n] = merged
        if not fixes:
            # an objection with no usable fix must not pass as a silent ✓
            objection = (clean.splitlines() or ["no verdict"])[0][:160]
            warnings.append(f"step {idx}: foreman flagged an issue that "
                            f"was not auto-fixed: {objection}")
            _mem_note(mem, f"Step {idx}: foreman objected but produced no "
                           f"usable fix: {objection}")
            break
        files_before, changed_before = dict(files), dict(changed)
        changed.update(fixes)
        files.update(fixes)
        if len(_static_check_files(changed)) > len(issues):
            # never let a "fix" ship a step worse than the draft it judged
            files.clear()
            files.update(files_before)
            changed = changed_before
            warnings.append(f"step {idx}: a foreman fix made the syntax "
                            f"checks worse and was reverted")
            _mem_note(mem, f"Step {idx}: a foreman fix touching "
                           + ", ".join(sorted(fixes))
                           + " made the syntax checks WORSE and was REVERTED "
                             "— that approach failed, do not retry it")
            break
        fn = re.search(r"FIXED:\s*(.+)", clean)
        note = fn.group(1).strip() if fn else "foreman revised the step"
        _mem_note(mem, f"Step {idx} foreman round: edited "
                       + ", ".join(sorted(fixes)) + f" — {note}")
    step_notes.append(f"{idx}. {step['title']} "
                      + ("⚠ partial — " + ", ".join(still_bad) + " not applied"
                         if still_bad else "✓")
                      + (f" — 🔧 {note}" if note else ""))
    return True


def run_code_pipeline(job_id: str, question: str, session_id: str):
    """Agentic code route: spec -> execute steps -> foreman-check each ->
    delivery inspection -> browser runtime check.

    The planner writes a spec (goal, promised features, visual direction)
    and the smallest steps that deliver ALL of it; each step is coded in its
    own fresh call and foreman-checked before the next builds on it. The
    inspector then compares the finished build against the promised features
    and builds what's missing (bounded by extra_steps()), and a headless-
    Firefox pass catches runtime errors static checks can't see. Every role
    runs on coder_url() — one model, no mixing."""
    with SESSIONS_LOCK:
        sess = sessions().get(session_id, {})
        proj_files = dict(sess.get("files") or {})
        history = list(sess.get("history") or [])
        prev_mem = sess.get("job_memory")
    prefs = _load_prefs()

    # Follow-up on an existing build → patch mode: minimal plan, and the job
    # memory carries what the previous job on this project already tried.
    followup = bool(proj_files)
    mem = _mem_new(question, followup)
    if followup and prev_mem:
        mem["prev"] = prev_mem
    try:
        _coder_model = (requests.get(
            coder_url().rsplit("/v1/", 1)[0] + "/props", timeout=4
        ).json().get("model_path") or "?").rsplit("/", 1)[-1]
    except Exception:
        _coder_model = "?"
    _trace(job_id, "job_start", question=question[:500], session=session_id,
           coder_model=_coder_model,
           mode="PATCH" if followup else "BUILD",
           project_files=sorted(proj_files),
           project_chars=sum(len(v) for v in proj_files.values()),
           whole_project_budget=small_project_chars(), coder=coder_url())
    print(f"[code] job {job_id} start: "
          f"{'PATCH' if followup else 'BUILD'} mode, "
          f"{len(proj_files)} files / {sum(len(v) for v in proj_files.values())} chars "
          f"-> trace: {os.path.join(TRACE_DIR, job_id)}")

    # ---- patch mode: REPRODUCE the report before touching anything ----------
    # SWE-agent pattern: get a failing check on the CURRENT build first, so
    # "fixed" later means that exact check passes — ground truth, not vibes.
    behavior_checks = []
    if followup and any(n.endswith((".html", ".htm")) for n in proj_files):
        _set_stage(job_id, "writing repro checks")
        behavior_checks = _write_checks(job_id, None, proj_files,
                                        patch_bug=question, mem=mem)
        if behavior_checks:
            _set_stage(job_id, "reproducing the report")
            results, fatal = _run_behavior_checks(proj_files, behavior_checks)
            if fatal:
                behavior_checks = []
            else:
                # a check whose ACTION crashed (bad selector) can never pass —
                # drop it so the fix loop doesn't chase a broken check
                bad = {r["desc"] for r in results if r["action_error"]}
                if bad:
                    behavior_checks = [c for c in behavior_checks
                                       if _check_desc(c) not in bad]
                failing = [r for r in results
                           if not r["ok"] and not r["action_error"]]
                if failing:
                    _mem_note(mem, "REPRODUCED on the current build — these "
                              "checks FAIL (ground truth; the fix must make "
                              "them pass): " + "; ".join(
                                  f"{r['desc']} ({r['detail']})"
                                  for r in failing[:4]))
                elif behavior_checks:
                    _mem_note(mem, "Behavior checks all PASS on the current "
                              "build — the reported problem was not captured "
                              "by them; they serve as regression guards only")

    # ---- plan ----
    _set_stage(job_id, "planning")
    planner_system = PLANNER_SYSTEM
    if not proj_files:
        # A fresh session has NO files. Users re-paste accumulated prompts
        # full of "already done ✓" claims; believing them plans features onto
        # a base that does not exist (observed: an app whose sidebar, canvas
        # and palette were all "✓" and none were ever built — shipped with no
        # way to add an item at all).
        planner_system += (
            "\n\nGROUND TRUTH: the project currently has ZERO files. Ignore "
            "any claim in the request that a part is already done, checked "
            "off, or implemented — NOTHING exists yet. Plan every feature "
            "the app needs to actually work, including the base skeleton, "
            "and the interactions that CREATE content (e.g. if items can be "
            "placed, plan the UI that adds them).")
    if prefs:
        planner_system += ("\n\nStanding design standards the spec must "
                           "honor:\n" + prefs)
    if followup:
        planner_system += PATCH_PLANNER_NOTE
    plan_user = ""
    recent_qs = [t.get("q", "")[:200] for t in history[-3:] if t.get("q")]
    if recent_qs:
        plan_user += ("Recent requests in this session (context only):\n"
                      + "\n".join(f"- {q}" for q in recent_qs) + "\n\n")
    if proj_files:
        # smallest-first with entry html leading, sized to the step budget —
        # the planner used to get a hardcoded 30k slice of arbitrary order
        # and silently planned against files it never saw
        ordered = dict(sorted(proj_files.items(),
                              key=lambda kv: (not kv[0].endswith((".html", ".htm")),
                                              len(kv[1]))))
        plan_user += _files_context(ordered, limit=step_file_chars(),
                                    header="Existing project files:") + "\n\n"
    plan_user += (("Follow-up request on the existing build (change ONLY "
                   "what it asks):\n" if followup else "Build request:\n")
                  + question)
    try:
        # Generous budget: a reasoning model's think-text counts against
        # max_tokens, and a starved cap truncates the JSON mid-object —
        # the same failure mode that once broke the classifier.
        plan_text = _stream_chat(
            coder_url(),
            [{"role": "system", "content": planner_system},
             {"role": "user", "content": plan_user}],
            job_id, temperature=0.15, max_tokens=16384, label="planner")
    except Exception as e:
        print(f"[code] planner error: {e}")
        plan_text = ""
    plan = _parse_plan(plan_text, question)
    steps = plan["steps"]
    total = len(steps)
    with SESSIONS_LOCK:
        sessions().setdefault(session_id, {})["last_plan"] = plan

    # Standards ride in the system prompt: stable across a build's calls
    # (good for llama.cpp prefix cache), and the model treats system-level
    # rules as law rather than as one more user suggestion.
    standards = ""
    if prefs:
        standards += f"\n\nStanding design standards:\n{prefs}"
    if plan.get("design"):
        standards += f"\n\nProject visual direction:\n{plan['design']}"
    if plan.get("features"):
        standards += ("\n\nPromised features (the finished build must have "
                      "all of these):\n"
                      + "\n".join(f"- {f}" for f in plan["features"]))
    step_system = STEP_SYSTEM + standards
    foreman_system = FOREMAN_SYSTEM + standards

    outline = "\n".join(f"{i}. {s['title']}" for i, s in enumerate(steps, 1))
    files = dict(proj_files)
    step_notes, warnings = [], []
    _mem_note(mem, ("PATCH-mode plan" if followup else "Build plan")
                   + f" ({total} step{'s' if total != 1 else ''}): "
                   + "; ".join(s["title"] for s in steps))

    for idx, step in enumerate(steps, 1):
        if _job_cancelled(job_id):
            step_notes.append("⏹ build stopped by user")
            break
        if not _execute_step(job_id, step, idx, total, plan["goal"], outline,
                             files, step_notes, warnings,
                             step_system, foreman_system, mem):
            break

    # ---- delivery inspection: did the build keep the spec's promises? ----
    if plan.get("features") and files and not _job_cancelled(job_id):
        _set_stage(job_id, "final inspection")
        feat_lines = "\n".join(f"- {f}" for f in plan["features"])
        # Pack smallest-first (html leading) so the most files fit, and TELL
        # the inspector what didn't fit — it cannot tell "not shown" from
        # "doesn't exist", and phantom 'missing file' verdicts burn real
        # build steps recreating files that are already there.
        shown, omitted, tot = {}, [], 0
        for n, c in sorted(files.items(),
                           key=lambda kv: (not kv[0].endswith((".html", ".htm")),
                                           len(kv[1]))):
            if tot + len(c) <= step_file_chars():
                shown[n] = c
                tot += len(c)
            else:
                omitted.append(n)
        insp_user = (f"Promised features:\n{feat_lines}\n\n"
                     + _files_context(shown, limit=step_file_chars() * 2,
                                      header="Project files:"))
        if omitted:
            insp_user += ("\n\nThese files also EXIST in the build but are "
                          "not shown for length. NEVER report them or their "
                          "contents as missing: " + ", ".join(omitted))
        try:
            verdict = _stream_chat(
                coder_url(),
                [{"role": "system", "content": COMPLETE_SYSTEM + standards},
                 {"role": "user", "content": insp_user}],
                job_id, temperature=0.1, max_tokens=8192, label="inspector")
            clean = _strip_think(verdict).strip()
        except Exception as e:
            print(f"[code] inspector error: {e}")
            clean = "COMPLETE"
        missing = _inspector_missing(clean)
        if missing or not clean.upper().startswith("COMPLETE"):
            if not missing:
                _mem_note(mem, "Delivery inspector returned prose with no "
                               "parseable missing-feature lines — treated "
                               "as complete")
            for k, m in enumerate(missing, 1):
                xstep = {"title": m[:60],
                         "detail": ("The delivery inspector found this "
                                    "missing or broken — implement a WORKING "
                                    "version (a described error is a defect "
                                    "to fix, never a spec to preserve): "
                                    + m),
                         "files": [], "done_when": m}
                step_notes.append(f"🔍 inspector: missing — {m[:120]}")
                _mem_note(mem, f"Delivery inspector: '{m[:120]}' judged "
                               f"missing/hollow — building it now")
                if not _execute_step(job_id, xstep, total + k,
                                     total + len(missing), plan["goal"],
                                     outline, files, step_notes, warnings,
                                     step_system, foreman_system, mem):
                    break

    AUTOFIX = os.environ.get("CREWE_AUTOFIX", "doa")
    VERIFY_BUDGET = int(os.environ.get("CREWE_VERIFY_BUDGET", "600"))
    verify_t0 = time.time()

    def _verify_spent():
        return time.time() - verify_t0

    def _budget_left(stage_name):
        if _verify_spent() < VERIFY_BUDGET:
            return True
        msg = (f"verification budget exhausted after "
               f"{_verify_spent()/60:.0f} min — {stage_name} findings ship "
               f"as warnings instead of more fix rounds")
        if msg not in warnings:
            warnings.append(msg)
            _trace(job_id, "gate", gate="budget", phase="exhausted",
                   stage=stage_name, spent_s=round(_verify_spent()))
            _mem_note(mem, msg)
        return False

    # ---- runtime gate: cross-file bundle check + real browser, with up to
    # two fix rounds. One round proved not enough when several errors share
    # a root cause the first fix only half-lands.
    if (files and any(n.endswith((".html", ".htm")) for n in files)
            and not _job_cancelled(job_id)):
        prev_gate: list = []
        for attempt in range(3):
            if _job_cancelled(job_id):
                break
            _set_stage(job_id, "browser check" if attempt == 0
                       else "browser re-check")
            gate_errors = _dup_symbol_check(files) + _bundle_check(files)
            runtime_errors, _shot = _browser_check(files)
            _trace(job_id, "gate", gate="browser", phase="result",
                   attempt=attempt, screenshot=bool(_shot),
                   static_errors=len(gate_errors), runtime_errors=runtime_errors[:6])
            gate_errors += [e for e in runtime_errors if e not in gate_errors]
            survived = [e for e in gate_errors if e in prev_gate]
            if survived:
                _mem_note(mem, "Browser errors that SURVIVED the previous "
                               "fix (its diagnosis was wrong — a DIFFERENT "
                               "approach is required): "
                               + "; ".join(survived[:6]))
            if not gate_errors:
                step_notes.append("🌐 browser check: loads clean"
                                  if attempt == 0 else
                                  "🌐 browser re-check: loads clean")
                break
            if AUTOFIX != "full" and not runtime_errors:
                # page loads and runs; remaining findings are quality, not
                # death — deliver now, fix on request
                warnings.extend(f"check finding: {e}" for e in gate_errors[:6])
                step_notes.append("🌐 browser check: loads, with findings")
                break
            if AUTOFIX == "off":
                warnings.extend(f"check finding: {e}" for e in gate_errors[:6])
                break
            if attempt == 2:
                warnings.append("browser check still failing: "
                                + "; ".join(gate_errors[:3]))
                _mem_note(mem, "Job ended with browser errors UNFIXED after "
                               "2 attempts: " + "; ".join(gate_errors[:4]))
                break
            if attempt == 0:
                step_notes.append("🌐 browser check found problems")
            detail = ("The finished build was loaded in a real browser "
                      "and checked as a whole; these problems were found:\n"
                      + "\n".join(f"- {e}" for e in gate_errors[:8]))
            if survived:
                detail += ("\nThe problems marked in working memory as "
                           "SURVIVED were not cured by the previous fix — "
                           "re-diagnose them from scratch instead of "
                           "repeating that edit.")
            if attempt == 1:
                detail += ("\nTargeted edits already failed once — rewrite "
                           "the ENTIRE smallest file that owns the problem "
                           "(output the full file, not SEARCH/REPLACE edits).")
            detail += ("\nFix every problem without removing features. "
                       "Prefer targeted SEARCH/REPLACE edits."
                       if attempt == 0 else
                       "\nFix every problem without removing features.")
            fix_step = {
                "title": "fix runtime errors",
                "detail": detail,
                "files": [], "done_when": "the page loads and draws without errors"}
            if not _budget_left("browser gate"):
                break
            prev_gate = list(gate_errors)
            if not _execute_step(job_id, fix_step, 1, 1, plan["goal"],
                                 outline, files, step_notes, warnings,
                                 step_system, foreman_system, mem,
                                 temp=0.35 if attempt == 1 else None):
                warnings.append("browser fix aborted — specialist unreachable")
                break

    # ---- behavior verification: the build must PASS its checks --------------
    # Patch mode re-runs the repro checks written before the fix; build mode
    # writes checks from the promised features now that selectors exist.
    if (files and any(n.endswith((".html", ".htm")) for n in files)
            and not _job_cancelled(job_id)):
        if not behavior_checks and plan.get("features"):
            _set_stage(job_id, "writing behavior checks")
            feats = "Features the finished build promises:\n- " \
                    + "\n- ".join(plan["features"][:4])
            behavior_checks = _write_checks(job_id, feats, files, mem=mem)
        for vround in range(3):
            if not behavior_checks or _job_cancelled(job_id):
                break
            _set_stage(job_id, "verifying behavior")
            results, fatal = _run_behavior_checks(files, behavior_checks)
            failing = [r for r in results if not r["ok"]]
            _trace(job_id, "gate", gate="behavior", phase="result", attempt=vround,
                   checks=len(results), failed=len(failing), fatal=bool(fatal),
                   failing=[r["desc"][:60] for r in failing[:4]])
            if fatal:
                break
            # In PATCH mode the checks are derived from the reported bug,
            # so ANY failure means the requested fix is unverified — that is
            # a mission failure regardless of the pass ratio. (Observed: a
            # drag-fix job shipped with its repro check failing because 1/4
            # was a "minority".)
            functionally_dead = (len(failing) > len(results) // 2
                                 or (followup and bool(failing)))
            if (AUTOFIX == "doa" and functionally_dead and vround == 0
                    and _budget_left("behavior checks")):
                warnings.append(f"build is functionally dead — "
                                f"{len(failing)}/{len(results)} behavior "
                                f"checks fail; running one repair round")
                _mem_note(mem, f"{len(failing)}/{len(results)} behavior checks "
                               f"FAILED — the app cannot be used as built. One "
                               f"repair round runs; further findings ship as "
                               f"warnings.")
            elif AUTOFIX != "full" and failing:
                step_notes.append(f"✔ behavior checks: "
                                  f"{len(results)-len(failing)}/{len(results)} pass")
                warnings.extend(
                    f"behavior check failed: {r['desc']} ({str(r.get('detail'))[:80]})"
                    for r in failing[:5])
                _mem_note(mem, "Behavior checks reported (deliver-first): "
                          + "; ".join(r["desc"] for r in failing[:4]))
                break
            if not failing:
                step_notes.append(f"✅ behavior checks: "
                                  f"{len(results)}/{len(results)} pass")
                _mem_note(mem, f"All {len(results)} behavior checks PASS — "
                               f"the change is verified working")
                break
            _mem_note(mem, "Behavior checks FAILING: " + "; ".join(
                f"{r['desc']} ({r['detail']})" for r in failing[:4]))
            if vround == 2:
                warnings.append("behavior checks still failing: " + "; ".join(
                    r["desc"] for r in failing[:3]))
                step_notes.append(f"⚠ behavior checks: "
                                  f"{len(results) - len(failing)}"
                                  f"/{len(results)} pass")
                break
            detail = ("These checks ran in a real browser against the "
                      "current build and FAILED:\n"
                      + "\n".join(f"- {r['desc']} — {r['detail']}"
                                  for r in failing[:6])
                      + "\nMake them pass without breaking working behavior.")
            if vround == 1:
                detail += ("\nTargeted edits already failed once — rewrite "
                           "the ENTIRE smallest file that owns the failing "
                           "behavior (output the full file, not "
                           "SEARCH/REPLACE edits).")
            vfix = {"title": "make behavior checks pass", "detail": detail,
                    "files": [], "done_when": "all behavior checks pass"}
            if not _budget_left("behavior checks"):
                break
            if not _execute_step(job_id, vfix, 1, 1, plan["goal"], outline,
                                 files, step_notes, warnings, step_system,
                                 foreman_system, mem,
                                 temp=0.35 if vround == 1 else None):
                break

    # ---- visual verification: does it LOOK like what was asked for? --------
    # Runs LAST, once the build loads clean and behaves: selector checks pass
    # happily on a page that rendered as an unstyled skeleton (observed — a
    # build passed 8/8 checks while showing nothing but default-styled text),
    # so the only thing that can judge "did a UI actually appear" is an eye.
    if (files and any(n.endswith((".html", ".htm")) for n in files)
            and not _job_cancelled(job_id)):
        prev_vis: list = []
        for vattempt in range(3):
            if _job_cancelled(job_id):
                break
            _set_stage(job_id, "visual check" if vattempt == 0
                       else "visual re-check")
            _errs, shot = _browser_check(files)
            ready = _vision_ready(_vision_url())
            _trace(job_id, "gate", gate="vision", phase="run", attempt=vattempt,
                   screenshot=bool(shot), vision_ready=ready,
                   url=_vision_url())
            defects = _vision_critique(shot, plan["goal"],
                                       plan.get("features") or [], job_id)
            _trace(job_id, "gate", gate="vision", phase="result",
                   attempt=vattempt, defects=defects,
                   verdict=("skipped" if not (shot and ready)
                            else "looks right" if not defects else "defects"))
            if not defects:
                if shot and _vision_ready(_vision_url()):
                    step_notes.append("👁 visual check: looks right"
                                      if vattempt == 0 else
                                      "👁 visual re-check: looks right")
                break
            if AUTOFIX != "full":
                step_notes.append("👁 visual check: findings reported")
                warnings.extend(f"visual finding: {d}" for d in defects[:4])
                _mem_note(mem, "Visual review (deliver-first): "
                          + "; ".join(defects[:4]))
                break
            survived = [d for d in defects if d in prev_vis]
            if survived:
                _mem_note(mem, "Visual problems that SURVIVED the previous "
                               "fix (that diagnosis was wrong — try a "
                               "DIFFERENT approach): " + "; ".join(survived[:4]))
            if vattempt == 2:
                warnings.append("visual check still flags: "
                                + "; ".join(defects[:3]))
                step_notes.append("👁 visual check: still flagged")
                break
            if vattempt == 0:
                step_notes.append("👁 visual check found problems")
            _mem_note(mem, "Visual review of the rendered page: "
                           + "; ".join(defects[:5]))
            detail = ("The finished build was rendered in a real browser and "
                      "a REVIEWER LOOKED AT THE SCREENSHOT. These problems "
                      "are visible on the page:\n"
                      + "\n".join(f"- {d}" for d in defects[:5])
                      + "\nFix what is visibly wrong. The page must actually "
                        "render its interface — styled, laid out, and "
                        "visible — not merely contain the right elements. "
                        "Do not remove working features.")
            if survived:
                detail += ("\nThe problems marked SURVIVED were not cured by "
                           "the last edit — re-diagnose from scratch.")
            if vattempt == 1:
                detail += ("\nTargeted edits already failed once — rewrite "
                           "the ENTIRE smallest file that owns the layout.")
            if not _budget_left("visual check"):
                break
            prev_vis = list(defects)
            vstep = {"title": "fix what the page visibly gets wrong",
                     "detail": detail, "files": [],
                     "done_when": "the rendered page visibly shows the "
                                  "requested interface"}
            if not _execute_step(job_id, vstep, 1, 1, plan["goal"], outline,
                                 files, step_notes, warnings, step_system,
                                 foreman_system, mem,
                                 temp=0.35 if vattempt == 1 else None):
                warnings.append("visual fix aborted — specialist unreachable")
                break

    # ---- assemble the answer ----
    changed_overall = {n: c for n, c in files.items() if proj_files.get(n) != c}
    final_warnings = _static_check_files(changed_overall) + warnings
    prose = f"**Build:** {plan['goal']}\n\n" + "\n".join(step_notes)
    if changed_overall:
        unchanged = [n for n in proj_files if n not in changed_overall]
        final_answer = _files_to_markdown(changed_overall, prose=prose,
                                          unchanged=unchanged,
                                          warnings=final_warnings)
        _store_files(session_id, changed_overall, force=True)
    else:
        detail = f" — {warnings[-1]}" if warnings else ""
        final_answer = (prose
                        + f"\n\n[router error: the build produced no files{detail}]")

    # stream-level findings (caps, loops) surface in the answer + jobmem —
    # a degraded call must never be invisible to the user
    with JOBS_LOCK:
        stream_notes = list(JOBS.get(job_id, {}).get("notes", []))
    for n in stream_notes:
        final_warnings.append(n)
        _mem_note(mem, f"Stream note: {n}")
    _trace(job_id, "job_end", files_changed=sorted(changed_overall),
           warnings=len(final_warnings), steps=len(step_notes))
    _mem_note(mem, "Job finished. Files changed overall: "
                   + (", ".join(sorted(changed_overall)) or "none"))
    with SESSIONS_LOCK:
        sessions().setdefault(session_id, {})["job_memory"] = _mem_digest(mem)
    try:  # full log for humans: ~/llama_logs/jobmem/<job_id>.md
        memdir = os.path.expanduser("~/llama_logs/jobmem")
        os.makedirs(memdir, exist_ok=True)
        with open(os.path.join(memdir, f"{job_id}.md"), "w") as fh:
            fh.write(_mem_render(mem, limit=100000))
    except Exception:
        pass

    _append_history(session_id, question, final_answer, "code")
    with JOBS_LOCK:
        JOBS[job_id]["answer"] = final_answer
        JOBS[job_id]["done"] = True
    spawn_owned(target=summarize, args=(session_id, question, final_answer),
                     daemon=True).start()


def _specialist_error(route, url, exc):
    """A failure the USER can act on, instead of a requests traceback.

    On a fresh install nobody has a model server running yet, so this is the
    very first thing most people ever see Crewe say. "HTTPConnectionPool(...)
    Max retries exceeded with url: /v1/chat/completions" tells them nothing;
    "nothing is listening, here is where to fix it" tells them everything."""
    host = ""
    try:
        p = urlparse(url)
        host = f"{p.hostname}:{p.port}" if p.port else (p.hostname or "")
    except Exception:
        pass
    kind = exc.__class__.__name__
    unreachable = isinstance(exc, (requests.ConnectionError, requests.Timeout))
    if unreachable:
        return (f"\n\n**No model backend is reachable for the `{route}` route.**\n\n"
                f"Crewe tried `{host or url}` and nothing answered.\n\n"
                "- If your model server isn't running yet, start it.\n"
                "- If it lives somewhere else, open **⚙ Settings** and point "
                f"the `{route}` route at the right backend — *Scan localhost* "
                "will find servers already running on this machine.\n\n"
                f"*({kind})*")
    return (f"\n\n**The `{route}` backend returned an error.**\n\n"
            f"`{host or url}` — {kind}: {exc}\n\n"
            "If you've just changed models, check the server is finished "
            "loading; llama.cpp answers `/health` with 503 until it is.")


def run_job(job_id: str, route: str, question: str, session_id: str,
            effort: str = DEFAULT_EFFORT):
    """Stream from the specialist, counting tokens as they arrive, store final answer."""
    # Capture the effort level onto THIS thread, the same way spawn_owned
    # captures the owner — the pipeline reads it from a thread-local.
    set_effort(effort)
    if route == "audio":
        run_audio_job(job_id, question, session_id)
        return
    if route == "code":
        try:
            run_code_pipeline(job_id, question, session_id)
        except Exception as e:
            # a bug in the pipeline must never leave the UI polling forever
            print(f"[code] pipeline crashed: {e}")
            crashed = False
            with JOBS_LOCK:
                if not JOBS[job_id]["done"]:
                    JOBS[job_id]["answer"] = f"\n\n[router error in code pipeline: {e}]"
                    JOBS[job_id]["done"] = True
                    crashed = True
            if crashed:
                # keep the conversation record whole — a follow-up must see
                # that this request happened and failed
                _append_history(session_id, question,
                                f"[router error in code pipeline: {e}]", "code")
        return
    url = SPECIALISTS[route]

    # inject conversation memory into the system prompt if available
    system = SYSTEM_PROMPTS[route]
    with SESSIONS_LOCK:
        sess = sessions().get(session_id, {})
        summary = sess.get("summary", "")
        history = list(sess.get("history") or [])
    if summary:
        system += f"\n\nConversation context (older exchanges, summarized):\n{summary}"

    user_content = question
    search_sources: list = []

    # Attached documents go to WHATEVER route the question landed on — the file
    # belongs to the session, not to one question, so "summarise this" and
    # "what does clause 4 say?" both work off a single upload.
    ups = session_uploads(session_id)
    if ups:
        with JOBS_LOCK:
            JOBS[job_id]["stage"] = "reading"
        JOBS[job_id]["stage_ts"] = time.time()
        # Half the window for the document, leaving room for the system prompt,
        # conversation memory, the question and the whole answer.
        ctx = _backend_ctx(url)
        budget = int(ctx * 0.5) * 4 if ctx else 40000
        per_doc = max(2000, budget // len(ups))
        blocks = []
        for up in ups:
            body = _upload_body(up)
            if not body:
                continue
            if len(body) > per_doc:
                if route == "review" and "=== FILE: " in body:
                    # A summary of code cannot be reviewed. Keep WHOLE files
                    # until the budget is spent and say which were left out —
                    # a reviewer told what it has not seen beats one working
                    # from a paraphrase of everything.
                    head, _, rest = body.partition("\n\n=== FILE: ")
                    parts = ("=== FILE: " + rest).split("\n\n=== FILE: ")
                    kept, omitted, used = [], [], len(head)
                    for p in parts:
                        p = p if p.startswith("=== FILE: ") else "=== FILE: " + p
                        fname = p.split("===")[1].replace("FILE:", "").strip()
                        if used + len(p) <= per_doc:
                            kept.append(p)
                            used += len(p)
                        else:
                            omitted.append(fname)
                    note = (f"\n\n[{len(omitted)} files omitted for size: "
                            + ", ".join(omitted[:40])
                            + ("…" if len(omitted) > 40 else "") + "]"
                            if omitted else "")
                    body = head + "\n\n" + "\n\n".join(kept) + note
                else:
                    # Too big for the window: condense rather than truncate,
                    # so the tail of a long document is represented at all.
                    with JOBS_LOCK:
                        JOBS[job_id]["stage"] = "condensing"
                    JOBS[job_id]["stage_ts"] = time.time()
                    body = _condense(body, per_doc, url, job_id)
            blocks.append(f"--- attached file: {up['name']} ---\n{body}")
        if blocks:
            joined = "\n\n".join(blocks)
            user_content = (
                f"The user has attached {len(blocks)} file(s), shown below. "
                "Answer using their contents; if the answer is not in them, "
                f"say so rather than inventing it.\n\n{joined}\n\n"
                f"---\nUser's message: {question}")
        with JOBS_LOCK:
            JOBS[job_id]["stage"] = "drafting"
        JOBS[job_id]["stage_ts"] = time.time()

    # Search route: query the web, fetch top pages, hand results to the model.
    if route == "search":
        with JOBS_LOCK:
            JOBS[job_id]["stage"] = "searching"
        JOBS[job_id]["stage_ts"] = time.time()
        results = web_search(question)
        with JOBS_LOCK:
            JOBS[job_id]["stage"] = "reading"
        JOBS[job_id]["stage_ts"] = time.time()
        context, search_sources = build_search_context(results)
        if context:
            user_content = (
                "Live web search results for the question are below. Use them to "
                "answer and cite sources inline as [n] matching the numbered "
                f"results.\n\n{context}\n\n---\nQuestion: {question}"
            )
        else:
            user_content = (
                "(Web search returned no usable results — answer from your own "
                "knowledge and tell the user you could not verify this online.)\n\n"
                f"{question}"
            )
        # results are in hand; the model is now drafting the synthesized answer
        with JOBS_LOCK:
            JOBS[job_id]["stage"] = "drafting"
        JOBS[job_id]["stage_ts"] = time.time()

    # Recent turns go in as real chat history so follow-ups ("no", "the second
    # one", a 20-questions answer) land mid-conversation, not cold.
    messages = [{"role": "system", "content": system}]
    for turn in history:
        messages.append({"role": "user",      "content": turn.get("q", "")})
        messages.append({"role": "assistant", "content": turn.get("a", "")})
    messages.append({"role": "user", "content": user_content})

    payload = {"messages": messages, "stream": True}
    answer_parts = []
    usage = None
    try:
        payload = _with_usage(_inject_model(payload, url, route), url)
        with requests.post(url, json=payload, headers=_hdrs(url),
                           stream=True, timeout=600) as r:
            r.raise_for_status()
            for line in _cancellable_lines(r, job_id):
                if not line:
                    continue
                line = line.decode("utf-8")
                if line.startswith("data: "):
                    data = line[6:]
                    if data.strip() == "[DONE]":
                        break
                    try:
                        chunk = json.loads(data)
                        # The usage-bearing final chunk has an EMPTY choices
                        # list, so read it before touching choices[0] — the
                        # except below would otherwise swallow it silently.
                        if chunk.get("usage"):
                            usage = chunk["usage"]
                        if not chunk.get("choices"):
                            continue
                        d = chunk["choices"][0]["delta"]
                        delta = d.get("content", "")
                        if delta:
                            answer_parts.append(delta)
                        # count hidden reasoning too, so the UI counter moves
                        # while a thinking model is still in its think phase
                        if delta or d.get("reasoning_content"):
                            with JOBS_LOCK:
                                JOBS[job_id]["tokens"] += 1
                    except Exception:
                        pass
    except Exception as e:
        answer_parts.append(_specialist_error(route, url, e))

    final_answer = _scrub_tokens("".join(answer_parts))
    cost = record_spend(url, route, usage, user_content, final_answer)
    if cost:
        with JOBS_LOCK:
            JOBS[job_id]["cost"] = round(cost, 6)
    if _job_cancelled(job_id):
        final_answer += "\n\n*(stopped by user)*"

    # Search route: append the sources the model was given, as a clickable list,
    # so the user can follow the [n] citations back to the originals.
    if route == "search" and search_sources and not final_answer.startswith("\n\n[router error"):
        lines = "\n".join(f"{i}. [{title or url}]({url})" for i, title, url in search_sources)
        final_answer += f"\n\n---\n**Sources**\n{lines}"

    if final_answer:
        # Non-code routes can still produce code — capture real files into
        # memory so a later code-route follow-up has them as context.
        got = {k: v for k, v in _extract_files(final_answer).items()
               if k.rsplit(".", 1)[-1].lower() in CODE_EXTS}
        _store_files(session_id, got)

    # record the exchange before the answer is handed to the client, so an
    # immediate follow-up request always sees this turn
    _append_history(session_id, question, final_answer, route)

    with JOBS_LOCK:
        JOBS[job_id]["answer"] = final_answer
        JOBS[job_id]["done"] = True

    # auto-save to cookbook if this was a recipe response — one card per dish
    if route == "recipes":
        ts = time.time()
        entries = [
            {"id": uuid.uuid4().hex, "question": question,
             "title": rec["title"], "answer": rec["body"], "timestamp": ts}
            for rec in split_recipes(final_answer, question)
        ]
        with COOKBOOK_LOCK:
            cookbook().extend(entries)
        _save_cookbook()
        # photo + meal/protein classification for each new dish, in the
        # background — the answer is already delivered, never delays anything
        spawn_owned(target=_backfill_photos, args=(entries, 0.5),
                         daemon=True).start()
        spawn_owned(target=_backfill_meta, args=(entries, 0.3),
                         daemon=True).start()

    # fire-and-forget: update session memory in background
    spawn_owned(
        target=summarize, args=(session_id, question, final_answer), daemon=True
    ).start()


# ---- panel ------------------------------------------------------------------
# Default panel line-up. A user may now pick any subset of their CONFIGURED
# routes instead; this is only the fallback when none is chosen.
PANEL_ROUTES = ["recipes", "creative", "code", "general", "reasoning"]
PANEL_MAX_ROUTES = 8


def panel_routes():
    """Routes a panel may use, with the backend each one would reach.

    Routes, not backends — a panel compares configured setups (model + persona),
    which is the thing a user actually experiences. /compare picks raw backends
    when you want the models themselves. Names only: no URLs, no keys, because
    /panel is open to any logged-in user while /settings is admin-only."""
    with CONFIG_LOCK:
        routes = ROUTER_CONFIG.get("routes", {})
        by_id = {b["id"]: b for b in ROUTER_CONFIG.get("backends", [])}
        out = []
        for name in sorted(routes):
            b = by_id.get(routes[name].get("backend")) or {}
            out.append({"route": name, "backend": b.get("id", ""),
                        "model": b.get("model", ""),
                        "paid": bool(b.get("paid")),
                        "default": name in PANEL_ROUTES})
    return out

JUDGE_SYSTEM = (
    "You are a neutral judge evaluating multiple AI responses to the same question. "
    "You will see the question and responses labelled by model type. "
    "In 3-5 sentences: identify which response best answers the question and why, "
    "then note one notable strength or weakness of each response. Be direct and specific."
)


# Retained for the panel endpoints that outlived the panel page (/panel now
# redirects to /compare). Left working rather than half-removed; delete the
# whole panel surface in one deliberate pass if it is ever worth the churn.
PANEL_JOBS: dict = {}
PANEL_JOBS_LOCK = threading.Lock()


def run_panel_judge(job_id: str, question: str):
    with PANEL_JOBS_LOCK:
        specialists = {k: dict(v) for k, v in PANEL_JOBS[job_id]["specialists"].items()}
    if len(specialists) < 2:
        with PANEL_JOBS_LOCK:
            PANEL_JOBS[job_id]["judge"]["answer"] = (
                "_Only one route was selected, so there is nothing to compare._")
            PANEL_JOBS[job_id]["judge"]["done"] = True
            PANEL_JOBS[job_id]["done"] = True
        return

    responses_text = "\n\n".join(
        f"=== {route.upper()} ===\n{data['answer'][:1500]}"
        for route, data in specialists.items()
    )
    payload = {
        "messages": [
            {"role": "system", "content": JUDGE_SYSTEM},
            {"role": "user",   "content": f"Question: {question}\n\n{responses_text}"},
        ],
        "stream": True,
        "max_tokens": 1200,
        # See the compare judge: thinking on burns the whole budget and returns
        # nothing visible.
        "chat_template_kwargs": {"enable_thinking": False},
    }
    parts = []
    try:
        _rurl = SPECIALISTS["reasoning"]
        with requests.post(_rurl, json=_inject_model(payload, _rurl, "reasoning"),
                           headers=_hdrs(_rurl), stream=True, timeout=300) as r:
            r.raise_for_status()
            for line in _lines(r):
                if not line:
                    continue
                line = line.decode("utf-8")
                if line.startswith("data: "):
                    data = line[6:]
                    if data.strip() == "[DONE]":
                        break
                    try:
                        chunk = json.loads(data)
                        delta = chunk["choices"][0]["delta"].get("content", "")
                        if delta:
                            parts.append(delta)
                            with PANEL_JOBS_LOCK:
                                PANEL_JOBS[job_id]["judge"]["tokens"] += 1
                    except Exception:
                        pass
    except Exception as e:
        parts = [f"[judge error: {e}]"]

    with PANEL_JOBS_LOCK:
        PANEL_JOBS[job_id]["judge"]["answer"] = "".join(parts)
        PANEL_JOBS[job_id]["judge"]["done"] = True
        PANEL_JOBS[job_id]["done"] = True


# ---- backend health (route status dots in the UI) ---------------------------
_health_cache = {"t": 0.0, "data": {}}
_health_lock = threading.Lock()

def _backend_alive(url):
    base = url.rsplit("/v1/", 1)[0]
    if _kind(url) == "anthropic":
        # Anthropic has no /health; /v1/models needs the key, and a 401 still
        # proves the endpoint is up — only a connection failure means dead.
        try:
            return requests.get(base + "/v1/models", headers=_hdrs(url),
                                timeout=4).status_code in (200, 401, 403)
        except Exception:
            return False
    for probe in ("/health", "/v1/models"):   # llama.cpp /health; /v1/models fallback
        try:
            if requests.get(base + probe, headers=_hdrs(url),
                            timeout=1.2).status_code == 200:
                return True
        except Exception:
            pass
    return False

# ---- settings API ------------------------------------------------------------
def _mask_config(cfg):
    out = json.loads(json.dumps(cfg))
    for b in out["backends"]:
        b["has_key"] = bool(b.get("key"))
        b["key"] = "__KEEP__" if b.get("key") else ""
    return out


@app.route("/settings")
def settings_page():
    return Response(SETTINGS_PAGE, mimetype="text/html")


@app.route("/settings/config")
def settings_get():
    with CONFIG_LOCK:
        return jsonify(_mask_config(ROUTER_CONFIG))


@app.route("/settings/config", methods=["POST"])
def settings_save():
    cfg = request.json or {}
    errs = []
    backends = cfg.get("backends") or []
    routes = cfg.get("routes") or {}
    roles = cfg.get("roles") or {}
    ids = [str(b.get("id", "")).strip() for b in backends]
    if len(ids) != len(set(ids)) or not all(ids):
        errs.append("backend names must be unique and non-empty")
    for b in backends:
        u = str(b.get("url", "")).strip().rstrip("/")
        if not re.match(r"^https?://\S+$", u):
            errs.append(f"backend '{b.get('id')}': bad URL")
        b["url"] = u
        # __KEEP__ = the UI never saw the real key; restore the stored one
        if b.get("key") == "__KEEP__":
            with CONFIG_LOCK:
                old = next((x for x in ROUTER_CONFIG["backends"]
                            if x["id"] == b["id"]), None)
            b["key"] = old.get("key", "") if old else ""
        # Cost settings. Prices are advisory only — they never gate a request,
        # they only feed the spend estimate, so bad input is coerced not rejected.
        b["kind"] = b.get("kind") if b.get("kind") in BACKEND_KINDS else "openai"
        if b["kind"] == "anthropic" and not b.get("model"):
            errs.append(f"backend '{b.get('id')}': Anthropic needs a model name "
                        "(e.g. claude-opus-5) — it has no default")
        b["paid"] = bool(b.get("paid"))
        for k in ("price_in", "price_out"):
            try:
                b[k] = max(0.0, float(b.get(k) or 0))
            except (TypeError, ValueError):
                b[k] = 0.0
        if b["paid"] and not (b["price_in"] or b["price_out"]):
            errs.append(f"backend '{b.get('id')}': marked paid but both prices "
                        "are 0 — spend would always estimate as free")
        b.pop("has_key", None)
    for n in _BUILTIN_ROUTES:
        if n not in routes:
            errs.append(f"built-in route '{n}' is missing")
    for n, r in routes.items():
        if not _ROUTE_NAME_RE.fullmatch(n) or n == "audio":
            errs.append(f"bad route name '{n}' (one lowercase word, 2-16 chars)")
        if r.get("backend") not in set(ids):
            errs.append(f"route '{n}': unknown backend")
        r["custom"] = n not in _BUILTIN_ROUTES
        if r["custom"]:
            if not str(r.get("subject", "")).strip():
                errs.append(f"route '{n}': subject is required (the classifier "
                            "routes by it)")
            if not r.get("color"):
                used = {x.get("color") for x in routes.values()}
                r["color"] = next((c for c in _CUSTOM_COLOR_POOL
                                   if c not in used), "#666666")
    for role in ("classifier", "summarizer", "coder"):
        if roles.get(role) not in set(ids):
            errs.append(f"role '{role}': unknown backend")
    # Optional: absent means "use the deep coder for both effort levels".
    if roles.get("coder_fast") and roles["coder_fast"] not in set(ids):
        errs.append("role 'coder_fast': unknown backend")
    if roles.get("coder_quick") and roles["coder_quick"] not in set(ids):
        errs.append("role 'coder_quick': unknown backend")
    if errs:
        return jsonify({"ok": False, "errors": errs}), 400
    clean = {"backends": backends, "routes": routes, "roles": roles}
    _save_router_config(clean)
    _apply_router_config(clean)
    return jsonify({"ok": True})


@app.route("/settings/backend_check", methods=["POST"])
def settings_backend_check():
    body = request.json or {}
    url = str(body.get("url", "")).strip().rstrip("/")
    key = body.get("key") or ""
    if key == "__KEEP__":
        with CONFIG_LOCK:
            b = next((x for x in ROUTER_CONFIG["backends"]
                      if x["url"].rstrip("/") == url), None)
        key = b.get("key", "") if b else ""
    kind = body.get("kind") if body.get("kind") in BACKEND_KINDS else "openai"
    if kind == "anthropic":
        hdr = {"anthropic-version": ANTHROPIC_VERSION}
        if key:
            hdr["x-api-key"] = key
    else:
        hdr = {"Authorization": f"Bearer {key}"} if key else {}
    try:
        r = requests.get(url + "/v1/models", headers=hdr, timeout=8)
        if r.status_code in (401, 403) and kind == "anthropic":
            return jsonify({"alive": False,
                            "error": "reachable, but the API key was rejected"})
        r.raise_for_status()
        models = [m.get("id", "?") for m in r.json().get("data", [])][:20]
        return jsonify({"alive": True, "models": models})
    except Exception as e:
        return jsonify({"alive": False, "error": str(e)[:120]})


@app.route("/settings/scan", methods=["POST"])
def settings_scan():
    """Probe well-known localhost ports for OpenAI-compatible servers."""
    found = []
    for port in (11434, 1234, 8000, 8001, 8080, 8081, 8082, 8083, 8084,
                 8085, 8086, 8087, 8088, 8090):
        url = f"http://127.0.0.1:{port}"
        try:
            r = requests.get(url + "/v1/models", timeout=1.0)
            if r.status_code == 200:
                models = [m.get("id", "?") for m in r.json().get("data", [])][:5]
                found.append({"url": url, "models": models})
        except Exception:
            pass
    return jsonify({"found": found})


@app.route("/settings/draft_route", methods=["POST"])
def settings_draft_route():
    body = request.json or {}
    name = str(body.get("name", "route")).strip()
    purpose = str(body.get("purpose", "")).strip()
    if not purpose:
        return jsonify({"error": "describe the route's purpose first"}), 400
    sys_p = (
        "You configure routes for an LLM router. Given a route name and its "
        "purpose, respond with ONLY a JSON object — no fences, no commentary:\n"
        '{"subject": "<1-2 sentences listing what belongs to this route, '
        "concrete and discriminative, written like: 'D&D campaigns, character "
        "builds, encounter design, rules questions'>\", "
        '"persona": "<a 2-3 sentence system prompt for the specialist that '
        'answers these questions>"}')
    payload = {"messages": [
        {"role": "system", "content": sys_p},
        {"role": "user", "content": f"Route name: {name}\nPurpose: {purpose}"}],
        "stream": False, "temperature": 0.3, "max_tokens": 512,
        "chat_template_kwargs": {"enable_thinking": False}}
    try:
        u = SPECIALISTS.get("general", CLASSIFIER)
        r = requests.post(u, json=_inject_model(payload, u, "general"),
                          headers=_hdrs(u), timeout=90)
        content = _resp_json(r)["choices"][0]["message"]["content"] or ""
        content = _strip_think(content)
        m = re.search(r"\{[\s\S]*\}", content)
        d = json.loads(m.group(0)) if m else {}
        return jsonify({"subject": str(d.get("subject", ""))[:400],
                        "persona": str(d.get("persona", ""))[:800]})
    except Exception as e:
        return jsonify({"error": f"draft failed: {e}"}), 502


@app.route("/settings/test_route", methods=["POST"])
def settings_test_route():
    """Run example questions through the REAL classifier with the candidate
    route included — ground truth for 'will traffic route correctly'."""
    body = request.json or {}
    name = str(body.get("name", "")).strip()
    subject = str(body.get("subject", "")).strip()
    questions = [q for q in (body.get("questions") or []) if str(q).strip()][:5]
    if not (_ROUTE_NAME_RE.fullmatch(name) and subject and questions):
        return jsonify({"error": "need a valid name, a subject, and at least "
                                 "one test question"}), 400
    with CONFIG_LOCK:
        cand = json.loads(json.dumps(ROUTER_CONFIG))
    cand["routes"].setdefault(name, {"custom": True})
    cand["routes"][name]["subject"] = subject
    cand["routes"][name]["custom"] = name not in _BUILTIN_ROUTES
    system = _build_classifier_system(cand)
    valid = set(cand["routes"]) | {"audio"}
    results = [{"q": q, "route": _classify_with(system, valid, str(q))}
               for q in questions]
    return jsonify({"results": results})


@app.route("/settings/routes_meta")
def settings_routes_meta():
    with CONFIG_LOCK:
        return jsonify([{"name": n, "color": r.get("color", ""),
                         "custom": bool(r.get("custom"))}
                        for n, r in ROUTER_CONFIG["routes"].items()])


@app.route("/health/routes")
def health_routes():
    """Liveness of each specialist backend, cached 10s (dedup'd by URL)."""
    with _health_lock:
        if time.time() - _health_cache["t"] > 10:
            by_url = {}
            for _route, url in SPECIALISTS.items():
                if url not in by_url:
                    by_url[url] = _backend_alive(url)
            _health_cache["data"] = {r: by_url[u] for r, u in SPECIALISTS.items()}
            _health_cache["t"] = time.time()
    return jsonify(_health_cache["data"])

@app.route("/")
def index():
    # AGPL s13: the source link must point at the version actually running, so
    # it is substituted at serve time from CREWE_SOURCE_URL rather than baked
    # into the template.
    return Response(PAGE.replace("__SOURCE_URL__", SOURCE_URL),
                    mimetype="text/html")


@app.route("/scratch")
def scratch():
    return Response(SCRATCH_PAGE, mimetype="text/html")


@app.route("/upload", methods=["POST"])
def upload():
    """Attach a document to a session.

    The extracted TEXT is what the model ever sees; the original is kept only
    so the user can download back what they sent. Both live under the caller's
    own crewe_userdata root, named by a generated id — the client's filename is
    never used as a path."""
    if request.is_json:
        body = request.json or {}
        session_id = (body.get("session_id") or "anonymous").strip()
        gh = (body.get("github_url") or "").strip()
        if not gh:
            return jsonify({"error": "no github_url"}), 400
        try:
            name, text = _github_repo_text(gh)
        except ValueError as e:
            return jsonify({"error": str(e)}), 400
        up_id = uuid.uuid4().hex[:16]
        try:
            with open(os.path.join(uploads_dir(), f"{up_id}.txt"), "w",
                      encoding="utf-8") as fh:
                fh.write(text)
        except OSError as e:
            print(f"[upload] repo write failed: {e}")
            return jsonify({"error": "could not store that repository"}), 500
        meta = {"id": up_id, "name": name, "ext": "repo",
                "chars": len(text), "ts": time.time()}
        with SESSIONS_LOCK:
            sess = sessions().setdefault(session_id, {})
            ups = sess.setdefault("uploads", [])
            ups.append(meta)
            if len(ups) > MAX_SESSION_UPLOADS:
                del ups[:-MAX_SESSION_UPLOADS]
        _save_memory()
        return jsonify({"ok": True, "upload": meta})

    session_id = (request.form.get("session_id") or "anonymous").strip()
    f = request.files.get("file")
    if not f or not f.filename:
        return jsonify({"error": "no file"}), 400
    name = os.path.basename(f.filename)[:180]
    raw = f.read()
    if not raw:
        return jsonify({"error": "that file is empty"}), 400
    try:
        text = _upload_text(name, raw)
    except ValueError as e:
        return jsonify({"error": str(e)}), 400

    up_id = uuid.uuid4().hex[:16]
    ext = name.rsplit(".", 1)[-1].lower() if "." in name else "bin"
    d = uploads_dir()
    try:
        with open(os.path.join(d, f"{up_id}.{ext}"), "wb") as fh:
            fh.write(raw)
        with open(os.path.join(d, f"{up_id}.txt"), "w", encoding="utf-8") as fh:
            fh.write(text)
    except OSError as e:
        print(f"[upload] write failed: {e}")
        return jsonify({"error": "could not store that file"}), 500

    meta = {"id": up_id, "name": name, "ext": ext,
            "chars": len(text), "ts": time.time()}
    with SESSIONS_LOCK:
        sess = sessions().setdefault(session_id, {})
        ups = sess.setdefault("uploads", [])
        ups.append(meta)
        if len(ups) > MAX_SESSION_UPLOADS:      # oldest attachments fall off
            del ups[:-MAX_SESSION_UPLOADS]
    _save_memory()
    return jsonify({"ok": True, "upload": meta})


@app.route("/spend")
def spend():
    """Estimated spend for the logged-in user, plus which routes cost money.

    Estimates only: they use the prices YOU entered and the token counts the
    provider reported (or a 4-chars-per-token guess when it reported none).
    Treat them as a running indication, never as a bill."""
    with CONFIG_LOCK:
        paid_routes = sorted(r for r, u in SPECIALISTS.items() if _is_paid(u))
        any_paid = any(_is_paid(u) for u in set(SPECIALISTS.values())) \
            or _is_paid(CODE_AGENT_URL) or _is_paid(CODE_AGENT_URL_FAST) \
            or _is_paid(CODE_AGENT_URL_QUICK)
    s = spend_summary()
    s["paid_routes"] = paid_routes
    s["any_paid"] = bool(any_paid)
    s["coder_paid"] = {"fast": _is_paid(CODE_AGENT_URL_QUICK),
                       "normal": _is_paid(CODE_AGENT_URL_FAST),
                       "extra": _is_paid(CODE_AGENT_URL)}
    return jsonify(s)


@app.route("/uploads/<session_id>")
def uploads_list(session_id):
    return jsonify({"uploads": session_uploads(session_id)})


@app.route("/uploads/<session_id>/<up_id>", methods=["DELETE"])
def upload_remove(session_id, up_id):
    """Detach a file. The stored copies are removed too — a user who clears an
    attachment expects it gone, not merely hidden."""
    with SESSIONS_LOCK:
        sess = sessions().get(session_id, {})
        ups = sess.get("uploads") or []
        keep = [u for u in ups if u["id"] != up_id]
        removed = len(keep) != len(ups)
        sess["uploads"] = keep
    if removed:
        d = uploads_dir()
        for fn in os.listdir(d):
            if fn.split(".")[0] == up_id:
                try:
                    os.remove(os.path.join(d, fn))
                except OSError:
                    pass
        _save_memory()
    return jsonify({"ok": True, "removed": removed})


@app.route("/ask", methods=["POST"])
def ask():
    body       = request.json or {}
    question   = body.get("question", "").strip()
    session_id = body.get("session_id", "anonymous")
    if not question:
        return jsonify({"error": "empty question"}), 400

    # Seed file memory from the client's recent code answers — covers sessions
    # that predate server-side file tracking. Oldest first, so newer complete
    # versions win and elided fragments never replace fuller ones.
    seeds = body.get("seed_code") or []
    if isinstance(seeds, str):
        seeds = [seeds]
    with SESSIONS_LOCK:
        has_files = bool(sessions().get(session_id, {}).get("files"))
    if not has_files and seeds:
        merged = {}
        for chunk in seeds:
            for k, v in _extract_files(chunk).items():
                if (k.rsplit(".", 1)[-1].lower() in CODE_EXTS
                        and not _looks_elided(v, merged.get(k))):
                    merged[k] = v
        _store_files(session_id, merged)
        has_files = bool(merged)

    history, last_route = _recent_context(session_id)
    prev_q = history[-1]["q"] if history else ""
    prev_a = history[-1]["a"] if history else ""

    # Pasted errors about an in-progress build must reach the code route even
    # when the classifier would call them 'general'.
    if has_files and _ERRORISH.search(question):
        route = "code"
    # Short continuation replies ("no", "keep going") stay with the previous
    # specialist — its history has the conversation state (e.g. a game in
    # progress), and the classifier can't route a bare "no" reliably.
    elif last_route in SPECIALISTS and _CONTINUATION.match(question):
        route = last_route
    else:
        route = classify(question, prev_q, prev_a, last_route)

    # The code route is the only one that EXECUTES model-written code on this box
    # (headless browser + node). Limit it to admins. Non-admins get a clean
    # message delivered through the normal job/poll UI — returning a raw error
    # here would leave the client spinner polling a job that never exists.
    if route == "code" and not (current_user() or {}).get("is_admin"):
        job_id = uuid.uuid4().hex
        with JOBS_LOCK:
            JOBS[job_id] = {"owner": _owner(), "tokens": 0, "done": True, "route": "code",
                            "answer": "🔒 The code builder is limited to admins right "
                                      "now. Ask a regular question, or have an admin run "
                                      "the build.",
                            "stage": "done", "audio": None, "stage_ts": time.time()}
            _prune_jobs(JOBS)
        return jsonify({"job_id": job_id, "route": "code"})

    # The UI always sends effort, but it is only ever consulted for the code
    # route — every other specialist has exactly one backend to go to.
    effort = str(body.get("effort", "")).lower()
    effort = EFFORT_ALIASES.get(effort, effort)
    effort = effort if effort in EFFORTS else DEFAULT_EFFORT

    job_id = uuid.uuid4().hex
    with JOBS_LOCK:
        JOBS[job_id] = {"owner": _owner(), "tokens": 0, "done": False, "route": route, "answer": "",
                        "stage": "drafting", "audio": None, "stage_ts": time.time(),
                        "effort": effort}
        _prune_jobs(JOBS)

    spawn_owned(target=run_job, args=(job_id, route, question, session_id, effort),
                daemon=True).start()
    return jsonify({"job_id": job_id, "route": route, "effort": effort})


@app.route("/progress/<job_id>")
def progress(job_id):
    with JOBS_LOCK:
        job = owned_job(JOBS, job_id)
        if not job:
            return jsonify({"error": "unknown job"}), 404
        resp = {"tokens": job["tokens"], "done": job["done"], "route": job["route"], "stage": job.get("stage", "drafting"),
                "stage_age": round(time.time() - job.get("stage_ts", time.time()))}
        if "partial" in job:   # live text for jobs that stream it (doc writer)
            resp["partial"] = job["partial"]
        if job["done"]:
            resp["answer"] = job["answer"]
            resp["audio"] = job.get("audio")
    return jsonify(resp)


def _job_cancelled(job_id):
    with JOBS_LOCK:
        j = JOBS.get(job_id)
        return bool(j and j.get("cancel"))


@app.route("/stop/<job_id>", methods=["POST"])
def stop_job(job_id):
    """Flag a running job for cancellation; streams and the step loop check it."""
    with JOBS_LOCK:
        job = owned_job(JOBS, job_id)
        if not job:
            return jsonify({"error": "unknown job"}), 404
        job["cancel"] = True
    return jsonify({"ok": True})


_MIME_BY_EXT = {"html": "text/html", "htm": "text/html", "css": "text/css",
                "js": "application/javascript", "mjs": "application/javascript",
                "json": "application/json", "svg": "image/svg+xml",
                "png": "image/png", "wav": "audio/wav"}


@app.route("/play/<session_id>/")
@app.route("/play/<session_id>/<path:filename>")
def play_build(session_id, filename="index.html"):
    """Serve a session's built project over HTTP, the way a real host would.
    file:// blocks ES modules and local fetch(), so 'download and double-
    click' can black-screen a build that actually works — this route is the
    no-setup answer (no more hand-rolled python -m http.server)."""
    with SESSIONS_LOCK:
        files = dict(sessions().get(session_id, {}).get("files") or {})
    if filename not in files and filename == "index.html":
        # entry fallback: first html file the build produced
        filename = next((n for n in files if n.endswith((".html", ".htm"))),
                        filename)
    if filename not in files:
        return jsonify({"error": "not found"}), 404
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext in ("html", "htm"):
        # sandbox: generated code runs with an opaque origin so it can never
        # read or wipe the chat UI's localStorage (routerSessions etc.).
        # The shim gives games that use localStorage an in-memory stand-in
        # instead of a SecurityError.
        resp = Response(_inject_html(files[filename], _PLAY_SHIM),
                        mimetype="text/html")
        resp.headers["Content-Security-Policy"] = "sandbox allow-scripts allow-pointer-lock"
        return resp
    return Response(files[filename],
                    mimetype=_MIME_BY_EXT.get(ext, "text/plain"))


_PLAY_SHIM = (
    "<script>try{window.localStorage}catch(e){(function(){var m={};"
    "var s={getItem:function(k){return m.hasOwnProperty(k)?m[k]:null},"
    "setItem:function(k,v){m[k]=String(v)},removeItem:function(k){delete m[k]},"
    "clear:function(){m={}},key:function(i){return Object.keys(m)[i]||null}};"
    "Object.defineProperty(s,'length',{get:function(){return Object.keys(m).length}});"
    "Object.defineProperty(window,'localStorage',{value:s});})()}</script>")

SHOTS_SHELL = """<!doctype html><html><head><meta charset="utf-8">
<title>build checks</title>
<link rel="icon" href="data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 64 64'%3E%3Crect width='64' height='64' rx='14' fill='%238b4a18'/%3E%3Cpath d='M32 57V34M32 34 16 13M32 34 48 13' stroke='%23faf6f0' stroke-width='7' stroke-linecap='round' fill='none'/%3E%3Ccircle cx='32' cy='34' r='6' fill='%23faf6f0'/%3E%3C/svg%3E"><meta name="viewport" content="width=device-width, initial-scale=1">
<link href="/static/fonts.css" rel="stylesheet">
<script>if(localStorage.getItem('creweTheme')==='dark')document.documentElement.dataset.theme='dark';</script>
<style>
:root{--bg:#faf6f0;--panel:#ffffff;--chip:#f2ebe0;--ink:#1e160a;--muted:#8a7560;
      --border:#ddd0b8;--link:#2a4e8c;--warm:#8b4a18}
:root[data-theme="dark"]{--bg:#191410;--panel:#221c15;--chip:#2b241b;--ink:#ece2d3;
      --muted:#9c8a75;--border:#3a3125;--link:#8aaee8;--warm:#d29a62;color-scheme:dark}
body{background:var(--bg);color:var(--ink);font:15px/1.5 'DM Sans',system-ui,sans-serif;margin:0;padding:24px}
h1{font-size:20px;margin:0 0 18px;font-family:'Lora',Georgia,serif;color:var(--warm)}
.run{background:var(--panel);border:1px solid var(--border);border-radius:10px;padding:14px 16px;margin-bottom:18px;max-width:760px}
.hd{display:flex;justify-content:space-between;align-items:baseline;margin-bottom:10px}
.hd b{color:var(--link);font-family:monospace}
.hd span{color:var(--muted);font-size:13px}
img{max-width:100%;border-radius:6px;border:1px solid var(--border);display:block}
.files{margin-top:10px;display:flex;flex-wrap:wrap;gap:6px}
.f{background:var(--chip);border:1px solid var(--border);color:var(--ink);text-decoration:none;
   padding:2px 9px;border-radius:20px;font:12px monospace}
.f:hover{background:var(--border)}
.none{color:var(--muted)}
a.top{color:var(--link);text-decoration:none;font-size:13px}
.theme-toggle{position:fixed;top:22px;right:24px;display:inline-flex;align-items:center;gap:5px;cursor:pointer;user-select:none}
.theme-toggle input{position:absolute;opacity:0;pointer-events:none}
.tt-track{width:34px;height:18px;border-radius:999px;background:var(--panel);border:1px solid var(--border);position:relative}
.tt-knob{position:absolute;top:1px;left:1px;width:14px;height:14px;border-radius:50%;background:var(--warm);transition:left .2s}
.theme-toggle input:checked + .tt-track .tt-knob{left:17px}
.tt-icon{font-size:11px;color:var(--muted)}
</style></head><body>
<h1>🌐 build checks <a class="top" href="/">← chat</a></h1>
<label class="theme-toggle" title="Dark mode"><span class="tt-icon">☀</span><input type="checkbox" id="themeSw"><span class="tt-track"><span class="tt-knob"></span></span><span class="tt-icon">🌙</span></label>
<script>(function(){const t=document.getElementById('themeSw');t.checked=localStorage.getItem('creweTheme')==='dark';t.addEventListener('change',()=>{if(t.checked){document.documentElement.dataset.theme='dark';localStorage.setItem('creweTheme','dark');}else{delete document.documentElement.dataset.theme;localStorage.setItem('creweTheme','light');}});})();</script>
__BODY__
</body></html>"""


@app.route("/shots")
def shots_page():
    """Build-check history: every browser-gate run dir, newest first, with
    its screenshot inline and a browsable listing of the exact files that
    were loaded — what the gate saw, viewable from any machine."""
    rows = []
    try:
        dirs = sorted(
            (d for d in os.listdir(checks_dir())
             if os.path.isdir(os.path.join(checks_dir(), d))),
            key=lambda d: os.path.getmtime(os.path.join(checks_dir(), d)),
            reverse=True)[:40]
    except OSError:
        dirs = []
    for d in dirs:
        base = os.path.join(checks_dir(), d)
        ts = time.strftime("%Y-%m-%d %H:%M",
                           time.localtime(os.path.getmtime(base)))
        entries = []
        for root, _dd, fns in os.walk(base):
            for fn in fns:
                entries.append(os.path.relpath(os.path.join(root, fn), base))
        entries.sort()
        shot = (f'<a href="/shots/{d}/screenshot.png" target="_blank">'
                f'<img src="/shots/{d}/screenshot.png" loading="lazy"></a>'
                if "screenshot.png" in entries else
                '<p class="none">no screenshot</p>')
        links = "".join(
            f'<a class="f" href="/shots/{d}/{e}" target="_blank">{e}</a>'
            for e in entries if e != "screenshot.png")
        rows.append(f'<div class="run"><div class="hd"><b>{d}</b>'
                    f'<span>{ts}</span></div>{shot}'
                    f'<div class="files">{links}</div></div>')
    body = "\n".join(rows) or '<p class="none">no build checks yet</p>'
    return Response(SHOTS_SHELL.replace("__BODY__", body), mimetype="text/html")


@app.route("/shots/<run>/<path:filename>")
def shots_file(run, filename):
    """Serve one file from a build-check run dir. Screenshots render as
    images; code is served as plain text (a viewer, not a runner — /play is
    the runner)."""
    if not re.fullmatch(r"[A-Za-z0-9_-]+", run):
        return jsonify({"error": "not found"}), 404
    base = os.path.abspath(os.path.join(checks_dir(), run))
    path = os.path.abspath(os.path.join(base, filename))
    if not path.startswith(base + os.sep) or not os.path.isfile(path):
        return jsonify({"error": "not found"}), 404
    if filename.lower().endswith(".png"):
        with open(path, "rb") as f:
            return Response(f.read(), mimetype="image/png")
    with open(path, "r", errors="replace") as f:
        return Response(f.read(), mimetype="text/plain; charset=utf-8")


@app.route("/audio/<path:filename>")
def serve_audio(filename):
    """Serve a recorded wav from audio_dir(). Basename-only, .wav-only — no traversal."""
    safe = os.path.basename(filename)
    if not safe.endswith(".wav"):
        return jsonify({"error": "not found"}), 404
    path = os.path.join(audio_dir(), safe)
    if not os.path.isfile(path):
        return jsonify({"error": "not found"}), 404
    with open(path, "rb") as f:
        return Response(f.read(), mimetype="audio/wav")


@app.route("/memory/<session_id>")
def get_memory(session_id):
    with SESSIONS_LOCK:
        sess = sessions().get(session_id, {})
        return jsonify({
            "summary": sess.get("summary", ""),
            "turns": sess.get("turns", 0),
            # names only — file contents can be large
            "files": sorted((sess.get("files") or {}).keys()),
        })


@app.route("/memory/<session_id>", methods=["DELETE"])
def clear_memory(session_id):
    with SESSIONS_LOCK:
        sessions().pop(session_id, None)
    _save_memory()
    return jsonify({"ok": True})


@app.route("/cookbook")
def cookbook_page():
    return Response(COOKBOOK_PAGE, mimetype="text/html")


@app.route("/cookbook/recipes")
def cookbook_recipes():
    with COOKBOOK_LOCK:
        return jsonify(list(reversed(cookbook())))


def _image_search(query, n=8):
    """SearXNG image search -> candidate image URLs (best-effort, [] on any
    failure — the book must still build without photos)."""
    try:
        r = requests.get(SEARXNG_URL,
                         params={"q": query, "categories": "images",
                                 "format": "json"}, timeout=10)
        results = r.json().get("results", [])
    except Exception:
        return []
    out = []
    for res in results:
        u = res.get("img_src") or ""
        if u.startswith("//"):
            u = "https:" + u
        if u.startswith("http"):
            out.append(u)
        if len(out) >= n:
            break
    return out


_IMG_EXT_BY_CT = {"image/jpeg": "jpg", "image/png": "png", "image/webp": "webp",
                  "image/gif": "gif", "image/avif": "avif"}


def _fetch_image_bytes(url):
    """Download one image; returns (bytes, content_type) or None."""
    try:
        r = requests.get(url, timeout=8, headers={
            "User-Agent": ("Mozilla/5.0 (X11; Linux x86_64; rv:127.0) "
                           "Gecko/20100101 Firefox/127.0"),
            "Accept": "image/*,*/*;q=0.8"})
        ct = r.headers.get("content-type", "").split(";")[0].strip().lower()
        if (r.status_code != 200 or ct not in _IMG_EXT_BY_CT):
            return None
        data = r.content
        if not (8_000 <= len(data) <= 6_000_000):
            return None
        return data, ct
    except Exception:
        return None


def _fetch_image_datauri(url):
    """Download one image and return it as a data URI (None on any problem)."""
    got = _fetch_image_bytes(url)
    if not got:
        return None
    data, ct = got
    return f"data:{ct};base64,{base64.b64encode(data).decode()}"


# ---- per-card recipe photos --------------------------------------------------
# Every cookbook card gets a photo file in photos_dir(), fetched in background
# threads (never blocks answering). card['photo'] = filename ('' = tried, none
# found; key absent = not tried yet). The designer reuses these for books.
_BACKFILL_STARTED = threading.Event()


def _card_photo(card):
    """Search + download a photo for one card; returns filename or ''."""
    title = ((card.get("title") or "").strip()
             or (card.get("question") or "").strip() or "dish")
    for u in _image_search(f"{title} dish food photography")[:5]:
        got = _fetch_image_bytes(u)
        if got:
            data, ct = got
            os.makedirs(photos_dir(), exist_ok=True)
            fname = f"{card['id']}.{_IMG_EXT_BY_CT[ct]}"
            with open(os.path.join(photos_dir(), fname), "wb") as f:
                f.write(data)
            return fname
    return ""


# meal-type + protein metadata, model-classified in the background (same
# lifecycle as photos: new cards on save, old cards via one-time startup sweep)
MEAL_TYPES = ["breakfast", "lunch", "dinner", "dessert", "snack", "drink", "side"]
PROTEINS = ["chicken", "beef", "pork", "seafood", "turkey", "lamb",
            "vegetarian", "other"]
_META_BACKFILL_STARTED = threading.Event()


def _classify_card(card):
    """One model call → {'meal': ..., 'protein': ...} from the fixed vocab
    (falls back to 'other'/'dinner' on any failure — never blocks)."""
    title = (card.get("title") or "").strip() or card.get("question", "")
    body = (card.get("answer") or "")[:800]
    sys_p = ("Classify a recipe. Respond with ONLY a JSON object, no "
             "commentary:\n"
             f'{{"meal": one of {MEAL_TYPES}, "protein": one of {PROTEINS}}}\n'
             "protein is the MAIN protein; use 'vegetarian' when there is "
             "none, 'other' when unsure.")
    payload = {"messages": [{"role": "system", "content": sys_p},
                            {"role": "user", "content": f"{title}\n\n{body}"}],
               "stream": False, "temperature": 0.0, "max_tokens": 128,
               "chat_template_kwargs": {"enable_thinking": False}}
    meal, protein = "dinner", "other"
    try:
        u = SPECIALISTS.get("general", CLASSIFIER)
        r = requests.post(u, json=_inject_model(payload, u, "general"),
                          headers=_hdrs(u), timeout=60)
        content = _resp_json(r)["choices"][0]["message"]["content"] or ""
        content = _strip_think(content)
        m = re.search(r"\{[\s\S]*?\}", content)
        d = json.loads(m.group(0)) if m else {}
        if str(d.get("meal", "")).lower() in MEAL_TYPES:
            meal = str(d["meal"]).lower()
        if str(d.get("protein", "")).lower() in PROTEINS:
            protein = str(d["protein"]).lower()
    except Exception:
        pass
    return meal, protein


def _backfill_meta(cards, pace=0.5):
    """Classify meal/protein for cards missing the 'meal' key."""
    todo = [c["id"] for c in cards if c.get("id") and "meal" not in c]
    done = 0
    for cid in todo:
        with COOKBOOK_LOCK:
            card = next((c for c in cookbook() if c.get("id") == cid), None)
            if card is None or "meal" in card:
                continue
            snap = dict(card)
        meal, protein = _classify_card(snap)
        with COOKBOOK_LOCK:
            card = next((c for c in cookbook() if c.get("id") == cid), None)
            if card is not None:
                card["meal"] = meal
                card["protein"] = protein
        done += 1
        if done % 8 == 0:
            _save_cookbook()
        time.sleep(pace)
    if done:
        _save_cookbook()
        print(f"[cookbook] meal/protein backfill: {done} card(s) classified")


def _backfill_photos(cards, pace=1.5):
    """Fetch photos for the given cards (those without a 'photo' key),
    updating the cookbook as results land. Runs in a daemon thread."""
    todo = [c["id"] for c in cards if c.get("id") and "photo" not in c]
    done = 0
    for cid in todo:
        with COOKBOOK_LOCK:
            card = next((c for c in cookbook() if c.get("id") == cid), None)
            if card is None or "photo" in card:
                continue
            snap = dict(card)
        fname = _card_photo(snap)
        with COOKBOOK_LOCK:
            card = next((c for c in cookbook() if c.get("id") == cid), None)
            if card is not None:
                card["photo"] = fname
        done += 1
        if done % 5 == 0:
            _save_cookbook()
        time.sleep(pace)
    if done:
        _save_cookbook()
        print(f"[cookbook] photo backfill: {done} card(s) processed")


@app.route("/cookbook/delete/<cid>", methods=["POST"])
def cookbook_delete(cid):
    with COOKBOOK_LOCK:
        idx = next((i for i, c in enumerate(cookbook())
                    if c.get("id") == cid), None)
        card = cookbook().pop(idx) if idx is not None else None
    if card:
        _save_cookbook()
        for f in (card.get("photo"), card.get("scan")):
            if f and os.path.basename(f) == f:
                try:
                    os.remove(os.path.join(photos_dir(), f))
                except OSError:
                    pass
    return jsonify({"ok": card is not None})


@app.route("/cookbook/photo/<name>")
def cookbook_photo(name):
    if os.path.basename(name) != name:
        return "bad name", 400
    p = os.path.join(photos_dir(), name)
    if not os.path.isfile(p):
        return "not found", 404
    ext = name.rsplit(".", 1)[-1].lower()
    mime = {"jpg": "image/jpeg", "png": "image/png", "webp": "image/webp",
            "gif": "image/gif", "avif": "image/avif"}.get(ext, "application/octet-stream")
    with open(p, "rb") as f:
        resp = Response(f.read(), mimetype=mime)
    resp.headers["Cache-Control"] = "public, max-age=86400"
    return resp


# ---- docs: a native Google-Docs-style editor with an AI writing bar ----------
# Documents are JSON files in docs_dir() ({id,title,html,created,updated}),
# autosaved from the editor. /docs/write streams a model answer into
# JOBS[job]["partial"] so the editor can render it typing live into the page.
_DOC_ID = re.compile(r"[0-9a-f]{32}$")

DOC_WRITER_SYSTEM = (
    "You are writing content directly INTO a document editor. Respond with "
    "clean, well-structured Markdown for the requested content ONLY — no "
    "chat preamble, no 'Here is…', no closing questions. Use headings, "
    "lists, bold, and tables only where they genuinely serve the document."
)


def _doc_path(did):
    return os.path.join(docs_dir(), f"{did}.json")


def _load_doc(did):
    try:
        with open(_doc_path(did)) as f:
            return json.load(f)
    except Exception:
        return None


def _store_doc(doc):
    os.makedirs(docs_dir(), exist_ok=True)
    tmp = _doc_path(doc["id"]) + ".tmp"
    with open(tmp, "w") as f:
        json.dump(doc, f)
    os.replace(tmp, _doc_path(doc["id"]))


def run_doc_write_job(job_id, prompt, doc_title):
    """Stream a document-writing answer, mirroring partial text into the job
    so the editor can show it typing."""
    user = prompt
    if doc_title and doc_title.lower() != "untitled document":
        user = f'Document title: "{doc_title}"\n\nRequest: {prompt}'
    # thinking off so visible text starts streaming immediately (better live
    # typing) — quality is unchanged for document prose on the 26B
    payload = {"messages": [
        {"role": "system", "content": DOC_WRITER_SYSTEM},
        {"role": "user", "content": user}], "stream": True, "temperature": 0.7,
        "chat_template_kwargs": {"enable_thinking": False}}
    parts = []
    try:
        _gurl = SPECIALISTS["general"]
        payload = _inject_model(payload, _gurl, "general")
        with requests.post(_gurl, json=payload, headers=_hdrs(_gurl),
                           stream=True, timeout=600) as r:
            r.raise_for_status()
            for line in _cancellable_lines(r, job_id):
                if not line:
                    continue
                line = line.decode("utf-8")
                if not line.startswith("data: "):
                    continue
                data = line[6:]
                if data.strip() == "[DONE]":
                    break
                try:
                    d = json.loads(data)["choices"][0]["delta"]
                    delta = d.get("content", "")
                    if delta:
                        parts.append(delta)
                        with JOBS_LOCK:
                            JOBS[job_id]["tokens"] += 1
                            JOBS[job_id]["partial"] = "".join(parts)
                    elif d.get("reasoning_content"):
                        with JOBS_LOCK:
                            JOBS[job_id]["tokens"] += 1
                except Exception:
                    pass
    except Exception as e:
        parts.append(f"\n\n*(error talking to the model: {e})*")
    ans = "".join(parts)
    with JOBS_LOCK:
        JOBS[job_id]["partial"] = ans
        JOBS[job_id]["answer"] = ans
        JOBS[job_id]["done"] = True


@app.route("/docs")
def docs_page():
    return Response(DOCS_PAGE, mimetype="text/html")


@app.route("/docs/list")
def docs_list():
    out = []
    if os.path.isdir(docs_dir()):
        for n in os.listdir(docs_dir()):
            if not n.endswith(".json"):
                continue
            d = _load_doc(n[:-5])
            if d:
                snippet = re.sub(r"<[^>]+>", " ", d.get("html", ""))
                snippet = re.sub(r"\s+", " ", snippet).strip()[:120]
                out.append({"id": d["id"], "title": d.get("title", "Untitled"),
                            "updated": d.get("updated", 0), "snippet": snippet})
    out.sort(key=lambda x: -x["updated"])
    return jsonify(out)


@app.route("/docs/new", methods=["POST"])
def docs_new():
    doc = {"id": uuid.uuid4().hex, "title": "Untitled document",
           "html": "", "created": time.time(), "updated": time.time()}
    _store_doc(doc)
    return jsonify({"id": doc["id"]})


@app.route("/docs/get/<did>")
def docs_get(did):
    if not _DOC_ID.fullmatch(did):
        return jsonify({"error": "bad id"}), 400
    doc = _load_doc(did)
    return (jsonify(doc), 200) if doc else (jsonify({"error": "not found"}), 404)


@app.route("/docs/save/<did>", methods=["POST"])
def docs_save(did):
    if not _DOC_ID.fullmatch(did):
        return jsonify({"error": "bad id"}), 400
    doc = _load_doc(did)
    if not doc:
        return jsonify({"error": "not found"}), 404
    body = request.json or {}
    if "title" in body:
        doc["title"] = str(body["title"])[:200] or "Untitled document"
    if "html" in body:
        doc["html"] = str(body["html"])[:2_000_000]
    doc["updated"] = time.time()
    _store_doc(doc)
    return jsonify({"ok": True, "updated": doc["updated"]})


@app.route("/docs/delete/<did>", methods=["POST"])
def docs_delete(did):
    if not _DOC_ID.fullmatch(did):
        return jsonify({"error": "bad id"}), 400
    try:
        os.remove(_doc_path(did))
    except FileNotFoundError:
        pass
    return jsonify({"ok": True})


# ---- sheets: a native spreadsheet with formulas and an AI fill bar ----------
# Same pattern as docs: JSON files in sheets_dir() ({id,title,cells:{A1:raw},
# rows,cols,...}). The formula engine lives client-side; the AI writer returns
# a JSON cells map (values + formulas) that the grid applies with animation.

SHEET_WRITER_SYSTEM = (
    "You fill in a spreadsheet. Respond with ONLY a JSON object mapping "
    "A1-style cell references to values — no commentary, no code fences. "
    "Values may be text labels, numbers, or formulas starting with '='. "
    "Formulas may use + - * / ^ ( ), cell refs (B2), ranges inside functions, "
    "and ONLY these functions: SUM, AVERAGE, COUNT, MIN, MAX, ROUND, ABS. "
    'Example: {"A1":"Item","B1":"Cost","A2":"Rent","B2":1200,'
    '"B6":"=SUM(B2:B5)"}. Put headers in row 1. Stay within columns A-J and '
    "rows 1-40 unless the request needs more.")

_CELL_REF = re.compile(r"[A-Z]{1,2}[1-9][0-9]{0,2}$")


def _sheet_path(sid):
    return os.path.join(sheets_dir(), f"{sid}.json")


def _load_sheet(sid):
    try:
        with open(_sheet_path(sid)) as f:
            return json.load(f)
    except Exception:
        return None


def _store_sheet(sh):
    os.makedirs(sheets_dir(), exist_ok=True)
    tmp = _sheet_path(sh["id"]) + ".tmp"
    with open(tmp, "w") as f:
        json.dump(sh, f)
    os.replace(tmp, _sheet_path(sh["id"]))


def run_sheet_write_job(job_id, prompt, title, existing):
    user = prompt
    if title and title.lower() != "untitled sheet":
        user = f'Spreadsheet title: "{title}"\n\nRequest: {prompt}'
    if existing:
        user += ("\n\nCells already filled (do not clobber unless the request "
                 "says to): " + json.dumps(existing)[:2000])
    # thinking off: structured JSON output, measured 5x faster on the 26B
    # with identical quality (same finding as the code pipeline)
    payload = {"messages": [
        {"role": "system", "content": SHEET_WRITER_SYSTEM},
        {"role": "user", "content": user}], "stream": True, "temperature": 0.3,
        "chat_template_kwargs": {"enable_thinking": False}}
    parts = []
    try:
        _gurl = SPECIALISTS["general"]
        payload = _inject_model(payload, _gurl, "general")
        with requests.post(_gurl, json=payload, headers=_hdrs(_gurl),
                           stream=True, timeout=600) as r:
            r.raise_for_status()
            for line in _cancellable_lines(r, job_id):
                if not line:
                    continue
                line = line.decode("utf-8")
                if not line.startswith("data: "):
                    continue
                data = line[6:]
                if data.strip() == "[DONE]":
                    break
                try:
                    d = json.loads(data)["choices"][0]["delta"]
                    if d.get("content"):
                        parts.append(d["content"])
                    if d.get("content") or d.get("reasoning_content"):
                        with JOBS_LOCK:
                            JOBS[job_id]["tokens"] += 1
                except Exception:
                    pass
    except Exception as e:
        with JOBS_LOCK:
            JOBS[job_id]["answer"] = f"[sheet error: model unreachable — {e}]"
            JOBS[job_id]["done"] = True
        return
    if _job_cancelled(job_id):
        with JOBS_LOCK:
            JOBS[job_id]["answer"] = "[sheet fill stopped]"
            JOBS[job_id]["done"] = True
        return
    out = _strip_think("".join(parts))
    m = re.search(r"\{[\s\S]*\}", out)
    cells = {}
    if m:
        try:
            raw = json.loads(m.group(0))
            for k, v in (raw.items() if isinstance(raw, dict) else []):
                k = str(k).strip().upper()
                if _CELL_REF.fullmatch(k) and isinstance(v, (str, int, float)):
                    cells[k] = str(v)[:500]
        except Exception:
            pass
    with JOBS_LOCK:
        if cells:
            JOBS[job_id]["answer"] = json.dumps({"cells": cells})
        else:
            JOBS[job_id]["answer"] = ("[sheet error: the model returned no "
                                      "usable cell data]")
        JOBS[job_id]["done"] = True


@app.route("/sheets")
def sheets_page():
    return Response(SHEETS_PAGE, mimetype="text/html")


@app.route("/sheets/list")
def sheets_list():
    out = []
    if os.path.isdir(sheets_dir()):
        for n in os.listdir(sheets_dir()):
            if not n.endswith(".json"):
                continue
            s = _load_sheet(n[:-5])
            if s:
                out.append({"id": s["id"], "title": s.get("title", "Untitled"),
                            "updated": s.get("updated", 0),
                            "filled": len(s.get("cells") or {})})
    out.sort(key=lambda x: -x["updated"])
    return jsonify(out)


@app.route("/sheets/new", methods=["POST"])
def sheets_new():
    sh = {"id": uuid.uuid4().hex, "title": "Untitled sheet", "cells": {},
          "rows": 60, "cols": 26, "created": time.time(),
          "updated": time.time()}
    _store_sheet(sh)
    return jsonify({"id": sh["id"]})


@app.route("/sheets/get/<sid>")
def sheets_get(sid):
    if not _DOC_ID.fullmatch(sid):
        return jsonify({"error": "bad id"}), 400
    sh = _load_sheet(sid)
    return (jsonify(sh), 200) if sh else (jsonify({"error": "not found"}), 404)


@app.route("/sheets/save/<sid>", methods=["POST"])
def sheets_save(sid):
    if not _DOC_ID.fullmatch(sid):
        return jsonify({"error": "bad id"}), 400
    sh = _load_sheet(sid)
    if not sh:
        return jsonify({"error": "not found"}), 404
    body = request.json or {}
    if "title" in body:
        sh["title"] = str(body["title"])[:200] or "Untitled sheet"
    if "cells" in body and isinstance(body["cells"], dict):
        sh["cells"] = {k: str(v)[:500] for k, v in body["cells"].items()
                       if _CELL_REF.fullmatch(str(k)) and str(v) != ""}
    if "fmt" in body and isinstance(body["fmt"], dict):
        clean = {}
        for k, v in body["fmt"].items():
            if _CELL_REF.fullmatch(str(k)) and isinstance(v, dict):
                vv = {kk: str(sv)[:20] for kk, sv in v.items()
                      if kk in ("b", "i", "al", "fc", "bg", "nf")
                      and isinstance(sv, (str, int))}
                if vv:
                    clean[k] = vv
        sh["fmt"] = clean
    if "colw" in body and isinstance(body["colw"], dict):
        sh["colw"] = {k: max(40, min(500, int(v)))
                      for k, v in body["colw"].items()
                      if re.fullmatch(r"[A-Z]{1,2}", str(k))
                      and str(v).lstrip("-").isdigit()}
    if "rows" in body:
        sh["rows"] = max(10, min(500, int(body["rows"])))
    sh["updated"] = time.time()
    _store_sheet(sh)
    return jsonify({"ok": True})


@app.route("/sheets/delete/<sid>", methods=["POST"])
def sheets_delete(sid):
    if not _DOC_ID.fullmatch(sid):
        return jsonify({"error": "bad id"}), 400
    try:
        os.remove(_sheet_path(sid))
    except FileNotFoundError:
        pass
    return jsonify({"ok": True})


@app.route("/sheets/write", methods=["POST"])
def sheets_write():
    body = request.json or {}
    prompt = (body.get("prompt") or "").strip()
    if not prompt:
        return jsonify({"error": "empty prompt"}), 400
    existing = body.get("cells") if isinstance(body.get("cells"), dict) else {}
    job_id = uuid.uuid4().hex
    with JOBS_LOCK:
        JOBS[job_id] = {"owner": _owner(), "tokens": 0, "done": False, "route": "general",
                        "answer": "", "stage": "filling", "audio": None,
                        "stage_ts": time.time()}
        _prune_jobs(JOBS)
    spawn_owned(target=run_sheet_write_job,
                     args=(job_id, prompt, body.get("title") or "", existing),
                     daemon=True).start()
    return jsonify({"job_id": job_id})


@app.route("/docs/write", methods=["POST"])
def docs_write():
    body = request.json or {}
    prompt = (body.get("prompt") or "").strip()
    if not prompt:
        return jsonify({"error": "empty prompt"}), 400
    job_id = uuid.uuid4().hex
    with JOBS_LOCK:
        JOBS[job_id] = {"owner": _owner(), "tokens": 0, "done": False, "route": "general",
                        "answer": "", "partial": "", "stage": "writing",
                        "audio": None, "stage_ts": time.time()}
        _prune_jobs(JOBS)
    spawn_owned(target=run_doc_write_job,
                     args=(job_id, prompt, body.get("title") or ""),
                     daemon=True).start()
    return jsonify({"job_id": job_id})


@app.route("/panel")
def panel_page_route():
    """Merged into /compare. The panel compared a fixed list of routes; compare
    does that as its "routes" mode, alongside raw-model comparison, with one
    picker and one judge. Kept as a redirect so bookmarks keep working."""
    return redirect("/compare", code=302)


@app.route("/panel/routes")
def panel_routes_route():
    return jsonify({"routes": panel_routes()})


@app.route("/panel/progress/<job_id>")
def panel_progress(job_id):
    with PANEL_JOBS_LOCK:
        job = owned_job(PANEL_JOBS, job_id)
        if not job:
            return jsonify({"error": "unknown job"}), 404
        return jsonify({
            "specialists": job["specialists"],
            "judge": job["judge"],
            "done": job["done"],
        })


# ---- compare (A/B: direct code vs creative-spec → code) ---------------------
COMPARE_JOBS: dict = {}
COMPARE_JOBS_LOCK = threading.Lock()

COMPARE_JUDGE_SYSTEM = (
    "You are judging two code implementations of the same request. "
    "Track A was coded directly by a code specialist. "
    "Track B was first designed by a creative director (brief shown), then coded by an engineer. "
    "Pick the better result on: visual quality, completeness, correctness, and wow factor. "
    "In 3-5 sentences: name the winner clearly, explain the key difference, and note one weakness of each."
)


# (The old direct-vs-creative-brief A/B prompts lived here. /compare is now a
# backend shoot-out; recover them from git history if that experiment returns.)


def _stream_into(url: str, payload: dict, job_id: str, field: str,
                 jobs=None, lock=None, route: str = None) -> str:
    """Stream into a COMPARE_JOBS-shaped pane. `jobs`/`lock` default to the
    compare tables; passing them lets other pages reuse this."""
    jobs = COMPARE_JOBS if jobs is None else jobs
    lock = COMPARE_JOBS_LOCK if lock is None else lock
    parts, usage = [], None
    think_n = 0
    try:
        payload = _with_usage(_inject_model(payload, url, route), url)
        with requests.post(url, json=payload, headers=_hdrs(url),
                           stream=True, timeout=600) as r:
            r.raise_for_status()
            for line in _lines(r):
                if not line:
                    continue
                line = line.decode("utf-8")
                if line.startswith("data: "):
                    data = line[6:]
                    if data.strip() == "[DONE]":
                        break
                    try:
                        chunk = json.loads(data)
                        if chunk.get("usage"):
                            usage = chunk["usage"]
                        if not chunk.get("choices"):
                            continue
                        d = chunk["choices"][0]["delta"]
                        delta = d.get("content", "")
                        if d.get("reasoning_content"):
                            think_n += 1
                        if delta:
                            parts.append(delta)
                        # count hidden reasoning too, so the pane's counter
                        # moves during a thinking model's think phase instead
                        # of sitting at 0 looking like a dead backend
                        if delta or d.get("reasoning_content"):
                            with lock:
                                jobs[job_id][field]["tokens"] += 1
                    except Exception:
                        pass
    except Exception as e:
        parts.append(_specialist_error(field, url, e))
    out = "".join(parts)
    # Belt and braces: some chat templates ignore enable_thinking:false. When
    # that happens the whole budget goes to reasoning_content and this pane is
    # blank (or one truncated sentence), which reads as a broken backend. Say
    # what actually happened rather than showing an empty box.
    if think_n and len(out.strip()) < 200:
        out += (f"\n\n*(this backend spent its budget on hidden reasoning "
                f"({think_n} thinking chunks) and returned little or no visible "
                f"answer — its chat template appears to ignore "
                f"`enable_thinking: false`)*")
    # A shoot-out can include paid backends — bill them like anything else.
    record_spend(url, f"compare:{field}", usage,
                 json.dumps(payload.get("messages", ""))[:20000], out)
    return out


# The worked example must NOT be parseable as a real file block. It used to
# show a literal path + fence + "// file contents"; the model echoed it back,
# the parser dutifully created src/components/App.jsx containing that comment,
# and every later step was then shown a phantom React file in a vanilla-JS
# project. Describe the shape in prose instead of demonstrating it verbatim.
_FILE_FORMAT_NOTE = (
    "\n\nIMPORTANT — output format: return every file as its own fenced code block, "
    "with the file's path on the line immediately BEFORE its opening fence "
    "(the path in backticks, e.g. bold-backtick style), then the fence with the "
    "language, then the file's REAL contents, then a closing fence. "
    "Never write an example or placeholder file — only real project files.\n"
    "Do not combine multiple files into one block. Do not show terminal commands — just the files. "
    "Every file must be COMPLETE from first line to last. NEVER stand in for code with "
    "placeholder comments like '// ... rest unchanged' or '// lines 1-349' — that destroys the file."
)


COMPARE_SYSTEM = (
    "You are a capable assistant. Answer the request directly and completely. "
    "If the request calls for code, return complete, working code with each "
    "file in its own fenced block, the filename on the line above it."
)


def compare_backends():
    """Backends a shoot-out may use — NAMES ONLY.

    /compare is open to any logged-in user while /settings is admin-only, so
    this must never leak URLs or API keys. It reports whether a backend costs
    money, because picking four models where three of them bill you is a
    decision the user should make with their eyes open."""
    with CONFIG_LOCK:
        out = []
        for b in ROUTER_CONFIG.get("backends", []):
            out.append({"id": b.get("id"), "kind": b.get("kind") or "openai",
                        "model": b.get("model") or "",
                        "paid": bool(b.get("paid"))})
    return out


def run_compare_backend(job_id: str, bid: str, url: str, question: str,
                        system: str = None, route: str = None):
    """One entrant in a comparison.

    `system` differs by mode and is the whole point of having two: comparing
    MODELS gives every entrant the same system prompt so you judge the models;
    comparing ROUTES gives each its own persona so you judge the setups."""
    # Thinking OFF, same as both judges. An entrant is capped at max_tokens, and
    # a reasoning model spends that entire budget on reasoning_content — the
    # pane then shows nothing, or the one sentence it managed before the cap.
    # Measured on Qwen3.6-27B (.53): 800-token budget → 800 thinking tokens and
    # a half-finished opening sentence; with the flag → a complete answer in 243.
    payload = {"messages": [{"role": "system", "content": system or COMPARE_SYSTEM},
                            {"role": "user", "content": question}],
               "stream": True, "max_tokens": COMPARE_MAX_TOKENS,
               "chat_template_kwargs": {"enable_thinking": False}}
    answer = _stream_into(url, payload, job_id, bid, route=route)
    with COMPARE_JOBS_LOCK:
        pane = COMPARE_JOBS.get(job_id, {}).get(bid)
        if pane is not None:
            pane["answer"] = answer
            pane["done"] = True
            pane["stage"] = "done"
    _check_compare_judge(job_id, question)


def _check_compare_judge(job_id: str, question: str):
    with COMPARE_JOBS_LOCK:
        job = COMPARE_JOBS.get(job_id)
        if not job:
            return
        ids = job.get("order") or []
        if not all(job.get(b, {}).get("done") for b in ids):
            return
        if job.get("judge", {}).get("started"):
            return                      # last finisher wins; only judge once
        job["judge"]["started"] = True
        answers = [(b, job[b]["answer"]) for b in ids]

    if len(answers) < 2:
        with COMPARE_JOBS_LOCK:
            COMPARE_JOBS[job_id]["judge"]["done"] = True
            COMPARE_JOBS[job_id]["judge"]["answer"] = (
                "_Only one model was selected, so there is nothing to compare._")
            COMPARE_JOBS[job_id]["done"] = True
        return

    blocks = "\n\n".join(
        f"=== {bid} ===\n{ans[:COMPARE_JUDGE_EXCERPT]}" for bid, ans in answers)
    payload = {
        "messages": [
            {"role": "system", "content": COMPARE_JUDGE_SYSTEM},
            {"role": "user", "content": (
                f"Request: {question}\n\n"
                "Several models answered the same request. Compare them and say "
                "which is strongest and why, naming them exactly as labelled.\n\n"
                + blocks
            )},
        ],
        "stream": True,
        "max_tokens": 1200,
        # Thinking OFF. The judge is a structured writer like the doc/sheet
        # writers, and on a reasoning model it will otherwise spend the ENTIRE
        # token budget on reasoning_content and emit no visible answer at all:
        # measured 697 thinking deltas and 0 content deltas on a realistic
        # payload. The page then sat on "deliberating" forever.
        "chat_template_kwargs": {"enable_thinking": False},
    }
    parts = []
    try:
        _rurl = SPECIALISTS["reasoning"]
        with requests.post(_rurl, json=_inject_model(payload, _rurl, "reasoning"),
                           headers=_hdrs(_rurl), stream=True, timeout=300) as r:
            r.raise_for_status()
            for line in _lines(r):
                if not line:
                    continue
                line = line.decode("utf-8")
                if line.startswith("data: "):
                    data = line[6:]
                    if data.strip() == "[DONE]":
                        break
                    try:
                        chunk = json.loads(data)
                        delta = chunk["choices"][0]["delta"].get("content", "")
                        if delta:
                            parts.append(delta)
                            with COMPARE_JOBS_LOCK:
                                COMPARE_JOBS[job_id]["judge"]["tokens"] += 1
                    except Exception:
                        pass
    except Exception as e:
        parts = [f"[judge error: {e}]"]

    verdict = "".join(parts).strip()
    if not verdict:
        # Defence in depth: an empty verdict used to render as an eternal
        # "deliberating…" because the UI only swapped in text when the answer
        # was non-empty. Say what happened instead.
        verdict = ("_The judge returned nothing. This usually means its backend "
                   "spent the whole token budget on hidden reasoning — check "
                   "that the reasoning route has thinking disabled._")
    with COMPARE_JOBS_LOCK:
        COMPARE_JOBS[job_id]["judge"]["answer"] = verdict
        COMPARE_JOBS[job_id]["judge"]["done"] = True
        COMPARE_JOBS[job_id]["done"] = True


@app.route("/compare")
def compare_page_route():
    return Response(COMPARE_PAGE, mimetype="text/html")


COMPARE_MODES = ("backends", "routes")


@app.route("/compare/backends")
def compare_backends_route():
    """Superseded by /compare/targets; kept so an older page still works."""
    return jsonify({"backends": compare_backends()})


@app.route("/compare/targets")
def compare_targets_route():
    mode = request.args.get("mode", "backends")
    if mode not in COMPARE_MODES:
        mode = "backends"
    if mode == "routes":
        items = [{"id": r["route"], "label": r["route"], "sub": r["backend"],
                  "paid": r["paid"], "default": r["default"]}
                 for r in panel_routes()]
    else:
        items = [{"id": b["id"], "label": b["id"], "sub": b["model"],
                  "paid": b["paid"], "default": False}
                 for b in compare_backends()]
    return jsonify({"mode": mode, "targets": items})


COMPARE_MAX_BACKENDS = 6
# Output budget per entrant. 1400 was too tight: anything asking for code or a
# structured multi-part answer stopped mid-sentence (measured — a "design a rate
# limiter, then implement it" prompt hit the cap with the answer half-written).
# Env-tunable because the picker can include PAID backends, where a bigger budget
# is real money: +2600 tokens on a $10/Mtok model is ~2.6c per entrant per run.
COMPARE_MAX_TOKENS = int(os.environ.get("CREWE_COMPARE_MAX_TOKENS", "4000"))
# How much of each answer the judge is shown. Must scale WITH the budget above —
# raising the entrant cap alone would just mean the judge grades a smaller
# fraction of what it is comparing. 6 entrants x 6000 chars is ~9k tokens, well
# inside the judge backend's context (the 26B on 8087 reports n_ctx 65536).
COMPARE_JUDGE_EXCERPT = int(os.environ.get("CREWE_COMPARE_JUDGE_EXCERPT", "6000"))


@app.route("/compare/ask", methods=["POST"])
def compare_ask():
    body = request.json or {}
    question = body.get("question", "").strip()
    if not question:
        return jsonify({"error": "empty question"}), 400

    mode = body.get("mode", "backends")
    if mode not in COMPARE_MODES:
        return jsonify({"error": "unknown mode"}), 400

    wanted = body.get("targets") or body.get("backends") or []
    if not isinstance(wanted, list):
        return jsonify({"error": "targets must be a list"}), 400

    # Names in, URLs resolved SERVER-SIDE — the browser never supplies a URL,
    # so a crafted request cannot aim a run at an arbitrary host.
    chosen, seen = [], set()
    if mode == "routes":
        for name in wanted:
            if name in SPECIALISTS and name not in seen:
                seen.add(name)
                chosen.append((name, SPECIALISTS[name],
                               SYSTEM_PROMPTS.get(name, SYSTEM_PROMPTS["general"]),
                               name))
    else:
        with CONFIG_LOCK:
            by_id = {b["id"]: b for b in ROUTER_CONFIG.get("backends", [])}
        for bid in wanted:
            b = by_id.get(bid)
            if b and bid not in seen:
                seen.add(bid)
                chosen.append((bid, _chat_url(b), COMPARE_SYSTEM, None))
    if not chosen:
        return jsonify({"error": "pick at least one"}), 400
    if len(chosen) > COMPARE_MAX_BACKENDS:
        return jsonify({"error": f"at most {COMPARE_MAX_BACKENDS} at once"}), 400

    job_id = uuid.uuid4().hex
    with COMPARE_JOBS_LOCK:
        job = {"owner": _owner(), "order": [k for k, _u, _s, _r in chosen],
               "mode": mode,
               "judge": {"tokens": 0, "done": False, "answer": "",
                         "started": False},
               "done": False}
        for k, _u, _s, _r in chosen:
            job[k] = {"tokens": 0, "done": False, "answer": "",
                      "stage": "drafting"}
        COMPARE_JOBS[job_id] = job
        _prune_jobs(COMPARE_JOBS)

    for k, u, sysp, rt in chosen:
        spawn_owned(target=run_compare_backend,
                    args=(job_id, k, u, question, sysp, rt), daemon=True).start()
    return jsonify({"job_id": job_id, "mode": mode,
                    "targets": [k for k, _u, _s, _r in chosen]})


@app.route("/compare/progress/<job_id>")
def compare_progress(job_id):
    with COMPARE_JOBS_LOCK:
        job = owned_job(COMPARE_JOBS, job_id)
        if not job:
            return jsonify({"error": "unknown job"}), 404
        return jsonify({k: dict(v) if isinstance(v, dict) else v for k, v in job.items()})


PAGE = r"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>Crewe</title>
<link rel="icon" href="data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 64 64'%3E%3Crect width='64' height='64' rx='14' fill='%238b4a18'/%3E%3Cpath d='M32 57V34M32 34 16 13M32 34 48 13' stroke='%23faf6f0' stroke-width='7' stroke-linecap='round' fill='none'/%3E%3Ccircle cx='32' cy='34' r='6' fill='%23faf6f0'/%3E%3C/svg%3E">
<link href="/static/fonts.css" rel="stylesheet">
<script src="/static/marked.min.js"></script>
<script src="/static/jszip.min.js"></script>
<style>
  :root {
    --bg:        #faf6f0;
    --panel:     #ffffff;
    --panel2:    #f2ebe0;
    --ink:       #1e160a;
    --muted:     #8a7560;
    --border:    #ddd0b8;
    --recipes:   #a85e20;
    --creative:  #903828;
    --code:      #2e6e2a;
    --general:   #226660;
    --reasoning: #2a4e8c;
    --audio:     #7a4ba0;
    --search:    #1f6f6a;
    --accent:    #8b4a18;
    --sw:        240px;
  }
  * { box-sizing:border-box; margin:0; padding:0; }
  html,body { height:100%; overflow:hidden; }
  body { font:15px/1.6 'DM Sans',system-ui,sans-serif;
         background:var(--bg); color:var(--ink); display:flex; flex-direction:column; }

  /* ---- header ---- */
  header { flex:0 0 auto; padding:18px 28px 14px; border-bottom:1px solid var(--border);
           background:var(--panel); }
  .hdr-top { display:flex; align-items:baseline; gap:14px; flex-wrap:wrap; }
  header h1 { font-family:'Lora',Georgia,serif; font-size:21px; font-weight:600;
              color:var(--accent); letter-spacing:-.01em; }
  header .sub { color:var(--muted); font-size:12px; }
  header a { color:var(--reasoning); text-decoration:none; }
  header a:hover { text-decoration:underline; }
  .mem-row { display:flex; align-items:center; gap:10px; margin-top:8px; }
  .mem-badge { font-size:12px; color:var(--muted); cursor:pointer;
               user-select:none; transition:color .15s; }
  .mem-badge:hover { color:var(--ink); }
  .mem-btn { background:none; border:1px solid var(--border); color:var(--muted);
             border-radius:5px; padding:2px 9px; font:12px 'DM Sans',sans-serif;
             cursor:pointer; transition:all .15s; }
  .mem-btn:hover { border-color:var(--accent); color:var(--accent); }

  /* ---- memory panel ---- */
  #memPanelWrap { flex:0 0 auto; background:var(--panel2); border-bottom:1px solid var(--border); }
  .mem-panel { max-width:820px; margin:0 auto; padding:14px 28px; }
  .mem-panel-head { display:flex; justify-content:space-between; align-items:center;
                    font-size:11px; font-weight:600; color:var(--muted);
                    text-transform:uppercase; letter-spacing:.08em; margin-bottom:10px; }
  .mem-close { background:none; border:0; color:var(--muted); cursor:pointer;
               font-size:16px; line-height:1; transition:color .15s; }
  .mem-close:hover { color:var(--ink); }
  .mem-text { font-size:13px; color:var(--ink); white-space:pre-wrap; line-height:1.65; }

  /* ---- layout ---- */
  .main { flex:1; display:flex; min-height:0; }
  .chat-col { flex:1; display:flex; flex-direction:column; min-width:0; }
  #chat { flex:1; overflow-y:auto; scroll-behavior:smooth; }
  #chat::-webkit-scrollbar { width:6px; }
  #chat::-webkit-scrollbar-track { background:transparent; }
  #chat::-webkit-scrollbar-thumb { background:var(--border); border-radius:3px; }
  .inner { max-width:820px; margin:0 auto; padding:30px 28px;
           display:flex; flex-direction:column; gap:24px; }

  /* ---- messages ---- */
  .msg { background:var(--panel); border:1px solid var(--border); border-radius:14px;
         padding:20px 24px; box-shadow:0 1px 4px rgba(80,50,20,.06); }
  .msg.user { background:transparent; border-color:transparent;
              box-shadow:none; padding:0 4px; }
  .msg.user .q { color:var(--ink); font-weight:500; font-size:15px; }
  .badge { display:inline-flex; align-items:center; gap:5px; font-size:10px; font-weight:700;
           text-transform:uppercase; letter-spacing:.08em; padding:3px 10px;
           border-radius:999px; margin-bottom:12px; }
  .badge::before { content:''; width:5px; height:5px; border-radius:50%;
                   background:currentColor; opacity:.7; flex:0 0 auto; }
  .badge.recipes  { background:rgba(168,94,32,.10);  color:var(--recipes); }
  .badge.creative { background:rgba(144,56,40,.10);  color:var(--creative); }
  .badge.code     { background:rgba(46,110,42,.10);  color:var(--code); }
  .badge.general  { background:rgba(34,102,96,.10);  color:var(--general); }
  .badge.reasoning{ background:rgba(42,78,140,.10);  color:var(--reasoning); }
  .badge.audio    { background:rgba(122,75,160,.10); color:var(--audio); }
  .badge.search   { background:rgba(31,111,106,.10); color:var(--search); }
  .audwrap { margin:2px 0 12px; }
  .aud { display:block; width:100%; max-width:420px; height:38px; }
  .auddl { display:inline-flex; align-items:center; gap:5px; margin-top:6px; font-size:12px;
           font-weight:600; color:var(--audio); text-decoration:none;
           background:rgba(122,75,160,.10); padding:4px 10px; border-radius:8px; }
  .auddl:hover { background:rgba(122,75,160,.18); }
  .pending { display:flex; align-items:center; gap:12px; color:var(--muted); font-size:14px; }
  .spinner { width:16px; height:16px; border:2px solid var(--border);
             border-top-color:var(--accent); border-radius:50%;
             animation:spin .7s linear infinite; flex:0 0 auto; }
  @keyframes spin { to { transform:rotate(360deg); } }
  .answer { word-wrap:break-word; line-height:1.8; font-size:14.5px; }
  .counter { font-variant-numeric:tabular-nums; font-size:13px; }

  /* ---- file chips under code answers ---- */
  .filebar { display:flex; flex-wrap:wrap; gap:6px; margin-top:12px;
             padding-top:10px; border-top:1px dashed var(--border); align-items:center; }
  .fchip { display:inline-flex; align-items:center; gap:4px; background:var(--panel2);
           border:1px solid var(--border); border-radius:8px; padding:3px 6px 3px 10px;
           font-size:12px; }
  .fchip .fname { font-family:ui-monospace,SFMono-Regular,Menlo,Consolas,monospace;
                  font-size:11.5px; color:var(--ink); }
  .fchip button, .fall { background:none; border:0; cursor:pointer; color:var(--muted);
                         font-size:12px; padding:2px 4px; border-radius:4px; height:auto; }
  .fchip button:hover, .fall:hover { color:var(--accent); background:rgba(139,74,24,.08); }
  .fall { border:1px solid var(--border); border-radius:8px; padding:3px 10px; font-weight:600; }

  /* ---- markdown rendering ---- */
  .md p { margin-bottom:.9em; }
  .md p:last-child { margin-bottom:0; }
  .md h1,.md h2,.md h3,.md h4 { font-family:'Lora',Georgia,serif; font-weight:600;
    margin:1.3em 0 .5em; line-height:1.3; }
  .md h1 { font-size:1.35em; color:var(--accent); }
  .md h2 { font-size:1.18em; color:var(--recipes); }
  .md h3 { font-size:1.04em; color:var(--general); letter-spacing:.01em; }
  .md h4 { font-size:1em;    color:var(--muted); }
  .md ul,.md ol { padding-left:1.6em; margin-bottom:.9em; }
  .md li { margin-bottom:.35em; }
  .md li p { margin-bottom:.2em; }
  .md li::marker { color:var(--code); font-weight:700; }
  .md strong { font-weight:600; color:var(--recipes); }
  .md em { font-style:italic; color:var(--creative); }
  .md code { font-family:ui-monospace,SFMono-Regular,Menlo,Consolas,monospace;
             font-size:.86em; background:var(--panel2); padding:2px 6px;
             border-radius:4px; color:var(--code); }
  .md pre { background:var(--panel2); border:1px solid var(--border);
            border-radius:10px; padding:16px 18px; overflow-x:auto; margin-bottom:.9em; }
  .md pre code { background:none; padding:0; border-radius:0; color:var(--ink); font-size:.86em; }
  .md blockquote { border-left:3px solid var(--creative); margin:.9em 0;
                   padding-left:1.1em; color:var(--muted); font-style:italic; }
  .md a { color:var(--reasoning); text-decoration:none; }
  .md a:hover { text-decoration:underline; }
  .md hr { border:0; border-top:1px solid var(--border); margin:1.2em 0; }
  .md table { border-collapse:collapse; width:100%; margin-bottom:.9em; font-size:13.5px; }
  .md th,.md td { border:1px solid var(--border); padding:8px 12px; text-align:left; }
  .md th { background:var(--panel2); font-weight:600; color:var(--accent); }

  /* ---- input ---- */
  footer { flex:0 0 auto; border-top:1px solid var(--border);
           background:var(--panel); padding:16px 28px; }
  .bar { max-width:820px; margin:0 auto; display:flex; gap:10px; align-items:flex-end; }
  textarea { flex:1; resize:none; max-height:160px; min-height:46px;
             background:var(--bg); color:var(--ink);
             border:1px solid var(--border); border-radius:12px;
             padding:12px 16px; font:15px/1.5 'DM Sans',sans-serif;
             outline:none; transition:border-color .2s; }
  textarea:focus { border-color:var(--accent); }
  textarea::placeholder { color:var(--muted); }
  button { background:var(--accent); color:#fff; border:0; border-radius:12px;
           padding:0 24px; height:46px; font:600 14px 'DM Sans',sans-serif;
           cursor:pointer; transition:opacity .15s; letter-spacing:.01em; }
  button:hover { opacity:.85; }
  button:disabled { opacity:.3; cursor:default; }
  #go { color:var(--btn-ink); }

  .spend-tag { font:500 11.5px 'DM Sans',sans-serif; color:var(--muted);
               border:1px solid var(--border); border-radius:999px;
               padding:3px 9px; cursor:default; }

  /* ---- attachments ---- */
  .attach-row { max-width:820px; margin:0 auto 8px; display:flex; flex-wrap:wrap;
                gap:6px; }
  .attach-row:empty { display:none; }
  .chip { display:inline-flex; align-items:center; gap:7px;
          background:var(--panel2); border:1px solid var(--border);
          border-radius:999px; padding:5px 10px 5px 12px;
          font:500 12px 'DM Sans',sans-serif; color:var(--ink); }
  .chip .sz { color:var(--muted); font-weight:400; }
  .chip .x { cursor:pointer; color:var(--muted); font-size:14px; line-height:1;
             padding:0 2px; }
  .chip .x:hover { color:var(--accent); }
  .chip.err { border-color:#c0392b; color:#c0392b; }
  .attach-btn { background:var(--bg); color:var(--muted);
                border:1px solid var(--border); border-radius:12px;
                height:46px; padding:0 14px; font-size:17px; cursor:pointer; }
  .attach-btn:hover { border-color:var(--accent); opacity:1; }
  .attach-btn.busy { opacity:.45; cursor:default; }

  /* ---- effort selector (code route only) ---- */
  #effort { height:46px; background:var(--bg); color:var(--ink);
            border:1px solid var(--border); border-radius:12px;
            padding:0 10px; font:500 13px 'DM Sans',sans-serif;
            cursor:pointer; outline:none; }
  #effort:focus { border-color:var(--accent); }
  .effort-help { max-width:820px; margin:6px auto 0; color:var(--muted);
                 font:400 11.5px/1.4 'DM Sans',sans-serif; text-align:right; }

  /* ---- sidebar ---- */
  .sidebar { width:var(--sw); border-left:1px solid var(--border);
             display:flex; flex-direction:column; overflow:hidden;
             background:var(--panel2); }
  .sb-top { flex:0 0 auto; padding:12px; border-bottom:1px solid var(--border); }
  .new-btn { width:100%; background:var(--panel); border:1px solid var(--border);
             color:var(--muted); border-radius:9px; padding:8px 12px;
             font:500 13px 'DM Sans',sans-serif; cursor:pointer;
             text-align:left; display:flex; align-items:center;
             gap:7px; transition:all .15s; height:auto; box-shadow:0 1px 2px rgba(80,50,20,.05); }
  .new-btn:hover { border-color:var(--accent); color:var(--accent); }
  .sess-list { flex:1; overflow-y:auto; padding:6px;
               display:flex; flex-direction:column; gap:2px; }
  .sess-list::-webkit-scrollbar { width:4px; }
  .sess-list::-webkit-scrollbar-thumb { background:var(--border); border-radius:2px; }
  .sess { padding:9px 11px; border-radius:9px; cursor:pointer;
          position:relative; border:1px solid transparent; transition:all .15s; }
  .sess:hover { background:var(--panel); border-color:var(--border); }
  .sess.active { background:var(--panel); border-color:var(--border);
                 box-shadow:0 1px 3px rgba(80,50,20,.07); }
  .sess-name { font-size:13px; font-weight:500; white-space:nowrap; overflow:hidden;
               text-overflow:ellipsis; padding-right:18px; color:var(--ink); }
  .sess-meta { font-size:11px; color:var(--muted); margin-top:2px; }
  .sess-del { position:absolute; right:7px; top:50%; transform:translateY(-50%);
              background:none; border:0; color:var(--muted); cursor:pointer;
              font-size:13px; padding:2px 4px; display:none; height:auto;
              line-height:1; border-radius:4px; transition:all .15s; }
  .sess:hover .sess-del { display:block; }
  .sess-del:hover { color:var(--creative); }
  .sb-bot { flex:0 0 auto; padding:10px; border-top:1px solid var(--border); }
  .cb-link { display:flex; align-items:center; gap:9px; color:var(--muted);
             font:13px 'DM Sans',sans-serif; text-decoration:none;
             padding:9px 11px; border-radius:9px; transition:all .15s; }
  .cb-link:hover { background:var(--panel); color:var(--recipes); }

  /* ---- header links + memory (right-aligned) ---- */
  .hdr-links { display:flex; align-items:baseline; gap:16px; font-size:12px; }
  .hdr-links a { color:var(--muted); text-decoration:none; transition:color .15s; }
  .hdr-links a:hover { color:var(--accent); }
  .mem-row { margin-top:0; margin-left:auto; }

  /* ---- empty / welcome state ---- */
  #emptyState { display:flex; flex-direction:column; align-items:center;
                justify-content:center; min-height:100%; padding:48px 28px; text-align:center; }
  .es-card { max-width:600px; width:100%; }
  .es-title { font-family:'Lora',Georgia,serif; font-weight:600; font-size:30px;
              color:var(--accent); letter-spacing:-.01em; margin-bottom:8px; }
  .es-sub { color:var(--muted); font-size:14.5px; margin-bottom:24px; }
  .es-legend { display:flex; flex-wrap:wrap; gap:8px; justify-content:center; margin-bottom:30px; }
  .es-legend .badge { margin-bottom:0; }
  .es-examples { display:grid; grid-template-columns:1fr 1fr; gap:10px; }
  .es-ex { display:flex; align-items:center; justify-content:space-between; gap:12px;
           background:var(--panel); border:1px solid var(--border); border-radius:12px;
           padding:13px 15px; cursor:pointer; height:auto; width:100%; text-align:left;
           box-shadow:0 1px 3px rgba(80,50,20,.05); transition:all .15s; }
  .es-ex:hover { border-color:var(--accent); opacity:1; transform:translateY(-1px);
                 box-shadow:0 4px 12px rgba(80,50,20,.10); }
  .es-ex-q { font:500 13.5px 'DM Sans',sans-serif; color:var(--ink); line-height:1.4; }
  .es-ex .badge { margin-bottom:0; flex:0 0 auto; }
  @media (max-width:640px){ .es-examples { grid-template-columns:1fr; } }

  /* ---- dark theme (manual slider, defaults to light) ---- */
  :root { --btn-ink:#fff; --down:#c0392b; }
  :root[data-theme="dark"] {
    --bg:#191410; --panel:#221c15; --panel2:#2b241b;
    --ink:#ece2d3; --muted:#9c8a75; --border:#3a3125;
    --recipes:#d99a5b; --creative:#d98276; --code:#7cbf72; --general:#63b8ab;
    --reasoning:#8aaee8; --audio:#bb95dd; --search:#5fb8ad; --accent:#d29a62;
    --btn-ink:#241b11; --down:#e06552;
    color-scheme:dark;
  }
  .theme-toggle { display:inline-flex; align-items:center; gap:5px; cursor:pointer; user-select:none; }
  .theme-toggle input { position:absolute; opacity:0; pointer-events:none; }
  .tt-track { width:34px; height:18px; border-radius:999px; background:var(--panel2);
              border:1px solid var(--border); position:relative; transition:background .2s; }
  .tt-knob { position:absolute; top:1px; left:1px; width:14px; height:14px; border-radius:50%;
             background:var(--accent); transition:left .2s; }
  .theme-toggle input:checked + .tt-track .tt-knob { left:17px; }
  .tt-icon { font-size:11px; color:var(--muted); }

  /* ---- route health dots ---- */
  .badge.down { opacity:.45; }
  .badge.down::before { background:var(--down); opacity:1; }
  .hstrip { display:inline-flex; gap:5px; align-items:center; margin-left:4px; }
  .hdot { width:8px; height:8px; border-radius:50%; opacity:.8; cursor:default; }
  .hdot.down { background:var(--down) !important; opacity:1;
               animation:hpulse 1.2s ease-in-out infinite; }
  @keyframes hpulse { 50% { opacity:.35; } }

  /* ---- copy answer ---- */
  .msg { position:relative; }
  .msg-copy { position:absolute; top:12px; right:12px; background:none;
              border:1px solid var(--border); color:var(--muted); border-radius:6px;
              padding:3px 9px; font:500 12px 'DM Sans',sans-serif; height:auto;
              opacity:0; transition:all .15s; cursor:pointer; }
  .msg:hover .msg-copy { opacity:1; }
  .msg-copy:hover { color:var(--accent); border-color:var(--accent); }

  /* ---- stop button state ---- */
  #go.stop { background:var(--creative); }

  /* ---- mobile ---- */
  #sbToggle { display:none; background:none; border:1px solid var(--border);
              color:var(--muted); border-radius:8px; padding:5px 11px; height:auto;
              font-size:16px; line-height:1; cursor:pointer; }
  #sbBackdrop { position:fixed; inset:0; background:rgba(20,12,4,.35);
                z-index:55; display:none; }
  #sbBackdrop.show { display:block; }
  @media (max-width:700px){
    #sbToggle { display:inline-flex; }
    header { padding:12px 16px 10px; }
    .inner { padding:20px 14px; }
    footer { padding:12px 14px; }
    .sidebar { position:fixed; top:0; right:0; bottom:0; width:min(80vw,300px);
               transform:translateX(105%); transition:transform .25s ease;
               z-index:60; box-shadow:-8px 0 30px rgba(0,0,0,.25); }
    .sidebar.open { transform:none; }

    /* The composer grew an attach button and an effort select, which squeezed
       the textarea and pushed Send off the right edge. Give the textarea its
       own row and let the controls sit underneath it. */
    .bar { flex-wrap:wrap; gap:8px; }
    .bar textarea { flex:1 1 100%; min-height:44px; }
    #attach { order:2; padding:0 12px; }
    #attachRepo { order:2; padding:0 12px; }
    #effort { order:3; flex:1 1 auto; min-width:0; }
    #go     { order:4; padding:0 18px; }
    .effort-help { text-align:center; }

    /* Eight header links wrapped into a three-row block with the last one
       clipped. Tighten them and drop the decorative arrows. */
    .hdr-links { flex-wrap:wrap; gap:2px 12px; font-size:12.5px;
                 max-width:100%; }
    .hdr-links a { white-space:nowrap; }
    .hdr-links .arr { display:none; }
    .spend-tag { font-size:11px; padding:2px 7px; }
  }
</style>
<script>if(localStorage.getItem('creweTheme')==='dark')document.documentElement.dataset.theme='dark';</script>
</head>
<body>
<header>
  <div class="hdr-top">
    <h1>Crewe</h1>
    <span id="healthStrip" class="hstrip" title="route backend health"></span>
    <nav class="hdr-links">
      <a href="/docs" target="_blank">docs<span class="arr"> ↗</span></a>
      <a href="/sheets" target="_blank">sheets<span class="arr"> ↗</span></a>
      <a href="/scratch" target="_blank">scratchpad<span class="arr"> ↗</span></a>
      <a href="/shots" target="_blank">build checks<span class="arr"> ↗</span></a>
      <a href="/panel" target="_blank">panel<span class="arr"> ↗</span></a>
      <a href="/compare" target="_blank">compare<span class="arr"> ↗</span></a>
      <a href="/admin" class="admin-only" style="display:none">👥 people</a>
      <a href="/settings" class="admin-only" style="display:none">⚙ settings</a>
      <span id="spendTag" class="spend-tag" style="display:none"></span>
      <a href="/help">? help</a>
      <a href="/account">👤 account</a>
      <a href="__SOURCE_URL__" target="_blank" rel="noopener"
         title="Crewe is AGPL-3.0 — this links to the source of the version you are using">&lt;/&gt; source</a>
    </nav>
    <script>
    // Hide admin-only nav items from non-admins (the routes are server-gated
    // regardless; this just avoids showing links that 403).
    fetch('/whoami').then(r=>r.json()).then(w=>{
      if(w && w.is_admin) document.querySelectorAll('.admin-only').forEach(e=>e.style.display='');
    }).catch(()=>{});
    </script>
    <div class="mem-row">
      <span id="memBadge" class="mem-badge">💭 loading…</span>
      <button id="memClear" class="mem-btn">Clear memory</button>
      <label class="theme-toggle" title="Toggle dark mode">
        <span class="tt-icon">☀</span>
        <input type="checkbox" id="themeSw">
        <span class="tt-track"><span class="tt-knob"></span></span>
        <span class="tt-icon">🌙</span>
      </label>
      <button id="sbToggle" title="Chats">☰</button>
    </div>
  </div>
</header>
<div id="memPanelWrap" style="display:none">
  <div class="mem-panel">
    <div class="mem-panel-head">
      <span>Conversation memory</span>
      <button class="mem-close" id="memClose">✕</button>
    </div>
    <div id="memText" class="mem-text">No memory yet.</div>
  </div>
</div>
<div id="sbBackdrop"></div>
<div class="main">
  <div class="chat-col">
    <div id="chat">
      <div id="emptyState">
        <div class="es-card">
          <h2 class="es-title">Ask anything.</h2>
          <p class="es-sub">Crewe reads your question and routes it to the right specialist.</p>
          <div class="es-legend">
            <span class="badge recipes"   data-route="recipes">recipes</span>
            <span class="badge creative"  data-route="creative">creative</span>
            <span class="badge code"      data-route="code">code</span>
            <span class="badge reasoning" data-route="reasoning">reasoning</span>
            <span class="badge general"   data-route="general">general</span>
            <span class="badge search"    data-route="search">search</span>
          </div>
          <div class="es-examples" id="esExamples"></div>
        </div>
      </div>
      <div class="inner" id="inner"></div>
    </div>
    <footer>
      <div class="attach-row" id="attachRow"></div>
      <div class="bar">
        <textarea id="q" rows="1" placeholder="Ask anything…  (Enter to send, Shift+Enter for newline)"></textarea>
        <input type="file" id="fileIn" hidden>
        <button id="attach" class="attach-btn" title="Attach a document — pdf, docx, xlsx, txt, md, csv…">&#128206;</button>
        <button id="attachRepo" class="attach-btn" title="Attach a public GitHub repository for review">&#128230;</button>
        <select id="effort" title="Only applies to coding answers. Higher tiers use larger, slower models.">
          <option value="fast">Effort: Fast</option>
          <option value="normal" selected>Effort: Normal</option>
          <option value="extra">Effort: Extra</option>
        </select>
        <button id="go">Send</button>
      </div>
      <div class="effort-help" id="effortHelp">Effort applies only to coding answers — more effort takes longer</div>
    </footer>
  </div>
  <div class="sidebar">
    <div class="sb-top">
      <button class="new-btn" id="newChat">＋&nbsp; New chat</button>
    </div>
    <div class="sess-list" id="sessList"></div>
    <div class="sb-bot">
      <a href="/cookbook" class="cb-link" target="_blank">📖&nbsp; Cookbook</a>
    </div>
  </div>
</div>
<script>
const q = document.getElementById('q');
const go = document.getElementById('go');
const chat = document.getElementById('chat');
const inner = document.getElementById('inner');
const emptyState = document.getElementById('emptyState');
let activeJob = null;   // job_id while a reply is generating (Send ⇄ Stop)

function updateEmptyState(){
  const show = !inner.children.length;
  emptyState.style.display = show ? 'flex' : 'none';
  if(show) renderExamples();
}
function scrollDown(){ chat.scrollTop = chat.scrollHeight; }

// ---- example prompts: a pool, 4 route-diverse picks refreshed each time the
//      welcome screen appears (so the same four don't keep reappearing) --------
const EXAMPLE_POOL = [
  {route:"reasoning", label:"Explain the CAP theorem with an example", q:"Explain the CAP theorem with a concrete example."},
  {route:"reasoning", label:"Monty Hall, step by step",               q:"Walk me through the Monty Hall problem step by step, and why switching wins."},
  {route:"reasoning", label:"Mergesort vs quicksort",                 q:"Compare mergesort and quicksort on speed, memory use, and stability."},
  {route:"code",      label:"A Python retry decorator with backoff",  q:"Write a Python decorator that retries a function with exponential backoff."},
  {route:"code",      label:"A debounce function in vanilla JS",      q:"Write a debounce function in vanilla JavaScript and explain how it works."},
  {route:"code",      label:"Parse a CSV string without libraries",   q:"Write a function to parse a CSV string into rows and columns without using any libraries."},
  {route:"recipes",   label:"A weeknight dinner with chickpeas & spinach", q:"A quick weeknight dinner using chickpeas, spinach, and whatever's likely in the pantry."},
  {route:"recipes",   label:"Use up three overripe bananas",          q:"What can I bake with three very overripe bananas?"},
  {route:"recipes",   label:"A cozy soup for a cold night",           q:"Suggest a hearty soup to simmer on a cold evening, with a short shopping list."},
  {route:"creative",  label:"A short poem about a railway junction",  q:"Write a short poem about a railway junction at night."},
  {route:"creative",  label:"A 100-word sci-fi story",                q:"Write a 100-word science fiction story about a lighthouse keeper on Mars."},
  {route:"creative",  label:"Name a neighborhood coffee shop",        q:"Brainstorm ten warm, memorable names for a neighborhood coffee shop."},
  {route:"general",   label:"Why is the sky blue?",                   q:"Why is the sky blue, and why do sunsets turn red?"},
  {route:"general",   label:"Weather vs climate",                     q:"What's the difference between weather and climate?"},
  {route:"general",   label:"How noise-cancelling headphones work",   q:"How do noise-cancelling headphones actually work?"},
  {route:"search",    label:"Latest on NASA's Artemis program",       q:"What's the latest news on NASA's Artemis program?"},
  {route:"search",    label:"Recent advances in battery tech",        q:"What are the most recent advances in battery technology?"},
  {route:"search",    label:"The most recent Nobel Prize in Physics", q:"Who won the most recent Nobel Prize in Physics, and for what?"},
];
function shuffle(a){ for(let i=a.length-1;i>0;i--){ const j=Math.floor(Math.random()*(i+1)); const t=a[i];a[i]=a[j];a[j]=t; } return a; }
function renderExamples(){
  const box=document.getElementById('esExamples'); if(!box) return;
  const byRoute={}; EXAMPLE_POOL.forEach(e=>{ (byRoute[e.route]=byRoute[e.route]||[]).push(e); });
  const routes=shuffle(Object.keys(byRoute)).slice(0,4);
  box.innerHTML='';
  routes.forEach(r=>{
    const arr=byRoute[r]; const e=arr[Math.floor(Math.random()*arr.length)];
    const b=document.createElement('button'); b.className='es-ex'; b.dataset.q=e.q;
    b.innerHTML=`<span class="es-ex-q">${escapeHtml(e.label)}</span><span class="badge ${e.route}">${e.route}</span>`;
    b.addEventListener('click',()=>{ q.value=e.q; q.style.height='auto'; q.style.height=Math.min(q.scrollHeight,160)+'px'; ask(); });
    box.appendChild(b);
  });
}
function escapeHtml(s){ const d=document.createElement('div'); d.textContent=s; return d.innerHTML; }
marked.use({ breaks: true, gfm: true });
// Raw HTML in a model answer (e.g. a website the code model wrote) must be shown
// as TEXT, not injected live into the page. Without this, marked passes raw tags
// straight through and the browser renders them — making the message look cut off
// at the first '<' tag. Escape block- and inline-level raw HTML tokens so they
// display verbatim. Fenced ``` code blocks are separate tokens and still render
// as proper code blocks.
marked.use({ renderer: {
  html(token){ const s = typeof token === 'string' ? token : (token && token.text) || ''; return escapeHtml(s); }
}});
function relTime(ts){
  const d=Date.now()-ts, m=60e3, h=36e5, day=864e5;
  if(d<m) return 'just now';
  if(d<h) return Math.floor(d/m)+'m ago';
  if(d<day) return Math.floor(d/h)+'h ago';
  return Math.floor(d/day)+'d ago';
}
function newId(){
  return crypto.randomUUID ? crypto.randomUUID()
       : Math.random().toString(36).slice(2)+Date.now().toString(36);
}

// (file extraction & per-file actions moved above session restore — they run
// during renderStoredMsg, so they must be initialized before it)
const EXTS = {html:'html',htm:'html',css:'css',javascript:'js',js:'js',
              typescript:'ts',ts:'ts',python:'py',py:'py',
              json:'json',bash:'sh',shell:'sh',sql:'sql',c:'c',cpp:'cpp',rust:'rs'};
function isCommandBlock(lang, code){
  if(!['bash','shell','sh','zsh','console','terminal','cmd'].includes(lang)) return false;
  return !code.trim().startsWith('#!/');
}
function extractFiles(mdText){
  const files=[]; const seen=new Set();
  const re=/```(\w*) *\n([\s\S]*?)```/g;
  let m,lastIndex=0;
  while((m=re.exec(mdText))!==null){
    const lang=m[1].trim()||''; const code=m[2];
    const before=mdText.slice(lastIndex,m.index); lastIndex=re.lastIndex;
    if(isCommandBlock(lang,code)) continue;
    const beforeLines=before.trimEnd().split('\n');
    let filename=null;
    for(let i=beforeLines.length-1;i>=Math.max(0,beforeLines.length-4);i--){
      const line=beforeLines[i].trim().replace(/[`*_#>]+/g,' ').trim();
      const fnm=line.match(/([a-zA-Z0-9_\-]+(?:\/[a-zA-Z0-9_\-.]+)*\.[a-zA-Z0-9]+)/);
      if(fnm&&!fnm[1].match(/^https?:/)&&fnm[1].length<80){filename=fnm[1];break;}
    }
    if(!filename){
      const firstLine=code.split('\n')[0];
      const cm=firstLine.match(/(?:\/\/|#|<!--|\/\*)\s*([a-zA-Z0-9_\-./]+\.[a-zA-Z0-9]+)/);
      if(cm&&!cm[1].match(/^https?:/)) filename=cm[1];
    }
    if(!filename){
      const ext=EXTS[lang]||(lang||'txt');
      filename=files.length===0?`index.${ext}`:`file${files.length+1}.${ext}`;
    }
    let final=filename,n=2;
    while(seen.has(final)){
      const dot=filename.lastIndexOf('.');
      final=dot>0?filename.slice(0,dot)+`_${n}`+filename.slice(dot):filename+`_${n}`;
      n++;
    }
    seen.add(final);
    files.push({name:final,code,lang});
  }
  return files;
}
const FILE_STORE=[];
// clipboard with fallback — navigator.clipboard is unavailable over plain-http
// LAN/tailnet origins (non-secure context), so fall back to execCommand there
function writeClipboard(t){
  if(navigator.clipboard && window.isSecureContext) return navigator.clipboard.writeText(t);
  const ta=document.createElement('textarea'); ta.value=t;
  ta.style.position='fixed'; ta.style.opacity='0';
  document.body.appendChild(ta); ta.select();
  try{ document.execCommand('copy'); }finally{ ta.remove(); }
  return Promise.resolve();
}
const ANSWER_STORE=[];
function copyBtn(text){
  const i=ANSWER_STORE.push(text)-1;
  return `<button class="msg-copy" onclick="copyAnswer(${i},this)" title="Copy answer as markdown">⧉ copy</button>`;
}
async function copyAnswer(i,btn){
  try{ await writeClipboard(ANSWER_STORE[i]); btn.textContent='✓ copied'; }
  catch(e){ btn.textContent='✗ failed'; }
  setTimeout(()=>btn.textContent='⧉ copy',1200);
}
function fileBar(text){
  if(!text||!text.includes('```')) return '';
  const files=extractFiles(text).filter(f=>f.name!=='output.txt');
  if(!files.length) return '';
  const idx=FILE_STORE.push(files)-1;
  const chips=files.map((f,i)=>`<span class="fchip"><span class="fname">${escapeHtml(f.name)}</span><button onclick="copyFile(${idx},${i},this)" title="Copy">⧉</button><button onclick="dlFile(${idx},${i})" title="Download">⬇</button></span>`).join('');
  const zip=files.length>1?`<button class="fall" onclick="dlZip(${idx})">⬇ all (.zip)</button>`:'';
  const sp=files.some(f=>f.name.endsWith('.html'))?`<button class="fall" onclick="toScratch(${idx})">▶ scratchpad</button>`:'';
  const pl=files.some(f=>f.name.endsWith('.html'))?`<button class="fall" onclick="window.open('/play/'+sessionId+'/index.html','_blank')">▶ play</button>`:'';
  return `<div class="filebar">${chips}${zip}${sp}${pl}</div>`;
}
async function copyFile(idx,i,btn){
  try{
    await writeClipboard(FILE_STORE[idx][i].code);
    const t=btn.textContent; btn.textContent='✓'; setTimeout(()=>btn.textContent=t,900);
  }catch(e){ alert('copy failed: '+e); }
}
function dlFile(idx,i){
  const f=FILE_STORE[idx][i];
  const a=document.createElement('a');
  a.href=URL.createObjectURL(new Blob([f.code],{type:'text/plain'}));
  a.download=f.name.split('/').pop();
  a.click(); URL.revokeObjectURL(a.href);
}
async function dlZip(idx){
  const zip=new JSZip();
  FILE_STORE[idx].forEach(f=>zip.file(f.name,f.code));
  const blob=await zip.generateAsync({type:'blob'});
  const a=document.createElement('a');
  a.href=URL.createObjectURL(blob); a.download='project.zip';
  a.click(); URL.revokeObjectURL(a.href);
}
function toScratch(idx){
  const files=FILE_STORE[idx];
  const f=files.find(x=>x.name==='index.html')||files.find(x=>x.name.endsWith('.html'))||files[0];
  localStorage.setItem('scratchLoad',f.code);
  window.open('/scratch','_blank');
}

// ---- session management -----------------------------------------------------
let sessions = JSON.parse(localStorage.getItem('routerSessions')||'[]');
let sessionId = localStorage.getItem('routerCurrentSession');

function saveSessions(){ localStorage.setItem('routerSessions', JSON.stringify(sessions)); }
function getMsgs(sid){ return JSON.parse(localStorage.getItem('routerMsgs_'+sid)||'[]'); }
function saveMsgs(sid,msgs){ localStorage.setItem('routerMsgs_'+sid, JSON.stringify(msgs)); }

function createSession(){
  const s={id:newId(), name:'New chat', createdAt:Date.now()};
  sessions.unshift(s); saveSessions(); return s;
}

if(!sessions.length){ const s=createSession(); sessionId=s.id; }
if(!sessionId||!sessions.find(s=>s.id===sessionId)) sessionId=sessions[0].id;
localStorage.setItem('routerCurrentSession', sessionId);

let currentMsgs = getMsgs(sessionId);

function switchSession(id){
  sessionId=id; localStorage.setItem('routerCurrentSession',id);
  currentMsgs=getMsgs(id);
  inner.innerHTML='';
  currentMsgs.forEach(renderStoredMsg);
  scrollDown(); renderSidebar(); loadMemory(); updateEmptyState(); closeSb();
  refreshAttachments();
}

function deleteSession(id){
  sessions=sessions.filter(s=>s.id!==id);
  localStorage.removeItem('routerMsgs_'+id); saveSessions();
  if(sessionId===id){
    if(!sessions.length){ const s=createSession(); sessionId=s.id; }
    else sessionId=sessions[0].id;
    localStorage.setItem('routerCurrentSession',sessionId);
    currentMsgs=getMsgs(sessionId);
    inner.innerHTML=''; currentMsgs.forEach(renderStoredMsg); scrollDown(); loadMemory(); updateEmptyState();
  }
  renderSidebar();
}

function renderSidebar(){
  const list=document.getElementById('sessList');
  list.innerHTML='';
  sessions.forEach(s=>{
    const el=document.createElement('div');
    el.className='sess'+(s.id===sessionId?' active':'');
    el.innerHTML=`<div class="sess-name">${escapeHtml(s.name)}</div>
      <div class="sess-meta">${relTime(s.createdAt)}</div>
      <button class="sess-del" title="Delete">✕</button>`;
    el.addEventListener('click',e=>{if(!e.target.classList.contains('sess-del'))switchSession(s.id);});
    el.querySelector('.sess-del').addEventListener('click',e=>{e.stopPropagation();deleteSession(s.id);});
    list.appendChild(el);
  });
}

document.getElementById('newChat').addEventListener('click',()=>{
  const s=createSession(); sessionId=s.id;
  localStorage.setItem('routerCurrentSession',sessionId);
  currentMsgs=[]; inner.innerHTML=''; renderSidebar(); loadMemory(); updateEmptyState(); closeSb(); q.focus();
});

// ---- message persistence ----------------------------------------------------
function pushMsg(msg){
  currentMsgs.push(msg); saveMsgs(sessionId,currentMsgs);
  const s=sessions.find(s=>s.id===sessionId);
  if(s&&s.name==='New chat'){
    const first=currentMsgs.find(m=>m.type==='user');
    if(first){s.name=first.content.slice(0,28)+(first.content.length>28?'…':'');saveSessions();}
  }
  renderSidebar();
}

function renderStoredMsg(msg){
  if(msg.type==='user'){
    const el=document.createElement('div');
    el.className='msg user';
    el.innerHTML=`<div class="q">${escapeHtml(msg.content)}</div>`;
    inner.appendChild(el);
  } else {
    const el=document.createElement('div');
    el.className='msg';
    el.innerHTML=`<span class="badge ${msg.route}">→ ${msg.route}</span>
      ${copyBtn(msg.content)}
      ${audioPlayer(msg.audio)}
      <div class="answer md">${marked.parse(msg.content)}</div>
      ${fileBar(msg.content)}`;
    inner.appendChild(el);
  }
}

// restore current session on load
currentMsgs.forEach(renderStoredMsg);
scrollDown(); updateEmptyState();

// ---- memory -----------------------------------------------------------------
async function loadMemory(){
  try{
    const m=await fetch('/memory/'+sessionId).then(r=>r.json());
    updateMemoryUI(m);
  }catch(e){}
}
function updateMemoryUI(m){
  const turns=(m&&m.turns)||0, summary=(m&&m.summary)||'', files=(m&&m.files)||[];
  let badge=turns>0?`💭 ${turns} turn${turns!==1?'s':''}`:'💭 no memory';
  if(files.length) badge+=` · 📁 ${files.length}`;
  document.getElementById('memBadge').textContent=badge;
  let text=summary||'No memory yet.';
  if(files.length) text+=`\n\nProject files: ${files.join(', ')}`;
  document.getElementById('memText').textContent=text;
}
document.getElementById('memBadge').addEventListener('click',async()=>{
  const w=document.getElementById('memPanelWrap');
  if(w.style.display==='none'){await loadMemory();w.style.display='block';}
  else w.style.display='none';
});
document.getElementById('memClose').addEventListener('click',()=>{
  document.getElementById('memPanelWrap').style.display='none';
});
document.getElementById('memClear').addEventListener('click',async()=>{
  if(!confirm('Clear conversation memory for this session?'))return;
  await fetch('/memory/'+sessionId,{method:'DELETE'});
  updateMemoryUI({turns:0,summary:''});
  document.getElementById('memPanelWrap').style.display='none';
});
loadMemory();

// ---- chat -------------------------------------------------------------------
function addUser(question){
  const el=document.createElement('div');
  el.className='msg user';
  el.innerHTML=`<div class="q">${escapeHtml(question)}</div>`;
  inner.appendChild(el); scrollDown(); updateEmptyState();
}
function addReply(){
  const el=document.createElement('div');
  el.className='msg'; inner.appendChild(el); scrollDown(); return el;
}
function renderPending(el,route,tokens,stage,stageAge){
  let label = 'routing…';
  if(route){
    const stageLabels = {improving:'improving…', verifying:'verifying…', scripting:'writing script…', recording:'recording…', searching:'searching web…', reading:'reading pages…', planning:'planning the build…'};
    const stageLabel = stageLabels[stage] || (stage && stage !== 'drafting' ? escapeHtml(stage) + '…' : 'drafting…');
    label = stage === 'recording' ? '🎙 recording…' : `${stageLabel} ${tokens} tokens`;
    // slow-backend honesty: distinguish "computing on a slow model" from dead
    if(stageAge>60){
      const m=Math.floor(stageAge/60);
      label += (tokens===0)
        ? ` · model is reading the prompt (${m}m — slow backends take a while)`
        : ` · ${m}m in this stage`;
    }
  }
  const badge=route?`<span class="badge ${route}">→ ${route}</span>`:'';
  el.innerHTML=`${badge}<div class="pending"><div class="spinner"></div>
    <span class="counter">${label}</span></div>`;
  scrollDown();
}
function audioPlayer(audio){
  if(!audio) return '';
  return `<div class="audwrap">
    <audio class="aud" controls preload="metadata" src="${audio}"></audio>
    <a class="auddl" href="${audio}" download="crewe-recording.wav">⤓ Download .wav</a>
  </div>`;
}

function renderAnswer(el,route,answer,audio){
  el.innerHTML=`<span class="badge ${route}">→ ${route}</span>
    ${copyBtn(answer)}
    ${audioPlayer(audio)}
    <div class="answer md">${marked.parse(answer)}</div>
    ${fileBar(answer)}`;
  scrollDown();
}

function lastCodeMsg(){
  // last few assistant messages containing code (oldest first) — used to seed
  // server-side file memory; multiple messages so one truncated/lazy answer
  // can't poison the seed
  const picks=[];
  for(let i=currentMsgs.length-1;i>=0&&picks.length<3;i--){
    const m=currentMsgs[i];
    if(m.type==='assistant'&&m.content&&m.content.includes('```'))
      picks.push(m.content.slice(0,300000));
  }
  return picks.length?picks.reverse():null;
}

async function ask(){
  const question=q.value.trim();
  if(!question||activeJob)return;
  q.value=''; q.style.height='auto'; go.disabled=true;

  addUser(question);
  pushMsg({type:'user',content:question,timestamp:Date.now()});
  const reply=addReply();
  renderPending(reply,null,0);

  let res;
  try{
    res=await(await fetch('/ask',{
      method:'POST',headers:{'Content-Type':'application/json'},
      body:JSON.stringify({question,session_id:sessionId,seed_code:lastCodeMsg(),
                           effort:(document.getElementById('effort')||{}).value||'hard'})
    })).json();
  }catch(e){
    reply.innerHTML=`<div class="answer">router unreachable: ${escapeHtml(String(e))}</div>`;
    go.disabled=false; return;
  }

  const route=res.route, jobId=res.job_id;
  activeJob=jobId; go.textContent='Stop'; go.classList.add('stop'); go.disabled=false;
  renderPending(reply,route,0);

  const poll=setInterval(async()=>{
    let p;
    try{p=await(await fetch('/progress/'+jobId)).json();}catch(e){return;}
    if(p.done){
      clearInterval(poll);
      renderAnswer(reply,route,p.answer,p.audio);
      pushMsg({type:'assistant',content:p.answer,route,audio:p.audio,timestamp:Date.now()});
      activeJob=null; go.textContent='Send'; go.classList.remove('stop');
      go.disabled=false; q.focus();
      setTimeout(loadMemory,3000);
    } else {
      renderPending(reply,route,p.tokens,p.stage,p.stage_age);
    }
  },400);
}

q.addEventListener('input',()=>{ q.style.height='auto'; q.style.height=Math.min(q.scrollHeight,160)+'px'; });
go.addEventListener('click',()=>{
  if(activeJob){ fetch('/stop/'+activeJob,{method:'POST'}).catch(()=>{}); }
  else ask();
});
q.addEventListener('keydown',e=>{ if(e.key==='Enter'&&!e.shiftKey){e.preventDefault();ask();} });

// ---- dark mode slider (defaults to light; persisted per-browser) ------------
const themeSw=document.getElementById('themeSw');
themeSw.checked = localStorage.getItem('creweTheme')==='dark';
themeSw.addEventListener('change',()=>{
  if(themeSw.checked){ document.documentElement.dataset.theme='dark'; localStorage.setItem('creweTheme','dark'); }
  else { delete document.documentElement.dataset.theme; localStorage.setItem('creweTheme','light'); }
});

// ---- attachments ------------------------------------------------------------
// A file attaches to the SESSION, so it stays available across follow-up
// questions and whatever route each one lands on — not just "summarise this".
const attachRow=document.getElementById('attachRow');
const attachBtn=document.getElementById('attach');
const fileIn=document.getElementById('fileIn');

function fmtSize(n){ return n>=1e6?(n/1e6).toFixed(1)+'M chars'
                          :n>=1e3?Math.round(n/1e3)+'k chars':n+' chars'; }

function renderAttachments(list){
  attachRow.innerHTML=(list||[]).map(u=>`
    <span class="chip" data-id="${u.id}">${escapeHtml(u.name)}
      <span class="sz">${fmtSize(u.chars)}</span>
      <span class="x" title="Remove">&times;</span></span>`).join('');
  attachRow.querySelectorAll('.chip .x').forEach(x=>{
    x.addEventListener('click',async()=>{
      const id=x.closest('.chip').dataset.id;
      await fetch(`/uploads/${encodeURIComponent(sessionId)}/${id}`,{method:'DELETE'});
      refreshAttachments();
    });
  });
}

async function refreshAttachments(){
  try{
    const r=await(await fetch('/uploads/'+encodeURIComponent(sessionId))).json();
    renderAttachments(r.uploads);
  }catch(e){ /* an unreachable router is already reported on send */ }
}

if(attachBtn){
  attachBtn.addEventListener('click',()=>fileIn.click());
  fileIn.addEventListener('change',async()=>{
    const f=fileIn.files[0];
    if(!f) return;
    attachBtn.classList.add('busy'); attachBtn.disabled=true;
    const fd=new FormData();
    fd.append('file',f); fd.append('session_id',sessionId);
    try{
      const r=await(await fetch('/upload',{method:'POST',body:fd})).json();
      if(r.error){
        attachRow.insertAdjacentHTML('beforeend',
          `<span class="chip err">${escapeHtml(f.name)} — ${escapeHtml(r.error)}</span>`);
      } else {
        refreshAttachments();
      }
    }catch(e){
      attachRow.insertAdjacentHTML('beforeend',
        `<span class="chip err">upload failed: ${escapeHtml(String(e))}</span>`);
    }finally{
      attachBtn.classList.remove('busy'); attachBtn.disabled=false;
      fileIn.value='';
    }
  });
}

const repoBtn=document.getElementById('attachRepo');
if(repoBtn){
  repoBtn.addEventListener('click',async()=>{
    const gh=prompt('Public GitHub repository URL:\n(e.g. https://github.com/owner/repo — a /tree/branch suffix works too)');
    if(!gh) return;
    repoBtn.classList.add('busy'); repoBtn.disabled=true;
    try{
      const r=await(await fetch('/upload',{method:'POST',
        headers:{'Content-Type':'application/json'},
        body:JSON.stringify({session_id:sessionId,github_url:gh})})).json();
      if(r.error){
        attachRow.insertAdjacentHTML('beforeend',
          `<span class="chip err">${escapeHtml(r.error)}</span>`);
      } else refreshAttachments();
    }catch(e){
      attachRow.insertAdjacentHTML('beforeend',
        `<span class="chip err">repo attach failed: ${escapeHtml(String(e))}</span>`);
    }finally{ repoBtn.classList.remove('busy'); repoBtn.disabled=false; }
  });
}

// ---- effort selector (persisted per-browser, like the theme) ----------------
// Always visible, but the server only consults it when a question routes to
// code — every other specialist has exactly one backend.
const effortSel=document.getElementById('effort');
if(effortSel){
  // Browsers that saved a choice before the rename hold "easy"/"hard".
  const _alias={easy:'normal',hard:'extra'};
  let _saved=localStorage.getItem('creweEffort')||'normal';
  effortSel.value=_alias[_saved]||_saved;
  if(!effortSel.value||![...effortSel.options].some(o=>o.value===effortSel.value))
    effortSel.value='normal';
  effortSel.addEventListener('change',()=>{
    localStorage.setItem('creweEffort', effortSel.value);
  });
}

// ---- cost awareness ---------------------------------------------------------
// The effort switch is only an honest money control if it SAYS so, and only
// when the backend behind it actually bills. Everything here is driven by
// /spend, never hardcoded — a local-only install should never mention money.
let spendState=null;
function money(n){ return (n<0.01&&n>0) ? '<$0.01' : '$'+n.toFixed(2); }

async function refreshSpend(){
  try{
    const s=await(await fetch('/spend')).json();
    spendState=s;
    const help=document.getElementById('effortHelp');
    const sel=document.getElementById('effort');
    if(sel&&s.coder_paid){
      const LBL={fast:'Fast',normal:'Normal',extra:'Extra'};
      [...sel.options].forEach(o=>{
        o.textContent='Effort: '+(LBL[o.value]||o.value)+(s.coder_paid[o.value]?' 💵':'');
      });
      // Any tier can be local or paid — no level implies cost. Say only what
      // this install's config actually is.
      if(help){
        const paid=Object.keys(LBL).filter(k=>s.coder_paid[k]);
        const base='Effort applies only to coding answers — ';
        if(!paid.length)             help.textContent=base+'higher tiers take longer';
        else if(paid.length===3)     help.innerHTML=base+'<b>every tier costs money</b>';
        else help.innerHTML=base+'<b>'+paid.map(k=>LBL[k]).join(' and ')+
                            (paid.length>1?' cost':' costs')+' money</b>, the rest are free';
      }
    }
    const tag=document.getElementById('spendTag');
    if(tag){
      if(s.any_paid){
        tag.style.display='';
        tag.textContent='💵 '+money(s.month)+' / 30d';
        tag.title=s.requests+' paid requests · '+money(s.day)+' today · '+money(s.all)+' all time'
                  +(s.estimated_any?' · some figures estimated from text length':'');
      } else tag.style.display='none';
    }
  }catch(e){}
}
refreshSpend(); setInterval(refreshSpend,60000);

// ---- route health dots ------------------------------------------------------
async function pollHealth(){
  try{
    const h=await(await fetch('/health/routes')).json();
    document.querySelectorAll('.es-legend .badge').forEach(b=>{
      const r=b.dataset.route; if(!(r in h)) return;
      b.classList.toggle('down', !h[r]);
      b.title = r+' backend '+(h[r]?'online':'offline');
    });
    // always-visible header strip — one dot per route, red when down
    const strip=document.getElementById('healthStrip');
    if(strip){
      strip.innerHTML='';
      Object.keys(h).sort().forEach(r=>{
        const d=document.createElement('span');
        d.className='hdot'+(h[r]?'':' down');
        d.style.background=`var(--${r}, var(--muted))`;
        d.title=r+': '+(h[r]?'online':'OFFLINE');
        strip.appendChild(d);
      });
    }
  }catch(e){}
}
pollHealth(); setInterval(pollHealth,30000);

// ---- custom-route badge colors (from /settings) -----------------------------
fetch('/settings/routes_meta').then(r=>r.json()).then(meta=>{
  const customs=meta.filter(m=>m.custom&&m.color);
  if(!customs.length)return;
  const css=customs.map(m=>
    `:root{--${m.name}:${m.color}}`+
    `.badge.${m.name}{background:color-mix(in srgb, ${m.color} 12%, transparent);color:var(--${m.name})}`
  ).join('\n');
  const st=document.createElement('style'); st.textContent=css;
  document.head.appendChild(st);
}).catch(()=>{});

// ---- mobile sidebar ---------------------------------------------------------
function closeSb(){
  document.querySelector('.sidebar').classList.remove('open');
  document.getElementById('sbBackdrop').classList.remove('show');
}
document.getElementById('sbToggle').addEventListener('click',()=>{
  const sb=document.querySelector('.sidebar');
  sb.classList.toggle('open');
  document.getElementById('sbBackdrop').classList.toggle('show', sb.classList.contains('open'));
});
document.getElementById('sbBackdrop').addEventListener('click',closeSb);

renderSidebar();
refreshAttachments();
q.focus();
</script>
</body>
</html>"""


COOKBOOK_PAGE = r"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>Cookbook</title>
<link rel="icon" href="data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 64 64'%3E%3Crect width='64' height='64' rx='14' fill='%238b4a18'/%3E%3Cpath d='M32 57V34M32 34 16 13M32 34 48 13' stroke='%23faf6f0' stroke-width='7' stroke-linecap='round' fill='none'/%3E%3Ccircle cx='32' cy='34' r='6' fill='%23faf6f0'/%3E%3C/svg%3E">
<link href="/static/fonts.css" rel="stylesheet">
<script>if(localStorage.getItem('creweTheme')==='dark')document.documentElement.dataset.theme='dark';</script>
<style>
  :root { --bg:#faf6f0; --panel:#ffffff; --ink:#1e160a; --muted:#8a7560;
          --border:#ddd0b8; --recipes:#a85e20; --accent:#2a4e8c;
          --btn-ink:#fff; }
  :root[data-theme="dark"] {
    --bg:#191410; --panel:#221c15; --ink:#ece2d3; --muted:#9c8a75;
    --border:#3a3125; --recipes:#d99a5b; --accent:#8aaee8;
    --btn-ink:#241b11; color-scheme:dark; }
  .theme-toggle { display:inline-flex; align-items:center; gap:5px;
                  cursor:pointer; user-select:none; flex:0 0 auto; }
  .theme-toggle input { position:absolute; opacity:0; pointer-events:none; }
  .tt-track { width:34px; height:18px; border-radius:999px; background:var(--bg);
              border:1px solid var(--border); position:relative; }
  .tt-knob { position:absolute; top:1px; left:1px; width:14px; height:14px;
             border-radius:50%; background:var(--recipes); transition:left .2s; }
  .theme-toggle input:checked + .tt-track .tt-knob { left:17px; }
  .tt-icon { font-size:11px; color:var(--muted); }
  header h1 { font-family:'Lora',Georgia,serif; color:var(--recipes); }
  * { box-sizing:border-box; }
  html,body { margin:0; font:15px/1.5 'DM Sans',-apple-system,Segoe UI,Roboto,sans-serif;
              background:var(--bg); color:var(--ink); }
  header { padding:14px 24px; border-bottom:1px solid var(--border);
           display:flex; align-items:center; gap:14px; }
  .back { color:var(--accent); text-decoration:none; font-size:13px; flex:0 0 auto; }
  header h1 { margin:0; font-size:18px; flex:0 0 auto; }
  input[type=search] { flex:1; background:var(--panel); border:1px solid var(--border);
                       color:var(--ink); border-radius:8px; padding:8px 12px;
                       font:inherit; outline:none; max-width:400px; }
  input[type=search]::placeholder { color:var(--muted); }
  .grid { display:grid; grid-template-columns:repeat(auto-fill,minmax(270px,1fr));
          gap:16px; padding:24px; }
  .card { background:var(--panel); border:1px solid var(--border); border-radius:12px;
          cursor:pointer; transition:border-color .15s; }
  .card:hover { border-color:var(--recipes); }
  .card-head { padding:14px 16px; border-bottom:1px solid var(--border); }
  .card-title { font-size:14px; font-weight:600; margin:0 0 4px; }
  .card-date { font-size:11px; color:var(--muted); }
  .card-tags { font-size:11px; color:var(--recipes); margin-top:3px;
               text-transform:capitalize; }
  .csel { background:var(--panel); color:var(--ink);
          border:1px solid var(--border); border-radius:8px; padding:7px 8px;
          font:13px inherit; outline:none; cursor:pointer; }
  .csel:focus { border-color:var(--recipes); }
  .grp-head { grid-column:1/-1; font:600 15px 'Lora',Georgia,serif;
              color:var(--recipes); text-transform:capitalize;
              padding-top:6px; margin-bottom:-6px;
              border-bottom:1px solid var(--border); padding-bottom:6px; }
  .card-del { position:absolute; top:10px; left:10px; background:var(--panel);
              border:1px solid var(--border); border-radius:6px;
              color:var(--muted); cursor:pointer; font-size:12px;
              width:24px; height:24px; line-height:1; display:none; z-index:2; }
  .card:hover .card-del { display:block; }
  .card-del:hover { color:#c0392b; border-color:#c0392b; }
  .card-preview { padding:12px 16px; font-size:13px; color:var(--muted);
                  display:-webkit-box; -webkit-line-clamp:2; -webkit-box-orient:vertical;
                  overflow:hidden; }
  .empty { text-align:center; color:var(--muted); padding:80px 24px; font-size:14px; }
  .overlay { display:none; position:fixed; inset:0; background:rgba(0,0,0,.75);
             z-index:100; align-items:center; justify-content:center; padding:24px; }
  .overlay.open { display:flex; }
  .modal { background:var(--panel); border:1px solid var(--border); border-radius:14px;
           max-width:680px; width:100%; max-height:82vh; display:flex; flex-direction:column; }
  .modal-head { padding:16px 20px; border-bottom:1px solid var(--border);
                display:flex; justify-content:space-between; align-items:flex-start; gap:12px; }
  .modal-title { font-size:16px; font-weight:600; margin:0; }
  .modal-close { background:none; border:0; color:var(--muted); font-size:20px;
                 cursor:pointer; padding:0; line-height:1; flex:0 0 auto; }
  .modal-body { padding:20px; overflow-y:auto; flex:1; }
  .modal-q { font-size:12px; color:var(--muted); margin-bottom:14px; font-style:italic; }
  .recipe-text { white-space:pre-wrap; font-size:14px; line-height:1.7; }
  .card { position:relative; }
  .card-img { width:100%; height:130px; object-fit:cover; display:block;
              border-radius:12px 12px 0 0;
              border-bottom:1px solid var(--border); }
  .modal-img { width:100%; max-height:260px; object-fit:cover;
               border-radius:10px; margin-bottom:14px; display:block; }
</style>
</head>
<body>
<header>
  <a href="/" class="back">← Crewe</a>
  <h1>📖 Cookbook</h1>
  <input type="search" id="search" placeholder="Search recipes…">
  <select id="groupSel" class="csel" title="Group by">
    <option value="meal">Group: meal</option>
    <option value="protein">Group: protein</option>
    <option value="none">Newest first</option>
  </select>
  <select id="mealSel" class="csel" title="Filter by meal">
    <option value="">All meals</option>
    <option>breakfast</option><option>lunch</option><option>dinner</option>
    <option>dessert</option><option>snack</option><option>drink</option>
    <option>side</option>
  </select>
  <select id="protSel" class="csel" title="Filter by protein">
    <option value="">All proteins</option>
    <option>chicken</option><option>beef</option><option>pork</option>
    <option>seafood</option><option>turkey</option><option>lamb</option>
    <option>vegetarian</option><option>other</option>
  </select>
  <label class="theme-toggle" title="Dark mode"><span class="tt-icon">☀</span><input type="checkbox" id="themeSw"><span class="tt-track"><span class="tt-knob"></span></span><span class="tt-icon">🌙</span></label>
</header>
<div id="grid" class="grid"></div>
<div class="empty" id="empty" style="display:none">
  No recipes saved yet.<br>Ask about food in the router and they'll appear here automatically.
</div>
<div class="overlay" id="overlay">
  <div class="modal">
    <div class="modal-head">
      <div class="modal-title" id="modalTitle"></div>
      <button class="modal-close" id="modalClose">✕</button>
    </div>
    <div class="modal-body">
      <img class="modal-img" id="modalImg" alt="" style="display:none">
      <div class="modal-q" id="modalQ"></div>
      <div class="recipe-text" id="modalBody"></div>
    </div>
  </div>
</div>
<script>
function escapeHtml(s){const d=document.createElement('div');d.textContent=s;return d.innerHTML;}
function fmtDate(ts){return new Date(ts*1000).toLocaleDateString(undefined,{month:'short',day:'numeric',year:'numeric'});}
function titleFromQ(q){
  return q.replace(/^(how (do i|to) (make|cook|bake|prepare|do)|recipe for|give me( a)?( recipe for)?|make me( a)?|can you make( a)?)\s+/i,'')
          .replace(/[?!.]+$/,'').trim()
          .replace(/\b\w/g,c=>c.toUpperCase());
}
// Prefer the dish name captured at save time; fall back for legacy entries.
function recipeTitle(r){ return (r.title && r.title.trim()) || titleFromQ(r.question); }

let recipes=[];

const MEALS=['breakfast','lunch','dinner','dessert','snack','drink','side'];
const PROTS=['chicken','beef','pork','seafood','turkey','lamb','vegetarian','other'];

function makeCard(r){
  const title=recipeTitle(r);
  const preview=r.answer.split('\n').find(l=>l.trim()&&!l.trim().startsWith('#'))||r.answer.slice(0,100);
  const card=document.createElement('div');
  card.className='card';
  const photo=r.photo?`<img class="card-img" loading="lazy" alt="" src="/cookbook/photo/${encodeURIComponent(r.photo)}">`:'';
  const tags=(r.meal||r.protein)?`<div class="card-tags">${escapeHtml([r.meal,r.protein].filter(Boolean).join(' · '))}</div>`:'';
  card.innerHTML=`${photo}<div class="card-head">
    <div class="card-title">${escapeHtml(title)}</div>
    <div class="card-date">${fmtDate(r.timestamp)}</div>
    ${tags}
  </div>
  <div class="card-preview">${escapeHtml(preview)}</div>`;
  if(r.id){
    const del=document.createElement('button');
    del.className='card-del'; del.title='Delete recipe'; del.textContent='✕';
    del.addEventListener('click',async e=>{
      e.stopPropagation();
      if(!confirm(`Delete "${title}"? This also removes its photo.`))return;
      await fetch('/cookbook/delete/'+r.id,{method:'POST'});
      recipes=recipes.filter(x=>x.id!==r.id);
      applyFilters();
    });
    card.appendChild(del);
  }
  card.addEventListener('click',()=>openModal(r));
  return card;
}

function render(list){
  const grid=document.getElementById('grid'), empty=document.getElementById('empty');
  grid.innerHTML='';
  if(!list.length){empty.style.display='block';return;}
  empty.style.display='none';
  const mode=document.getElementById('groupSel').value;
  if(mode==='none'){
    list.forEach(r=>grid.appendChild(makeCard(r)));
    return;
  }
  const keys=mode==='meal'?MEALS:PROTS;
  const other=mode==='meal'?'protein':'meal';
  const grouped={}, uncat=[];
  list.forEach(r=>{
    const k=r[mode];
    if(k&&keys.includes(k)) (grouped[k]=grouped[k]||[]).push(r);
    else uncat.push(r);
  });
  keys.forEach(k=>{
    const items=grouped[k];
    if(!items||!items.length)return;
    items.sort((a,b)=>(a[other]||'zz').localeCompare(b[other]||'zz')
                      ||recipeTitle(a).localeCompare(recipeTitle(b)));
    const h=document.createElement('div');
    h.className='grp-head';
    h.textContent=`${k} (${items.length})`;
    grid.appendChild(h);
    items.forEach(r=>grid.appendChild(makeCard(r)));
  });
  if(uncat.length){
    const h=document.createElement('div');
    h.className='grp-head';
    h.textContent=`still categorizing… (${uncat.length})`;
    grid.appendChild(h);
    uncat.forEach(r=>grid.appendChild(makeCard(r)));
  }
}

function applyFilters(){
  const s=document.getElementById('search').value.toLowerCase();
  const meal=document.getElementById('mealSel').value;
  const prot=document.getElementById('protSel').value;
  let list=recipes;
  if(s)list=list.filter(r=>r.question.toLowerCase().includes(s)||r.answer.toLowerCase().includes(s));
  if(meal)list=list.filter(r=>r.meal===meal);
  if(prot)list=list.filter(r=>r.protein===prot);
  render(list);
}

function openModal(r){
  document.getElementById('modalTitle').textContent=recipeTitle(r);
  document.getElementById('modalQ').textContent='Asked: '+r.question;
  document.getElementById('modalBody').textContent=r.answer;
  const mi=document.getElementById('modalImg');
  if(r.photo){mi.src='/cookbook/photo/'+encodeURIComponent(r.photo);mi.style.display='block';}
  else mi.style.display='none';
  document.getElementById('overlay').classList.add('open');
}
document.getElementById('modalClose').addEventListener('click',()=>{
  document.getElementById('overlay').classList.remove('open');
});
document.getElementById('overlay').addEventListener('click',e=>{
  if(e.target===document.getElementById('overlay'))
    document.getElementById('overlay').classList.remove('open');
});
document.getElementById('search').addEventListener('input',applyFilters);
document.getElementById('groupSel').addEventListener('change',applyFilters);
document.getElementById('mealSel').addEventListener('change',applyFilters);
document.getElementById('protSel').addEventListener('change',applyFilters);


fetch('/cookbook/recipes').then(r=>r.json()).then(data=>{recipes=data;applyFilters();});
// refresh once the background classifier has had a chance to tag more cards
setTimeout(()=>{fetch('/cookbook/recipes').then(r=>r.json()).then(data=>{recipes=data;applyFilters();});},45000);
const _tsw=document.getElementById('themeSw');
_tsw.checked=localStorage.getItem('creweTheme')==='dark';
_tsw.addEventListener('change',()=>{
  if(_tsw.checked){document.documentElement.dataset.theme='dark';localStorage.setItem('creweTheme','dark');}
  else{delete document.documentElement.dataset.theme;localStorage.setItem('creweTheme','light');}
});
</script>
</body>
</html>"""


DOCS_PAGE = r"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>Docs</title>
<link rel="icon" href="data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 64 64'%3E%3Crect width='64' height='64' rx='14' fill='%238b4a18'/%3E%3Cpath d='M32 57V34M32 34 16 13M32 34 48 13' stroke='%23faf6f0' stroke-width='7' stroke-linecap='round' fill='none'/%3E%3Ccircle cx='32' cy='34' r='6' fill='%23faf6f0'/%3E%3C/svg%3E">
<script src="/static/marked.min.js"></script>
<link href="/static/fonts.css" rel="stylesheet">
<script>if(localStorage.getItem('creweTheme')==='dark')document.documentElement.dataset.theme='dark';</script>
<style>
  :root { --bg:#faf6f0; --panel:#ffffff; --panel2:#f2ebe0; --ink:#1e160a;
          --muted:#8a7560; --border:#ddd0b8; --accent:#2a4e8c;
          --gold:#8b4a18; --btn-ink:#fff; }
  :root[data-theme="dark"] {
    --bg:#191410; --panel:#221c15; --panel2:#14100d; --ink:#ece2d3;
    --muted:#9c8a75; --border:#3a3125; --accent:#8aaee8;
    --gold:#d29a62; --btn-ink:#241b11; color-scheme:dark; }
  .theme-toggle { display:inline-flex; align-items:center; gap:5px;
                  cursor:pointer; user-select:none; flex:0 0 auto; }
  .theme-toggle input { position:absolute; opacity:0; pointer-events:none; }
  .tt-track { width:34px; height:18px; border-radius:999px; background:var(--bg);
              border:1px solid var(--border); position:relative; }
  .tt-knob { position:absolute; top:1px; left:1px; width:14px; height:14px;
             border-radius:50%; background:var(--gold); transition:left .2s; }
  .theme-toggle input:checked + .tt-track .tt-knob { left:17px; }
  .tt-icon { font-size:11px; color:var(--muted); }
  header h1 { font-family:'Lora',Georgia,serif; color:var(--gold); }
  * { box-sizing:border-box; }
  html,body { margin:0; height:100%; }
  body { font:15px/1.5 'DM Sans',-apple-system,Segoe UI,Roboto,sans-serif;
         background:var(--bg); color:var(--ink); display:flex;
         flex-direction:column; }
  header { padding:12px 24px; border-bottom:1px solid var(--border);
           display:flex; align-items:center; gap:14px; flex:0 0 auto; }
  .back { color:var(--accent); text-decoration:none; font-size:13px; }
  header h1 { margin:0; font-size:18px; }
  #savedTag { color:var(--muted); font-size:12px; margin-left:auto; }
  /* ---- list view ---- */
  #listView { flex:1; overflow-y:auto; }
  .grid { display:grid; grid-template-columns:repeat(auto-fill,minmax(230px,1fr));
          gap:16px; padding:24px; }
  .doc-card { background:var(--panel); border:1px solid var(--border);
              border-radius:12px; padding:16px; cursor:pointer;
              transition:border-color .15s; position:relative; min-height:120px; }
  .doc-card:hover { border-color:var(--gold); }
  .doc-card.new { display:flex; align-items:center; justify-content:center;
                  font-size:15px; color:var(--muted); border-style:dashed; }
  .doc-card.new:hover { color:var(--gold); }
  .doc-title { font-weight:600; font-size:14px; margin-bottom:6px;
               padding-right:20px; }
  .doc-date { font-size:11px; color:var(--muted); margin-bottom:8px; }
  .doc-snip { font-size:12px; color:var(--muted); line-height:1.5;
              display:-webkit-box; -webkit-line-clamp:3;
              -webkit-box-orient:vertical; overflow:hidden; }
  .doc-del { position:absolute; top:10px; right:10px; background:none;
             border:0; color:var(--muted); cursor:pointer; font-size:14px;
             display:none; }
  .doc-card:hover .doc-del { display:block; }
  .doc-del:hover { color:#e06552; }
  /* ---- editor view ---- */
  #editView { flex:1; display:none; flex-direction:column; min-height:0; }
  .ebar { display:flex; align-items:center; gap:12px; padding:10px 24px;
          border-bottom:1px solid var(--border); flex:0 0 auto; }
  #docTitle { background:none; border:1px solid transparent; color:var(--ink);
              font:600 16px inherit; padding:5px 10px; border-radius:7px;
              min-width:280px; outline:none; }
  #docTitle:hover,#docTitle:focus { border-color:var(--border);
                                    background:var(--panel); }
  .toolbar { display:flex; gap:2px; padding:7px 24px; flex-wrap:wrap;
             border-bottom:1px solid var(--border); flex:0 0 auto;
             background:var(--panel); }
  .tb { background:none; border:0; color:var(--ink); border-radius:6px;
        min-width:32px; height:30px; font:14px inherit; cursor:pointer;
        padding:0 8px; }
  .tb:hover { background:var(--border); }
  .tb.sep { width:1px; min-width:1px; padding:0; background:var(--border);
            margin:3px 7px; cursor:default; }
  .tsel { background:var(--bg); color:var(--ink); border:1px solid var(--border);
          border-radius:6px; height:30px; font:13px inherit; padding:0 6px;
          outline:none; cursor:pointer; }
  .tb input[type=color] { width:22px; height:18px; border:0; padding:0;
                          background:none; cursor:pointer; vertical-align:middle; }
  #pageWrap { flex:1; overflow-y:auto; padding:34px 16px 120px;
              background:var(--panel2); }
  #page { width:min(816px, 100%); min-height:1056px; margin:0 auto;
          background:#ffffff; color:#202124; border-radius:3px;
          box-shadow:0 3px 16px rgba(60,40,10,.25);
          padding:96px 86px; outline:none;
          font:12.5pt/1.7 Georgia,'Times New Roman',serif; }
  #page :is(h1,h2,h3) { font-family:inherit; line-height:1.3; color:#111; }
  #page h1 { font-size:24pt; margin:.7em 0 .35em; }
  #page h2 { font-size:17pt; margin:.8em 0 .3em; }
  #page h3 { font-size:13.5pt; margin:.8em 0 .25em; }
  #page p { margin:0 0 .55em; }
  #page ul,#page ol { margin:0 0 .6em; padding-left:1.7em; }
  #page blockquote { border-left:3px solid #c9a35c; margin:.6em 0;
                     padding-left:1em; color:#555; }
  #page table { border-collapse:collapse; margin:.6em 0; }
  #page th,#page td { border:1px solid #bbb; padding:5px 10px; }
  #page code { font-family:ui-monospace,Menlo,monospace; font-size:.85em;
               background:#f1f1ec; padding:1px 5px; border-radius:3px; }
  #page pre { background:#f6f6f2; border:1px solid #ddd; border-radius:6px;
              padding:12px 14px; overflow-x:auto; }
  .ai-writing { border-left:3px solid var(--gold); padding-left:12px;
                animation:aiPulse 1.6s ease-in-out infinite; }
  .ai-writing::after { content:'▍'; color:#c9861f; animation:blink 1s step-end infinite; }
  @keyframes blink { 50% { opacity:0; } }
  @keyframes aiPulse { 50% { border-left-color:#f3cf8e; } }
  /* ---- ai bar ---- */
  #aiBar { position:fixed; left:0; right:0; bottom:0; background:var(--panel);
           border-top:1px solid var(--border); padding:12px 24px;
           display:none; gap:10px; align-items:center; z-index:50; }
  #aiBar.show { display:flex; }
  #aiQ { flex:1; max-width:820px; background:var(--bg); color:var(--ink);
         border:1px solid var(--border); border-radius:9px; padding:10px 14px;
         font:14px inherit; outline:none; }
  #aiQ:focus { border-color:var(--gold); }
  #aiGo { background:var(--gold); color:var(--btn-ink); border:0; border-radius:9px;
          padding:10px 18px; font:600 13px inherit; cursor:pointer; }
  #aiGo:disabled { opacity:.5; }
  #aiStatus { color:var(--muted); font-size:12px; }
  @media print {
    header,.ebar,.toolbar,#aiBar,#listView { display:none !important; }
    body { background:#fff; }
    #editView { display:block !important; }
    #pageWrap { padding:0; background:#fff; overflow:visible; }
    #page { box-shadow:none; border-radius:0; width:100%; min-height:0;
            padding:0.4in 0.2in; }
  }
</style>
</head>
<body>
<header>
  <a href="/" class="back">← Crewe</a>
  <h1>📄 Docs</h1>
  <span id="savedTag"></span>
  <label class="theme-toggle" title="Dark mode"><span class="tt-icon">☀</span><input type="checkbox" id="themeSw"><span class="tt-track"><span class="tt-knob"></span></span><span class="tt-icon">🌙</span></label>
</header>

<div id="listView"><div class="grid" id="docGrid"></div></div>

<div id="editView">
  <div class="ebar">
    <a href="#" class="back" id="backBtn">← All docs</a>
    <input id="docTitle" placeholder="Untitled document">
    <button class="tb" id="printBtn" title="Print / save as PDF"
            style="border:1px solid var(--border);font-size:12px;">🖨 Print</button>
  </div>
  <div class="toolbar">
    <select class="tsel" id="styleSel" title="Paragraph style">
      <option value="P">Normal text</option>
      <option value="H1">Heading 1</option>
      <option value="H2">Heading 2</option>
      <option value="H3">Heading 3</option>
      <option value="BLOCKQUOTE">Quote</option>
      <option value="PRE">Code block</option>
    </select>
    <select class="tsel" id="sizeSel" title="Text size">
      <option value="3" selected>Normal</option>
      <option value="2">Small</option>
      <option value="5">Large</option>
      <option value="7">Huge</option>
    </select>
    <span class="tb sep"></span>
    <button class="tb" data-cmd="bold" title="Bold (Ctrl+B)"><b>B</b></button>
    <button class="tb" data-cmd="italic" title="Italic (Ctrl+I)"><i>I</i></button>
    <button class="tb" data-cmd="underline" title="Underline (Ctrl+U)"><u>U</u></button>
    <button class="tb" data-cmd="strikeThrough" title="Strikethrough"><s>S</s></button>
    <span class="tb sep"></span>
    <label class="tb" title="Text color">A<input type="color" id="fcDoc" value="#202124"></label>
    <label class="tb" title="Highlight">▨<input type="color" id="hlDoc" value="#fff3bf"></label>
    <span class="tb sep"></span>
    <button class="tb" data-cmd="justifyLeft" title="Align left">⇤</button>
    <button class="tb" data-cmd="justifyCenter" title="Center">≡</button>
    <button class="tb" data-cmd="justifyRight" title="Align right">⇥</button>
    <span class="tb sep"></span>
    <button class="tb" data-cmd="insertUnorderedList" title="Bulleted list">• list</button>
    <button class="tb" data-cmd="insertOrderedList" title="Numbered list">1. list</button>
    <span class="tb sep"></span>
    <button class="tb" data-cmd="removeFormat" title="Clear formatting">Tx</button>
    <button class="tb" data-cmd="undo" title="Undo (Ctrl+Z)">↶</button>
    <button class="tb" data-cmd="redo" title="Redo">↷</button>
  </div>
  <div id="pageWrap"><div id="page" contenteditable="true" spellcheck="true"></div></div>
</div>

<div id="aiBar">
  <input id="aiQ" placeholder="Ask Crewe to write into this doc…  (e.g. 'a help document for basic Linux commands')">
  <button id="aiGo">✍ Write</button>
  <span id="aiStatus"></span>
</div>

<script>
marked.use({ breaks:true, gfm:true });
function escapeHtml(s){const d=document.createElement('div');d.textContent=s;return d.innerHTML;}
marked.use({ renderer:{ html(t){const s=typeof t==='string'?t:(t&&t.text)||'';return escapeHtml(s);} }});
function fmtDate(ts){return new Date(ts*1000).toLocaleString(undefined,{month:'short',day:'numeric',hour:'numeric',minute:'2-digit'});}

let curDoc=null, saveTimer=null, aiBusy=false, aiJob=null;
const page=document.getElementById('page'), titleEl=document.getElementById('docTitle');

// ---- list view --------------------------------------------------------------
async function loadList(){
  const docs=await(await fetch('/docs/list')).json();
  const g=document.getElementById('docGrid'); g.innerHTML='';
  const nc=document.createElement('div');
  nc.className='doc-card new'; nc.textContent='＋  New document';
  nc.addEventListener('click',async()=>{
    const r=await(await fetch('/docs/new',{method:'POST'})).json();
    openDoc(r.id,true);
  });
  g.appendChild(nc);
  docs.forEach(d=>{
    const c=document.createElement('div'); c.className='doc-card';
    c.innerHTML=`<div class="doc-title">${escapeHtml(d.title)}</div>
      <div class="doc-date">${fmtDate(d.updated)}</div>
      <div class="doc-snip">${escapeHtml(d.snippet||'Empty document')}</div>
      <button class="doc-del" title="Delete">✕</button>`;
    c.querySelector('.doc-del').addEventListener('click',async e=>{
      e.stopPropagation();
      if(confirm(`Delete "${d.title}"?`)){
        await fetch('/docs/delete/'+d.id,{method:'POST'}); loadList();
      }
    });
    c.addEventListener('click',()=>openDoc(d.id));
    g.appendChild(c);
  });
}

// ---- editor -----------------------------------------------------------------
async function openDoc(id,isNew){
  const d=await(await fetch('/docs/get/'+id)).json();
  if(d.error) return;
  curDoc=id;
  titleEl.value=d.title==='Untitled document'&&isNew?'':d.title;
  titleEl.placeholder='Untitled document';
  page.innerHTML=d.html||'<p><br></p>';
  document.getElementById('listView').style.display='none';
  document.getElementById('editView').style.display='flex';
  document.getElementById('aiBar').classList.add('show');
  document.getElementById('savedTag').textContent='';
  (isNew?titleEl:page).focus();
}
function closeDoc(){
  saveNow().then(()=>{
    curDoc=null;
    document.getElementById('editView').style.display='none';
    document.getElementById('aiBar').classList.remove('show');
    document.getElementById('listView').style.display='block';
    document.getElementById('savedTag').textContent='';
    loadList();
  });
}
document.getElementById('backBtn').addEventListener('click',e=>{e.preventDefault();closeDoc();});

async function saveNow(){
  if(!curDoc) return;
  document.getElementById('savedTag').textContent='Saving…';
  try{
    await fetch('/docs/save/'+curDoc,{method:'POST',
      headers:{'Content-Type':'application/json'},
      body:JSON.stringify({title:titleEl.value.trim()||'Untitled document',
                           html:page.innerHTML})});
    document.getElementById('savedTag').textContent='Saved ✓';
  }catch(e){ document.getElementById('savedTag').textContent='⚠ not saved'; }
}
function queueSave(){
  document.getElementById('savedTag').textContent='…';
  clearTimeout(saveTimer); saveTimer=setTimeout(saveNow,1200);
}
page.addEventListener('input',queueSave);
titleEl.addEventListener('input',queueSave);
document.getElementById('printBtn').addEventListener('click',async()=>{
  await saveNow(); window.print();
});
window.addEventListener('beforeunload',()=>{ if(curDoc) navigator.sendBeacon&&saveNow(); });

// toolbar — mousedown preventDefault keeps the text selection alive
document.querySelectorAll('.tb[data-cmd],.tb[data-block]').forEach(b=>{
  b.addEventListener('mousedown',e=>e.preventDefault());
  b.addEventListener('click',()=>{
    if(b.dataset.cmd) document.execCommand(b.dataset.cmd,false,null);
    else document.execCommand('formatBlock',false,b.dataset.block);
    page.focus(); queueSave();
  });
});
// selects & color pickers steal focus, so save/restore the page selection
let savedRange=null;
function keepRange(){
  const s=window.getSelection();
  if(s.rangeCount&&page.contains(s.anchorNode))
    savedRange=s.getRangeAt(0).cloneRange();
}
function restoreRange(){
  if(!savedRange)return;
  const s=window.getSelection(); s.removeAllRanges(); s.addRange(savedRange);
}
page.addEventListener('keyup',keepRange);
page.addEventListener('mouseup',keepRange);
function execWithRange(cmd,val){
  page.focus(); restoreRange();
  document.execCommand(cmd,false,val);
  keepRange(); queueSave();
}
document.getElementById('styleSel').addEventListener('change',e=>{
  execWithRange('formatBlock',e.target.value);
});
document.getElementById('sizeSel').addEventListener('change',e=>{
  execWithRange('fontSize',e.target.value);
});
document.getElementById('fcDoc').addEventListener('input',e=>{
  execWithRange('foreColor',e.target.value);
});
document.getElementById('hlDoc').addEventListener('input',e=>{
  execWithRange('hiliteColor',e.target.value);
});
// keep the style dropdown in sync with wherever the caret is
document.addEventListener('selectionchange',()=>{
  const s=window.getSelection();
  if(!s.rangeCount||!page.contains(s.anchorNode))return;
  keepRange();
  let n=s.anchorNode;
  while(n&&n!==page&&!(n.tagName&&/^(H1|H2|H3|BLOCKQUOTE|PRE)$/.test(n.tagName)))
    n=n.parentNode;
  document.getElementById('styleSel').value=
    (n&&n!==page&&n.tagName)?n.tagName:'P';
});

// ---- AI writes into the doc -------------------------------------------------
async function aiWrite(){
  const q=document.getElementById('aiQ'), go=document.getElementById('aiGo'),
        st=document.getElementById('aiStatus');
  const prompt=q.value.trim();
  if(!prompt||aiBusy||!curDoc) return;
  aiBusy=true; go.disabled=true; st.textContent='starting…';
  // insertion point: caret inside the page if there is one, else end of doc
  const block=document.createElement('div'); block.className='ai-writing';
  const sel=window.getSelection();
  let placed=false;
  if(sel.rangeCount&&page.contains(sel.anchorNode)){
    try{ const r=sel.getRangeAt(0); r.collapse(false); r.insertNode(block); placed=true; }
    catch(e){}
  }
  if(!placed) page.appendChild(block);
  block.scrollIntoView({block:'center',behavior:'smooth'});
  let job;
  try{
    job=(await(await fetch('/docs/write',{method:'POST',
      headers:{'Content-Type':'application/json'},
      body:JSON.stringify({prompt,title:titleEl.value.trim()})})).json()).job_id;
  }catch(e){ block.remove(); st.textContent='failed to start'; aiBusy=false; go.disabled=false; return; }
  aiJob=job; go.disabled=false; go.textContent='⏹ Stop';
  q.value='';
  const poll=setInterval(async()=>{
    let p; try{p=await(await fetch('/progress/'+job)).json();}catch(e){return;}
    if(p.partial) block.innerHTML=marked.parse(p.partial);
    st.textContent=p.done?'':'writing… '+(p.tokens||0)+' tokens';
    if(p.done){
      clearInterval(poll);
      block.innerHTML=marked.parse(p.answer||p.partial||'');
      block.classList.remove('ai-writing');
      if(!titleEl.value.trim()) titleEl.value=prompt.slice(0,60);
      st.textContent='done ✓'; setTimeout(()=>{st.textContent='';},3000);
      aiBusy=false; aiJob=null; go.disabled=false; go.textContent='✍ Write';
      queueSave();
    }
  },700);
}
document.getElementById('aiGo').addEventListener('click',()=>{
  if(aiBusy&&aiJob){ fetch('/stop/'+aiJob,{method:'POST'}).catch(()=>{}); return; }
  aiWrite();
});
document.getElementById('aiQ').addEventListener('keydown',e=>{
  if(e.key==='Enter'){e.preventDefault();aiWrite();}
});

loadList();
const _tsw=document.getElementById('themeSw');
_tsw.checked=localStorage.getItem('creweTheme')==='dark';
_tsw.addEventListener('change',()=>{
  if(_tsw.checked){document.documentElement.dataset.theme='dark';localStorage.setItem('creweTheme','dark');}
  else{delete document.documentElement.dataset.theme;localStorage.setItem('creweTheme','light');}
});
</script>
</body>
</html>"""


SHEETS_PAGE = r"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>Sheets</title>
<link rel="icon" href="data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 64 64'%3E%3Crect width='64' height='64' rx='14' fill='%238b4a18'/%3E%3Cpath d='M32 57V34M32 34 16 13M32 34 48 13' stroke='%23faf6f0' stroke-width='7' stroke-linecap='round' fill='none'/%3E%3Ccircle cx='32' cy='34' r='6' fill='%23faf6f0'/%3E%3C/svg%3E">
<link href="/static/fonts.css" rel="stylesheet">
<script>if(localStorage.getItem('creweTheme')==='dark')document.documentElement.dataset.theme='dark';</script>
<style>
  :root { --bg:#faf6f0; --panel:#ffffff; --panel2:#f2ebe0; --ink:#1e160a;
          --muted:#8a7560; --border:#ddd0b8; --accent:#2a4e8c;
          --gold:#8b4a18; --btn-ink:#fff; --grid:#e4dac4; --cellbg:#fffdf8;
          --insel:#f3ead7; --selbg:#f8f0dd; --aiflash:#f3dfae;
          --grn:#2e6e2a; --down:#c0392b; }
  :root[data-theme="dark"] {
    --bg:#191410; --panel:#221c15; --panel2:#2b241b; --ink:#ece2d3;
    --muted:#9c8a75; --border:#3a3125; --accent:#8aaee8; --gold:#d29a62;
    --btn-ink:#241b11; --grid:#33291d; --cellbg:#1d1712; --insel:#2b241b;
    --selbg:#241d15; --aiflash:#3a2f16; --grn:#7cbf72; --down:#e06552;
    color-scheme:dark; }
  .theme-toggle { display:inline-flex; align-items:center; gap:5px;
                  cursor:pointer; user-select:none; flex:0 0 auto; }
  .theme-toggle input { position:absolute; opacity:0; pointer-events:none; }
  .tt-track { width:34px; height:18px; border-radius:999px; background:var(--bg);
              border:1px solid var(--border); position:relative; }
  .tt-knob { position:absolute; top:1px; left:1px; width:14px; height:14px;
             border-radius:50%; background:var(--gold); transition:left .2s; }
  .theme-toggle input:checked + .tt-track .tt-knob { left:17px; }
  .tt-icon { font-size:11px; color:var(--muted); }
  header h1 { font-family:'Lora',Georgia,serif; color:var(--gold); }
  * { box-sizing:border-box; }
  html,body { margin:0; height:100%; }
  body { font:14px/1.4 'DM Sans',-apple-system,Segoe UI,Roboto,sans-serif;
         background:var(--bg); color:var(--ink); display:flex;
         flex-direction:column; }
  header { padding:12px 24px; border-bottom:1px solid var(--border);
           display:flex; align-items:center; gap:14px; flex:0 0 auto; }
  .back { color:var(--accent); text-decoration:none; font-size:13px; }
  header h1 { margin:0; font-size:18px; }
  #savedTag { color:var(--muted); font-size:12px; margin-left:auto; }
  #listView { flex:1; overflow-y:auto; }
  .grid-cards { display:grid;
                grid-template-columns:repeat(auto-fill,minmax(230px,1fr));
                gap:16px; padding:24px; }
  .sh-card { background:var(--panel); border:1px solid var(--border);
             border-radius:12px; padding:16px; cursor:pointer;
             position:relative; min-height:96px; transition:border-color .15s; }
  .sh-card:hover { border-color:var(--grn); }
  .sh-card.new { display:flex; align-items:center; justify-content:center;
                 color:var(--muted); border-style:dashed; }
  .sh-card.new:hover { color:var(--grn); }
  .sh-title { font-weight:600; font-size:14px; margin-bottom:6px; padding-right:20px; }
  .sh-meta { font-size:11px; color:var(--muted); }
  .sh-del { position:absolute; top:10px; right:10px; background:none; border:0;
            color:var(--muted); cursor:pointer; display:none; font-size:14px; }
  .sh-card:hover .sh-del { display:block; }
  .sh-del:hover { color:#e06552; }
  /* editor */
  #editView { flex:1; display:none; flex-direction:column; min-height:0; }
  .ebar { display:flex; align-items:center; gap:12px; padding:9px 24px;
          border-bottom:1px solid var(--border); flex:0 0 auto; }
  #shTitle { background:none; border:1px solid transparent; color:var(--ink);
             font:600 15px inherit; padding:5px 10px; border-radius:7px;
             min-width:240px; outline:none; }
  #shTitle:hover,#shTitle:focus { border-color:var(--border); background:var(--panel); }
  .ebtn { background:var(--panel); border:1px solid var(--border); color:var(--muted);
          border-radius:7px; padding:6px 12px; font:12px inherit; cursor:pointer; }
  .ebtn:hover { color:var(--ink); border-color:var(--muted); }
  .fbar { display:flex; align-items:center; gap:10px; padding:7px 24px;
          border-bottom:1px solid var(--border); background:var(--panel);
          flex:0 0 auto; }
  #selRef { font:600 12px ui-monospace,monospace; color:var(--gold);
            min-width:44px; text-align:center; background:var(--bg);
            border:1px solid var(--border); border-radius:6px; padding:6px 4px; }
  #fx { flex:1; background:var(--bg); border:1px solid var(--border);
        color:var(--ink); border-radius:6px; padding:6px 12px;
        font:13px ui-monospace,monospace; outline:none; }
  #fx:focus { border-color:var(--gold); }
  #gridWrap { flex:1; overflow:auto; padding-bottom:90px; }
  table.sheet { border-collapse:collapse; user-select:none; }
  .sheet th { background:var(--panel2); color:var(--muted); font:600 11px inherit;
              border:1px solid var(--grid); min-width:92px; max-width:92px;
              padding:4px 6px; position:sticky; top:0; z-index:3; }
  .sheet th.rn { min-width:44px; max-width:44px; left:0; z-index:4; }
  .sheet td.rn { background:var(--panel2); color:var(--muted); font:11px inherit;
                 text-align:center; position:sticky; left:0; z-index:2;
                 border:1px solid var(--grid); min-width:44px; max-width:44px; }
  .sheet td.cell { border:1px solid var(--grid); min-width:92px; max-width:92px;
                   padding:4px 7px; font:13px inherit; height:26px;
                   white-space:nowrap; overflow:hidden; cursor:cell;
                   background:var(--cellbg); }
  .sheet td.cell.num { text-align:right; font-variant-numeric:tabular-nums; }
  .sheet td.cell.err { color:var(--down); font-size:11px; }
  .sheet td.cell.sel { outline:2px solid var(--gold); outline-offset:-2px;
                       background:var(--selbg); }
  .sheet td.cell.aiflash { background:var(--aiflash); transition:background .6s; }
  .sheet td.cell.insel { background:var(--insel); }
  .sheet th { position:relative; }
  .colgrip { position:absolute; top:0; right:-3px; width:7px; height:100%;
             cursor:col-resize; z-index:5; }
  .stb { display:flex; gap:3px; align-items:center; padding:6px 24px;
         border-bottom:1px solid var(--border); background:var(--panel);
         flex:0 0 auto; flex-wrap:wrap; }
  .tbx { background:none; border:0; color:var(--ink); border-radius:6px;
         min-width:30px; height:28px; font:13px inherit; cursor:pointer;
         padding:0 8px; }
  .tbx:hover { background:var(--border); }
  .tbx.sep { width:1px; min-width:1px; padding:0; background:var(--border);
             margin:2px 7px; cursor:default; }
  .tbx input[type=color] { width:22px; height:18px; border:0; padding:0;
                           background:none; cursor:pointer; vertical-align:middle; }
  #selStats { margin-left:auto; color:var(--muted);
              font:12px ui-monospace,monospace; }
  #cellEd { position:absolute; display:none; z-index:10;
            border:2px solid var(--gold); background:var(--cellbg); color:var(--ink);
            font:13px inherit; padding:3px 6px; outline:none; }
  #aiBar { position:fixed; left:0; right:0; bottom:0; background:var(--panel);
           border-top:1px solid var(--border); padding:12px 24px; display:none;
           gap:10px; align-items:center; z-index:50; }
  #aiBar.show { display:flex; }
  #aiQ { flex:1; max-width:820px; background:var(--bg); color:var(--ink);
         border:1px solid var(--border); border-radius:9px; padding:10px 14px;
         font:14px inherit; outline:none; }
  #aiQ:focus { border-color:var(--grn); }
  #aiGo { background:var(--grn); color:var(--btn-ink); border:0; border-radius:9px;
          padding:10px 18px; font:600 13px inherit; cursor:pointer; }
  #aiGo:disabled { opacity:.5; }
  #aiStatus { color:var(--muted); font-size:12px; }
</style>
</head>
<body>
<header>
  <a href="/" class="back">← Crewe</a>
  <h1>🧮 Sheets</h1>
  <span id="savedTag"></span>
  <label class="theme-toggle" title="Dark mode"><span class="tt-icon">☀</span><input type="checkbox" id="themeSw"><span class="tt-track"><span class="tt-knob"></span></span><span class="tt-icon">🌙</span></label>
</header>

<div id="listView"><div class="grid-cards" id="shGrid"></div></div>

<div id="editView">
  <div class="ebar">
    <a href="#" class="back" id="backBtn">← All sheets</a>
    <input id="shTitle" placeholder="Untitled sheet">
    <button class="ebtn" id="csvBtn">⬇ CSV</button>
    <button class="ebtn" id="moreRows">+20 rows</button>
  </div>
  <div class="fbar">
    <span id="selRef">A1</span>
    <input id="fx" placeholder="Type a value or formula…  =SUM(B2:B9)  =ROUND(C3*1.07, 2)">
    <span id="selStats"></span>
  </div>
  <div class="stb">
    <button class="tbx" data-f="b" title="Bold"><b>B</b></button>
    <button class="tbx" data-f="i" title="Italic"><i>I</i></button>
    <span class="tbx sep"></span>
    <button class="tbx" data-al="l" title="Align left">⇤</button>
    <button class="tbx" data-al="c" title="Center">≡</button>
    <button class="tbx" data-al="r" title="Align right">⇥</button>
    <span class="tbx sep"></span>
    <label class="tbx" title="Text color">A<input type="color" id="fcPick" value="#e6e8ec"></label>
    <label class="tbx" title="Fill color">▧<input type="color" id="bgPick" value="#3a2f16"></label>
    <span class="tbx sep"></span>
    <button class="tbx" data-nf="cur" title="Currency">$</button>
    <button class="tbx" data-nf="pct" title="Percent">%</button>
    <button class="tbx" data-nf="2dp" title="Two decimals">.00</button>
    <span class="tbx sep"></span>
    <button class="tbx" id="clearFmt" title="Clear formatting">Tx</button>
  </div>
  <div id="gridWrap"></div>
  <input id="cellEd" autocomplete="off">
</div>

<div id="aiBar">
  <input id="aiQ" placeholder="Ask Crewe to fill this sheet…  (e.g. 'a monthly household budget with totals')">
  <button id="aiGo">▦ Fill</button>
  <span id="aiStatus"></span>
</div>

<script>
'use strict';
function escapeHtml(s){const d=document.createElement('div');d.textContent=s;return d.innerHTML;}
function fmtDate(ts){return new Date(ts*1000).toLocaleString(undefined,{month:'short',day:'numeric',hour:'numeric',minute:'2-digit'});}
function colName(i){let s='';i++;while(i>0){i--;s=String.fromCharCode(65+i%26)+s;i=Math.floor(i/26);}return s;}
function refOf(c,r){return colName(c)+(r+1);}
function parseRef(ref){const m=/^([A-Z]{1,2})([0-9]{1,3})$/.exec(ref);if(!m)return null;
  let c=0;for(const ch of m[1])c=c*26+(ch.charCodeAt(0)-64);return {c:c-1,r:+m[2]-1};}

// ================= formula engine =================
const FUNCS={
  SUM:a=>a.reduce((x,y)=>x+y,0),
  AVERAGE:a=>a.length?a.reduce((x,y)=>x+y,0)/a.length:0,
  AVG:a=>FUNCS.AVERAGE(a),
  COUNT:a=>a.length,
  MIN:a=>a.length?Math.min.apply(null,a):0,
  MAX:a=>a.length?Math.max.apply(null,a):0,
  ROUND:a=>{const n=a[1]===undefined?0:a[1];const f=Math.pow(10,n);return Math.round(a[0]*f)/f;},
  ABS:a=>Math.abs(a[0]||0),
};
function tokenize(src){
  const out=[]; let i=0;
  const re=/^(?:([A-Z]{1,2}[0-9]{1,3})|([A-Z]+)(?=\()|([0-9]*\.?[0-9]+)|(:|,|\(|\)|\+|-|\*|\/|\^|%)|(\s+))/;
  src=src.toUpperCase();
  while(i<src.length){
    const m=re.exec(src.slice(i));
    if(!m) throw '#ERR';
    if(m[1])out.push({t:'ref',v:m[1]});
    else if(m[2])out.push({t:'fn',v:m[2]});
    else if(m[3])out.push({t:'num',v:parseFloat(m[3])});
    else if(m[4])out.push({t:'op',v:m[4]});
    i+=m[0].length;
  }
  return out;
}
function evalFormula(src, getRef){
  const toks=tokenize(src); let p=0;
  const peek=()=>toks[p], eat=()=>toks[p++];
  function expr(){
    let v=term();
    while(peek()&&peek().t==='op'&&(peek().v==='+'||peek().v==='-')){
      const o=eat().v, r=term(); v=o==='+'?v+r:v-r;
    }
    return v;
  }
  function term(){
    let v=factor();
    while(peek()&&peek().t==='op'&&(peek().v==='*'||peek().v==='/')){
      const o=eat().v, r=factor();
      if(o==='/'){ if(r===0) throw '#DIV0'; v=v/r; } else v=v*r;
    }
    return v;
  }
  function factor(){
    let neg=false;
    while(peek()&&peek().t==='op'&&(peek().v==='-'||peek().v==='+')){
      if(eat().v==='-')neg=!neg;
    }
    let v=atom();
    if(peek()&&peek().t==='op'&&peek().v==='^'){ eat(); v=Math.pow(v,factor()); }
    if(peek()&&peek().t==='op'&&peek().v==='%'){ eat(); v=v/100; }
    return neg?-v:v;
  }
  function rangeVals(a,b){
    const A=parseRef(a),B=parseRef(b); if(!A||!B)throw '#REF';
    const vals=[];
    for(let r=Math.min(A.r,B.r);r<=Math.max(A.r,B.r);r++)
      for(let c=Math.min(A.c,B.c);c<=Math.max(A.c,B.c);c++){
        const v=getRef(refOf(c,r),true);
        if(typeof v==='number')vals.push(v);
      }
    return vals;
  }
  function args(){
    const vals=[];
    for(;;){
      if(peek()&&peek().t==='ref'&&toks[p+1]&&toks[p+1].t==='op'&&toks[p+1].v===':'){
        const a=eat().v; eat(); const bTok=eat();
        if(!bTok||bTok.t!=='ref')throw '#REF';
        vals.push.apply(vals,rangeVals(a,bTok.v));
      } else vals.push(expr());
      if(peek()&&peek().t==='op'&&peek().v===','){eat();continue;}
      break;
    }
    return vals;
  }
  function atom(){
    const t=eat(); if(!t)throw '#ERR';
    if(t.t==='num')return t.v;
    if(t.t==='ref'){const v=getRef(t.v,false);return typeof v==='number'?v:(v===''?0:(()=>{throw '#VAL';})());}
    if(t.t==='fn'){
      const fn=FUNCS[t.v]; if(!fn)throw '#NAME';
      const o=eat(); if(!o||o.v!=='(')throw '#ERR';
      let a=[];
      if(peek()&&!(peek().t==='op'&&peek().v===')')) a=args();
      const c=eat(); if(!c||c.v!==')')throw '#ERR';
      return fn(a);
    }
    if(t.t==='op'&&t.v==='('){const v=expr();const c=eat();if(!c||c.v!==')')throw '#ERR';return v;}
    throw '#ERR';
  }
  const v=expr();
  if(p<toks.length)throw '#ERR';
  return v;
}
// computed-value resolver with cycle detection + memo per recompute pass
let cells={}, compCache={};
function valueOf(ref, forRange, stack){
  stack=stack||new Set();
  if(ref in compCache)return compCache[ref];
  const raw=cells[ref];
  if(raw===undefined||raw==='')return '';
  if(raw[0]!=='='){
    const n=parseFloat(raw);
    const v=(isFinite(n)&&/^\s*-?[0-9.,]+\s*$/.test(raw.replace(/,/g,'')))?parseFloat(raw.replace(/,/g,'')):raw;
    compCache[ref]=v; return v;
  }
  if(stack.has(ref)){compCache[ref]='#CIRC';return '#CIRC';}
  stack.add(ref);
  let v;
  try{
    v=evalFormula(raw.slice(1),(r,inRange)=>{
      const x=valueOf(r,inRange,stack);
      if(typeof x==='string'&&x.startsWith('#'))throw x;
      return x;
    });
    if(typeof v==='number'&&!isFinite(v))v='#DIV0';
    if(typeof v==='number')v=Math.round(v*1e10)/1e10;
  }catch(e){ v=typeof e==='string'&&e.startsWith('#')?e:'#ERR'; }
  stack.delete(ref);
  compCache[ref]=v; return v;
}

// ================= grid =================
let COLS=26, ROWS=60, sel='A1', selEnd='A1', curSheet=null, saveTimer=null,
    aiBusy=false, aiJob=null, fmt={}, colw={}, draggingSel=false;
const wrap=document.getElementById('gridWrap'), ed=document.getElementById('cellEd'),
      fx=document.getElementById('fx'), titleEl=document.getElementById('shTitle');

function buildGrid(){
  let h='<table class="sheet" style="table-layout:fixed"><colgroup><col style="width:44px">';
  for(let c=0;c<COLS;c++)h+=`<col id="col-${colName(c)}" style="width:${colw[colName(c)]||92}px">`;
  h+='</colgroup><tr><th class="rn"></th>';
  for(let c=0;c<COLS;c++)h+=`<th>${colName(c)}<div class="colgrip" data-col="${colName(c)}"></div></th>`;
  h+='</tr>';
  for(let r=0;r<ROWS;r++){
    h+=`<tr><td class="rn">${r+1}</td>`;
    for(let c=0;c<COLS;c++)h+=`<td class="cell" id="c-${refOf(c,r)}" data-ref="${refOf(c,r)}"></td>`;
    h+='</tr>';
  }
  wrap.innerHTML=h+'</table>';
  wrap.querySelectorAll('td.cell').forEach(td=>{
    td.addEventListener('mousedown',e=>{
      if(e.button!==0)return;
      e.preventDefault();commitEdit();
      if(e.shiftKey){selEnd=td.dataset.ref;paint();}
      else{draggingSel=true;selectRange(td.dataset.ref,td.dataset.ref);}
    });
    td.addEventListener('mouseover',()=>{if(draggingSel){selEnd=td.dataset.ref;paint();}});
    td.addEventListener('dblclick',()=>startEdit(td.dataset.ref));
  });
  wrap.querySelectorAll('.colgrip').forEach(g=>{
    g.addEventListener('mousedown',e=>{
      e.preventDefault();e.stopPropagation();
      const col=g.dataset.col,startX=e.clientX,startW=colw[col]||92;
      function mv(ev){colw[col]=Math.max(40,Math.min(500,startW+ev.clientX-startX));
        const el=document.getElementById('col-'+col);if(el)el.style.width=colw[col]+'px';}
      function up(){document.removeEventListener('mousemove',mv);
        document.removeEventListener('mouseup',up);queueSave();}
      document.addEventListener('mousemove',mv);
      document.addEventListener('mouseup',up);
    });
  });
}
document.addEventListener('mouseup',()=>{draggingSel=false;});

function rangeRefs(){
  const A=parseRef(sel),B=parseRef(selEnd);
  if(!A||!B)return [sel];
  const out=[];
  for(let r=Math.min(A.r,B.r);r<=Math.max(A.r,B.r);r++)
    for(let c=Math.min(A.c,B.c);c<=Math.max(A.c,B.c);c++)out.push(refOf(c,r));
  return out;
}
function fmtVal(v,f){
  if(typeof v!=='number')return v===''?'':String(v);
  const nf=f&&f.nf;
  if(nf==='cur')return '$'+v.toLocaleString(undefined,{minimumFractionDigits:2,maximumFractionDigits:2});
  if(nf==='pct')return (v*100).toLocaleString(undefined,{maximumFractionDigits:2})+'%';
  if(nf==='2dp')return v.toLocaleString(undefined,{minimumFractionDigits:2,maximumFractionDigits:2});
  return String(v);
}
function recompute(){
  compCache={};
  wrap.querySelectorAll('td.cell').forEach(td=>{
    const ref=td.dataset.ref, v=valueOf(ref), f=fmt[ref];
    td.textContent=fmtVal(v,f);
    td.classList.toggle('num',typeof v==='number');
    td.classList.toggle('err',typeof v==='string'&&v.startsWith('#'));
    td.style.fontWeight=f&&f.b?'700':'';
    td.style.fontStyle=f&&f.i?'italic':'';
    td.style.color=f&&f.fc?f.fc:'';
    td.style.backgroundColor=f&&f.bg?f.bg:'';
    td.style.textAlign=f&&f.al?({l:'left',c:'center',r:'right'}[f.al]):'';
  });
}
function fmtNum(v){return String(Math.round(v*100)/100);}
function paint(){
  wrap.querySelectorAll('td.cell.sel,td.cell.insel')
      .forEach(td=>td.classList.remove('sel','insel'));
  const refs=rangeRefs();
  refs.forEach(r=>{const td=document.getElementById('c-'+r);
                   if(td)td.classList.add('insel');});
  const a=document.getElementById('c-'+sel);
  if(a){a.classList.add('sel');a.classList.remove('insel');}
  document.getElementById('selRef').textContent=
    refs.length>1?sel+':'+selEnd:sel;
  fx.value=cells[sel]||'';
  const st=document.getElementById('selStats');
  if(refs.length>1){
    const nums=refs.map(r=>valueOf(r)).filter(v=>typeof v==='number');
    st.textContent=nums.length
      ?`Sum ${fmtNum(nums.reduce((x,y)=>x+y,0))} · Avg ${fmtNum(nums.reduce((x,y)=>x+y,0)/nums.length)} · n ${nums.length}`
      :refs.length+' cells';
  } else st.textContent='';
}
function selectRange(a,b){
  sel=a; selEnd=b;
  const td=document.getElementById('c-'+a);
  if(td)td.scrollIntoView({block:'nearest',inline:'nearest'});
  paint();
}
function select(ref){selectRange(ref,ref);}
function startEdit(ref,replaceWith){
  select(ref);
  const td=document.getElementById('c-'+ref);
  const rect=td.getBoundingClientRect();
  ed.style.left=(rect.left+window.scrollX)+'px';
  ed.style.top=(rect.top+window.scrollY)+'px';
  ed.style.width=rect.width+'px'; ed.style.height=rect.height+'px';
  ed.style.display='block';
  ed.value=replaceWith!==undefined?replaceWith:(cells[ref]||'');
  ed.focus();
  if(replaceWith===undefined)ed.select();
}
function commitEdit(){
  if(ed.style.display==='none'||ed.style.display==='')return;
  setCell(sel,ed.value);
  ed.style.display='none';
}
function setCell(ref,val){
  val=String(val).trim();
  if(val==='')delete cells[ref]; else cells[ref]=val;
  recompute(); fx.value=cells[ref]||''; queueSave();
}
function move(dr,dc,extend){
  const p=parseRef(extend?selEnd:sel); if(!p)return;
  const r=Math.max(0,Math.min(ROWS-1,p.r+dr)), c=Math.max(0,Math.min(COLS-1,p.c+dc));
  if(extend){
    selEnd=refOf(c,r);
    const td=document.getElementById('c-'+selEnd);
    if(td)td.scrollIntoView({block:'nearest',inline:'nearest'});
    paint();
  } else select(refOf(c,r));
}
function copyRange(){
  const A=parseRef(sel),B=parseRef(selEnd); if(!A||!B)return;
  const rows=[];
  for(let r=Math.min(A.r,B.r);r<=Math.max(A.r,B.r);r++){
    const row=[];
    for(let c=Math.min(A.c,B.c);c<=Math.max(A.c,B.c);c++){
      const v=valueOf(refOf(c,r)); row.push(v===''?'':String(v));
    }
    rows.push(row.join('\t'));
  }
  const txt=rows.join('\n');
  if(navigator.clipboard&&window.isSecureContext)navigator.clipboard.writeText(txt);
  else{
    const ta=document.createElement('textarea'); ta.value=txt;
    ta.style.position='fixed'; ta.style.opacity='0';
    document.body.appendChild(ta); ta.select();
    try{document.execCommand('copy');}finally{ta.remove();}
  }
  document.getElementById('selStats').textContent='copied ✓';
  setTimeout(paint,900);
}
function applyFmt(patch){
  rangeRefs().forEach(r=>{
    fmt[r]=Object.assign({},fmt[r]||{},patch);
    Object.keys(fmt[r]).forEach(k=>{
      if(fmt[r][k]===''||fmt[r][k]===0)delete fmt[r][k];});
    if(!Object.keys(fmt[r]).length)delete fmt[r];
  });
  recompute(); paint(); queueSave();
}
document.querySelectorAll('.tbx[data-f]').forEach(b=>{
  b.addEventListener('mousedown',e=>e.preventDefault());
  b.addEventListener('click',()=>{const k=b.dataset.f;
    applyFmt({[k]:(fmt[sel]&&fmt[sel][k])?0:1});});
});
document.querySelectorAll('.tbx[data-al]').forEach(b=>{
  b.addEventListener('mousedown',e=>e.preventDefault());
  b.addEventListener('click',()=>{const v=b.dataset.al;
    applyFmt({al:(fmt[sel]&&fmt[sel].al===v)?'':v});});
});
document.querySelectorAll('.tbx[data-nf]').forEach(b=>{
  b.addEventListener('mousedown',e=>e.preventDefault());
  b.addEventListener('click',()=>{const v=b.dataset.nf;
    applyFmt({nf:(fmt[sel]&&fmt[sel].nf===v)?'':v});});
});
document.getElementById('fcPick').addEventListener('input',e=>applyFmt({fc:e.target.value}));
document.getElementById('bgPick').addEventListener('input',e=>applyFmt({bg:e.target.value}));
document.getElementById('clearFmt').addEventListener('click',()=>{
  rangeRefs().forEach(r=>{delete fmt[r];});
  recompute(); paint(); queueSave();
});
ed.addEventListener('keydown',e=>{
  if(e.key==='Enter'){e.preventDefault();commitEdit();move(1,0);}
  else if(e.key==='Tab'){e.preventDefault();commitEdit();move(0,1);}
  else if(e.key==='Escape'){ed.style.display='none';}
});
ed.addEventListener('blur',commitEdit);
document.addEventListener('keydown',e=>{
  if(document.getElementById('editView').style.display!=='flex')return;
  const inEd=document.activeElement===ed, inFx=document.activeElement===fx,
        inTitle=document.activeElement===titleEl, inAi=document.activeElement===document.getElementById('aiQ');
  if(inEd||inFx||inTitle||inAi)return;
  if(e.key==='ArrowUp'){e.preventDefault();move(-1,0,e.shiftKey);}
  else if(e.key==='ArrowDown'){e.preventDefault();move(1,0,e.shiftKey);}
  else if(e.key==='Enter'){e.preventDefault();move(1,0,false);}
  else if(e.key==='ArrowLeft'){e.preventDefault();move(0,-1,e.shiftKey);}
  else if(e.key==='ArrowRight'){e.preventDefault();move(0,1,e.shiftKey);}
  else if(e.key==='Tab'){e.preventDefault();move(0,e.shiftKey?-1:1,false);}
  else if(e.key==='Delete'||e.key==='Backspace'){e.preventDefault();
    rangeRefs().forEach(r=>{delete cells[r];});
    recompute(); paint(); queueSave();}
  else if(e.key==='F2'){e.preventDefault();startEdit(sel);}
  else if((e.ctrlKey||e.metaKey)&&e.key.toLowerCase()==='c'){
    e.preventDefault();copyRange();}
  else if(e.key.length===1&&!e.ctrlKey&&!e.metaKey&&!e.altKey){
    e.preventDefault();startEdit(sel,e.key);
  }
});
fx.addEventListener('keydown',e=>{
  if(e.key==='Enter'){e.preventDefault();setCell(sel,fx.value);move(1,0);}
  else if(e.key==='Escape'){fx.value=cells[sel]||'';fx.blur();}
});
// paste TSV/CSV block starting at the selected cell
document.addEventListener('paste',e=>{
  if(document.getElementById('editView').style.display!=='flex')return;
  if(document.activeElement===ed||document.activeElement===fx
     ||document.activeElement===titleEl
     ||document.activeElement===document.getElementById('aiQ'))return;
  const txt=(e.clipboardData||{}).getData?e.clipboardData.getData('text'):'';
  if(!txt)return;
  e.preventDefault();
  const p=parseRef(sel); if(!p)return;
  txt.replace(/\r/g,'').split('\n').forEach((line,dr)=>{
    if(line==='')return;
    line.split('\t').forEach((v,dc)=>{
      if(p.r+dr<ROWS&&p.c+dc<COLS&&v!=='')cells[refOf(p.c+dc,p.r+dr)]=v;
    });
  });
  recompute(); queueSave();
});

// ================= persistence =================
async function loadList(){
  const l=await(await fetch('/sheets/list')).json();
  const g=document.getElementById('shGrid'); g.innerHTML='';
  const nc=document.createElement('div');
  nc.className='sh-card new'; nc.textContent='＋  New spreadsheet';
  nc.addEventListener('click',async()=>{
    const r=await(await fetch('/sheets/new',{method:'POST'})).json();
    openSheet(r.id,true);
  });
  g.appendChild(nc);
  l.forEach(s=>{
    const c=document.createElement('div'); c.className='sh-card';
    c.innerHTML=`<div class="sh-title">${escapeHtml(s.title)}</div>
      <div class="sh-meta">${fmtDate(s.updated)} · ${s.filled} filled cell${s.filled!==1?'s':''}</div>
      <button class="sh-del" title="Delete">✕</button>`;
    c.querySelector('.sh-del').addEventListener('click',async e=>{
      e.stopPropagation();
      if(confirm(`Delete "${s.title}"?`)){await fetch('/sheets/delete/'+s.id,{method:'POST'});loadList();}
    });
    c.addEventListener('click',()=>openSheet(s.id));
    g.appendChild(c);
  });
}
async function openSheet(id,isNew){
  const s=await(await fetch('/sheets/get/'+id)).json();
  if(s.error)return;
  curSheet=id; cells=s.cells||{}; fmt=s.fmt||{}; colw=s.colw||{};
  ROWS=s.rows||60; COLS=s.cols||26;
  titleEl.value=s.title==='Untitled sheet'&&isNew?'':s.title;
  document.getElementById('listView').style.display='none';
  document.getElementById('editView').style.display='flex';
  document.getElementById('aiBar').classList.add('show');
  buildGrid(); recompute(); select('A1');
  if(isNew)titleEl.focus();
}
function closeSheet(){
  commitEdit();
  saveNow().then(()=>{
    curSheet=null;
    document.getElementById('editView').style.display='none';
    document.getElementById('aiBar').classList.remove('show');
    document.getElementById('listView').style.display='block';
    document.getElementById('savedTag').textContent='';
    loadList();
  });
}
document.getElementById('backBtn').addEventListener('click',e=>{e.preventDefault();closeSheet();});
async function saveNow(){
  if(!curSheet)return;
  document.getElementById('savedTag').textContent='Saving…';
  try{
    await fetch('/sheets/save/'+curSheet,{method:'POST',
      headers:{'Content-Type':'application/json'},
      body:JSON.stringify({title:titleEl.value.trim()||'Untitled sheet',
                           cells:cells,fmt:fmt,colw:colw,rows:ROWS})});
    document.getElementById('savedTag').textContent='Saved ✓';
  }catch(e){document.getElementById('savedTag').textContent='⚠ not saved';}
}
function queueSave(){
  document.getElementById('savedTag').textContent='…';
  clearTimeout(saveTimer); saveTimer=setTimeout(saveNow,1200);
}
titleEl.addEventListener('input',queueSave);
document.getElementById('moreRows').addEventListener('click',()=>{
  ROWS=Math.min(500,ROWS+20); buildGrid(); recompute(); select(sel); queueSave();
});
document.getElementById('csvBtn').addEventListener('click',()=>{
  commitEdit();
  let maxR=0,maxC=0;
  Object.keys(cells).forEach(k=>{const p=parseRef(k);if(p){maxR=Math.max(maxR,p.r);maxC=Math.max(maxC,p.c);}});
  const rows=[];
  for(let r=0;r<=maxR;r++){
    const row=[];
    for(let c=0;c<=maxC;c++){
      let v=valueOf(refOf(c,r)); v=v===undefined?'':String(v);
      row.push(/[",\n]/.test(v)?'"'+v.replace(/"/g,'""')+'"':v);
    }
    rows.push(row.join(','));
  }
  const blob=new Blob([rows.join('\n')],{type:'text/csv'});
  const a=document.createElement('a');
  a.href=URL.createObjectURL(blob);
  a.download=(titleEl.value.trim()||'sheet')+'.csv';
  a.click();
});

// ================= AI fill =================
async function aiFill(){
  const q=document.getElementById('aiQ'), go=document.getElementById('aiGo'),
        st=document.getElementById('aiStatus');
  const prompt=q.value.trim();
  if(!prompt||aiBusy||!curSheet)return;
  aiBusy=true; go.disabled=true; st.textContent='thinking…';
  let job;
  try{
    job=(await(await fetch('/sheets/write',{method:'POST',
      headers:{'Content-Type':'application/json'},
      body:JSON.stringify({prompt,title:titleEl.value.trim(),cells})})).json()).job_id;
  }catch(e){st.textContent='failed to start';aiBusy=false;go.disabled=false;return;}
  aiJob=job; go.disabled=false; go.textContent='⏹ Stop';
  q.value='';
  const poll=setInterval(async()=>{
    let p; try{p=await(await fetch('/progress/'+job)).json();}catch(e){return;}
    if(!p.done){st.textContent='filling… '+(p.tokens||0)+' tokens';return;}
    clearInterval(poll);
    aiJob=null; go.textContent='▦ Fill';
    let data=null;
    try{data=JSON.parse(p.answer);}catch(e){}
    if(data&&data.cells){
      const entries=Object.entries(data.cells).sort((a,b)=>{
        const A=parseRef(a[0]),B=parseRef(b[0]);
        return (A.r-B.r)||(A.c-B.c);
      });
      // grow the grid if the fill goes past it
      let needR=ROWS;
      entries.forEach(([k])=>{const pr=parseRef(k);if(pr)needR=Math.max(needR,pr.r+1);});
      if(needR>ROWS){ROWS=Math.min(500,needR);buildGrid();recompute();select(sel);}
      let i=0;
      const anim=setInterval(()=>{
        if(i>=entries.length){
          clearInterval(anim); recompute(); queueSave();
          if(!titleEl.value.trim())titleEl.value=prompt.slice(0,60);
          st.textContent='done ✓ '+entries.length+' cells';
          setTimeout(()=>{st.textContent='';},4000);
          aiBusy=false; go.disabled=false;
          return;
        }
        const [k,v]=entries[i++];
        cells[k]=String(v); recompute();
        const td=document.getElementById('c-'+k);
        if(td){td.classList.add('aiflash');
               td.scrollIntoView({block:'nearest',inline:'nearest'});
               setTimeout(()=>td.classList.remove('aiflash'),600);}
      },60);
    } else {
      st.textContent=(p.answer||'failed').replace(/\[|\]/g,'').slice(0,140);
      aiBusy=false; go.disabled=false;
    }
  },900);
}
document.getElementById('aiGo').addEventListener('click',()=>{
  if(aiBusy&&aiJob){ fetch('/stop/'+aiJob,{method:'POST'}).catch(()=>{}); return; }
  aiFill();
});
document.getElementById('aiQ').addEventListener('keydown',e=>{
  if(e.key==='Enter'){e.preventDefault();aiFill();}
});

loadList();
const _tsw=document.getElementById('themeSw');
_tsw.checked=localStorage.getItem('creweTheme')==='dark';
_tsw.addEventListener('change',()=>{
  if(_tsw.checked){document.documentElement.dataset.theme='dark';localStorage.setItem('creweTheme','dark');}
  else{delete document.documentElement.dataset.theme;localStorage.setItem('creweTheme','light');}
});
</script>
</body>
</html>"""


SETTINGS_PAGE = r"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>Settings</title>
<link rel="icon" href="data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 64 64'%3E%3Crect width='64' height='64' rx='14' fill='%238b4a18'/%3E%3Cpath d='M32 57V34M32 34 16 13M32 34 48 13' stroke='%23faf6f0' stroke-width='7' stroke-linecap='round' fill='none'/%3E%3Ccircle cx='32' cy='34' r='6' fill='%23faf6f0'/%3E%3C/svg%3E">
<link href="/static/fonts.css" rel="stylesheet">
<script>if(localStorage.getItem('creweTheme')==='dark')document.documentElement.dataset.theme='dark';</script>
<style>
  :root { --bg:#faf6f0; --panel:#ffffff; --panel2:#f2ebe0; --ink:#1e160a;
          --muted:#8a7560; --border:#ddd0b8; --accent:#2a4e8c;
          --gold:#8b4a18; --btn-ink:#fff; --ok:#2e6e2a; --down:#c0392b; }
  :root[data-theme="dark"] {
    --bg:#191410; --panel:#221c15; --panel2:#2b241b; --ink:#ece2d3;
    --muted:#9c8a75; --border:#3a3125; --accent:#8aaee8; --gold:#d29a62;
    --btn-ink:#241b11; --ok:#7cbf72; --down:#e06552; color-scheme:dark; }
  .theme-toggle { display:inline-flex; align-items:center; gap:5px;
                  cursor:pointer; user-select:none; flex:0 0 auto; }
  .theme-toggle input { position:absolute; opacity:0; pointer-events:none; }
  .tt-track { width:34px; height:18px; border-radius:999px; background:var(--bg);
              border:1px solid var(--border); position:relative; }
  .tt-knob { position:absolute; top:1px; left:1px; width:14px; height:14px;
             border-radius:50%; background:var(--gold); transition:left .2s; }
  .theme-toggle input:checked + .tt-track .tt-knob { left:17px; }
  .tt-icon { font-size:11px; color:var(--muted); }
  * { box-sizing:border-box; }
  body { margin:0; font:14px/1.5 'DM Sans',system-ui,sans-serif;
         background:var(--bg); color:var(--ink); padding-bottom:90px; }
  header { padding:12px 24px; border-bottom:1px solid var(--border);
           display:flex; align-items:center; gap:14px; background:var(--panel); }
  .back { color:var(--accent); text-decoration:none; font-size:13px; }
  header h1 { margin:0; font-size:18px; font-family:'Lora',Georgia,serif;
              color:var(--gold); flex:1; }
  main { max-width:900px; margin:0 auto; padding:20px 24px; }
  h2 { font:600 15px 'Lora',Georgia,serif; color:var(--gold); margin:26px 0 4px; }
  .hint { color:var(--muted); font-size:12px; margin:0 0 12px; }
  .card { background:var(--panel); border:1px solid var(--border);
          border-radius:12px; padding:14px 16px; margin-bottom:10px; }
  .row { display:flex; gap:8px; flex-wrap:wrap; align-items:center;
         margin-bottom:8px; }
  .row:last-child { margin-bottom:0; }
  label.f { font-size:11px; color:var(--muted); display:block; margin-bottom:2px; }
  input[type=text],input[type=password],select,textarea {
    background:var(--bg); color:var(--ink); border:1px solid var(--border);
    border-radius:7px; padding:7px 10px; font:13px inherit; outline:none; }
  input:focus,select:focus,textarea:focus { border-color:var(--gold); }
  textarea { width:100%; resize:vertical; min-height:44px; }
  .grow { flex:1; min-width:140px; }
  button { background:var(--panel2); color:var(--ink); border:1px solid var(--border);
           border-radius:8px; padding:7px 13px; font:600 12px inherit;
           cursor:pointer; }
  button:hover { border-color:var(--gold); }
  button.primary { background:var(--gold); color:var(--btn-ink); border:0; }
  button.danger:hover { border-color:var(--down); color:var(--down); }
  .tag { font-size:10px; font-weight:700; text-transform:uppercase;
         letter-spacing:.06em; padding:2px 8px; border-radius:999px;
         background:var(--panel2); color:var(--muted); }
  .dot { display:inline-block; width:9px; height:9px; border-radius:50%;
         background:var(--muted); }
  .dot.ok { background:var(--ok); } .dot.bad { background:var(--down); }
  .status { font-size:12px; color:var(--muted); }
  .status.ok { color:var(--ok); } .status.bad { color:var(--down); }
  .rname { font-weight:700; font-size:14px; }
  .verdict { font:12px ui-monospace,monospace; margin:2px 0; }
  .verdict .yes { color:var(--ok); } .verdict .no { color:var(--down); }
  #saveBar { position:fixed; left:0; right:0; bottom:0; background:var(--panel);
             border-top:1px solid var(--border); padding:12px 24px;
             display:flex; gap:14px; align-items:center; z-index:50; }
  #saveMsg { font-size:13px; color:var(--muted); white-space:pre-wrap; }
  #saveMsg.ok { color:var(--ok); } #saveMsg.bad { color:var(--down); }
</style>
</head>
<body>
<header>
  <a href="/" class="back">← Crewe</a>
  <h1>⚙ Settings — router &amp; models</h1>
  <label class="theme-toggle" title="Dark mode"><span class="tt-icon">☀</span><input type="checkbox" id="themeSw"><span class="tt-track"><span class="tt-knob"></span></span><span class="tt-icon">🌙</span></label>
</header>
<main>
  <h2>Backends</h2>
  <p class="hint">Anywhere a model lives — llama.cpp, Ollama, vLLM, LM Studio,
  or a hosted OpenAI-compatible API (OpenRouter, Groq…) with an API key.
  Routes point at these.</p>
  <div id="backends"></div>
  <div class="row">
    <button id="addBackend">＋ Add backend</button>
    <button id="scanBtn">🔎 Scan localhost</button>
    <span class="status" id="scanMsg"></span>
  </div>

  <h2>Routes</h2>
  <p class="hint">The classifier reads each route's subject to decide where a
  question goes. Built-in routes have tuned routing rules and special
  machinery; you can repoint their backend and rewrite their persona, but not
  delete them.</p>
  <p class="hint" id="routeWarn" style="display:none"></p>
  <div id="routes"></div>

  <div class="card" id="addRouteCard">
    <div class="row"><span class="rname">＋ Add a route</span></div>
    <div class="row">
      <div><label class="f">name (one word)</label>
        <input type="text" id="nrName" placeholder="dnd" size="10"></div>
      <div class="grow"><label class="f">what's it for? (plain English)</label>
        <input type="text" id="nrPurpose" class="grow" style="width:100%"
               placeholder="my D&D campaign — characters, encounters, rules"></div>
      <button id="draftBtn">✨ Draft it</button>
      <span class="status" id="draftMsg"></span>
    </div>
    <div class="row"><div class="grow"><label class="f">subject — what the classifier routes here</label>
      <textarea id="nrSubject" placeholder="(drafted or written by you)"></textarea></div></div>
    <div class="row"><div class="grow"><label class="f">persona — system prompt for the specialist</label>
      <textarea id="nrPersona"></textarea></div></div>
    <div class="row">
      <div><label class="f">backend</label><select id="nrBackend"></select></div>
      <div><label class="f">model override (optional)</label>
        <input type="text" id="nrModel" size="18"></div>
    </div>
    <div class="row"><div class="grow"><label class="f">test questions that SHOULD route here (one per line; add a
      counter-example that shouldn't)</label>
      <textarea id="nrTests" placeholder="how do I build a level 5 ranger?&#10;design a boss encounter for four level 3 players&#10;recipe for beef stew"></textarea></div></div>
    <div class="row">
      <button id="testBtn">🧪 Test routing</button>
      <button class="primary" id="addRouteBtn">Add route</button>
      <span class="status" id="testMsg"></span>
    </div>
    <div id="testResults"></div>
  </div>

  <h2>System models</h2>
  <p class="hint">The router's own machinery: the classifier that routes,
  the summarizer that maintains conversation memory, and the model the code
  pipeline uses.</p>
  <div class="card" id="rolesCard"><div class="row" id="rolesRow"></div></div>
</main>
<div id="saveBar">
  <button class="primary" id="saveBtn">💾 Save &amp; apply</button>
  <span id="saveMsg">changes apply live — no restart</span>
</div>

<script>
function esc(s){const d=document.createElement('div');d.textContent=s;return d.innerHTML;}
let cfg=null;
const BUILTINS=["recipes","creative","code","general","reasoning","search"];

async function load(){
  cfg=await(await fetch('/settings/config')).json();
  render();
}
function backendOptions(sel){
  return cfg.backends.map(b=>`<option value="${esc(b.id)}"${b.id===sel?' selected':''}>${esc(b.id)}</option>`).join('');
}
function render(){
  // ---- backends
  const bx=document.getElementById('backends'); bx.innerHTML='';
  cfg.backends.forEach((b,i)=>{
    const c=document.createElement('div'); c.className='card';
    c.innerHTML=`<div class="row">
      <span class="dot" id="bdot-${i}"></span>
      <div><label class="f">type</label>
        <select data-k="kind">${['openai','llamacpp','ollama','anthropic'].map(k=>
          `<option value="${k}"${(b.kind||'openai')===k?' selected':''}>${
            {openai:'OpenAI-compatible',llamacpp:'llama.cpp',ollama:'Ollama',anthropic:'Claude (Anthropic)'}[k]
          }</option>`).join('')}</select></div>
      <div><label class="f">name</label><input type="text" size="13" value="${esc(b.id)}" data-k="id"></div>
      <div class="grow"><label class="f">base URL</label><input type="text" style="width:100%" value="${esc(b.url)}" data-k="url"></div>
      <div><label class="f">API key</label><input type="password" size="12" value="${esc(b.key||'')}" placeholder="${b.has_key?'(saved)':'none'}" data-k="key"></div>
      <div><label class="f">default model</label><input type="text" size="14" value="${esc(b.model||'')}" data-k="model"></div>
      <button data-chk="${i}">Check</button>
      <button class="danger" data-del="${i}">✕</button></div>
      <div class="row" style="margin-top:8px;align-items:center">
        <label class="f" style="display:flex;gap:6px;align-items:center;cursor:pointer">
          <input type="checkbox" data-k="paid"${b.paid?' checked':''}> this backend costs money
        </label>
        <div class="paidonly" style="${b.paid?'':'display:none'}"><label class="f">price in / 1M tokens</label>
          <input type="number" step="0.01" min="0" size="7" value="${b.price_in||0}" data-k="price_in"></div>
        <div class="paidonly" style="${b.paid?'':'display:none'}"><label class="f">price out / 1M tokens</label>
          <input type="number" step="0.01" min="0" size="7" value="${b.price_out||0}" data-k="price_out"></div>
        <span class="paidonly f" style="${b.paid?'':'display:none'};opacity:.7">
          from your provider's pricing page — used only to estimate spend</span>
      </div>
      <div class="kindhint f" style="opacity:.75;margin-top:6px"></div>
      <div class="status" id="bst-${i}"></div>`;
    const KIND_HINT={
      openai:"Any OpenAI-compatible server or hosted API. Base URL WITHOUT /v1 — Crewe appends it. Most hosted APIs need an API key and a model name.",
      llamacpp:"llama.cpp server. Base URL WITHOUT /v1 (e.g. http://127.0.0.1:8080). No key or model name needed — it serves one model. Crewe reads its context size from /props.",
      ollama:"Ollama. Base URL WITHOUT /v1 (e.g. http://127.0.0.1:11434). No key, but the model name is REQUIRED (e.g. qwen3:8b) — Ollama serves many.",
      anthropic:"Claude, direct. Base URL https://api.anthropic.com — Crewe uses /v1/messages, sends x-api-key, and translates the request and response for you. Model name is REQUIRED (e.g. claude-opus-5). Note: Claude ignores temperature — steer it with the route persona instead."};
    const hintEl=c.querySelector('.kindhint');
    const paintHint=()=>{hintEl.textContent=KIND_HINT[b.kind||'openai'];};
    paintHint();
    c.querySelector('[data-k="kind"]').addEventListener('change',e=>{
      b.kind=e.target.value; paintHint();
      if(b.kind==='anthropic'&&!b.url) { b.url='https://api.anthropic.com';
        const u=c.querySelector('[data-k="url"]'); if(u) u.value=b.url; }
    });
    c.querySelectorAll('input').forEach(inp=>inp.addEventListener('input',()=>{
      const k=inp.dataset.k;
      if(k==='paid'){
        b.paid=inp.checked;
        c.querySelectorAll('.paidonly').forEach(e=>e.style.display=inp.checked?'':'none');
        return;
      }
      b[k]=(k==='price_in'||k==='price_out')?parseFloat(inp.value||0):inp.value;
      if(k==='key') b.has_key=false;
    }));
    c.querySelector('[data-chk]').addEventListener('click',()=>checkBackend(i));
    c.querySelector('[data-del]').addEventListener('click',()=>{
      const used=Object.values(cfg.routes).some(r=>r.backend===b.id)
        || Object.values(cfg.roles).includes(b.id);
      if(used){alert('A route or system role still uses this backend.');return;}
      cfg.backends.splice(i,1); render();
    });
    bx.appendChild(c);
    checkBackend(i,true);
  });
  // ---- routes (+ soft capacity warning: the classifier is a small model)
  const rw=document.getElementById('routeWarn'),
        nRoutes=Object.keys(cfg.routes).length;
  if(nRoutes>9){
    rw.style.display='block'; rw.style.color='var(--down)';
    rw.textContent=`⚠ ${nRoutes} routes — the classifier is a small model and `
      +`one-word routing gets unreliable past ~9-10 routes. Use the routing `
      +`test on anything new, and consider merging niche routes.`;
  } else rw.style.display='none';
  const rx=document.getElementById('routes'); rx.innerHTML='';
  Object.entries(cfg.routes).forEach(([name,r])=>{
    const c=document.createElement('div'); c.className='card';
    const color=r.custom?(r.color||'#666'):`var(--${name})`;
    c.innerHTML=`<div class="row">
        <span class="dot" style="background:${color}"></span>
        <span class="rname">${esc(name)}</span>
        <span class="tag">${r.custom?'custom':'built-in'}</span>
        ${r.custom?'<button class="danger" data-delr>✕ delete</button>':''}
      </div>
      <div class="row">
        <div><label class="f">backend</label><select data-k="backend">${backendOptions(r.backend)}</select></div>
        <div><label class="f">model override (optional)</label>
          <input type="text" size="18" value="${esc(r.model||'')}" data-k="model"></div>
      </div>
      ${r.custom?`<div class="row"><div class="grow"><label class="f">subject (routing rule)</label>
        <textarea data-k="subject">${esc(r.subject||'')}</textarea></div></div>`:''}
      <div class="row"><div class="grow"><label class="f">persona — system prompt${r.custom?'':' (blank = built-in default)'}</label>
        <textarea data-k="persona" placeholder="${r.custom?'':'built-in persona in use'}">${esc(r.persona||'')}</textarea></div></div>`;
    c.querySelectorAll('[data-k]').forEach(el=>el.addEventListener('input',()=>{r[el.dataset.k]=el.value;}));
    const del=c.querySelector('[data-delr]');
    if(del)del.addEventListener('click',()=>{
      if(confirm(`Delete route "${name}"?`)){delete cfg.routes[name];render();}
    });
    rx.appendChild(c);
  });
  // ---- add-route backend select + roles
  document.getElementById('nrBackend').innerHTML=backendOptions(null);
  const rr=document.getElementById('rolesRow');
  // Older configs may lack the newer coder roles — seed each missing one from
  // the tier above it, mirroring the server's own fallback chain, so the
  // select and the saved config never disagree about what is selected.
  if(!cfg.roles.coder_fast)  cfg.roles.coder_fast  = cfg.roles.coder;
  if(!cfg.roles.coder_quick) cfg.roles.coder_quick = cfg.roles.coder_fast;
  const ROLE_INFO=[
    ['classifier','classifier',
     'Reads every question and picks the route. Small and fast — this one runs on each request.'],
    ['summarizer','summarizer',
     'Condenses older conversation turns. Small and fast is fine.'],
    ['coder_quick','fast coder',
     'Used when Effort is Fast. Put your SMALLEST, QUICKEST model here — instant drafts, simple asks.'],
    ['coder_fast','normal coder',
     'Used when Effort is Normal — the everyday default. A solid mid-size model belongs here.'],
    ['coder','extra coder',
     'Used when Effort is Extra. Your STRONGEST model — often a paid API. Budgets are sized automatically from whatever you choose.'],
  ];
  rr.innerHTML=ROLE_INFO.map(([role,label,help])=>`
    <div><label class="f">${label}</label>
    <select data-role="${role}">${backendOptions(cfg.roles[role])}</select>
    <div style="color:var(--muted);font:400 11.5px/1.4 'DM Sans',sans-serif;margin-top:5px">${help}</div></div>`).join('');
  rr.querySelectorAll('select').forEach(s=>s.addEventListener('change',()=>{
    cfg.roles[s.dataset.role]=s.value;}));
}
async function checkBackend(i,quiet){
  const b=cfg.backends[i], st=document.getElementById('bst-'+i),
        dot=document.getElementById('bdot-'+i);
  if(!quiet)st.textContent='checking…';
  try{
    const r=await(await fetch('/settings/backend_check',{method:'POST',
      headers:{'Content-Type':'application/json'},
      body:JSON.stringify({url:b.url,kind:b.kind||'openai',
        key:b.has_key?'__KEEP__':(b.key||'')})})).json();
    if(r.alive){dot.className='dot ok';
      st.className='status ok';
      st.textContent='online — serving: '+(r.models.join(', ')||'(no model list)');}
    else{dot.className='dot bad'; st.className='status bad';
      st.textContent='offline — '+(r.error||'');}
  }catch(e){dot.className='dot bad';}
}
document.getElementById('addBackend').addEventListener('click',()=>{
  cfg.backends.push({id:'backend-'+(cfg.backends.length+1),url:'http://',key:'',model:'',has_key:false});
  render();
});
document.getElementById('scanBtn').addEventListener('click',async()=>{
  const m=document.getElementById('scanMsg'); m.textContent='scanning…';
  const r=await(await fetch('/settings/scan',{method:'POST'})).json();
  const known=new Set(cfg.backends.map(b=>b.url.replace(/\/$/,'')));
  const fresh=r.found.filter(f=>!known.has(f.url));
  if(!fresh.length){m.textContent=`found ${r.found.length}, all already configured`;return;}
  fresh.forEach(f=>{
    if(confirm(`Found ${f.url} serving: ${f.models.join(', ')||'?'} — add it?`))
      cfg.backends.push({id:'local-'+f.url.split(':').pop(),url:f.url,key:'',model:'',has_key:false});
  });
  m.textContent=''; render();
});
document.getElementById('draftBtn').addEventListener('click',async()=>{
  const m=document.getElementById('draftMsg'); m.textContent='drafting…';
  const r=await(await fetch('/settings/draft_route',{method:'POST',
    headers:{'Content-Type':'application/json'},
    body:JSON.stringify({name:document.getElementById('nrName').value,
                         purpose:document.getElementById('nrPurpose').value})})).json();
  if(r.error){m.className='status bad';m.textContent=r.error;return;}
  document.getElementById('nrSubject').value=r.subject||'';
  document.getElementById('nrPersona').value=r.persona||'';
  m.className='status ok'; m.textContent='drafted — edit freely, then test';
});
document.getElementById('testBtn').addEventListener('click',async()=>{
  const m=document.getElementById('testMsg'), out=document.getElementById('testResults');
  const name=document.getElementById('nrName').value.trim();
  const qs=document.getElementById('nrTests').value.split('\n').filter(x=>x.trim());
  m.textContent='running the real classifier…'; out.innerHTML='';
  const r=await(await fetch('/settings/test_route',{method:'POST',
    headers:{'Content-Type':'application/json'},
    body:JSON.stringify({name,subject:document.getElementById('nrSubject').value,
                         questions:qs})})).json();
  if(r.error){m.className='status bad';m.textContent=r.error;return;}
  m.textContent='';
  out.innerHTML=r.results.map(x=>
    `<div class="verdict"><span class="${x.route===name?'yes':'no'}">→ ${esc(x.route)}</span>  ${esc(x.q)}</div>`).join('');
});
document.getElementById('addRouteBtn').addEventListener('click',()=>{
  const name=document.getElementById('nrName').value.trim().toLowerCase();
  if(!/^[a-z][a-z0-9]{1,15}$/.test(name)||BUILTINS.includes(name)||name==='audio'||cfg.routes[name]){
    alert('Route name must be one lowercase word (2-16 chars) and not already exist.');return;}
  if(!document.getElementById('nrSubject').value.trim()){
    alert('Subject is required — the classifier routes by it.');return;}
  cfg.routes[name]={backend:document.getElementById('nrBackend').value,
    model:document.getElementById('nrModel').value.trim(),
    subject:document.getElementById('nrSubject').value.trim(),
    persona:document.getElementById('nrPersona').value.trim(),
    custom:true,color:''};
  ['nrName','nrPurpose','nrSubject','nrPersona','nrModel','nrTests'].forEach(id=>document.getElementById(id).value='');
  document.getElementById('testResults').innerHTML='';
  render();
  document.getElementById('saveMsg').textContent='route staged — hit Save & apply';
});
document.getElementById('saveBtn').addEventListener('click',async()=>{
  const m=document.getElementById('saveMsg');
  m.className=''; m.textContent='saving…';
  const r=await(await fetch('/settings/config',{method:'POST',
    headers:{'Content-Type':'application/json'},body:JSON.stringify(cfg)})).json();
  if(r.ok){m.className='ok';m.textContent='saved & applied live ✓';load();}
  else{m.className='bad';m.textContent='not saved:\n- '+(r.errors||['unknown']).join('\n- ');}
});
const _tsw=document.getElementById('themeSw');
_tsw.checked=localStorage.getItem('creweTheme')==='dark';
_tsw.addEventListener('change',()=>{
  if(_tsw.checked){document.documentElement.dataset.theme='dark';localStorage.setItem('creweTheme','dark');}
  else{delete document.documentElement.dataset.theme;localStorage.setItem('creweTheme','light');}
});
load();
</script>
</body>
</html>"""


SCRATCH_PAGE = r"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>HTML Scratchpad</title>
<link rel="icon" href="data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 64 64'%3E%3Crect width='64' height='64' rx='14' fill='%238b4a18'/%3E%3Cpath d='M32 57V34M32 34 16 13M32 34 48 13' stroke='%23faf6f0' stroke-width='7' stroke-linecap='round' fill='none'/%3E%3Ccircle cx='32' cy='34' r='6' fill='%23faf6f0'/%3E%3C/svg%3E">
<link href="/static/fonts.css" rel="stylesheet">
<script>if(localStorage.getItem('creweTheme')==='dark')document.documentElement.dataset.theme='dark';</script>
<style>
  :root { --bg:#faf6f0; --panel:#ffffff; --ink:#1e160a; --muted:#8a7560;
          --border:#ddd0b8; --accent:#2a4e8c; --btn-ink:#fff;
          --edbg:#fffdf8; }
  :root[data-theme="dark"] {
    --bg:#191410; --panel:#221c15; --ink:#ece2d3; --muted:#9c8a75;
    --border:#3a3125; --accent:#8aaee8; --btn-ink:#241b11;
    --edbg:#14100d; color-scheme:dark; }
  .theme-toggle { display:inline-flex; align-items:center; gap:5px;
                  cursor:pointer; user-select:none; flex:0 0 auto; }
  .theme-toggle input { position:absolute; opacity:0; pointer-events:none; }
  .tt-track { width:34px; height:18px; border-radius:999px; background:var(--bg);
              border:1px solid var(--border); position:relative; }
  .tt-knob { position:absolute; top:1px; left:1px; width:14px; height:14px;
             border-radius:50%; background:var(--accent); transition:left .2s; }
  .theme-toggle input:checked + .tt-track .tt-knob { left:17px; }
  .tt-icon { font-size:11px; color:var(--muted); }
  * { box-sizing:border-box; }
  html,body { height:100%; margin:0; }
  body { font:14px/1.5 'DM Sans',-apple-system,Segoe UI,Roboto,sans-serif;
         background:var(--bg); color:var(--ink); display:flex; flex-direction:column; }
  header { display:flex; align-items:center; gap:12px; padding:10px 14px;
           border-bottom:1px solid var(--border); flex:0 0 auto; }
  header h1 { font-size:14px; font-weight:600; margin:0; }
  header .hint { color:var(--muted); font-size:12px; }
  header .spacer { flex:1; }
  header a { color:var(--accent); text-decoration:none; font-size:13px; }
  button { background:var(--accent); color:var(--btn-ink); border:0; border-radius:8px;
           padding:7px 16px; font-weight:600; cursor:pointer; font:inherit; }
  button.alt { background:var(--panel); border:1px solid var(--border); color:var(--ink); }
  main { flex:1; display:flex; min-height:0; }
  .pane { flex:1; display:flex; flex-direction:column; min-width:0; }
  .pane.left { border-right:1px solid var(--border); }
  .pane label { padding:6px 12px; color:var(--muted); font-size:12px;
                border-bottom:1px solid var(--border); }
  textarea { flex:1; border:0; resize:none; background:var(--edbg); color:var(--ink);
             padding:14px; font:13px/1.55 ui-monospace,SFMono-Regular,Menlo,Consolas,monospace;
             outline:none; }
  iframe { flex:1; border:0; background:#fff; width:100%; }
  kbd { background:var(--panel); border:1px solid var(--border); border-radius:4px;
        padding:1px 6px; font-size:11px; color:var(--muted); }
</style>
</head>
<body>
<header>
  <h1>HTML Scratchpad</h1>
  <span class="hint">paste code, hit <kbd>Ctrl/Cmd + Enter</kbd> or Run</span>
  <span class="spacer"></span>
  <a href="/">← back to Crewe</a>
  <label class="theme-toggle" title="Dark mode"><span class="tt-icon">☀</span><input type="checkbox" id="themeSw"><span class="tt-track"><span class="tt-knob"></span></span><span class="tt-icon">🌙</span></label>
  <script>(function(){const t=document.getElementById('themeSw');t.checked=localStorage.getItem('creweTheme')==='dark';t.addEventListener('change',()=>{if(t.checked){document.documentElement.dataset.theme='dark';localStorage.setItem('creweTheme','dark');}else{delete document.documentElement.dataset.theme;localStorage.setItem('creweTheme','light');}});})();</script>
  <button class="alt" id="auto">Auto-run: ON</button>
  <button class="alt" id="clear">Clear</button>
  <button id="run">Run &#9654;</button>
</header>
<main>
  <div class="pane left">
    <label>HTML / CSS / JS</label>
    <textarea id="src" placeholder="<!-- paste HTML here -->" spellcheck="false"></textarea>
  </div>
  <div class="pane">
    <label>Preview</label>
    <iframe id="out" sandbox="allow-scripts"></iframe>
  </div>
</main>
<script>
  const src = document.getElementById('src');
  const out = document.getElementById('out');
  const runBtn = document.getElementById('run');
  const clearBtn = document.getElementById('clear');
  const autoBtn = document.getElementById('auto');
  let autorun = true, timer = null;

  function render(){ out.srcdoc = src.value; }
  runBtn.addEventListener('click', render);
  clearBtn.addEventListener('click', () => { src.value=''; render(); src.focus(); });
  autoBtn.addEventListener('click', () => {
    autorun = !autorun;
    autoBtn.textContent = 'Auto-run: ' + (autorun ? 'ON' : 'OFF');
  });
  src.addEventListener('input', () => {
    if (!autorun) return;
    clearTimeout(timer); timer = setTimeout(render, 400);
  });
  src.addEventListener('keydown', (e) => {
    if ((e.metaKey || e.ctrlKey) && e.key === 'Enter') { e.preventDefault(); render(); }
  });
  const _saved = localStorage.getItem('scratchLoad');
  if (_saved) { localStorage.removeItem('scratchLoad'); src.value = _saved; }
  else { src.value = '<!DOCTYPE html>\n<html>\n<body style="font-family:sans-serif;padding:2rem">\n  <h1>It works \u2728</h1>\n  <p>Paste HTML on the left. It renders here.</p>\n</body>\n</html>'; }
  render();
</script>
</body>
</html>"""


COMPARE_PAGE = r"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>Code A/B</title>
<script>if(localStorage.getItem('creweTheme')==='dark')document.documentElement.dataset.theme='dark';</script>
<link rel="icon" href="data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 64 64'%3E%3Crect width='64' height='64' rx='14' fill='%238b4a18'/%3E%3Cpath d='M32 57V34M32 34 16 13M32 34 48 13' stroke='%23faf6f0' stroke-width='7' stroke-linecap='round' fill='none'/%3E%3Ccircle cx='32' cy='34' r='6' fill='%23faf6f0'/%3E%3C/svg%3E">
<link href="/static/fonts.css" rel="stylesheet">
<script src="/static/marked.min.js"></script>
<script src="/static/jszip.min.js"></script>
<style>
  :root {
    --bg:        #faf6f0;
    --panel:     #ffffff;
    --panel2:    #f2ebe0;
    --ink:       #1e160a;
    --muted:     #8a7560;
    --border:    #ddd0b8;
    --code:      #2e6e2a;
    --creative:  #903828;
    --reasoning: #2a4e8c;
    --accent:    #8b4a18;
    --a-color:   #2e6e2a;
    --b-color:   #903828;
    --btn-ink:   #fff;
  }
  :root[data-theme="dark"] {
    --bg:#191410; --panel:#221c15; --panel2:#2b241b; --ink:#ece2d3;
    --muted:#9c8a75; --border:#3a3125; --code:#7cbf72; --creative:#d98276;
    --reasoning:#8aaee8; --accent:#d29a62; --a-color:#7cbf72;
    --b-color:#d98276; --btn-ink:#241b11; color-scheme:dark; }
  .theme-toggle { display:inline-flex; align-items:center; gap:5px;
                  cursor:pointer; user-select:none; flex:0 0 auto; }
  .theme-toggle input { position:absolute; opacity:0; pointer-events:none; }
  .tt-track { width:34px; height:18px; border-radius:999px; background:var(--bg);
              border:1px solid var(--border); position:relative; }
  .tt-knob { position:absolute; top:1px; left:1px; width:14px; height:14px;
             border-radius:50%; background:var(--accent); transition:left .2s; }
  .theme-toggle input:checked + .tt-track .tt-knob { left:17px; }
  .tt-icon { font-size:11px; color:var(--muted); }
  * { box-sizing: border-box; margin: 0; padding: 0; }
  body { font: 15px/1.6 'DM Sans', system-ui, sans-serif;
         background: var(--bg); color: var(--ink);
         display: flex; flex-direction: column; min-height: 100vh; }

  header { flex: 0 0 auto; padding: 14px 28px; border-bottom: 1px solid var(--border);
           background: var(--panel); display: flex; align-items: center; gap: 18px; flex-wrap: wrap; }
  header h1 { font-family: 'Lora', Georgia, serif; font-size: 20px;
              font-weight: 600; color: var(--accent); }
  header .sub { color: var(--muted); font-size: 12px; flex: 1; min-width: 0; }
  header a { color: var(--reasoning); text-decoration: none; font-size: 13px; white-space: nowrap; }
  header a:hover { text-decoration: underline; }

  .ask-bar { background: var(--panel); border-bottom: 1px solid var(--border); padding: 14px 28px; }
  .ask-inner { max-width: 1240px; margin: 0 auto; display: flex; gap: 10px; align-items: flex-end; }
  textarea { flex: 1; resize: none; background: var(--bg); color: var(--ink);
             border: 1px solid var(--border); border-radius: 12px;
             padding: 12px 16px; font: 15px/1.5 'DM Sans', sans-serif;
             outline: none; transition: border-color .2s; max-height: 120px; min-height: 46px; }
  textarea:focus { border-color: var(--accent); }
  textarea::placeholder { color: var(--muted); }
  #go { background: var(--accent); color: var(--btn-ink); border: 0; border-radius: 12px;
        padding: 0 28px; height: 46px; font: 600 14px 'DM Sans', sans-serif;
        cursor: pointer; transition: opacity .15s; white-space: nowrap; }
  #go:hover { opacity: .85; }
  #go:disabled { opacity: .35; cursor: default; }

  .page { flex: 1; max-width: 1240px; width: 100%; margin: 0 auto; padding: 24px 28px;
          display: flex; flex-direction: column; gap: 20px; }

  .idle { text-align: center; padding: 80px 20px; color: var(--muted); font-size: 14px; }
  .idle .icon { font-size: 42px; margin-bottom: 16px; }

  /* ---- split hero ---- */
  .split { display: grid; grid-template-columns: 1fr 1fr; gap: 20px; align-items: start; }
  @media (max-width: 800px) { .split { grid-template-columns: 1fr; } }

  /* ---- backend picker ---- */
  .pickrow { max-width:1400px; margin:10px auto 0; display:flex; flex-wrap:wrap;
             gap:8px; align-items:center; }
  .pickhint { color:var(--muted); font-size:12px; }
  .pick { display:inline-flex; align-items:center; gap:7px; cursor:pointer;
          border:1px solid var(--border); border-radius:999px;
          padding:5px 12px; font-size:12.5px; background:var(--panel);
          user-select:none; }
  .pick input { margin:0; cursor:pointer; }
  .pick.on { border-color:var(--accent); }
  .pick .mdl { color:var(--muted); font-size:11px; }
  .pick .cost { color:var(--accent); font-size:11px; font-weight:600; }
  .pickinner { display:flex; flex-wrap:wrap; gap:8px; align-items:center; }
  .modesw { display:inline-flex; border:1px solid var(--border);
            border-radius:999px; overflow:hidden; }
  .modebtn { background:var(--panel); border:0; color:var(--muted);
             cursor:pointer; font:600 12px 'DM Sans',sans-serif; padding:6px 14px; }
  .modebtn.on { background:var(--accent); color:var(--btn-ink); }

  /* ---- track card ---- */
  .track { background: var(--panel); border: 1px solid var(--border);
           border-radius: 14px; display: flex; flex-direction: column;
           box-shadow: 0 1px 4px rgba(80,50,20,.06); overflow: hidden; }
  .track-head { padding: 12px 16px; border-bottom: 1px solid var(--border);
                background: var(--panel2); display: flex; align-items: center;
                justify-content: space-between; gap: 8px; flex-wrap: wrap; }
  .track-label { font-family: 'Lora', Georgia, serif; font-size: 14px; font-weight: 600; }
  .track-label.a { color: var(--a-color); }
  .track-label.b { color: var(--b-color); }
  .stage-info { font-size: 11px; color: var(--muted);
                display: flex; align-items: center; gap: 5px; }
  .spinner { width: 11px; height: 11px; border: 2px solid var(--border);
             border-top-color: var(--accent); border-radius: 50%;
             animation: spin .7s linear infinite; flex: 0 0 auto; }
  @keyframes spin { to { transform: rotate(360deg); } }
  .done-tick { color: var(--code); font-size: 11px; }

  /* ---- brief ---- */
  .brief-toggle { width: 100%; background: none; border: none; border-bottom: 1px solid var(--border);
                  padding: 7px 16px; font: 500 12px 'DM Sans', sans-serif; color: var(--muted);
                  cursor: pointer; text-align: left; display: flex; justify-content: space-between;
                  align-items: center; transition: color .15s; }
  .brief-toggle:hover { color: var(--b-color); background: rgba(144,56,40,.04); }
  .brief-body { padding: 14px 16px; background: #fdf8f2; border-bottom: 1px solid var(--border);
                max-height: 300px; overflow-y: auto; font-size: 13px; line-height: 1.7; }

  /* ---- body ---- */
  .track-body { padding: 16px; overflow-y: auto; min-height: 180px; max-height: 560px; }
  .pending { display: flex; align-items: center; gap: 10px;
             color: var(--muted); font-size: 13px; padding: 16px 0; }

  /* ---- markdown ---- */
  .md { font-size: 13.5px; line-height: 1.75; word-wrap: break-word; }
  .md p { margin-bottom: .8em; } .md p:last-child { margin-bottom: 0; }
  .md h1,.md h2,.md h3,.md h4 { font-family: 'Lora', Georgia, serif; font-weight: 600; margin: 1em 0 .4em; }
  .md h1 { font-size: 1.2em; color: var(--accent); }
  .md h2 { font-size: 1.05em; color: var(--creative); }
  .md h3 { font-size: 1em; color: var(--muted); }
  .md ul,.md ol { padding-left: 1.4em; margin-bottom: .8em; }
  .md li { margin-bottom: .3em; }
  .md li::marker { color: var(--code); font-weight: 700; }
  .md strong { font-weight: 600; color: var(--accent); }
  .md em { font-style: italic; color: var(--creative); }
  .md code { font-family: ui-monospace,SFMono-Regular,Menlo,Consolas,monospace;
             font-size: .84em; background: var(--panel2); padding: 2px 5px;
             border-radius: 4px; color: var(--code); }
  .md pre { background: #1a1d24; border: 1px solid #2a2f3a;
            border-radius: 10px; padding: 14px 16px; overflow-x: auto; margin-bottom: .8em; }
  .md pre code { background: none; padding: 0; color: #e6e8ec; font-size: .84em; }
  .md table { border-collapse: collapse; width: 100%; margin-bottom: .8em; font-size: 13px; }
  .md th,.md td { border: 1px solid var(--border); padding: 7px 10px; text-align: left; }
  .md th { background: var(--panel2); font-weight: 600; color: var(--accent); }
  .md blockquote { border-left: 3px solid var(--creative); margin: .8em 0;
                   padding-left: 1em; color: var(--muted); font-style: italic; }

  /* ---- action bar ---- */
  .track-foot { border-top: 1px solid var(--border); padding: 10px 14px;
                display: flex; gap: 8px; flex-wrap: wrap; }
  .btn { border: 1px solid var(--border); background: var(--panel2); color: var(--ink);
         border-radius: 8px; padding: 6px 14px; font: 600 12px 'DM Sans', sans-serif;
         cursor: pointer; transition: all .15s; display: inline-flex; align-items: center; gap: 5px; }
  .btn:hover { border-color: var(--accent); color: var(--accent); background: var(--panel); }
  .btn.hi { background: var(--accent); color: var(--btn-ink); border-color: transparent; }
  .btn.hi:hover { opacity: .85; color: var(--btn-ink); border-color: transparent; }

  /* ---- preview ---- */
  .preview-wrap { background: var(--panel); border: 1px solid var(--border);
                  border-radius: 14px; overflow: hidden;
                  box-shadow: 0 2px 12px rgba(80,50,20,.08); }
  .preview-head { padding: 10px 16px; border-bottom: 1px solid var(--border);
                  background: var(--panel2); display: flex; align-items: center; justify-content: space-between; }
  .preview-title { font-size: 13px; font-weight: 600; color: var(--muted); }
  .preview-close { background: none; border: none; color: var(--muted); cursor: pointer;
                   font-size: 16px; line-height: 1; padding: 2px; transition: color .15s; }
  .preview-close:hover { color: var(--ink); }
  #previewFrame { display: block; width: 100%; height: 560px; border: none; background: #fff; }

  /* ---- judge ---- */
  .judge-wrap { background: var(--panel); border: 2px solid var(--reasoning);
                border-radius: 14px; overflow: hidden;
                box-shadow: 0 2px 14px rgba(42,78,140,.10); }
  .judge-head { padding: 13px 20px; border-bottom: 1px solid var(--border);
                background: var(--panel2); display: flex; align-items: center; gap: 10px; }
  .judge-title { font-family: 'Lora', Georgia, serif; font-size: 15px;
                 font-weight: 600; color: var(--reasoning); }
  .judge-tok { font-size: 12px; color: var(--muted); margin-left: auto;
               font-variant-numeric: tabular-nums; }
  .judge-body { padding: 18px 20px; }
</style>
</head>
<body>
<header>
  <h1>Code A/B</h1>
  <label class="theme-toggle" title="Dark mode"><span class="tt-icon">☀</span><input type="checkbox" id="themeSw"><span class="tt-track"><span class="tt-knob"></span></span><span class="tt-icon">🌙</span></label>
  <script>(function(){const t=document.getElementById('themeSw');t.checked=localStorage.getItem('creweTheme')==='dark';t.addEventListener('change',()=>{if(t.checked){document.documentElement.dataset.theme='dark';localStorage.setItem('creweTheme','dark');}else{delete document.documentElement.dataset.theme;localStorage.setItem('creweTheme','light');}});})();</script>
  <span class="sub">Direct code vs creative brief → code · judge picks the winner</span>
  <a href="/">← Crewe</a>
  <a href="/scratch" target="_blank">scratchpad ↗</a>
  
</header>
<div class="ask-bar">
  <div class="ask-inner">
    <textarea id="q" rows="1" placeholder="Ask anything — every selected model gets the same prompt (Enter to run, Shift+Enter for newline)"></textarea>
    <button id="go">Compare ›</button>
  </div>
  <div class="pickrow">
    <span class="modesw"><button class="modebtn on" data-mode="backends" title="Same system prompt for everyone — judges the models">models</button><button class="modebtn" data-mode="routes" title="Each route brings its own persona — judges your configured setups">routes</button></span>
    <span id="pickRow" class="pickinner"><span class="pickhint">loading…</span></span>
  </div>
</div>
<div class="page" id="page">
  <div class="idle">
    <div class="icon">⚗️</div>
    Pick two or more above and ask a question. <b>models</b> gives everyone the
    same system prompt, so you judge the models; <b>routes</b> gives each its own
    persona, so you judge your setups.<br>
    They all get the <b>same prompt and the same system prompt</b>, so what you
    are comparing is the models themselves. A judge reads the answers afterwards.
  </div>
</div>
<script>
const q = document.getElementById('q');
const go = document.getElementById('go');
const page = document.getElementById('page');

marked.use({ breaks: true, gfm: true });
marked.use({ renderer: {
  html(t){ const s=typeof t==='string'?t:(t&&t.text)||'';
    const d=document.createElement('div'); d.textContent=s; return d.innerHTML; }
}});

let answers = {};

function esc(s){ const d=document.createElement('div'); d.textContent=s; return d.innerHTML; }
function md(s){ return `<div class="md">${marked.parse(s)}</div>`; }

const EXTS = {html:'html',htm:'html',css:'css',javascript:'js',js:'js',
              typescript:'ts',ts:'ts',python:'py',py:'py',
              json:'json',bash:'sh',shell:'sh',sql:'sql',c:'c',cpp:'cpp',rust:'rs'};

// Returns true for bash blocks that are terminal commands, not actual script files.
// Real shell scripts start with a shebang; command dumps don't.
function isCommandBlock(lang, code) {
  if (!['bash','shell','sh','zsh','console','terminal','cmd'].includes(lang)) return false;
  return !code.trim().startsWith('#!/');
}

// Extract all named code blocks from a model response.
// Tries to infer filenames from the prose immediately before each fence,
// then falls back to inline comment on the first line of the block.
// Skips terminal command blocks (npm install, mkdir, etc.) — those aren't files.
function extractFiles(md) {
  const files = [];
  const seen = new Set();
  const re = /```(\w*) *\n([\s\S]*?)```/g;
  let m, lastIndex = 0;
  while ((m = re.exec(md)) !== null) {
    const lang = m[1].trim() || '';
    const code = m[2];
    const before = md.slice(lastIndex, m.index);
    lastIndex = re.lastIndex;

    // Skip blocks that are just terminal commands
    if (isCommandBlock(lang, code)) continue;

    // Look for a filename in the lines immediately before the fence.
    // Handles: **`src/App.jsx`**, ### src/App.jsx, `index.html`, plain path/to/file.ext
    const beforeLines = before.trimEnd().split('\n');
    let filename = null;
    for (let i = beforeLines.length - 1; i >= Math.max(0, beforeLines.length - 4); i--) {
      const line = beforeLines[i].trim().replace(/[`*_#>]+/g, ' ').trim();
      const fnm = line.match(/([a-zA-Z0-9_\-]+(?:\/[a-zA-Z0-9_\-.]+)*\.[a-zA-Z0-9]+)/);
      if (fnm && !fnm[1].match(/^https?:/) && fnm[1].length < 80) {
        filename = fnm[1]; break;
      }
    }

    // Fallback: check first line of code for an inline comment with a filename
    if (!filename) {
      const firstLine = code.split('\n')[0];
      const cm = firstLine.match(/(?:\/\/|#|<!--|\/\*)\s*([a-zA-Z0-9_\-./]+\.[a-zA-Z0-9]+)/);
      if (cm && !cm[1].match(/^https?:/)) filename = cm[1];
    }

    // Final fallback: generate a name from lang/position
    if (!filename) {
      const ext = EXTS[lang] || (lang || 'txt');
      filename = files.length === 0 ? `index.${ext}` : `file${files.length + 1}.${ext}`;
    }

    // Deduplicate
    let final = filename;
    let n = 2;
    while (seen.has(final)) {
      const dot = filename.lastIndexOf('.');
      final = dot > 0
        ? filename.slice(0, dot) + `_${n}` + filename.slice(dot)
        : filename + `_${n}`;
      n++;
    }
    seen.add(final);
    files.push({ name: final, code, lang });
  }
  // Nothing extracted — wrap the whole answer as a text file
  if (!files.length) files.push({ name: 'output.txt', code: md, lang: 'txt' });
  return files;
}

function htmlFile(files) {
  return files.find(f => f.name === 'index.html')
      || files.find(f => f.name.endsWith('.html'))
      || files[0];
}

async function doDownload(track) {
  const raw = answers[track]; if (!raw) return;
  const files = extractFiles(raw);
  if (files.length === 1) {
    const f = files[0];
    const a = document.createElement('a');
    a.href = URL.createObjectURL(new Blob([f.code], { type: 'text/plain' }));
    a.download = f.name;
    a.click(); URL.revokeObjectURL(a.href);
  } else {
    const zip = new JSZip();
    files.forEach(f => zip.file(f.name, f.code));
    const blob = await zip.generateAsync({ type: 'blob' });
    const a = document.createElement('a');
    a.href = URL.createObjectURL(blob);
    a.download = 'track-' + track.toLowerCase() + '.zip';
    a.click(); URL.revokeObjectURL(a.href);
  }
}

function doPreview(track) {
  const raw = answers[track]; if (!raw) return;
  const files = extractFiles(raw);
  const f = htmlFile(files);
  document.getElementById('previewWrap').style.display = 'block';
  const label = files.length > 1
    ? `Preview — Track ${track} (${f.name} · ${files.length} files total)`
    : `Preview — Track ${track}`;
  document.getElementById('previewTitle').textContent = label;
  document.getElementById('previewFrame').srcdoc = f.code;
  document.getElementById('previewWrap').scrollIntoView({ behavior: 'smooth', block: 'start' });
}

function closePreview() {
  document.getElementById('previewWrap').style.display = 'none';
  document.getElementById('previewFrame').srcdoc = '';
}

function doScratch(track) {
  const raw = answers[track]; if (!raw) return;
  const files = extractFiles(raw);
  localStorage.setItem('scratchLoad', htmlFile(files).code);
  window.open('/scratch', '_blank');
}

let briefOpen = false;
function toggleBrief() {
  briefOpen = !briefOpen;
  document.getElementById('briefBody').style.display = briefOpen ? 'block' : 'none';
  document.getElementById('briefArrow').textContent = briefOpen ? '▴' : '▾';
}

function setStage(id, html) {
  const el = document.getElementById(id); if (el) el.innerHTML = html;
}

function buildUI(ids) {
  page.innerHTML =
    '<div class="split" id="split">' + ids.map(id => `
  <div class="track">
    <div class="track-head">
      <span class="track-label a">${esc(id)}</span>
      <span class="stage-info" id="stage-${esc(id)}"><div class="spinner"></div>&nbsp;starting…</span>
    </div>
    <div class="track-body" id="body-${esc(id)}">
      <div class="pending"><div class="spinner"></div> waiting…</div>
    </div>
    <div class="track-foot" id="foot-${esc(id)}" style="display:none">
      <button class="btn" onclick="doDownload('${esc(id)}')">⬇ Download</button>
      <button class="btn hi" onclick="doPreview('${esc(id)}')">▶ Preview</button>
      <button class="btn" onclick="doScratch('${esc(id)}')">↗ Scratchpad</button>
    </div>
  </div>`).join('') + '</div>' + `
<div class="judge-wrap" id="judgeWrap" style="display:none">
  <div class="judge-head"><span class="judge-title">⚖ Judge</span>
    <span class="stage-info" id="judgeTok"></span></div>
  <div class="judge-body" id="judgeBody"></div>
</div>`;
  // Side-by-side to three; wrap beyond that so panes stay readable.
  const split = document.getElementById('split');
  if (split) split.style.gridTemplateColumns =
    'repeat(' + Math.min(ids.length, 3) + ', minmax(0, 1fr))';
}

function poll(jobId, ids) {
  const timer = setInterval(async () => {
    let p;
    try { p = await (await fetch('/compare/progress/' + jobId)).json(); }
    catch (e) { return; }

    for (const id of ids) {
      const pane = p[id]; if (!pane) continue;
      if (pane.done && answers[id] !== pane.answer) {
        answers[id] = pane.answer;
        document.getElementById('body-' + id).innerHTML = md(pane.answer);
        document.getElementById('foot-' + id).style.display = 'flex';
        setStage('stage-' + id, '<span class="done-tick">✓ done</span>');
      } else if (!pane.done) {
        setStage('stage-' + id,
          `<div class="spinner"></div>&nbsp;drafting… ${pane.tokens} tok`);
      }
    }

    const j = p.judge || {};
    if (j.tokens > 0 || j.done) {
      document.getElementById('judgeWrap').style.display = 'block';
      document.getElementById('judgeTok').textContent = j.tokens + ' tok' + (j.done ? '' : '…');
      document.getElementById('judgeBody').innerHTML = (j.done && j.answer)
        ? md(j.answer)
        : `<div class="pending"><div class="spinner"></div> deliberating… ${j.tokens} tokens</div>`;
    }

    if (p.done) { clearInterval(timer); go.disabled = false; q.focus(); }
  }, 400);
}

async function run() {
  const question = q.value.trim(); if (!question) return;
  const ids = currentPick();
  if (ids.length < 2) {
    page.innerHTML = `<div class="idle">Pick at least two ${pickMode === 'routes' ? 'routes' : 'models'} to compare.</div>`;
    return;
  }
  go.disabled = true;
  answers = {};
  buildUI(ids);
  let res;
  try {
    res = await (await fetch('/compare/ask', {
      method: 'POST', headers: { 'Content-Type': 'application/json' },
      body: JSON.stringify({ question, mode: pickMode, targets: ids }),
    })).json();
  } catch (e) {
    page.innerHTML = `<div class="idle">⚠️ Router unreachable: ${esc(String(e))}</div>`;
    go.disabled = false; return;
  }
  if (res.error) {
    page.innerHTML = `<div class="idle">⚠️ ${esc(res.error)}</div>`;
    go.disabled = false; return;
  }
  poll(res.job_id, res.targets || ids);
}

// ---- target picker ---------------------------------------------------------
// Two modes because they answer different questions. "models" gives every
// entrant the SAME system prompt, so what differs is the model. "routes" gives
// each its own persona, so what differs is your configured setup. Names only —
// backend URLs and keys are admin-only and never reach this page.
let pickMode = localStorage.getItem('creweCompareMode') || 'backends';
function currentPick(){
  return [...document.querySelectorAll('#pickRow .pick input:checked')]
           .map(i => i.dataset.id);
}
function saveKey(){ return 'creweComparePick:' + pickMode; }
async function loadTargets(){
  const row = document.getElementById('pickRow');
  document.querySelectorAll('.modebtn').forEach(b =>
    b.classList.toggle('on', b.dataset.mode === pickMode));
  let list = [];
  try { list = (await (await fetch('/compare/targets?mode=' + pickMode)).json()).targets || []; }
  catch (e) { row.innerHTML = '<span class="pickhint">could not load</span>'; return; }
  if (!list.length) {
    row.innerHTML = '<span class="pickhint">nothing configured — see Settings</span>';
    return;
  }
  const savedRaw = localStorage.getItem(saveKey());
  const saved = savedRaw === null ? null : savedRaw.split(',').filter(Boolean);
  row.innerHTML = list.map(t => {
    const on = saved === null ? !!t.default : saved.includes(t.id);
    return `<label class="pick${on ? ' on' : ''}">
      <input type="checkbox" data-id="${esc(t.id)}"${on ? ' checked' : ''}>
      <span>${esc(t.label)}</span>
      ${t.sub ? `<span class="mdl">${esc(t.sub)}</span>` : ''}
      ${t.paid ? '<span class="cost">&#128181;</span>' : ''}
    </label>`;
  }).join('');
  row.querySelectorAll('.pick input').forEach(inp =>
    inp.addEventListener('change', () => {
      inp.closest('.pick').classList.toggle('on', inp.checked);
      localStorage.setItem(saveKey(), currentPick().join(','));
    }));
}
document.querySelectorAll('.modebtn').forEach(b =>
  b.addEventListener('click', () => {
    pickMode = b.dataset.mode;
    localStorage.setItem('creweCompareMode', pickMode);
    loadTargets();
  }));
loadTargets();

q.addEventListener('input', () => { q.style.height='auto'; q.style.height=Math.min(q.scrollHeight,120)+'px'; });
go.addEventListener('click', run);
q.addEventListener('keydown', e => { if (e.key==='Enter'&&!e.shiftKey){ e.preventDefault(); run(); } });
q.focus();
</script>
</body>
</html>"""


# These run at startup, where there is no request and therefore no owner. They
# loop over every account and set the owner explicitly per user, so each user's
# backfill reads and writes only their own cookbook.
def _start_photo_backfill():
    """One-time per process, per user: fetch photos for cards that never got one."""
    if _BACKFILL_STARTED.is_set():
        return
    _BACKFILL_STARTED.set()
    try:
        for u in _load_users():
            set_owner(u["id"])
            with COOKBOOK_LOCK:
                missing = [dict(c) for c in cookbook()
                           if c.get("id") and "photo" not in c]
            if missing:
                print(f"[cookbook] backfilling photos for {len(missing)} card(s) "
                      f"— {u.get('email')}")
                spawn_owned(_backfill_photos, (missing,)).start()
    finally:
        set_owner(None)      # don't leave the main thread owned by the last user


def _start_meta_backfill():
    if _META_BACKFILL_STARTED.is_set():
        return
    _META_BACKFILL_STARTED.set()
    try:
        for u in _load_users():
            set_owner(u["id"])
            with COOKBOOK_LOCK:
                missing = [dict(c) for c in cookbook()
                           if c.get("id") and "meal" not in c]
            if missing:
                print(f"[cookbook] classifying meal/protein for {len(missing)} "
                      f"card(s) — {u.get('email')}")
                spawn_owned(_backfill_meta, (missing,)).start()
    finally:
        set_owner(None)


if __name__ == "__main__":
    _start_photo_backfill()
    _start_meta_backfill()
    # Bind host: default 0.0.0.0 keeps LAN/tailnet direct access working. For a
    # pure public-via-tunnel deploy, set CREWE_BIND=127.0.0.1 so ONLY cloudflared
    # can reach Flask (closes any header-spoofing path and hides :5000 from LAN).
    bind = os.environ.get("CREWE_BIND", "0.0.0.0")
    print(f"Router on http://{bind}:5000  (specialists 8080-8082, classifier 8083)")
    print("Scratchpad at http://localhost:5000/scratch")
    try:
        # Production WSGI server for public exposure. app.run (Werkzeug dev
        # server) is fine locally but not hardened for the internet.
        from waitress import serve
        print("Serving with waitress (production).")
        serve(app, host=bind, port=5000, threads=16)
    except ImportError:
        print("⚠ waitress not installed — falling back to the Flask dev server. "
              "For public hosting: pip install --user waitress --break-system-packages")
        app.run(host=bind, port=5000, threaded=True)
